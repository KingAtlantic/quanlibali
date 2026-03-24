import streamlit as st
import pandas as pd
import pdfplumber
import re
from datetime import datetime
import time
import base64
import hashlib
import sqlite3
import os
import io
import sys
import subprocess
import random
import string
from PIL import Image, ImageEnhance
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
from reportlab.lib.colors import HexColor
from typing import Any, List, Optional, Union, Literal, overload, Dict
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from lunardate import LunarDate

# --- QUAN TRỌNG: CẤU HÌNH TRANG PHẢI Ở ĐẦU TIÊN ---
st.set_page_config(
    page_title="Quản Lý Hóa Đơn Pro", 
    page_icon="🌸", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- IMPORTS KHÁC (Không dùng auto_install nữa) ---
import gspread
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import cv2
import numpy as np

# --- OCR CONFIGURATION ---
try:
    import pytesseract
    # Cấu hình đường dẫn Tesseract
    if sys.platform.startswith('win'):
        # Đường dẫn cho Windows (Local)
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    else:
        # Chạy trên Streamlit Cloud (Linux) - Không cần set path, nó tự tìm
        # Nếu cần thiết: pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
        pass
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    pytesseract = None

# --- EXCEL & DOCX LIBS ---
import openpyxl
import xlsxwriter
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
HAS_OPENPYXL = True
HAS_XLSXWRITER = True
HAS_CV = True
HAS_DOCX = True

# --- CẤU HÌNH GOOGLE (Đã cập nhật theo thông tin của bạn) ---
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

# Tên file chìa khóa (Hãy đổi tên file bạn tải về thành tên này)
SERVICE_ACCOUNT_FILE = 'service_account.json'

# ID Google Drive (Nơi lưu ảnh/pdf)
# Link: https://drive.google.com/drive/folders/1PMCKIUirYwbacu0evnRyuF0xSq-bQtBv?usp=drive_link
DRIVE_FOLDER_ID = '1PMCKIUirYwbacu0evnRyuF0xSq-bQtBv'

# ID Google Sheet (Lấy từ link bạn gửi)
# Link: https://docs.google.com/spreadsheets/d/1coeIPogjKEJSKv1hW1dFBrSAwF6V7c-tkVCZPuPQjoc/edit?gid=0#gid=0
SPREADSHEET_ID = '1coeIPogjKEJSKv1hW1dFBrSAwF6V7c-tkVCZPuPQjoc'

def get_gspread_client():
    # Kiểm tra xem đang chạy trên Cloud (dùng secrets) hay Local (dùng file json)
    if "gcp_service_account" in st.secrets:
        creds_dict = dict(st.secrets["gcp_service_account"])
        creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    else:
        creds = Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    
    client = gspread.authorize(creds)
    return client

def get_drive_service():
    if "gcp_service_account" in st.secrets:
        creds_dict = dict(st.secrets["gcp_service_account"])
        creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    else:
        creds = Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    
    service = build('drive', 'v3', credentials=creds)
    return service

# --- CÁC HÀM XỬ LÝ DỮ LIỆU MỚI (Thay thế SQL) ---

def load_table(table_name):
    """Đọc dữ liệu từ Local SQLite (Thay thế Google Sheet)"""
    conn = get_connection()
    try:
        df = pd.read_sql_query(f"SELECT * FROM {table_name}", conn)
        return df
    except Exception as e:
        print(f"Lỗi đọc bảng {table_name}: {e}")
        return pd.DataFrame()

def add_row_to_table(table_name, row_dict):
    """Thêm dòng mới vào Local SQLite VÀ Google Sheet"""
    # 1. Ghi vào SQLite (Local)
    conn = get_connection()
    c = conn.cursor()
    success = False
    try:
        columns = ', '.join(row_dict.keys())
        placeholders = ', '.join(['?'] * len(row_dict))
        sql = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"
        c.execute(sql, list(row_dict.values()))
        conn.commit()
        success = True
    except Exception as e:
        st.error(f"Lỗi ghi dữ liệu vào {table_name}: {e}")
        return False

    # 2. Ghi vào Google Sheet (Cloud)
    if success:
        try:
            gc = get_gspread_client()
            sh = gc.open_by_key(SPREADSHEET_ID)
            try:
                wks = sh.worksheet(table_name)
            except:
                wks = sh.add_worksheet(title=table_name, rows=100, cols=20)
            
            # Xử lý header và map dữ liệu
            existing = wks.get_all_values()
            if not existing:
                headers = list(row_dict.keys())
                wks.append_row(headers)
            else:
                headers = existing[0]
            
            row_values = []
            for h in headers:
                val = row_dict.get(h, "")
                if val is None: val = ""
                row_values.append(val)
                
            wks.append_row(row_values)
        except Exception as e:
            # [DEBUG] Thay đổi để hiển thị lỗi chi tiết hơn
            st.error(f"⚠️ LỖI ĐỒNG BỘ GOOGLE SHEET (Đã lưu vào máy nhưng không đẩy lên cloud được)")
            st.exception(e)
            
    return success

def upload_to_drive(file_obj, file_name, mimetype=None):
    """Upload file lên Google Drive"""
    try:
        service = get_drive_service()
        file_metadata = {'name': file_name, 'parents': [DRIVE_FOLDER_ID]}
        
        if not mimetype and hasattr(file_obj, 'type'):
            mimetype = file_obj.type
            
        media = MediaIoBaseUpload(file_obj, mimetype=mimetype or 'application/octet-stream', resumable=True)
        file = service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()
        return file.get('webViewLink')
    except Exception as e:
        st.warning(f"⚠️ Lỗi upload Drive: {e}")
        return None

def sync_all_data_to_gsheet():
    """Đọc tất cả dữ liệu từ SQLite và ghi đè lên Google Sheet."""
    TABLES_TO_SYNC = [
        'users', 'invoices', 'projects', 'project_links', 'company_info', 
        'flight_tickets', 'flight_groups', 'flight_group_links', 
        'service_bookings', 'customers', 'tours', 'tour_items', 'ocr_learning',
        'transaction_history'
    ]

    try:
        gc = get_gspread_client()
        sh = gc.open_by_key(SPREADSHEET_ID)
        conn = get_connection()

        st.info(f"Bắt đầu đồng bộ {len(TABLES_TO_SYNC)} bảng...")
        status_placeholder = st.empty()
        progress_bar = st.progress(0)
        
        for i, table_name in enumerate(TABLES_TO_SYNC):
            status_placeholder.info(f"Đang xử lý bảng: **{table_name}**...")
            
            try:
                df = pd.read_sql_query(f"SELECT * FROM {table_name}", conn)
            except Exception:
                st.warning(f"Bảng '{table_name}' không có trong DB, bỏ qua.")
                progress_bar.progress((i + 1) / len(TABLES_TO_SYNC))
                continue

            try:
                wks = sh.worksheet(table_name)
                wks.clear()
            except gspread.WorksheetNotFound:
                wks = sh.add_worksheet(title=table_name, rows=1, cols=20)

            if not df.empty:
                df = df.astype(str).replace({'nan': '', 'NaT': ''})
                # [FIX] Truncate cells that are too long for Google Sheets API to prevent 400 error
                df = df.map(lambda x: x[:49999] if isinstance(x, str) and len(x) >= 50000 else x)
                data_to_upload = [df.columns.tolist()] + df.values.tolist()
                wks.update(data_to_upload, 'A1')
                st.toast(f"✅ Đồng bộ '{table_name}' ({len(df)} dòng) OK.")
            else:
                st.toast(f"ℹ️ Bảng '{table_name}' rỗng, đã dọn dẹp trên cloud.")

            progress_bar.progress((i + 1) / len(TABLES_TO_SYNC))

        status_placeholder.empty()
        st.success("🎉 Đồng bộ toàn bộ dữ liệu hoàn tất!")
    except Exception as e:
        st.error("❌ Lỗi nghiêm trọng khi đồng bộ:")
        st.exception(e)
        st.info("💡 Gợi ý: Hãy chắc chắn rằng email của tài khoản dịch vụ (`client_email` trong file .json) đã được cấp quyền 'Editor' (Người chỉnh sửa) cho file Google Sheet này.")

# --- QUẢN LÝ SESSION STATE ---
if "logged_in" not in st.session_state: st.session_state.logged_in = False
if "user_info" not in st.session_state: st.session_state.user_info = None
if "db_initialized" not in st.session_state: st.session_state.db_initialized = False

# Biến lưu trữ
if "ready_pdf_bytes" not in st.session_state: st.session_state.ready_pdf_bytes = None
if "ready_file_name" not in st.session_state: st.session_state.ready_file_name = None
if "uploader_key" not in st.session_state: st.session_state.uploader_key = 0
if "pdf_data" not in st.session_state: st.session_state.pdf_data = None
if "edit_lock" not in st.session_state: st.session_state.edit_lock = True
if "local_edit_count" not in st.session_state: st.session_state.local_edit_count = 0
if "current_doc_type" not in st.session_state: st.session_state.current_doc_type = "Hóa đơn"
if "invoice_view_page" not in st.session_state: st.session_state.invoice_view_page = 0

# Biến riêng cho Edit Mode
if "unc_edit_mode" not in st.session_state: st.session_state.unc_edit_mode = False
if "est_edit_mode" not in st.session_state: st.session_state.est_edit_mode = False
if "current_tour_id_est" not in st.session_state: st.session_state.current_tour_id_est = None
if "est_editor_key" not in st.session_state: st.session_state.est_editor_key = 0

# Initialize tab variables to avoid Pylance undefined errors
tab_est = tab_act = tab_rpt = None

# FIX LỖI OUT TÀI KHOẢN
UPLOAD_FOLDER = ".uploaded_invoices"
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

DB_FILE = "invoice_app.db"

# ==========================================
# 2. XỬ LÝ DATABASE (SQLite)
# ==========================================
@st.cache_resource
def get_connection():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def migrate_db_columns():
    conn = get_connection()
    c = conn.cursor()
    # Thêm các cột nếu chưa có cho Hóa đơn/Dự án cũ
    try: c.execute("ALTER TABLE invoices ADD COLUMN request_edit INTEGER DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE flight_tickets ADD COLUMN airline TEXT")
    except: pass
    try: c.execute("ALTER TABLE projects ADD COLUMN pending_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE projects ADD COLUMN type TEXT DEFAULT 'NORMAL'")
    except: pass
    try: c.execute("ALTER TABLE tour_items ADD COLUMN category TEXT")
    except: pass
    try: c.execute("ALTER TABLE tour_items ADD COLUMN times REAL DEFAULT 1")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN pending_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN request_delete INTEGER DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN request_edit_act INTEGER DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN tour_code TEXT")
    except: pass
    try: c.execute("ALTER TABLE invoices ADD COLUMN cost_code TEXT")
    except: pass
    
    # --- Cập nhật cho Bàn Giao Tour (Mới) ---
    try: c.execute("ALTER TABLE tours ADD COLUMN pickup_location TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN pickup_time TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN flight_code TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN driver_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN driver_phone TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN car_plate TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN car_type TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN itinerary_summary TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN guide_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN guide_phone TEXT")
    except: pass
    try: c.execute("CREATE TABLE IF NOT EXISTS ocr_learning (keyword TEXT UNIQUE, weight INTEGER DEFAULT 1)")
    except: pass

    # --- Bảng Điểm Tham Quan (Mới) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS tour_sightseeings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        name TEXT,
        address TEXT,
        quantity INTEGER,
        note TEXT
    )''')
    except: pass
    
    # --- Cập nhật cột Tài chính cho KS/NH (Mới) ---
    try: c.execute("ALTER TABLE tour_hotels ADD COLUMN total_amount REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tour_hotels ADD COLUMN deposit REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tour_restaurants ADD COLUMN total_amount REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tour_restaurants ADD COLUMN deposit REAL DEFAULT 0")
    except: pass
    # --- Cập nhật cột Ngày và Tài chính cho Nhà hàng/Tham quan (Mới nhất) ---
    try: c.execute("ALTER TABLE tour_restaurants ADD COLUMN date TEXT")
    except: pass
    try: c.execute("ALTER TABLE tour_sightseeings ADD COLUMN date TEXT")
    except: pass
    try: c.execute("ALTER TABLE tour_sightseeings ADD COLUMN total_amount REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tour_sightseeings ADD COLUMN deposit REAL DEFAULT 0")
    except: pass

    # --- Cập nhật cho Payment Reminders (Mới) ---
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN cc_email TEXT")
    except: pass
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN sender_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN bank_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN bank_account TEXT")
    except: pass
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN bank_holder TEXT")
    except: pass

    # --- Bảng Lịch Trình Tour (Mới) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS tour_itineraries (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        day_index INTEGER,
        content TEXT
    )''')
    except: pass

    # --- Bảng Chi Phí Phát Sinh (Mới) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS tour_incurred_costs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        name TEXT,
        unit TEXT,
        quantity REAL,
        price REAL,
        total_amount REAL,
        deposit REAL DEFAULT 0,
        note TEXT
    )''')
    except: pass

    # --- Bảng Booking Dịch Vụ (Mới) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS service_bookings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        code TEXT UNIQUE,
        name TEXT,
        created_at TEXT,
        status TEXT DEFAULT 'active'
    )''')
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN type TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN details TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN customer_info TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN net_price REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN tax_percent REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN selling_price REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN profit REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN sale_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN hotel_code TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN room_type TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN guest_list TEXT")
    except: pass

    # --- Bảng Khách Hàng (Mới) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS customers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        phone TEXT,
        email TEXT,
        address TEXT,
        notes TEXT,
        created_at TEXT
    )''')
    except: pass
    try: c.execute("ALTER TABLE customers ADD COLUMN sale_name TEXT")
    except: pass

    # --- Cập nhật cột mới cho Tour (Giá chốt, Giá trẻ em, Giá trị hợp đồng) ---
    try: c.execute("ALTER TABLE tours ADD COLUMN final_tour_price REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN child_price REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN contract_value REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN final_qty REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN child_qty REAL DEFAULT 0")
    except: pass

    # --- Cập nhật thông tin khách hàng cho Tour ---
    try: c.execute("ALTER TABLE tours ADD COLUMN customer_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN customer_phone TEXT")
    except: pass

    # --- Cập nhật cho phần Danh sách & Dịch vụ (Mới) ---
    try: c.execute("ALTER TABLE tours ADD COLUMN handover_checklist TEXT")
    except: pass

    c.execute('''CREATE TABLE IF NOT EXISTS tour_guests (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        name TEXT,
        dob TEXT,
        hometown TEXT,
        cccd TEXT,
        type TEXT
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS tour_hotels (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        hotel_name TEXT,
        address TEXT,
        phone TEXT,
        total_rooms TEXT,
        room_type TEXT
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS tour_restaurants (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        meal_name TEXT,
        restaurant_name TEXT,
        address TEXT,
        phone TEXT,
        menu TEXT
    )''')

    # --- Cập nhật mã tour cho dữ liệu cũ ---
    try:
        old_tours = c.execute("SELECT id FROM tours WHERE tour_code IS NULL OR tour_code = ''").fetchall()
        for t in old_tours:
            code = ''.join(random.choices(string.ascii_uppercase, k=5))
            c.execute("UPDATE tours SET tour_code=? WHERE id=?", (code, t['id'])) # type: ignore
    except: pass
    
    # --- Cập nhật dữ liệu cũ để hiện thị dự án ---
    try: 
        c.execute("UPDATE projects SET type='NORMAL' WHERE type IS NULL OR type=''")
    except: pass

    # --- FIX QUAN TRỌNG: ĐẢM BẢO BẢNG TOURS TỒN TẠI KHI CẬP NHẬT ---
    # Phần này giúp tạo bảng ngay cả khi DB đã tồn tại từ trước
    c.execute('''CREATE TABLE IF NOT EXISTS tours (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_name TEXT,
        sale_name TEXT,
        start_date TEXT,
        end_date TEXT,
        guest_count INTEGER,
        created_at TEXT,
        est_profit_percent REAL DEFAULT 10.0,
        est_tax_percent REAL DEFAULT 8.0,
        status TEXT DEFAULT 'running'
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS tour_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        item_type TEXT, 
        category TEXT,
        description TEXT,
        unit TEXT,
        quantity REAL,
        times REAL DEFAULT 1,
        unit_price REAL,
        total_amount REAL
    )''')
    
    # --- Bảng Công Nợ (Mới) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS transaction_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ref_code TEXT,
        type TEXT,
        amount REAL,
        payment_method TEXT,
        note TEXT,
        created_at TEXT
    )''')
    except: pass

    conn.commit()

def init_db():
    conn = get_connection()
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT UNIQUE, password TEXT, role TEXT, status TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS invoices (
        id INTEGER PRIMARY KEY AUTOINCREMENT, type TEXT, date TEXT, invoice_number TEXT, invoice_symbol TEXT, 
        seller_name TEXT, buyer_name TEXT, pre_tax_amount REAL, tax_amount REAL, total_amount REAL, 
        file_name TEXT, status TEXT, edit_count INTEGER, created_at TEXT, memo TEXT, file_path TEXT, request_edit INTEGER DEFAULT 0
    )''')
    # Thêm cột pending_name và type vào bảng projects
    c.execute('''CREATE TABLE IF NOT EXISTS projects (
        id INTEGER PRIMARY KEY AUTOINCREMENT, 
        project_name TEXT, 
        created_at TEXT,
        pending_name TEXT,
        type TEXT DEFAULT 'NORMAL'
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS project_links (id INTEGER PRIMARY KEY AUTOINCREMENT, project_id INTEGER, invoice_id INTEGER)''')
    c.execute('''CREATE TABLE IF NOT EXISTS company_info (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, address TEXT, phone TEXT, logo_base64 TEXT)''')
    
    # Bảng Vé máy bay
    c.execute('''CREATE TABLE IF NOT EXISTS flight_tickets (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ticket_code TEXT,
        flight_date TEXT,
        route TEXT,
        passenger_names TEXT,
        file_path TEXT,
        created_at TEXT,
        airline TEXT
    )''')
    
    # Bảng Đoàn bay (Cũ - Giữ nguyên để tương thích)
    c.execute('''CREATE TABLE IF NOT EXISTS flight_groups (id INTEGER PRIMARY KEY AUTOINCREMENT, group_name TEXT, created_at TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS flight_group_links (id INTEGER PRIMARY KEY AUTOINCREMENT, group_id INTEGER, ticket_id INTEGER)''')

    # --- BẢNG BOOKING DỊCH VỤ ---
    c.execute('''CREATE TABLE IF NOT EXISTS service_bookings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        code TEXT UNIQUE,
        name TEXT,
        created_at TEXT,
        status TEXT DEFAULT 'active'
    )''')

    # --- BẢNG QUẢN LÝ TOUR  ---
    c.execute('''CREATE TABLE IF NOT EXISTS tours (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_name TEXT,
        sale_name TEXT,
        start_date TEXT,
        end_date TEXT,
        guest_count INTEGER,
        created_at TEXT,
        est_profit_percent REAL DEFAULT 10.0,
        est_tax_percent REAL DEFAULT 8.0,
        status TEXT DEFAULT 'running'
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS tour_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        item_type TEXT, 
        category TEXT,
        description TEXT,
        unit TEXT,
        quantity REAL,
        times REAL DEFAULT 1,
        unit_price REAL,
        total_amount REAL
    )''')
    # item_type: 'EST' (Dự toán), 'ACT' (Quyết toán)

    # --- Bảng Công Nợ (Mới) ---
    c.execute('''CREATE TABLE IF NOT EXISTS transaction_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ref_code TEXT,
        type TEXT,
        amount REAL,
        payment_method TEXT,
        note TEXT,
        created_at TEXT
    )''')

    # --- THÊM VÀO TRONG HÀM init_db() ---
    c.execute('''CREATE TABLE IF NOT EXISTS payment_reminders (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ref_code TEXT,       -- Mã Booking/Tour
        ref_name TEXT,       -- Tên khách/Tour
        amount REAL,         -- Số tiền cần thu
        due_date TEXT,       -- Ngày hẹn thông báo lại
        receiver_email TEXT, -- Email người nhận (Nội bộ hoặc Khách)
        content TEXT,        -- Nội dung nhắc
        status TEXT,         -- 'sent_1': Đã gửi lần 1, 'sent_2': Đã gửi lần 2 (hoàn tất)
        created_at TEXT
    )''')

    c.execute("SELECT * FROM users WHERE username = 'admin'")
    if not c.fetchone():
        admin_pw = hashlib.sha256("admin123".encode()).hexdigest()
        c.execute("INSERT INTO users (username, password, role, status) VALUES (?, ?, ?, ?)", ('admin', admin_pw, 'admin', 'approved'))
    
    c.execute("SELECT * FROM company_info WHERE id = 1")
    if not c.fetchone():
        c.execute("INSERT INTO company_info (name, address, phone, logo_base64) VALUES (?, ?, ?, ?)", ('Tên Công Ty Của Bạn', 'Địa chỉ...', '090...', ''))

    conn.commit()

if not st.session_state.db_initialized:
    init_db()
    st.session_state.db_initialized = True

# Luôn chạy migration để đảm bảo cột mới được thêm vào (Fix lỗi Admin không nhận yêu cầu)
migrate_db_columns()

# --- CÁC HÀM HỖ TRỢ ---
@overload
def run_query(query: str, params: Any = ..., fetch_one: Literal[False] = ..., commit: Literal[False] = ...) -> List[Any]: ...

@overload
def run_query(query: str, params: Any, fetch_one: Literal[True], commit: Literal[False] = ...) -> Any: ...

@overload
def run_query(query: str, *, fetch_one: Literal[True], commit: Literal[False] = ...) -> Any: ...

@overload
def run_query(query: str, params: Any = ..., fetch_one: Any = ..., *, commit: Literal[True]) -> bool: ...

def run_query(query, params=(), fetch_one=False, commit=False):
    conn = get_connection()
    c = conn.cursor()
    try:
        c.execute(query, params)
        if commit:
            conn.commit()
            return True
        if fetch_one:
            return c.fetchone()
        return c.fetchall()
    except Exception as e:
        print(f"Lỗi truy vấn DB: {e}")
        if commit: return False
        if fetch_one: return None
        return []

def run_query_many(query, data):
    """Thực thi nhiều câu lệnh (thường là INSERT) cùng lúc."""
    conn = get_connection()
    c = conn.cursor()
    try:
        c.executemany(query, data)
        conn.commit()
        return True
    except Exception as e:
        print(f"Lỗi truy vấn DB (many): {e}")
        return False

def save_customer_check(name, phone, sale_name=None):
    """Lưu khách hàng mới nếu chưa tồn tại"""
    if not name: return
    try:
        exist = run_query("SELECT id FROM customers WHERE name=?", (name,), fetch_one=True)
        if not exist:
            data = {'name': name, 'phone': phone, 'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
            if sale_name:
                data['sale_name'] = sale_name
            add_row_to_table('customers', data)
    except: pass

def hash_pass(password):
    return hashlib.sha256(str.encode(password)).hexdigest()

def save_file_local(file_bytes, original_name):
    try:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_name = re.sub(r'[\\/*?:"<>|]', "", original_name)
        if not clean_name.lower().endswith('.pdf'):
            clean_name = os.path.splitext(clean_name)[0] + ".pdf"
            
        final_name = f"{ts}_{clean_name}"
        file_path = os.path.join(UPLOAD_FOLDER, final_name)
        
        with open(file_path, "wb") as f:
            f.write(file_bytes)
                
        return file_path, final_name
    except: return None, None

def format_vnd(amount):
    if amount is None: return "0"
    try: return "{:,.0f}".format(float(amount)).replace(",", ".")
    except: return "0"

@st.cache_data
def get_company_data():
    row = run_query("SELECT * FROM company_info WHERE id = 1", fetch_one=True)
    if isinstance(row, sqlite3.Row):
        return {'name': row['name'], 'address': row['address'], 'phone': row['phone'], 'logo_b64_str': row['logo_base64']}
    return {'name': 'Company', 'address': '...', 'phone': '...', 'logo_b64_str': ''}

def update_company_info(name, address, phone, logo_bytes=None):
    b64_str = base64.b64encode(logo_bytes).decode('utf-8') if logo_bytes else ""
    if not logo_bytes:
        old = run_query("SELECT logo_base64 FROM company_info WHERE id = 1", fetch_one=True)
        if isinstance(old, sqlite3.Row): b64_str = old['logo_base64'] # type: ignore
    run_query("UPDATE company_info SET name=?, address=?, phone=?, logo_base64=? WHERE id=1", (name, address, phone, b64_str), commit=True)
    get_company_data.clear()# type: ignore

# --- HÀM GỬI EMAIL ---
def send_email_notification(to_email, subject, body_html, cc_emails=None):
    """Hàm gửi email qua SMTP Gmail"""
    try:
        # Lấy cấu hình từ secrets.toml
        email_sender = st.secrets["email"]["sender"]
        email_password = st.secrets["email"]["password"]
        
        msg = MIMEMultipart()
        msg['From'] = f"Bali Tourist System <{email_sender}>"
        msg['To'] = to_email
        msg['Subject'] = subject
        if cc_emails:
            msg['Cc'] = cc_emails
        msg.attach(MIMEText(body_html, 'html'))

        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(email_sender, email_password)
        
        # Xử lý danh sách người nhận (To + Cc)
        recipients = [to_email]
        if cc_emails:
            if isinstance(cc_emails, str):
                recipients.extend([e.strip() for e in cc_emails.split(',') if e.strip()])
            elif isinstance(cc_emails, list):
                recipients.extend(cc_emails)
        
        server.send_message(msg, to_addrs=recipients)
        server.quit()
        return True, "Đã gửi mail thành công!"
    except Exception as e:
        return False, f"Lỗi gửi mail: {str(e)}"

# --- HÀM TỰ ĐỘNG QUÉT & GỬI LẦN 2 ---
def check_and_send_due_reminders():
    """Kiểm tra các lịch hẹn đến ngày hôm nay để gửi email lần 2"""
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    # Tìm các nhắc hẹn có ngày = hôm nay (hoặc quá khứ) VÀ mới chỉ gửi lần 1 ('sent_1')
    reminders = run_query("SELECT * FROM payment_reminders WHERE status='sent_1' AND due_date <= ?", (now_str,))
    
    count = 0
    if reminders:
        for row in reminders:
            r = dict(row)
            # Gửi email lần 2
            cc = r.get('cc_email', '')
            sender = r.get('sender_name', 'Bali Tourist System')
            
            # [FIX] Format ngày hiển thị trong mail (DD/MM/YYYY HH:MM)
            try:
                d_obj = datetime.strptime(r['due_date'], '%Y-%m-%d %H:%M:%S')
                date_display = d_obj.strftime('%H:%M %d/%m/%Y')
            except:
                try:
                    d_obj = datetime.strptime(r['due_date'], '%Y-%m-%d')
                    date_display = d_obj.strftime('%d/%m/%Y')
                except:
                    date_display = r['due_date']

            # [NEW] Bank Info for Automated Email
            bank_info_html = ""
            if r.get('bank_name') and r.get('bank_account'):
                bank_info_html = f"""
                <div style="background-color: #f8f9fa; padding: 15px; border-radius: 5px; margin: 15px 0;">
                    <h4 style="margin-top: 0;">🏦 THÔNG TIN CHUYỂN KHOẢN</h4>
                    <p><strong>Ngân hàng:</strong> {r['bank_name']}</p>
                    <p><strong>Số tài khoản:</strong> {r['bank_account']}</p>
                    <p><strong>Chủ tài khoản:</strong> {r.get('bank_holder', '')}</p>
                </div>
                """

            subject = f"🔔 [NHẮC HẸN LẦN 2] Thanh toán cho mã {r['ref_code']}"
            content = f"""
            <h3>🔔 NHẮC HẸN THANH TOÁN (LẦN 2)</h3>
            <p>Hệ thống tự động nhắc bạn về khoản thanh toán đã đến hẹn:</p>
            <ul>
                <li><strong>Mã hồ sơ:</strong> {r['ref_code']}</li>
                <li><strong>Tên:</strong> {r['ref_name']}</li>
                <li><strong>Số tiền:</strong> {format_vnd(r['amount'])} VND</li>
                <li><strong>Nội dung:</strong> {r['content']}</li>
                <li><strong>Ngày hẹn:</strong> {date_display}</li>
            </ul>
            {bank_info_html}
            <p>Vui lòng kiểm tra và xử lý.</p>
            <p>Trân trọng,<br>{sender}</p>
            """
            success, msg = send_email_notification(r['receiver_email'], subject, content, cc_emails=cc)
            if success:
                # Cập nhật trạng thái thành sent_2 (đã xong)
                run_query("UPDATE payment_reminders SET status='sent_2' WHERE id=?", (r['id'],), commit=True)
                count += 1
    return count

# --- HÀM HỖ TRỢ LỊCH ÂM/DƯƠNG ---
def convert_solar_to_lunar(solar_date):
    """Chuyển Dương lịch -> Âm lịch"""
    try:
        ld = LunarDate.fromSolarDate(solar_date.year, solar_date.month, solar_date.day)
        return f"{ld.day:02d}/{ld.month:02d}/{ld.year} (Âm lịch)"
    except:
        return "Không xác định"

def convert_lunar_to_solar(day, month, year, is_leap=False):
    """Chuyển Âm lịch -> Dương lịch"""
    try:
        sd = LunarDate(year, month, day, is_leap).toSolarDate()
        return sd # Trả về object date
    except ValueError:
        return None

def get_tour_financials(tour_id, tour_info):
    """
    Tính toán doanh thu và chi phí cho một tour.
    """
    # Lấy tổng chi phí quyết toán (ACT) từ bảng kê
    act_items = run_query("SELECT SUM(total_amount) as total FROM tour_items WHERE tour_id=? AND item_type='ACT'", (tour_id,), fetch_one=True)
    act_cost_items = act_items['total'] if act_items and act_items['total'] else 0

    # Lấy tổng chi phí từ hóa đơn đầu vào liên kết với tour (không tính UNC)
    inv_items = run_query("SELECT SUM(total_amount) as total FROM invoices WHERE cost_code=? AND status='active' AND type='IN' AND invoice_number NOT LIKE '%UNC%'", (tour_info['tour_code'],), fetch_one=True)
    inv_cost = inv_items['total'] if inv_items and inv_items['total'] else 0

    cost = (act_cost_items or 0) + (inv_cost or 0)

    # Lấy tổng chi phí dự toán (EST) để tính doanh thu nếu cần
    est_items = run_query("SELECT SUM(total_amount) as total FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id,), fetch_one=True)
    est_cost = est_items['total'] if est_items and est_items['total'] else 0

    # Tính doanh thu dựa trên giá chốt
    t_dict = dict(tour_info)
    final_price = float(t_dict.get('final_tour_price', 0) or 0)
    child_price = float(t_dict.get('child_price', 0) or 0)
    final_qty = float(t_dict.get('final_qty', 0) or 0)
    child_qty = float(t_dict.get('child_qty', 0) or 0)
    if final_qty == 0: final_qty = float(t_dict.get('guest_count', 1))
    
    revenue = (final_price * final_qty) + (child_price * child_qty)

    # Nếu chi phí quyết toán chưa có, dùng tạm chi phí dự toán
    if cost == 0 and est_cost > 0:
        cost = est_cost

    return revenue, cost
# ==========================================
# 3. CSS & GIAO DIỆN HIỆN ĐẠI
# ==========================================
comp = get_company_data()
st.markdown("""<style>
/* --- BASE & ANIMATION --- */
@keyframes fadeIn { 0% { opacity: 0; transform: translateY(10px); } 100% { opacity: 1; transform: translateY(0); } }
.stApp {
    background-color: #f8f9fa;
    font-family: 'Inter', 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
    animation: fadeIn 0.5s ease-in-out;
}

/* --- TYPOGRAPHY & LABELS --- */
h1, h2, h3, h4, h5, h6 { color: #2c3e50; }
div[data-testid="stMarkdownContainer"] p { font-weight: 400; white-space: normal; word-break: break-word; }
.company-info-text p, .report-card p { white-space: normal !important; }

/* --- MODERN INPUTS --- */
.stTextInput input, .stNumberInput input, .stSelectbox div[data-baseweb="select"], .stTextArea textarea, .stDateInput input {
    border-radius: 10px !important;
    border: 1px solid #e0e0e0 !important;
    padding: 10px 12px !important;
    background-color: #ffffff !important;
    transition: all 0.3s;
    font-size: 0.95rem;
}
.stTextInput input:focus, .stNumberInput input:focus, .stTextArea textarea:focus, .stDateInput input:focus {
    border-color: #56ab2f !important;
    box-shadow: 0 4px 12px rgba(86, 171, 47, 0.15) !important;
}

/* --- BUTTONS --- */
.stButton button {
    border-radius: 12px !important;
    font-weight: 600;
    font-size: 1rem;
    padding: 0.6rem 1.2rem !important;
    border: none !important;
    box-shadow: 0 4px 6px rgba(0,0,0,0.05);
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    white-space: normal !important;
    height: auto !important;
    min-height: 2.5rem;
}
.stButton button:hover {
    transform: translateY(-2px);
    box-shadow: 0 8px 15px rgba(0,0,0,0.1);
}
.stButton button[kind="primary"] {
    background: linear-gradient(90deg, #56ab2f 0%, #a8e063 100%);
    color: white;
}
.stButton button[kind="secondary"] {
    background-color: #f1f3f5;
    color: #333;
}

/* --- COMPANY HEADER --- */
.company-header-container {
    display: flex; align-items: center; justify-content: center; gap: 30px;
    padding: 25px 40px; background: rgba(255, 255, 255, 0.8);
    backdrop-filter: blur(10px); border-radius: 20px;
    box-shadow: 0 8px 32px rgba(0,0,0,0.05); margin-bottom: 30px;
    border: 1px solid rgba(255,255,255,0.3); flex-wrap: nowrap !important;
}
.company-logo-img { height: 70px; width: auto; object-fit: contain; flex-shrink: 0; }
.company-info-text { text-align: left; flex: 1; display: flex; flex-direction: column; justify-content: center; white-space: normal; }
.company-info-text h1 { margin: 0; font-size: 1.8rem; color: #2e7d32; font-weight: 800; line-height: 1.2; }
.company-info-text p { margin: 5px 0 0 0; color: #555; font-size: 0.9rem; font-weight: 500; display: flex; align-items: center; gap: 10px; }

/* --- CARD STYLES --- */
.report-card, .login-container {
    background-color: white; border: none; border-radius: 20px;
    padding: 25px; margin-bottom: 25px;
    box-shadow: 0 10px 30px rgba(0,0,0,0.04);
    transition: all 0.3s ease;
}
.report-card:hover { transform: translateY(-5px); box-shadow: 0 20px 40px rgba(0,0,0,0.08); }

/* --- MONEY BOX --- */
.money-box {
    background: linear-gradient(135deg, #00b09b, #96c93d) !important;
    color: #ffffff !important; padding: 25px; border-radius: 20px;
    box-shadow: 0 15px 30px -5px rgba(0, 176, 155, 0.3);
    font-size: clamp(1.2rem, 3vw, 2.5rem); font-weight: 800;
    text-align: center; margin: 1.5rem 0; width: 100%;
    text-shadow: 0 2px 4px rgba(0,0,0,0.1); letter-spacing: 1px;
    white-space: normal; word-wrap: break-word;
    transition: transform 0.3s ease;
}
.money-box:hover { transform: scale(1.02); }

/* --- MODERN TABS --- */
div[data-baseweb="tab-list"] { border-bottom: 2px solid #e0e0e0; }
button[data-baseweb="tab"] {
    background-color: transparent !important; border-bottom: 2px solid transparent !important;
    padding-bottom: 10px !important; margin-bottom: -2px !important; transition: all 0.3s !important;
}
button[data-baseweb="tab"]:hover { background-color: #f1f3f5 !important; }
button[aria-selected="true"] {
    border-bottom-color: #56ab2f !important; font-weight: 600; color: #56ab2f !important;
}

/* --- ENHANCED EXPANDER --- */
div[data-testid="stExpander"] {
    border: 1px solid #e0e0e0 !important; border-radius: 15px !important;
    overflow: hidden; box-shadow: none !important; background-color: #fff;
}
div[data-testid="stExpander"] > details > summary {
    font-weight: 600; font-size: 1.05rem; background-color: #fafafa;
    padding: 0.75rem 1rem !important;
}
div[data-testid="stExpander"] > details > summary:hover { background-color: #f1f3f5; }

/* --- DATA EDITOR --- */
div[data-testid="stDataEditor"] {
    border-radius: 15px; overflow: hidden;
    border: 1px solid #f0f0f0; box-shadow: 0 4px 12px rgba(0,0,0,0.03);
}

/* --- FINANCE SUMMARY CARDS --- */
.finance-summary-card {
    background-color: #ffffff; border: 1px solid #e9ecef; border-radius: 15px;
    padding: 20px; margin-top: 15px;
}
.finance-summary-card .row {
    display: flex; justify-content: space-between; align-items: center;
    padding: 8px 0; border-bottom: 1px solid #f1f3f5;
}
.finance-summary-card .row:last-child { border-bottom: none; }
.finance-summary-card .row span { color: #495057; }
.finance-summary-card .row b { color: #212529; }
.finance-summary-card .total-row {
    font-size: 1.2em; font-weight: bold; color: #2e7d32; padding-top: 15px;
}
.finance-summary-card .pax-price {
    text-align: right; font-size: 0.9em; color: #6c757d; margin-top: 5px;
}
.profit-summary-card {
    background-color: #e3f2fd; padding: 20px; border-radius: 15px;
    text-align: center; border: 1px solid #90caf9; margin-top: 10px;
}
.profit-summary-card h3 {
    margin: 0; color: #1565c0; font-size: 1.1rem; font-weight: 600;
}
.profit-summary-card .formula {
    font-size: 1.8em; font-weight: bold; color: #1e88e5; margin-top: 10px;
}
.profit-summary-card .formula .result { color: #d32f2f; }

/* --- RESPONSIVE --- */
@media only screen and (max-width: 600px) {
    .company-header-container { flex-direction: column; text-align: center; gap: 10px; flex-wrap: wrap !important; }
    .company-info-text { text-align: center; }
    .company-info-text p { justify-content: center; }
}
</style>""", unsafe_allow_html=True)

def convert_image_to_pdf(image_file):
    try:
        img = Image.open(image_file)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        img_width, img_height = img.size
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=(img_width, img_height))
        temp_img_path = f"temp_img_{int(time.time())}.jpg"
        img.save(temp_img_path)
        c.drawImage(temp_img_path, 0, 0, img_width, img_height)
        c.save()
        if os.path.exists(temp_img_path): os.remove(temp_img_path)
        pdf_buffer.seek(0)
        return pdf_buffer.getvalue()
    except Exception as e:
        return None

# --- HÀM OCR ---
def perform_ocr(image_input, lang='vie'):
    """
    Thực hiện OCR trên ảnh với các bước tiền xử lý nâng cao sử dụng OpenCV để cải thiện độ chính xác.
    """
    # Check for dependencies and provide clear feedback.
    # This also helps static analysis tools like Pylance understand that `np` and `cv2` are not None below.
    if not HAS_OCR or pytesseract is None:
        st.toast("⚠️ Tesseract OCR chưa được cài đặt.", icon="🚨")
        return ""
    if not HAS_CV or np is None or cv2 is None:
        st.toast("⚠️ OpenCV hoặc Numpy chưa được cài đặt.", icon="🚨")
        return ""
    try:
        # 1. Load ảnh từ input (có thể là file stream hoặc đối tượng PIL)
        if isinstance(image_input, Image.Image):
            img = image_input
        else:
            image_input.seek(0)
            img = Image.open(image_input)

        # 2. Chuyển đổi sang định dạng OpenCV
        # Chuyển sang ảnh xám (grayscale) và numpy array để xử lý
        img_np = np.array(img.convert('L'))

        # 3. Tăng kích thước ảnh (Upscaling)
        # OCR hoạt động tốt hơn với ảnh có DPI cao (khoảng 300). Việc upscale ảnh nhỏ giúp nhận diện ký tự tốt hơn.
        h, w = img_np.shape
        if w < 2000:
            scale = 2000 / w
            new_w, new_h = int(w * scale), int(h * scale)
            # Sử dụng Lanczos interpolation cho kết quả sắc nét khi phóng to
            img_np = cv2.resize(img_np, (new_w, new_h), interpolation=cv2.INTER_LANCZOS4)

        # 4. Giảm nhiễu (Noise Reduction)
        # Sử dụng Median Blur hiệu quả để loại bỏ nhiễu "muối tiêu" (salt-and-pepper noise) mà không làm mờ các cạnh quá nhiều.
        img_np = cv2.medianBlur(img_np, 3)

        # 5. Binarization thông minh (Adaptive Thresholding)
        # Đây là bước quan trọng nhất, thay thế cho việc tăng contrast và dùng ngưỡng cố định.
        # Nó tự động tính toán ngưỡng cho các vùng ảnh nhỏ, rất hiệu quả với ảnh có điều kiện sáng không đồng đều.
        img_processed = cv2.adaptiveThreshold(
            img_np,
            255,  # Giá trị tối đa cho pixel
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,  # Phương pháp tính ngưỡng dựa trên vùng lân cận theo phân phối Gaussian
            cv2.THRESH_BINARY, # Chuyển ảnh thành đen và trắng
            15,  # Kích thước vùng lân cận (block size), nên là số lẻ
            4    # Hằng số C, một giá trị được trừ đi từ giá trị trung bình tính được
        )

        # 6. Cấu hình Tesseract để có kết quả tốt nhất
        # --psm 4: Giả định văn bản là một cột duy nhất với kích thước thay đổi (tốt cho hóa đơn, UNC).
        # --oem 3: Sử dụng engine mặc định (kết hợp Legacy và LSTM), thường cho kết quả ổn định.
        config = '--psm 4 --oem 3'
        text = pytesseract.image_to_string(img_processed, lang='vie+eng', config=config) if pytesseract else ""
        return text
    except Exception as e:
        print(f"OCR Error: {e}")
        return ""

def extract_money_smart(line):
    cleaned = re.sub(r'[^\d.,]', '', line) 
    potential_numbers = []
    raw_digits = re.findall(r'\d+', cleaned)
    for rd in raw_digits:
        if len(rd) > 8 and str(rd).startswith('0'): continue
        if len(rd) >= 4: potential_numbers.append(float(rd))
    matches = re.findall(r'\d[\d.,\s]*\d', line) 
    for m in matches:
        s = m.replace('VND', '').replace('đ', '').replace(' ', '').strip()
        if len(s) > 8 and s.startswith('0'): continue
        try:
            val = 0.0
            if ',' in s and '.' not in s: val = float(s.replace(',', ''))
            elif '.' in s and ',' not in s: val = float(s.replace('.', ''))
            elif ',' in s and '.' in s:
                last_dot = s.rfind('.')
                last_comma = s.rfind(',')
                if last_dot > last_comma: val = float(s.replace(',', '')) 
                else: val = float(s.replace('.', '').replace(',', '.'))
            else: val = float(s)
            if (val > 2030 or val < 1900) and val > 1000:
                potential_numbers.append(val)
        except: pass
    return potential_numbers

def extract_numbers_from_line_basic(line):
    clean_line = line.replace("-", "").replace("VND", "").replace("đ", "").strip()
    raw_integers = re.findall(r'(?<!\d)\d{4,}(?!\d)', clean_line)
    results = []
    for n in raw_integers:
        try:
            val = float(n)
            if not (1990 <= val <= 2030): results.append(val)
        except: pass
    return results

# --- XỬ LÝ HÓA ĐƠN & UNC (LOGIC CŨ) ---
def extract_data_smart(file_obj, is_image, doc_type="Hóa đơn"):
    text_content = ""
    msg = None
    try:
        if is_image:
            if HAS_OCR:
                # Gọi hàm OCR đã sửa đổi
                text_content = perform_ocr(file_obj)
                if not text_content.strip(): msg = "Hic, ảnh mờ quá hoặc không tìm thấy chữ số nào 😭."
            else: msg = "⚠️ Tình yêu ơi, máy chưa cài Tesseract OCR nên không đọc được ảnh nè."
        else:
            # Xử lý PDF (Cả text và scan)
            file_obj.seek(0)
            with pdfplumber.open(file_obj) as pdf:
                for page in pdf.pages: 
                    extracted = page.extract_text()
                    if extracted and len(extracted.strip()) > 10: 
                        text_content += extracted + "\n"
                    else:
                        if HAS_OCR:
                            im = page.to_image(resolution=300).original
                            text_content += perform_ocr(im) + "\n"
            
            if not text_content.strip(): 
                if not HAS_OCR: msg = "⚠️ File PDF này là ảnh scan, cần cài Tesseract OCR để đọc."
                else: msg = "⚠️ File trắng tinh hoặc không đọc được nội dung."

    except Exception as e: return None, f"Lỗi xíu xiu: {str(e)}"
    
    info = {"date": "", "seller": "", "buyer": "", "inv_num": "", "inv_sym": "", "pre_tax": 0.0, "tax": 0.0, "total": 0.0, "content": ""}
    if not text_content: return info, msg

    lines = text_content.split('\n')
    all_found_numbers = set()

    # --- TÌM NGÀY THÁNG ---
    m_date = re.search(r'(?:Ngày|ngày)\s+(\d{1,2})\s+(?:tháng|Tháng|[/.-])\s+(\d{1,2})\s+(?:năm|Năm|[/.-])\s+(\d{4})', text_content)
    if m_date: 
        try: info["date"] = f"{int(m_date.group(1)):02d}/{int(m_date.group(2)):02d}/{m_date.group(3)}"
        except: pass
    else:
        m_date_alt = re.search(r'(\d{2}/\d{2}/\d{4})', text_content)
        if m_date_alt: info["date"] = m_date_alt.group(1)

    # --- LOGIC XỬ LÝ SỐ TIỀN ---
    if doc_type == "Hóa đơn":
        # ... (Giữ nguyên logic Hóa đơn cũ của bạn ở đây nếu cần, hoặc dùng đoạn dưới đây)
        m_no = re.search(r'(?:Số hóa đơn|Số HĐ|Số|No)[:\s\.]*(\d{1,8})\b', text_content, re.IGNORECASE)
        if m_no: info["inv_num"] = m_no.group(1).zfill(7)
        m_sym = re.search(r'(?:Ký hiệu|Mẫu số|Serial)[:\s\.]*([A-Z0-9]{1,2}[A-Z0-9/-]{3,10})', text_content, re.IGNORECASE)
        if m_sym: info["inv_sym"] = m_sym.group(1)
        
        for line in lines:
            line_l = line.lower()
            nums = extract_money_smart(line)
            for n in nums: all_found_numbers.add(n)
            if not nums: continue
            val = max(nums)
            if any(kw in line_l for kw in ["thanh toán", "tổng cộng", "cộng tiền hàng"]): info["total"] = val
            elif any(kw in line_l for kw in ["tiền hàng", "thành tiền", "trước thuế"]): info["pre_tax"] = val
            elif "thuế" in line_l and "suất" not in line_l: info["tax"] = val
        
        if info["total"] == 0 and all_found_numbers: info["total"] = max(all_found_numbers)
        if info["pre_tax"] == 0: info["pre_tax"] = round(info["total"] / 1.08)
        if info["tax"] == 0: info["tax"] = info["total"] - info["pre_tax"]
        
        # Tìm Buyer/Seller cho Hóa đơn
        for line in lines[:35]:
            l_c = line.strip()
            if re.search(r'^(Đơn vị bán|Người bán|Bên A|Nhà cung cấp)', l_c, re.IGNORECASE): 
                parts = l_c.split(':')
                if len(parts) > 1: info["seller"] = parts[-1].strip()
            elif re.search(r'^(Đơn vị mua|Người mua|Khách hàng|Bên B)', l_c, re.IGNORECASE): 
                parts = l_c.split(':')
                if len(parts) > 1: info["buyer"] = parts[-1].strip()

    else: # === UNC (NÂNG CẤP LOGIC) ===
        candidates_total = []
        BLOCK_KEYWORDS = ['số dư', 'balance', 'phí', 'fee', 'charge', 'vat', 'tax', 'điện thoại', 'tel', 'fax', 'mst', 'mã số thuế', 'lệ phí', 'so du', 'le phi']
        CONFIRM_KEYWORDS = ['số tiền', 'amount', 'thanh toán', 'chuyển khoản', 'transaction', 'giá trị', 'total', 'cộng', 'money', 'so tien', 'chuyen khoan', 'gia tri']
        
        # --- LOAD TỪ KHÓA ĐÃ HỌC TỪ DB ---
        learned_kws = run_query("SELECT keyword FROM ocr_learning")
        if learned_kws:
            CONFIRM_KEYWORDS.extend([r['keyword'] for r in learned_kws]) # type: ignore
            
        CURRENCY_KEYWORDS = ['vnd', 'đ', 'vnđ', 'usd']
        prev_line_score_boost = 0
        fallback_numbers = []

        for i, line in enumerate(lines):
            line_l = line.lower()
            
            is_label_line = False
            if any(kw in line_l for kw in CONFIRM_KEYWORDS):
                nums_in_line = extract_money_smart(line)
                if not nums_in_line: 
                    prev_line_score_boost = 15 
                    is_label_line = True
            
            if is_label_line: continue

            nums = extract_money_smart(line)
            if not nums: 
                prev_line_score_boost = 0
                continue
            
            max_val = max(nums)
            if max_val < 1000: 
                prev_line_score_boost = 0
                continue 
            
            is_blocked = any(bad in line_l for bad in BLOCK_KEYWORDS)
            if not is_blocked:
                fallback_numbers.append(max_val)
            
            score = 0
            score += prev_line_score_boost
            prev_line_score_boost = 0 
            
            if any(kw in line_l for kw in CONFIRM_KEYWORDS): score += 10
            if any(kw in line_l for kw in CURRENCY_KEYWORDS): score += 5
            if is_blocked and not any(good in line_l for good in CONFIRM_KEYWORDS):
                score -= 20
            if 'tài khoản' in line_l or 'account' in line_l or 'stk' in line_l: score -= 5

            val_str = "{:,.0f}".format(max_val) # 10,000,000
            val_str_dot = val_str.replace(",", ".") # 10.000.000
            
            if val_str in line or val_str_dot in line:
                score += 3
            elif max_val > 100000000: 
                score -= 3

            if score > -10: candidates_total.append((max_val, score))
        
        if candidates_total:
            candidates_total.sort(key=lambda x: (x[1], x[0]), reverse=True)
            info["total"] = candidates_total[0][0]
        elif fallback_numbers:
            info["total"] = max(fallback_numbers)
            
        info["pre_tax"] = info["total"]
        
        for line in lines:
            if re.search(r'(?:nội dung|diễn giải|lý do|remarks|narrative|description|message)', line, re.IGNORECASE):
                parts = re.split(r'[:\.\-]', line, 1)
                if len(parts) > 1: info["content"] = parts[1].strip()
                else: info["content"] = line.strip()
                break

        for i, line in enumerate(lines):
            line_clean = line.strip()
            if re.search(r'(?:người hưởng|đơn vị thụ hưởng|tài khoản nhận|tên người nhận|bên nhận|beneficiary)', line_clean, re.IGNORECASE):
                parts = line_clean.split(':')
                if len(parts) > 1 and len(parts[-1].strip()) > 3:
                    info["seller"] = parts[-1].strip()
                    break
                elif i + 1 < len(lines):
                    info["seller"] = lines[i+1].strip()
                    break

    info["raw_text"] = text_content
    return info, msg

def create_handover_docx(tour_info, guests, hotels, restaurants, sightseeings, checklist_str):
    if not HAS_DOCX: return None
    
    # Lấy thông tin công ty
    comp_data = get_company_data()

    doc = Document() # type: ignore
    
    # Styles
    style = doc.styles['Normal']
    font = style.font # type: ignore
    font.name = 'Times New Roman'
    font.size = Pt(11) # type: ignore
    
    # --- HEADER (COMPANY INFO) ---
    header_table = doc.add_table(rows=1, cols=1)
    header_table.autofit = False
    header_table.allow_autofit = False # type: ignore
    cell = header_table.cell(0, 0)
    cell.width = Inches(6.5) # type: ignore
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
    
    run_comp = p.add_run(f"{comp_data['name'].upper()}\n")
    run_comp.bold = True
    run_comp.font.size = Pt(14) # type: ignore
    run_comp.font.color.rgb = None # Black
    
    p.add_run(f"Địa chỉ: {comp_data['address']}\n")
    p.add_run(f"Mã Số Thuế: {comp_data['phone']}")
    
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
    doc.add_paragraph()
    
    # --- TITLE ---
    p_title = doc.add_heading('HỒ SƠ BÀN GIAO ĐOÀN', 0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
    p_title.style.font.name = 'Times New Roman' # type: ignore
    p_title.style.font.size = Pt(16) # type: ignore
    p_title.style.font.bold = True # type: ignore
    p_title.style.font.color.rgb = None # Black # type: ignore
    
    doc.add_paragraph()
    
    # --- I. THÔNG TIN CHUNG ---
    doc.add_heading('I. THÔNG TIN CHUNG', level=1)
    
    table_info = doc.add_table(rows=0, cols=4)
    table_info.style = 'Table Grid'
    table_info.autofit = True
    
    def add_kv(k1, v1, k2, v2):
        row = table_info.add_row()
        c = row.cells
        c[0].text = k1
        c[0].paragraphs[0].runs[0].bold = True
        c[1].text = str(v1)
        c[2].text = k2
        c[2].paragraphs[0].runs[0].bold = True
        c[3].text = str(v2)

    cust_info = f"{tour_info.get('customer_name','')} - {tour_info.get('customer_phone','')}"
    
    add_kv("Tên đoàn:", tour_info['tour_name'], "Mã đoàn:", tour_info['tour_code'])
    add_kv("Ngày đi:", tour_info['start_date'], "Ngày về:", tour_info['end_date'])
    add_kv("Số lượng khách:", str(tour_info['guest_count']), "Sales:", tour_info['sale_name'])
    
    # Row for Customer
    r = table_info.add_row()
    r.cells[0].text = "Khách hàng:"
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[1].merge(r.cells[3])
    r.cells[1].text = cust_info
    
    doc.add_paragraph()
    
    # --- II. DANH SÁCH ĐOÀN ---
    doc.add_heading('II. DANH SÁCH ĐOÀN', level=1)
    if guests:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['STT', 'Họ và tên', 'Ngày sinh', 'Số CCCD', 'Phân loại']
        for i, h in enumerate(headers): 
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
            
        for i, g in enumerate(guests):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = str(g['name'])
            row_cells[2].text = str(g['dob'])
            row_cells[3].text = str(g['cccd'])
            row_cells[4].text = str(g['type'])
    else: doc.add_paragraph("(Chưa có danh sách đoàn)")
    doc.add_paragraph()
        
    # --- III. KHÁCH SẠN ---
    doc.add_heading('III. THÔNG TIN LƯU TRÚ', level=1)
    if hotels:
        table_h = doc.add_table(rows=1, cols=7)
        table_h.style = 'Table Grid'
        hdr = table_h.rows[0].cells
        for i, h in enumerate(['Tên Khách sạn', 'Địa chỉ & Liên hệ', 'Tổng phòng', 'Loại phòng', 'Tổng tiền', 'Đã cọc', 'Còn lại']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            
        for h in hotels:
            total = float(h.get('total_amount', 0) or 0)
            dep = float(h.get('deposit', 0) or 0)
            rem = total - dep
            row = table_h.add_row().cells
            row[0].text = f"🏨 {h['hotel_name']}"
            row[1].text = f"{h['address']}\nSĐT: {h['phone']}"
            row[2].text = str(h['total_rooms'])
            row[3].text = str(h['room_type'])
            row[4].text = "{:,.0f}".format(total)
            row[5].text = "{:,.0f}".format(dep)
            row[6].text = "{:,.0f}".format(rem)
    else: doc.add_paragraph("(Chưa có thông tin khách sạn)")
    doc.add_paragraph()

    # --- IV. NHÀ HÀNG ---
    doc.add_heading('IV. ẨM THỰC & THỰC ĐƠN', level=1)
    if restaurants:
        table_r = doc.add_table(rows=1, cols=7)
        table_r.style = 'Table Grid'
        hdr = table_r.rows[0].cells
        for i, h in enumerate(['Bữa ăn', 'Nhà hàng', 'Liên hệ', 'Thực đơn', 'Tổng tiền', 'Đã cọc', 'Còn lại']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            
        for r in restaurants:
            total = float(r.get('total_amount', 0) or 0)
            dep = float(r.get('deposit', 0) or 0)
            rem = total - dep
            row = table_r.add_row().cells
            row[0].text = f"🍽️ {r['meal_name']}"
            row[1].text = str(r['restaurant_name'])
            row[2].text = f"{r['address']}\nSĐT: {r['phone']}"
            row[3].text = str(r['menu'])
            row[4].text = "{:,.0f}".format(total)
            row[5].text = "{:,.0f}".format(dep)
            row[6].text = "{:,.0f}".format(rem)
    else: doc.add_paragraph("(Chưa có thông tin nhà hàng)")
    doc.add_paragraph()

    # --- V. ĐIỂM THAM QUAN ---
    doc.add_heading('V. ĐIỂM THAM QUAN', level=1)
    if sightseeings:
        table_s = doc.add_table(rows=1, cols=4)
        table_s.style = 'Table Grid'
        hdr = table_s.rows[0].cells
        for i, h in enumerate(['Tên địa điểm', 'Địa chỉ', 'Số lượng', 'Lưu ý']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            
        for s in sightseeings:
            row = table_s.add_row().cells
            row[0].text = f"📍 {s['name']}"
            row[1].text = str(s['address'])
            row[2].text = str(s['quantity'])
            row[3].text = str(s['note'])
    else: doc.add_paragraph("(Chưa có thông tin điểm tham quan)")
    doc.add_paragraph()
        
    # --- VI. CHECKLIST ---
    doc.add_heading('VI. CHECKLIST BÀN GIAO', level=1)
    checked_items = checklist_str.split(',') if checklist_str else []
    all_items = ["Chương trình đóng mộc", "Danh sách đóng mộc", "Bảo hiểm du lịch", "Thực đơn đóng mộc", "Vé máy bay", "Xác nhận khu du lịch/nhà hàng (Nếu có)", "Hợp đồng hướng dẫn"]
    
    table_c = doc.add_table(rows=0, cols=2)
    for item in all_items:
        mark = "☑" if item in checked_items else "☐"
        row = table_c.add_row()
        row.cells[0].text = mark
        row.cells[0].width = Pt(20) # type: ignore
        row.cells[1].text = item
        
    # Footer
    doc.add_paragraph("\n")
    p_foot = doc.add_paragraph(f"Ngày xuất hồ sơ: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT # type: ignore
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- HÀM ĐỌC SỐ TIỀN BẰNG CHỮ (VIETNAMESE) ---
def read_money_vietnamese(amount):
    if amount == 0: return "Không đồng"
    
    digits = ["không", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
    units = ["", "nghìn", "triệu", "tỷ"]
    
    def read_group(n):
        res = ""
        h = n // 100
        t = (n % 100) // 10
        u = n % 10
        
        if h > 0:
            res += digits[h] + " trăm "
        elif n > 0: # Có hàng chục hoặc đơn vị nhưng hàng trăm = 0 (xử lý ở loop chính tốt hơn, đây là logic đơn giản)
            pass 
            
        if t > 1:
            res += digits[t] + " mươi "
            if u == 1: res += "mốt "
            elif u == 5: res += "lăm "
            elif u > 0: res += digits[u] + " "
        elif t == 1:
            res += "mười "
            if u == 1: res += "một "
            elif u == 5: res += "lăm "
            elif u > 0: res += digits[u] + " "
        else: # t = 0
            if h > 0 and u > 0: res += "lẻ "
            if u > 0: res += digits[u] + " "
        return res

    s_num = "{:.0f}".format(amount)
    groups = []
    while len(s_num) > 0:
        groups.append(int(s_num[-3:]))
        s_num = s_num[:-3]
    
    ret = ""
    for i, g in enumerate(groups):
        if g > 0:
            s_g = read_group(g)
            # Xử lý số 0 trăm
            if i < len(groups) - 1 and g < 100 and g > 0: 
                s_g = "không trăm " + s_g
                
            ret = s_g + units[i] + " " + ret
            
    ret = ret.strip()
    # Capitalize first letter
    if ret:
        ret = ret[0].upper() + ret[1:]
    
    return ret + " đồng"

def create_voucher_pdf(voucher_data):
    """Tạo file PDF phiếu thu/chi đẹp, có logo và màu sắc"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # --- SỬA LỖI FONT TIẾNG VIỆT TRÊN STREAMLIT CLOUD ---
    font_name = 'Helvetica' # Fallback mặc định
    font_name_bold = 'Helvetica-Bold'
    
    # Danh sách các cặp file font (Thường, Đậm) ưu tiên tìm kiếm
    # Bạn phải upload các file .ttf này lên cùng thư mục với app.py trên Streamlit Cloud
    font_candidates = [
        ("times.ttf", "timesbd.ttf", "TimesNewRoman"),  # Ưu tiên 1
        ("arial.ttf", "arialbd.ttf", "Arial"),          # Ưu tiên 2
        ("Roboto-Regular.ttf", "Roboto-Bold.ttf", "Roboto") # Ưu tiên 3 (Nếu dùng Google Fonts)
    ]
    
    font_registered = False

    # 1. Thử tìm font trong thư mục hiện tại (Dành cho Streamlit Cloud)
    for regular, bold, name in font_candidates:
        if os.path.exists(regular) and os.path.exists(bold):
            try:
                pdfmetrics.registerFont(TTFont(name, regular))
                pdfmetrics.registerFont(TTFont(f'{name}-Bold', bold))
                font_name = name
                font_name_bold = f'{name}-Bold'
                font_registered = True
                break
            except: pass

    # 2. Nếu chưa tìm thấy, thử tìm trong Windows (Dành cho chạy Local)
    if not font_registered:
        try:
            win_path = r"C:\Windows\Fonts\times.ttf"
            win_path_bd = r"C:\Windows\Fonts\timesbd.ttf"
            if os.path.exists(win_path) and os.path.exists(win_path_bd):
                pdfmetrics.registerFont(TTFont('TimesNewRoman', win_path))
                pdfmetrics.registerFont(TTFont('TimesNewRoman-Bold', win_path_bd))
                font_name = 'TimesNewRoman'
                font_name_bold = 'TimesNewRoman-Bold'
            else:
                # Fallback Arial trên Windows
                win_arial = r"C:\Windows\Fonts\arial.ttf"
                if os.path.exists(win_arial):
                    pdfmetrics.registerFont(TTFont('Arial', win_arial))
                    font_name = 'Arial'
                    font_name_bold = 'Arial' # Arial thường không tách file bold rõ trong logic đơn giản
        except: pass

    comp = get_company_data()
    
    # Màu sắc chủ đạo
    primary_color = "#2E7D32" if voucher_data['type'] == 'THU' else "#C62828" # Xanh cho Thu, Đỏ cho Chi
    text_color = "#212121"
    
    # --- HEADER ---
    # Logo
    logo_height = 60
    header_y = height - 50
    header_x_text = 50
    
    if comp['logo_b64_str']:
        try:
            logo_data = base64.b64decode(comp['logo_b64_str'])
            image_stream = io.BytesIO(logo_data)
            img_reader = ImageReader(image_stream)
            # Tính tỷ lệ ảnh
            iw, ih = img_reader.getSize()
            aspect = iw / float(ih)
            draw_w = logo_height * aspect
            
            c.drawImage(img_reader, 50, header_y - logo_height, width=draw_w, height=logo_height, mask='auto')
            header_x_text = 50 + draw_w + 20
        except: pass

    # --- LOGO CHÌM (WATERMARK) ---
    if comp['logo_b64_str']:
        try:
            c.saveState()
            logo_data = base64.b64decode(comp['logo_b64_str'])
            image_stream = io.BytesIO(logo_data)
            img_reader = ImageReader(image_stream)
            iw, ih = img_reader.getSize()
            aspect = iw / float(ih)
            wm_width = 300
            wm_height = wm_width / aspect
            c.setFillAlpha(0.1) # Độ mờ 10%
            # Vẽ chính giữa trang
            c.drawImage(img_reader, (width - wm_width)/2, (height - wm_height)/2, width=wm_width, height=wm_height, mask='auto')
            c.restoreState()
        except: pass

    # Thông tin công ty
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_name_bold, 16)
    c.drawString(header_x_text, header_y - 15, comp['name'].upper())
    
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 10)
    c.drawString(header_x_text, header_y - 35, f"ĐC: {comp['address']}")
    c.drawString(header_x_text, header_y - 50, f"MST: {comp['phone']}")
    
    # Đường kẻ trang trí
    c.setStrokeColor(HexColor(primary_color))
    c.setLineWidth(2)
    c.line(50, header_y - 70, width - 50, header_y - 70)
    
    # --- TIÊU ĐỀ ---
    title = "PHIẾU THU TIỀN" if voucher_data['type'] == 'THU' else "PHIẾU CHI TIỀN"
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_name_bold, 24)
    c.drawCentredString(width/2, height - 150, title)
    
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 11)
    c.drawCentredString(width/2, height - 170, f"Ngày: {voucher_data['date']}")
    
    # --- NỘI DUNG ---
    # Lấy tên khách hàng nếu có
    person_name = ""
    ref_code = voucher_data.get('ref_code', '')
    if ref_code:
        try:
            # Thử tìm trong Tours
            t = run_query("SELECT customer_name FROM tours WHERE tour_code=?", (ref_code,), fetch_one=True)
            if t and t['customer_name']: person_name = t['customer_name']
            else:
                # Thử tìm trong Bookings
                b = run_query("SELECT customer_info FROM service_bookings WHERE code=?", (ref_code,), fetch_one=True)
                if b and b['customer_info']: person_name = b['customer_info'].split(' - ')[0]
        except: pass

    # --- TÍNH TOÁN TÀI CHÍNH (MỚI) ---
    contract_val = 0.0
    total_paid = 0.0
    remaining = 0.0
    
    if ref_code:
        # 1. Lấy giá trị hợp đồng
        # Thử tìm Tour
        t_info = run_query("SELECT * FROM tours WHERE tour_code=?", (ref_code,), fetch_one=True)
        if t_info:
            t_dict = dict(t_info)
            final_price = float(t_dict.get('final_tour_price', 0) or 0)
            child_price = float(t_dict.get('child_price', 0) or 0)
            final_qty = float(t_dict.get('final_qty', 0) or 0)
            child_qty = float(t_dict.get('child_qty', 0) or 0)
            if final_qty == 0: final_qty = float(t_dict.get('guest_count', 1))
            contract_val = (final_price * final_qty) + (child_price * child_qty)
        else:
            # Thử tìm Booking
            b_info = run_query("SELECT selling_price FROM service_bookings WHERE code=?", (ref_code,), fetch_one=True)
            if b_info:
                contract_val = float(b_info['selling_price'] or 0)
        
        # 2. Lấy tổng đã thu (Bao gồm cả phiếu vừa tạo nếu đã lưu DB)
        txns = run_query("SELECT type, amount FROM transaction_history WHERE ref_code=?", (ref_code,))
        if txns:
            paid_sum = sum(r['amount'] for r in txns if r['type'] == 'THU')
            refund_sum = sum(r['amount'] for r in txns if r['type'] == 'CHI')
            total_paid = paid_sum - refund_sum
            
        remaining = contract_val - total_paid

    # --- VẼ PDF ---
    y = height - 220
    x_label = 70
    x_val = 200
    line_height = 30
    
    # Vẽ khung nền mờ
    bg_color = "#E8F5E9" if voucher_data['type'] == 'THU' else "#FFEBEE"
    c.setFillColor(HexColor(bg_color))
    # Tăng chiều cao khung để chứa thêm thông tin (210 -> 330)
    c.roundRect(50, y - 310, width - 100, 330, 10, fill=1, stroke=0)
    
    c.setFillColor(HexColor(text_color))
    
    def draw_line_content(label, value, y_pos, is_money=False):
        c.setFont(font_name, 12)
        c.drawString(x_label, y_pos, label)
        
        if is_money:
            c.setFont(font_name_bold, 14)
            c.setFillColor(HexColor(primary_color))
            c.drawString(x_val, y_pos, value)
            c.setFillColor(HexColor(text_color)) # Reset
        else:
            c.setFont(font_name, 12)
            if value:
                c.drawString(x_val, y_pos, value)
            else:
                # Vẽ dòng chấm
                c.setStrokeColor(HexColor("#BDBDBD"))
                c.setLineWidth(1)
                c.setDash(1, 3)
                c.line(x_val, y_pos - 3, width - 70, y_pos - 3)
                c.setDash([])

    label_person = "Người nộp tiền:" if voucher_data['type'] == 'THU' else "Người nhận tiền:"
    draw_line_content(label_person, person_name, y); y -= line_height
    draw_line_content("Địa chỉ/SĐT:", "", y); y -= line_height
    draw_line_content("Lý do:", f"{voucher_data['note']} (Mã: {voucher_data['ref_code']})", y); y -= line_height
    draw_line_content("Số tiền:", f"{format_vnd(voucher_data['amount'])} VND", y, is_money=True); y -= line_height
    draw_line_content("Bằng chữ:", read_money_vietnamese(voucher_data['amount']), y); y -= line_height
    
    # --- CÁC DÒNG MỚI ---
    draw_line_content("Tổng giá trị HĐ:", f"{format_vnd(contract_val)} VND", y); y -= line_height
    draw_line_content("Đã thanh toán:", f"{format_vnd(total_paid)} VND", y); y -= line_height
    draw_line_content("Còn lại:", f"{format_vnd(remaining)} VND", y); y -= line_height
    draw_line_content("Người xuất phiếu:", voucher_data.get('issuer', ''), y); y -= line_height

    draw_line_content("Kèm theo:", "", y); y -= line_height
    
    # --- CHỮ KÝ ---
    y_sig = y - 40
    sigs = ["Giám đốc", "Kế toán trưởng", "Người lập phiếu", "Người nộp/nhận"]
    x_positions = [50, 180, 310, 440]
    for i, sig in enumerate(sigs):
        c.setFont(font_name, 11)
        c.setFillColor(HexColor(text_color))
        c.drawCentredString(x_positions[i] + 40, y_sig, sig)
        c.setFont(font_name, 9)
        c.setFillColor(HexColor("#757575"))
        c.drawCentredString(x_positions[i] + 40, y_sig - 15, "(Ký, họ tên)")
        
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def create_voucher_docx(voucher_data):
    """Tạo file Word phiếu thu/chi"""
    doc = Document()
    # Styles
    style = doc.styles['Normal']
    font = style.font # type: ignore
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    comp = get_company_data()
    
    # Header
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(2.5)
    t.columns[1].width = Cm(14)
    
    # Logo
    if comp['logo_b64_str']:
        try:
            logo_data = base64.b64decode(comp['logo_b64_str'])
            image_stream = io.BytesIO(logo_data)
            cell = t.cell(0, 0)
            p = cell.paragraphs[0]
            r = p.add_run()
            r.add_picture(image_stream, width=Cm(2.5))
        except: pass
        
    # Company Info
    cell = t.cell(0, 1)
    p = cell.paragraphs[0]
    r = p.add_run(comp['name'].upper() + "\n")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(46, 125, 50)
    
    p.add_run(f"Địa chỉ: {comp['address']}\n")
    p.add_run(f"Hotline/MST: {comp['phone']}")
    
    doc.add_paragraph()
    
    # Title
    title = "PHIẾU THU TIỀN" if voucher_data['type'] == 'THU' else "PHIẾU CHI TIỀN"
    color = RGBColor(46, 125, 50) if voucher_data['type'] == 'THU' else RGBColor(198, 40, 40)
    
    p_title = doc.add_paragraph(title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.runs[0]
    r_t.bold = True
    r_t.font.size = Pt(20)
    r_t.font.color.rgb = color
    
    p_date = doc.add_paragraph(f"Ngày: {voucher_data['date']}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Data Logic
    person_name = ""
    ref_code = voucher_data.get('ref_code', '')
    if ref_code:
        try:
            t = run_query("SELECT customer_name FROM tours WHERE tour_code=?", (ref_code,), fetch_one=True)
            if t and t['customer_name']: person_name = t['customer_name']
            else:
                b = run_query("SELECT customer_info FROM service_bookings WHERE code=?", (ref_code,), fetch_one=True)
                if b and b['customer_info']: person_name = b['customer_info'].split(' - ')[0]
        except: pass

    # Content Table
    t_content = doc.add_table(rows=0, cols=2)
    t_content.autofit = False
    t_content.columns[0].width = Cm(4)
    t_content.columns[1].width = Cm(12)
    
    def add_row(label, value, bold_val=False, color_val=None):
        r = t_content.add_row()
        r.cells[0].text = label
        r.cells[1].text = str(value)
        if bold_val: 
            r.cells[1].paragraphs[0].runs[0].bold = True
        if color_val:
            r.cells[1].paragraphs[0].runs[0].font.color.rgb = color_val

    label_person = "Người nộp tiền:" if voucher_data['type'] == 'THU' else "Người nhận tiền:"
    add_row(label_person, person_name)
    add_row("Lý do:", f"{voucher_data['note']} (Mã: {voucher_data['ref_code']})")
    add_row("Số tiền:", f"{format_vnd(voucher_data['amount'])} VND", True, color)
    add_row("Bằng chữ:", read_money_vietnamese(voucher_data['amount']))
    add_row("Người xuất phiếu:", voucher_data.get('issuer', ''))
    
    doc.add_paragraph()
    
    # Signatures
    t_sig = doc.add_table(rows=1, cols=4)
    t_sig.autofit = True
    sigs = ["Giám đốc", "Kế toán trưởng", "Người lập phiếu", "Người nộp/nhận"]
    for i, s in enumerate(sigs):
        cell = t_sig.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(s + "\n\n\n\n")
        r.bold = True
        p.add_run("(Ký, họ tên)")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_booking_cfm_pdf(booking_info, company_info, lang='en'):
    """Tạo file PDF Booking Confirmation (CFM)"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # --- HELPER: WATERMARK & NEW PAGE ---
    def draw_watermark():
        if company_info['logo_b64_str']:
            try:
                c.saveState()
                logo_data = base64.b64decode(company_info['logo_b64_str'])
                image_stream = io.BytesIO(logo_data)
                img_reader = ImageReader(image_stream)
                iw, ih = img_reader.getSize()
                aspect = iw / float(ih)
                wm_width = 400
                wm_height = wm_width / aspect
                c.setFillAlpha(0.08) # Độ mờ 8%
                c.drawImage(img_reader, (width - wm_width)/2, (height - wm_height)/2, width=wm_width, height=wm_height, mask='auto')
                c.restoreState()
            except: pass

    draw_watermark()

    # --- CẤU HÌNH FONT (SỬA LẠI ĐỂ CHẠY TRÊN CLOUD) ---
    font_name = 'Helvetica'
    font_bold = 'Helvetica-Bold'
    
    # Danh sách các cặp file font (Thường, Đậm) ưu tiên tìm kiếm
    # Bạn phải upload các file .ttf này lên cùng thư mục với app.py trên Streamlit Cloud
    font_candidates = [
        ("times.ttf", "timesbd.ttf", "TimesNewRoman"),  # Ưu tiên 1
        ("arial.ttf", "arialbd.ttf", "Arial"),          # Ưu tiên 2
        ("Roboto-Regular.ttf", "Roboto-Bold.ttf", "Roboto"), # Ưu tiên 3 (Nếu dùng Google Fonts)
        # Thêm font hệ thống Linux (Streamlit Cloud)
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "DejaVuSans"),
        ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf", "LiberationSans"),
    ]
    
    font_registered = False

    # 1. Ưu tiên tìm font ngay tại thư mục gốc (Streamlit Cloud)
    for regular, bold, name in font_candidates:
        if os.path.exists(regular):
            try:
                pdfmetrics.registerFont(TTFont(name, regular))
                font_name = name
                
                if os.path.exists(bold):
                    pdfmetrics.registerFont(TTFont(f'{name}-Bold', bold))
                    font_bold = f'{name}-Bold'
                else:
                    # Nếu không có file đậm, dùng file thường cho đậm (để không lỗi font)
                    pdfmetrics.registerFont(TTFont(f'{name}-Bold', regular))
                    font_bold = f'{name}-Bold'

                font_registered = True
                break
            except: pass

    # 2. Nếu chưa tìm thấy, thử tìm trong Windows (Dành cho chạy Local)
    if not font_registered:
        try:
            win_path = r"C:\Windows\Fonts\times.ttf"
            win_path_bd = r"C:\Windows\Fonts\timesbd.ttf"
            if os.path.exists(win_path):
                pdfmetrics.registerFont(TTFont('TimesNewRoman', win_path))
                font_name = 'TimesNewRoman'
                
                if os.path.exists(win_path_bd):
                    pdfmetrics.registerFont(TTFont('TimesNewRoman-Bold', win_path_bd))
                    font_bold = 'TimesNewRoman-Bold'
                else:
                    pdfmetrics.registerFont(TTFont('TimesNewRoman-Bold', win_path))
                    font_bold = 'TimesNewRoman-Bold'
            else:
                # Fallback Arial trên Windows
                win_arial = r"C:\Windows\Fonts\arial.ttf"
                if os.path.exists(win_arial):
                    pdfmetrics.registerFont(TTFont('Arial', win_arial))
                    font_name = 'Arial'
                    # Arial thường tự xử lý bold hoặc dùng cùng file
                    pdfmetrics.registerFont(TTFont('Arial-Bold', win_arial))
                    font_bold = 'Arial-Bold' 
        except: pass

    # --- MÀU SẮC ---
    primary_color = "#1B5E20" # Xanh đậm thương hiệu
    text_color = "#212121"
    line_color = "#BDBDBD"

    # --- TỪ ĐIỂN NGÔN NGỮ ---
    labels = {
        'en': {
            'add': 'Add:', 'tax': 'Tax Code:',
            'title': 'BOOKING CONFIRMATION',
            'greeting': f"A warm greeting from {company_info['name']}!",
            'gen_info': 'I. GENERAL INFORMATION',
            'attn': 'Attention to:', 'bk_code': 'Booking Code:', 'svc_date': 'Service Date:',
            'date_created': 'Date Created:', 'checkout': 'Check-out:', 'created_by': 'Created By:',
            'status': 'Status:', 'hotel_code': 'Hotel Code:',
            'confirmed': 'Confirmed', 'cancelled': 'Cancelled',
            'svc_details': 'II. SERVICE DETAILS',
            'svc_details_cont': 'II. SERVICE DETAILS (Cont.)',
            'tbl_name': 'SERVICE NAME', 'tbl_det': 'DETAILS', 'tbl_note': 'NOTE',
            'guest_list': 'III. GUEST LIST',
            'guest_list_cont': 'III. GUEST LIST (Cont.)',
            'included': 'INCLUDED SERVICES',
            'inc_1': '- Tax and Service charges.',
            'inc_2': '- 24/7 Support from our team.',
            'confirmed_by': 'CONFIRMED BY',
            'signed': '[SIGNED]',
            'page': 'Page'
        },
        'vi': {
            'add': 'ĐC:', 'tax': 'MST:',
            'title': 'XÁC NHẬN ĐẶT DỊCH VỤ',
            'greeting': f"Lời chào trân trọng từ {company_info['name']}!",
            'gen_info': 'I. THÔNG TIN CHUNG',
            'attn': 'Kính gửi:', 'bk_code': 'Mã đặt chỗ:', 'svc_date': 'Ngày sử dụng:',
            'date_created': 'Ngày tạo:', 'checkout': 'Ngày trả phòng:', 'created_by': 'Người tạo:',
            'status': 'Trạng thái:', 'hotel_code': 'Mã khách sạn:',
            'confirmed': 'Đã xác nhận', 'cancelled': 'Đã hủy',
            'svc_details': 'II. CHI TIẾT DỊCH VỤ',
            'svc_details_cont': 'II. CHI TIẾT DỊCH VỤ (Tiếp)',
            'tbl_name': 'TÊN DỊCH VỤ', 'tbl_det': 'THÔNG TIN CHI TIẾT', 'tbl_note': 'GHI CHÚ',
            'guest_list': 'III. DANH SÁCH KHÁCH',
            'guest_list_cont': 'III. DANH SÁCH KHÁCH (Tiếp)',
            'included': 'DỊCH VỤ BAO GỒM',
            'inc_1': '- Thuế và phí phục vụ.',
            'inc_2': '- Hỗ trợ 24/7 từ đội ngũ của chúng tôi.',
            'confirmed_by': 'XÁC NHẬN BỞI',
            'signed': '[ĐÃ KÝ]',
            'page': 'Trang'
        }
    }
    txt = labels[lang]

    # --- HEADER ---
    y = height - 50
    # Logo
    if company_info['logo_b64_str']:
        try:
            logo_data = base64.b64decode(company_info['logo_b64_str'])
            image_stream = io.BytesIO(logo_data)
            img_reader = ImageReader(image_stream)
            iw, ih = img_reader.getSize()
            aspect = iw / float(ih)
            logo_h = 85 # Logo to hơn
            logo_w = logo_h * aspect
            c.drawImage(img_reader, 40, y - logo_h, width=logo_w, height=logo_h, mask='auto')
        except: pass

    # Xử lý địa chỉ (Dịch sơ bộ nếu là tiếng Anh)
    # Xử lý địa chỉ và tên công ty (Dịch sơ bộ nếu là tiếng Anh)
    comp_addr = company_info['address']
    comp_name = company_info['name']
    
    if lang == 'en':
        # [UPDATED] Hardcoded English details
        comp_name = "BALI TOURIST TRAVEL COMPANY LIMITED"
        comp_addr = "No. 46 Nguyen Oanh, Hanh Thong Ward, Ho Chi Minh City, Vietnam"
        txt['greeting'] = f"A warm greeting from {comp_name}!"

    # Thông tin công ty (Căn phải)
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 18)
    c.drawRightString(width - 40, y - 25, comp_name.upper())
    
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 10)
    c.drawRightString(width - 40, y - 45, f"{txt['add']} {comp_addr}")
    c.drawRightString(width - 40, y - 60, f"{txt['tax']} {company_info['phone']}")
    
    y -= 100
    c.setStrokeColor(HexColor(primary_color))
    c.setLineWidth(2)
    c.line(40, y, width - 40, y)
    
    # --- TITLE ---
    y -= 40
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 20)
    c.drawCentredString(width/2, y, txt['title'])
    
    y -= 25
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 11)
    c.drawCentredString(width/2, y, txt['greeting'])
    
    # --- XỬ LÝ DỮ LIỆU BOOKING ---
    # Parse Customer Info
    cust_raw = booking_info.get('customer_info', '')
    cust_name = cust_raw.split(' - ')[0] if ' - ' in cust_raw else cust_raw
    
    # Parse Dates from Details
    details = booking_info.get('details', '')
    dates = re.findall(r'\d{1,2}[/-]\d{1,2}[/-]\d{4}', details)
    check_in = dates[0] if len(dates) > 0 else booking_info.get('created_at', '')
    check_out = dates[1] if len(dates) > 1 else "N/A"
    
    # --- LOGIC MỚI: TRÍCH XUẤT MÃ ĐẶC THÙ (PHẦN I) ---
    bk_type = booking_info.get('type', '')
    specific_label = ""
    specific_value = ""

    if bk_type == 'HOTEL':
        if booking_info.get('hotel_code'):
            specific_label = txt['hotel_code']
            specific_value = booking_info.get('hotel_code', '')
    elif bk_type == 'TRANS':
        # Regex tìm Biển số xe: "Xe ...: [Biển số] |"
        m_car = re.search(r'(Xe\s*[^:]*:\s*)([^|]+)', details)
        if m_car:
            specific_label = "Biển số xe:" if lang == 'vi' else "Car Plate:"
            specific_value = m_car.group(2).strip()
        else:
            # Regex tìm Mã vé: "Vé: [Code] |"
            m_ticket = re.search(r'(Vé:\s*)([^|]+)', details)
            if m_ticket:
                specific_label = "Mã vé:" if lang == 'vi' else "Ticket Code:"
                specific_value = m_ticket.group(2).strip()
    
    # Nếu không tìm thấy trong details, thử tìm trong tên (cho trường hợp Vé máy bay)
    if bk_type == 'TRANS' and not specific_value and '[' in booking_info['name']:
         pass # Giữ nguyên logic cũ nếu không extract được

    # --- PHẦN 1: THÔNG TIN CHUNG ---
    y -= 40
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 12)
    c.drawString(40, y, txt['gen_info'])
    y -= 20
    
    # Danh sách thông tin (Cân đối lại layout)
    gen_items = [
        (txt['attn'], cust_name),
        (txt['bk_code'], booking_info['code']),
        (txt['date_created'], booking_info.get('created_at', '')),
        (txt['created_by'], booking_info.get('sale_name', '')),
        (txt['svc_date'], check_in),
    ]
    if check_out != "N/A":
        gen_items.append((txt['checkout'], check_out))
    
    status_txt = txt['confirmed'] if booking_info.get('status') != 'deleted' else txt['cancelled']
    gen_items.append((txt['status'], status_txt))
    
    if specific_label and specific_value:
        gen_items.append((specific_label, specific_value))

    # Tính toán chiều cao khung dựa trên số dòng (chia 2 cột)
    import math
    rows_needed = math.ceil(len(gen_items) / 2)
    row_h = 25
    box_height = rows_needed * row_h + 15

    # Vẽ khung
    c.setStrokeColor(HexColor(line_color))
    c.setLineWidth(1)
    c.rect(40, y - box_height, width - 80, box_height, fill=0)
    
    # Vẽ nội dung (2 cột)
    curr_y = y - 25
    col_x_positions = [50, 300] # Vị trí x cho cột 1 và cột 2
    
    for i, (label, val) in enumerate(gen_items):
        col_idx = i % 2
        x_pos = col_x_positions[col_idx]
        draw_y = curr_y - (i // 2) * row_h
        
        c.setFillColor(HexColor(text_color))
        c.setFont(font_name, 11)
        c.drawString(x_pos, draw_y, label)
        
        c.setFont(font_bold, 11)
        if label == txt['status']:
             c.setFillColor(HexColor("#2E7D32" if val == txt['confirmed'] else "#C62828"))
        
        # Căn chỉnh giá trị
        val_x_offset = 100 if col_idx == 0 else 110
        c.drawString(x_pos + val_x_offset, draw_y, str(val))

    # --- PHẦN 2: CHI TIẾT DỊCH VỤ ---
    y -= (box_height + 30)
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 12)
    c.drawString(40, y, txt['svc_details'])
    y -= 25
    
    # Header Bảng
    c.setFillColor(HexColor("#E8F5E9"))
    c.rect(40, y - 5, width - 80, 20, fill=1, stroke=0) # Header BG
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 10)
    c.drawString(50, y, txt['tbl_name'])
    c.drawString(250, y, txt['tbl_det'])
    c.drawString(450, y, txt['tbl_note'])
    
    y -= 20
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 10)
    
    # Nội dung bảng (Xử lý Combo tách dòng)
    items = []
    
    # [HELPER] Translate content if English
    def translate_content(text):
        if lang != 'en' or not text: return text
        replacements = {
            "Ngày:": "Date:", "SL:": "Qty:", "Lưu trú:": "Stay:",
            "Xe ": "Car ", "Vé:": "Ticket:", "Máy bay": "Flight", 
            "Tàu hỏa": "Train", "Du thuyền": "Cruise", "Cabin:": "Cabin:",
            "phòng": "rooms", "đêm": "nights", "khách": "pax",
            "[KS]": "[Hotel]", "[XE]": "[Car]", "[BAY]": "[Flight]", 
            "[TAU]": "[Train]", "[THUYEN]": "[Cruise]", "[CB]": "[Combo]"
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text

    if booking_info.get('type') == 'COMBO':
        # Tách các item trong combo (ngăn cách bởi | hoặc dòng mới)
        raw_items = re.split(r'[|\n]', details)
        for item in raw_items:
            if item.strip(): 
                display_item = translate_content(item.strip())
                items.append((display_item, ""))
    elif booking_info.get('type') == 'TRANS' and '[XE]' in booking_info.get('name', ''):
        # [NEW] Xử lý hiển thị chi tiết cho Booking Xe (Tách dòng)
        name_display = translate_content(booking_info['name'])
        raw_details = booking_info.get('details', '')
        
        # 1. Lịch trình (Lấy từ tên booking)
        itinerary = booking_info['name'].replace('[XE]', '').strip()
        
        # 2. Thông tin xe
        car_info = ""
        m_car = re.search(r'Xe\s*([^:]*):\s*([^|]+)', raw_details)
        if m_car: car_info = f"{m_car.group(1).strip()} - {m_car.group(2).strip()}"
            
        # 3. Tài xế & SĐT
        driver_name = ""
        driver_phone = ""
        m_drv_full = re.search(r'Tài xế:\s*([^|]+)', raw_details)
        if m_drv_full:
            drv_str = m_drv_full.group(1).strip()
            if ' - ' in drv_str:
                parts = drv_str.split(' - ')
                driver_name = parts[0]
                driver_phone = parts[1] if len(parts) > 1 else ""
            else:
                driver_name = drv_str

        lines = []
        if itinerary: lines.append(f"Lịch trình: {itinerary}" if lang == 'vi' else f"Itinerary: {itinerary}")
        if car_info: lines.append(f"Xe/Biển số: {car_info}" if lang == 'vi' else f"Car/Plate: {car_info}")
        if driver_name: lines.append(f"Tài xế: {driver_name}" if lang == 'vi' else f"Driver: {driver_name}")
        if driver_phone: lines.append(f"SĐT: {driver_phone}" if lang == 'vi' else f"Phone: {driver_phone}")
        
        details_display = "\n".join(lines)
        items.append((name_display, details_display))
    else:
        # Translate basic keywords for non-hotel types
        details_display = translate_content(details)
        name_display = translate_content(booking_info['name'])
        items.append((name_display, details_display))
        
    # [NEW] Xử lý hiển thị chi tiết cho Booking Khách sạn (Hotel Code, Room Type, Guest List)
    if booking_info.get('type') == 'HOTEL':
        # Override items list to show detailed info
        r_type = booking_info.get('room_type', '')
        g_list = booking_info.get('guest_list', '')
        
        # --- LOGIC MỚI: TÍNH SỐ ĐÊM & BỎ NGÀY CỤ THỂ ---
        new_details = ""
        
        # Tính số đêm từ dates đã extract ở trên
        nights_count = 0
        if len(dates) >= 2:
            try:
                d1 = datetime.strptime(dates[0], '%d/%m/%Y')
                d2 = datetime.strptime(dates[1], '%d/%m/%Y')
                nights_count = (d2 - d1).days
            except: pass
        
        if nights_count > 0:
            lbl_night = "nights" if lang == 'en' else "đêm"
            new_details += f"{nights_count} {lbl_night}"
        
        if r_type:
            lbl_room = "Room Type:" if lang == 'en' else "Loại phòng:"
            if new_details: new_details += "\n"
            new_details += f"{lbl_room} {r_type}"
        
        # Format phần Note hoặc thêm vào Details
        note_part = "" # Guest list moved to separate section
        
        name_display = translate_content(booking_info['name'])
        items = [(name_display, new_details, note_part)]

    for item in items:
        # Tự động xuống dòng nếu text quá dài (Logic đơn giản)
        if len(item) == 3:
            name, det, note = item
        else:
            name, det = item # type: ignore
            note = ""
            
        # Vẽ Name
        c.drawString(50, y, name[:45] + "..." if len(name)>45 else name)
        
        # Vẽ Details (Multi-line support basic)
        det_lines = det.split('\n')
        dy = y
        for line in det_lines:
            c.drawString(250, dy, line[:50] + "..." if len(line)>50 else line)
            dy -= 12
            
        # Vẽ Note (Guest List)
        note_lines = note.split('\n')
        ny = y
        for line in note_lines:
            c.drawString(450, ny, line[:40] + "..." if len(line)>40 else line)
            ny -= 12
            
        # Tính toán y tiếp theo dựa trên số dòng nhiều nhất
        max_lines = max(len(det_lines), len(note_lines), 1)
        row_height = max(25, max_lines * 12 + 10)
        
        # [NEW] Kiểm tra ngắt trang
        if y - row_height < 50:
            c.setFillColor(HexColor(text_color))
            c.setFont(font_name, 9)
            c.drawCentredString(width / 2, 15, f"Page {c.getPageNumber()}")
            c.showPage()
            y = height - 50
            draw_watermark()
            
            # Vẽ lại Header bảng
            c.setFillColor(HexColor(primary_color))
            c.setFont(font_bold, 12)
            c.drawString(40, y, txt['svc_details_cont'])
            y -= 25
            c.setFillColor(HexColor("#E8F5E9"))
            c.rect(40, y - 5, width - 80, 20, fill=1, stroke=0)
            c.setFillColor(HexColor(primary_color))
            c.setFont(font_bold, 10)
            c.drawString(50, y, txt['tbl_name'])
            c.drawString(250, y, txt['tbl_det'])
            c.drawString(450, y, txt['tbl_note'])
            y -= 20
            c.setFillColor(HexColor(text_color))
            c.setFont(font_name, 10)
        
        # Kẻ dòng dưới
        c.setStrokeColor(HexColor("#EEEEEE"))
        line_y = y - row_height + 15
        c.line(40, line_y, width - 40, line_y)
        y -= row_height

    # --- PHẦN 3: GUEST LIST (NẾU CÓ) ---
    g_list_content = booking_info.get('guest_list', '')
    next_section_idx = 3
    
    if g_list_content:
        if y < 100:
            c.setFillColor(HexColor(text_color))
            c.setFont(font_name, 9)
            c.drawCentredString(width / 2, 15, f"Page {c.getPageNumber()}")
            c.showPage()
            y = height - 50
            draw_watermark()
            
        y -= 20
        c.setFillColor(HexColor(primary_color))
        c.setFont(font_bold, 12)
        c.drawString(40, y, txt['guest_list'])
        y -= 20
        c.setFillColor(HexColor(text_color))
        c.setFont(font_name, 10)
        
        for line in g_list_content.split('\n'):
            if y < 50:
                c.setFillColor(HexColor(text_color))
                c.setFont(font_name, 9)
                c.drawCentredString(width / 2, 15, f"Page {c.getPageNumber()}")
                c.showPage()
                y = height - 50
                draw_watermark()
                c.setFillColor(HexColor(primary_color))
                c.setFont(font_bold, 12)
                c.drawString(40, y, txt['guest_list_cont'])
                y -= 20
                c.setFillColor(HexColor(text_color))
                c.setFont(font_name, 10)
                
            c.drawString(50, y, line)
            y -= 15
        next_section_idx = 4

    # --- PHẦN 4: INCLUDED & FOOTER ---
    if y < 150:
        c.setFillColor(HexColor(text_color))
        c.setFont(font_name, 9)
        c.drawCentredString(width / 2, 15, f"Page {c.getPageNumber()}")
        c.showPage()
        y = height - 50
        draw_watermark()
        
    y -= 20
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 12)
    roman_num = "IV" if next_section_idx == 4 else "III"
    c.drawString(40, y, f"{roman_num}. {txt['included']}")
    y -= 20
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 10)
    # c.drawString(50, y, txt['inc_1']) # Đã bỏ theo yêu cầu
    c.drawString(50, y, txt['inc_2'])
    

    # Signature
    y -= 45
    c.setFont(font_bold, 11)
    c.drawCentredString(width - 120, y, txt['confirmed_by'])
    c.setFont(font_name, 10)
    c.drawCentredString(width - 120, y - 15, comp_name)
    
    # Dấu mộc giả lập (Text)
    c.setFillColor(HexColor("#C62828"))
    c.setFont(font_bold, 14)
    c.saveState()
    c.translate(width - 120, y - 50)
    c.rotate(15)
    c.drawCentredString(0, 0, txt['signed'])
    c.restoreState()

    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 9)
    c.drawCentredString(width / 2, 15, f"{txt['page']} {c.getPageNumber()}")

    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def create_booking_cfm_docx(booking_info, company_info, lang='en'):
    """Tạo file Word Booking Confirmation"""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font # type: ignore
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Labels
    labels = {
        'en': {'title': 'BOOKING CONFIRMATION', 'greeting': f"A warm greeting from {company_info['name']}!", 'gen_info': 'I. GENERAL INFORMATION', 'svc_details': 'II. SERVICE DETAILS', 'guest_list': 'III. GUEST LIST', 'included': 'INCLUDED SERVICES'},
        'vi': {'title': 'XÁC NHẬN ĐẶT DỊCH VỤ', 'greeting': f"Lời chào trân trọng từ {company_info['name']}!", 'gen_info': 'I. THÔNG TIN CHUNG', 'svc_details': 'II. CHI TIẾT DỊCH VỤ', 'guest_list': 'III. DANH SÁCH KHÁCH', 'included': 'DỊCH VỤ BAO GỒM'}
    }
    txt = labels[lang]
    
    # Header
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(3)
    t.columns[1].width = Cm(13)
    
    if company_info['logo_b64_str']:
        try:
            logo_data = base64.b64decode(company_info['logo_b64_str'])
            image_stream = io.BytesIO(logo_data)
            t.cell(0, 0).paragraphs[0].add_run().add_picture(image_stream, width=Cm(2.5))
        except: pass
        
    cell = t.cell(0, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(company_info['name'].upper() + "\n")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(27, 94, 32)
    p.add_run(f"{company_info['address']}\nHotline: {company_info['phone']}")
    
    doc.add_paragraph("-" * 90).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    p_title = doc.add_paragraph(txt['title'])
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].bold = True; p_title.runs[0].font.size = Pt(18); p_title.runs[0].font.color.rgb = RGBColor(27, 94, 32)
    
    doc.add_paragraph(txt['greeting']).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # I. General Info
    doc.add_heading(txt['gen_info'], level=1)
    t_gen = doc.add_table(rows=0, cols=4)
    t_gen.style = 'Table Grid'
    
    cust_raw = booking_info.get('customer_info', '')
    cust_name = cust_raw.split(' - ')[0] if ' - ' in cust_raw else cust_raw
    
    def add_kv(k1, v1, k2, v2):
        r = t_gen.add_row()
        r.cells[0].text = k1; r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[1].text = str(v1)
        r.cells[2].text = k2; r.cells[2].paragraphs[0].runs[0].bold = True
        r.cells[3].text = str(v2)
        
    add_kv("Booking Code:", booking_info['code'], "Created Date:", booking_info.get('created_at', ''))
    add_kv("Customer:", cust_name, "Sales:", booking_info.get('sale_name', ''))
    add_kv("Status:", booking_info.get('status', ''), "Hotel Code:", booking_info.get('hotel_code', ''))
    
    doc.add_paragraph()
    
    # II. Service Details
    doc.add_heading(txt['svc_details'], level=1)
    t_svc = doc.add_table(rows=1, cols=2)
    t_svc.style = 'Table Grid'
    t_svc.rows[0].cells[0].text = "SERVICE NAME"
    t_svc.rows[0].cells[1].text = "DETAILS"
    t_svc.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    t_svc.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    r = t_svc.add_row()
    r.cells[0].text = booking_info['name']
    
    details = booking_info.get('details', '')
    if booking_info.get('type') == 'HOTEL':
        r_type = booking_info.get('room_type', '')
        if r_type: details += f"\nRoom Type: {r_type}"
    elif booking_info.get('type') == 'TRANS' and '[XE]' in booking_info.get('name', ''):
        # [NEW] Xử lý hiển thị chi tiết cho Booking Xe (Word)
        raw_details = details
        itinerary = booking_info['name'].replace('[XE]', '').strip()
        
        car_info = ""
        m_car = re.search(r'Xe\s*([^:]*):\s*([^|]+)', raw_details)
        if m_car: car_info = f"{m_car.group(1).strip()} - {m_car.group(2).strip()}"
            
        driver_name = ""
        driver_phone = ""
        m_drv_full = re.search(r'Tài xế:\s*([^|]+)', raw_details)
        if m_drv_full:
            drv_str = m_drv_full.group(1).strip()
            if ' - ' in drv_str:
                parts = drv_str.split(' - ')
                driver_name = parts[0]
                driver_phone = parts[1] if len(parts) > 1 else ""
            else:
                driver_name = drv_str

        lines = []
        if itinerary: lines.append(f"Lịch trình: {itinerary}" if lang == 'vi' else f"Itinerary: {itinerary}")
        if car_info: lines.append(f"Xe/Biển số: {car_info}" if lang == 'vi' else f"Car/Plate: {car_info}")
        if driver_name: lines.append(f"Tài xế: {driver_name}" if lang == 'vi' else f"Driver: {driver_name}")
        if driver_phone: lines.append(f"SĐT: {driver_phone}" if lang == 'vi' else f"Phone: {driver_phone}")
        
        if lines: details = "\n".join(lines)
    
    r.cells[1].text = details
    
    doc.add_paragraph()
    
    # III. Guest List
    g_list = booking_info.get('guest_list', '')
    if g_list:
        doc.add_heading(txt['guest_list'], level=1)
        doc.add_paragraph(g_list)
        doc.add_paragraph()
        
    # Included
    doc.add_heading(txt['included'], level=1)
    doc.add_paragraph("- Tax and Service charges.\n- 24/7 Support.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    p_sig = doc.add_paragraph("CONFIRMED BY")
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.runs[0].bold = True
    
    p_comp = doc.add_paragraph(company_info['name'])
    p_comp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ==========================================
# 4. GIAO DIỆN & LOGIC MODULES
# ==========================================

def render_notification_calendar():
    st.title("📅 Lịch Thông Báo & Nhắc Thanh Toán")
    
    # --- TỰ ĐỘNG CHẠY KIỂM TRA GỬI LẦN 2 ---
    if "auto_check_done" not in st.session_state:
        sent_count = check_and_send_due_reminders()
        if sent_count > 0:
            st.toast(f"🚀 Hệ thống vừa tự động gửi {sent_count} email nhắc hẹn đến hạn!", icon="✅")
        st.session_state.auto_check_done = True

    # Chia layout
    col_cal, col_form = st.columns([1, 1.5])

    # --- CỘT TRÁI: DANH SÁCH & CÔNG CỤ LỊCH ---
    with col_cal:
        # === 1. CÔNG CỤ CHUYỂN ĐỔI LỊCH (MỚI) ===
        with st.expander("☯️ Công cụ Chuyển đổi Âm / Dương", expanded=False):
            st.caption("Tra cứu nhanh ngày Âm/Dương lịch")
            cv_mode = st.radio("Chế độ:", ["Dương ➡ Âm", "Âm ➡ Dương"], horizontal=True, label_visibility="collapsed")
            
            if cv_mode == "Dương ➡ Âm":
                d_in = st.date_input("Chọn ngày Dương:", datetime.now(), format="DD/MM/YYYY")
                if d_in:
                    lunar_txt = convert_solar_to_lunar(d_in)
                    st.success(f"🗓️ **{lunar_txt}**")
            else:
                c_d, c_m, c_y = st.columns(3)
                l_day = c_d.number_input("Ngày", 1, 30, 1)
                l_month = c_m.number_input("Tháng", 1, 12, 1)
                l_year = c_y.number_input("Năm", 2024, 2030, datetime.now().year)
                is_leap = st.checkbox("Tháng nhuận")
                
                if st.button("Tra cứu Dương lịch"):
                    res_date = convert_lunar_to_solar(l_day, l_month, l_year, is_leap)
                    if res_date:
                        st.success(f"☀️ Ngày Dương: **{res_date.strftime('%d/%m/%Y')}**")
                        weekday_map = {0:"Thứ Hai", 1:"Thứ Ba", 2:"Thứ Tư", 3:"Thứ Năm", 4:"Thứ Sáu", 5:"Thứ Bảy", 6:"Chủ Nhật"}
                        st.caption(f"({weekday_map[res_date.weekday()]})")
                    else:
                        st.error("Ngày âm lịch không hợp lệ!")

        st.divider()
        
        # === 2. DANH SÁCH LỊCH HẸN ===
        st.markdown("### 🗓️ Lịch sắp tới")
        with st.container(border=True):
            # Lấy danh sách nhắc hẹn
            upcoming = run_query("SELECT * FROM payment_reminders WHERE status != 'sent_2' ORDER BY due_date ASC")
            
            if upcoming:
                for item in upcoming:
                    try:
                        d_obj = datetime.strptime(item['due_date'], '%Y-%m-%d %H:%M:%S')
                    except:
                        d_obj = datetime.strptime(item['due_date'], '%Y-%m-%d')
                        
                    days_left = (d_obj.date() - datetime.now().date()).days
                    
                    # Format ngày tháng năm
                    date_display = d_obj.strftime('%H:%M %d/%m/%Y')
                    lunar_display = convert_solar_to_lunar(d_obj).replace(" (Âm lịch)", "")
                    
                    color = "orange" if days_left == 0 else "green" if days_left > 0 else "red"
                    icon = "🔔" if days_left == 0 else "📅"
                    
                    with st.expander(f"{icon} {date_display} (Âm: {lunar_display}) | {item['ref_code']}"):
                        st.write(f"**Nội dung:** {item['content']}")
                        st.write(f"**Người nhận:** {item['receiver_email']}")
                        st.write(f"**Số tiền:** {format_vnd(item['amount'])} VND")
                        
                        status_txt = "Chờ gửi Lần 1" if item['status'] == 'pending' else "Đã gửi Lần 1, chờ Lần 2"
                        st.caption(f"Trạng thái: {status_txt}")
                        
                        if st.button("🗑️ Xóa", key=f"del_cal_{item['id']}"):
                            run_query("DELETE FROM payment_reminders WHERE id=?", (item['id'],), commit=True)
                            st.rerun()
            else:
                st.info("Không có lịch nhắc nào sắp tới.")

    # --- CỘT PHẢI: FORM TẠO (GIỮ NGUYÊN) ---
    with col_form:
        st.markdown("### ✍️ Tạo yêu cầu thanh toán mới")
        with st.container(border=True):
            # 1. Lấy dữ liệu Booking/Tour để liên kết
            # [UPDATED] Phân quyền xem Booking/Tour
            user_info = st.session_state.get("user_info", {})
            u_role = user_info.get('role')
            u_name = user_info.get('name')
            
            tour_q = "SELECT tour_code, tour_name FROM tours WHERE status='running'"
            tour_p = []
            bk_q = "SELECT code, name FROM service_bookings WHERE status='active'"
            bk_p = []
            
            if u_role not in ['admin', 'admin_f1']:
                tour_q += " AND sale_name=?"
                tour_p.append(u_name)
                bk_q += " AND sale_name=?"
                bk_p.append(u_name)
                
            tours = run_query(tour_q, tuple(tour_p))
            bookings = run_query(bk_q, tuple(bk_p))
            
            opts = ["-- Chọn mã liên kết --"]
            if tours: opts += [f"TOUR | {t['tour_code']} | {t['tour_name']}" for t in tours]
            if bookings: opts += [f"BOOK | {b['code']} | {b['name']}" for b in bookings]
            
            sel_ref = st.selectbox("Liên kết với Booking/Tour:", opts, key="notif_ref")
            
            # Tự động điền thông tin nếu chọn mã
            ref_code = ""
            ref_name = ""
            if sel_ref != "-- Chọn mã liên kết --":
                parts = sel_ref.split(" | ")
                ref_code = parts[1]
                ref_name = parts[2]

            c1, c2 = st.columns(2)
            
            # [CẬP NHẬT] Nhập số tiền có định dạng VND
            if "req_amount_val" not in st.session_state: st.session_state.req_amount_val = ""
            def fmt_req_amount():
                val = st.session_state.req_amount_val
                try:
                    v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    st.session_state.req_amount_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                except: pass

            amount_input = c1.text_input("Số tiền yêu cầu:", key="req_amount_val", on_change=fmt_req_amount, help="Nhập số tiền (VD: 1000000)")
            try: amount = float(amount_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
            except: amount = 0.0
            
            # [CẬP NHẬT] Chọn ngày và giờ với format DD/MM/YYYY
            with c2:
                c_d, c_t = st.columns(2)
                due_date = c_d.date_input("Ngày hẹn (Lần 2):", min_value=datetime.now(), format="DD/MM/YYYY", key="notif_date")
                due_time = c_t.time_input("Giờ hẹn:", value=datetime.now().time(), key="notif_time")
                due_datetime = datetime.combine(due_date, due_time)
            
            # Hiển thị ngày Âm lịch tương ứng ngay dưới để tiện theo dõi
            if due_date:
                st.caption(f"🗓️ Tương ứng Âm lịch: {convert_solar_to_lunar(due_date)}")
            
            # [CẬP NHẬT] Thông tin người gửi và CC
            current_user_name = st.session_state.user_info.get('name', '')
            sender_name = st.text_input("Người gửi (Hiển thị trong mail):", value=current_user_name, disabled=True)
            cc_email = st.text_input("CC Email (cách nhau dấu phẩy):", placeholder="boss@gmail.com, ketoan@gmail.com", key="notif_cc")
            
            # Email mặc định lấy từ secrets
            def_email = ""
            try: def_email = st.secrets["email"].get("receiver_default", "")
            except: pass
            
            receiver = st.text_input("Email người nhận thông báo:", value=def_email, help="Email của Kế toán hoặc Khách hàng", key="notif_receiver")
            content = st.text_area("Nội dung yêu cầu thanh toán:", height=100, placeholder="VD: Yêu cầu thanh toán đợt 1 cho đoàn...", key="notif_content")
            
            # [NEW] Bank Info Inputs
            st.markdown("##### 🏦 Thông tin chuyển khoản")
            c_b1, c_b2 = st.columns(2)
            bank_name = c_b1.text_input("Tên Ngân Hàng", placeholder="VD: Techcombank", key="notif_bank_name")
            bank_acc = c_b2.text_input("Số Tài Khoản", placeholder="VD: 1903...", key="notif_bank_acc")
            bank_holder = st.text_input("Chủ Tài Khoản", placeholder="VD: NGUYEN VAN A", key="notif_bank_holder")

            st.info("ℹ️ **Cơ chế:** Khi bấm nút dưới, hệ thống sẽ **GỬI NGAY 1 EMAIL** cho người nhận. Đến ngày hẹn ở trên, hệ thống sẽ **GỬI TIẾP 1 EMAIL NỮA**.")

            if st.button("🚀 Lưu & Gửi thông báo ngay", type="primary", use_container_width=True):
                if ref_code and receiver and content:
                    with st.spinner("Đang gửi email lần 1..."):
                        # 1. Gửi Email Lần 1 Ngay lập tức
                        bank_html = ""
                        if bank_name and bank_acc:
                            bank_html = f"""
                            <div style="background-color: #f8f9fa; padding: 15px; border-radius: 5px; margin: 15px 0;">
                                <h4 style="margin-top: 0;">🏦 THÔNG TIN CHUYỂN KHOẢN</h4>
                                <p><strong>Ngân hàng:</strong> {bank_name}</p>
                                <p><strong>Số tài khoản:</strong> {bank_acc}</p>
                                <p><strong>Chủ tài khoản:</strong> {bank_holder}</p>
                            </div>
                            """

                        subj = f"📢 [THÔNG BÁO] Yêu cầu thanh toán - {ref_code}"
                        html_body = f"""
                        <h3>📢 YÊU CẦU THANH TOÁN (LẦN 1)</h3>
                        <p>Kính gửi,</p>
                        <p>Chúng tôi gửi thông báo thanh toán cho dịch vụ <strong>{ref_name}</strong> (Mã: {ref_code}).</p>
                        <p><strong>Số tiền:</strong> {format_vnd(amount)} VND</p>
                        <p><strong>Nội dung:</strong> {content}</p>
                        {bank_html}
                        <p>Hệ thống sẽ gửi nhắc nhở lại vào lúc: <strong>{due_datetime.strftime('%H:%M %d/%m/%Y')}</strong>.</p>
                        <hr>
                        <p>Trân trọng,<br><strong>{sender_name}</strong><br><small>Bali Tourist Automated System</small></p>
                        """
                        
                        ok, msg = send_email_notification(receiver, subj, html_body, cc_emails=cc_email)
                        if ok:
                            # 2. Lưu vào DB để hẹn giờ gửi lần 2
                            run_query("""INSERT INTO payment_reminders 
                                (ref_code, ref_name, amount, due_date, receiver_email, content, status, created_at, cc_email, sender_name, bank_name, bank_account, bank_holder)
                                VALUES (?, ?, ?, ?, ?, ?, 'sent_1', ?, ?, ?, ?, ?, ?)""", 
                                (ref_code, ref_name, amount, due_datetime.strftime('%Y-%m-%d %H:%M:%S'), receiver, content, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), cc_email, sender_name, bank_name, bank_acc, bank_holder),
                                commit=True)
                            
                            # [FIX] Reset form fields
                            keys_to_reset = ["req_amount_val", "notif_receiver", "notif_content", "notif_bank_name", "notif_bank_acc", "notif_bank_holder", "notif_cc", "notif_date", "notif_time", "notif_ref"]
                            for k in keys_to_reset:
                                if k in st.session_state: del st.session_state[k]

                            st.success(f"✅ Đã gửi email lần 1 và lên lịch nhắc lần 2 vào ngày {due_date.strftime('%d/%m/%Y')}!"); time.sleep(1); st.rerun()
                        else: st.error(msg)
                else:
                    st.warning("Vui lòng chọn Mã liên kết, nhập Email và Nội dung.")

def render_dashboard():
    st.title("🏠 Trang Chủ - Tổng Quan Kinh Doanh")
    
    # User context
    user_info = st.session_state.get("user_info", {})
    role = user_info.get('role')
    username = user_info.get('name')
    
    # Time context
    now = datetime.now()
    current_month = now.month
    current_year = now.year
    
    st.markdown(f"### 📅 Số liệu tháng {current_month}/{current_year}")
    
    # Data fetching
    # 1. Tours
    tour_query = "SELECT * FROM tours WHERE status != 'deleted'"
    tour_params = []
    if role == 'sale':
        tour_query += " AND sale_name=?"
        tour_params.append(username)
    tours = run_query(tour_query, tuple(tour_params))
    
    # 2. Bookings
    bk_query = "SELECT * FROM service_bookings WHERE status != 'deleted'"
    bk_params = []
    if role == 'sale':
        bk_query += " AND sale_name=?"
        bk_params.append(username)
    bookings = run_query(bk_query, tuple(bk_params))
    
    # 3. Costs (for tours)
    all_items = run_query("SELECT tour_id, item_type, total_amount FROM tour_items")
    items_map = {}
    if all_items:
        for item in all_items:
            tid = item['tour_id']
            itype = item['item_type']
            amt = item['total_amount'] or 0
            if tid not in items_map: items_map[tid] = {'EST': 0, 'ACT': 0}
            items_map[tid][itype] += amt

    # Processing
    total_tour_rev = 0
    total_tour_profit = 0
    count_tours = 0
    tours_in_month = []
    
    total_bk_rev = 0
    total_bk_profit = 0
    count_bks = 0
    bks_in_month = []
    
    # Process Tours
    if tours:
        for t in tours:
            t = dict(t)
            try:
                s_date = datetime.strptime(t['start_date'], '%d/%m/%Y')
                if s_date.month == current_month and s_date.year == current_year:
                    count_tours += 1
                    
                    final_price = float(t.get('final_tour_price', 0) or 0)
                    child_price = float(t.get('child_price', 0) or 0)
                    final_qty = float(t.get('final_qty', 0) or 0)
                    child_qty = float(t.get('child_qty', 0) or 0)
                    if final_qty == 0: final_qty = float(t.get('guest_count', 1))
                    
                    rev = (final_price * final_qty) + (child_price * child_qty)
                    
                    costs = items_map.get(t['id'], {'EST': 0, 'ACT': 0})
                    est_cost = costs['EST']; act_cost = costs['ACT']
                    
                    if rev == 0:
                        p_pct = t.get('est_profit_percent', 0) or 0
                        t_pct = t.get('est_tax_percent', 0) or 0
                        profit_est_val = est_cost * (p_pct/100)
                        rev = (est_cost + profit_est_val) * (1 + t_pct/100)
                    
                    t_pct = t.get('est_tax_percent', 0) or 0
                    net_rev = rev / (1 + t_pct/100) if (1 + t_pct/100) != 0 else rev
                    prof = net_rev - act_cost
                    
                    total_tour_rev += rev; total_tour_profit += prof
                    
                    t_display = dict(t); t_display['revenue'] = rev; t_display['profit'] = prof
                    tours_in_month.append(t_display)
            except: pass

    # Process Bookings
    if bookings:
        for b in bookings:
            try:
                c_date = datetime.strptime(str(b['created_at']).split(' ')[0], '%Y-%m-%d')
                if c_date.month == current_month and c_date.year == current_year:
                    count_bks += 1
                    rev = float(b['selling_price'] or 0); prof = float(b['profit'] or 0)
                    total_bk_rev += rev; total_bk_profit += prof
                    b_display = dict(b); b_display['revenue'] = rev; b_display['profit'] = prof
                    bks_in_month.append(b_display)
            except: pass

    # Display Metrics
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Tổng Doanh Thu", format_vnd(total_tour_rev + total_bk_rev) + " VND")
    m2.metric("Tổng Lợi Nhuận", format_vnd(total_tour_profit + total_bk_profit) + " VND")
    m3.metric("Số lượng Tour", count_tours)
    m4.metric("Số lượng Booking", count_bks)
    
    st.divider()
    c_left, c_right = st.columns(2)
    with c_left:
        st.subheader("📦 Tour trong tháng")
        if tours_in_month:
            df_t = pd.DataFrame(tours_in_month)[['start_date', 'tour_name', 'revenue', 'profit']]
            df_t['revenue'] = df_t['revenue'].apply(lambda x: format_vnd(x) + " VND")
            df_t['profit'] = df_t['profit'].apply(lambda x: format_vnd(x) + " VND")
            st.dataframe(df_t, column_config={"start_date": "Ngày đi", "tour_name": "Tên đoàn", "revenue": "Doanh thu", "profit": "Lợi nhuận (TT)"}, use_container_width=True, hide_index=True)
        else: st.info("Không có tour nào.")
            
    with c_right:
        st.subheader("🔖 Booking trong tháng")
        if bks_in_month:
            df_b = pd.DataFrame(bks_in_month)[['created_at', 'name', 'revenue', 'profit']]
            df_b['revenue'] = df_b['revenue'].apply(lambda x: format_vnd(x) + " VND")
            df_b['profit'] = df_b['profit'].apply(lambda x: format_vnd(x) + " VND")
            st.dataframe(df_b, column_config={"created_at": "Ngày tạo", "name": "Tên dịch vụ", "revenue": "Doanh thu", "profit": "Lợi nhuận"}, use_container_width=True, hide_index=True)
        else: st.info("Không có booking nào.")

# ==========================================
# 4. GIAO DIỆN & LOGIC MODULES
# ==========================================

def render_login_page(comp):
    # 1. --- PHẦN STYLE (CSS) - LỚP VỎ ĐẸP NHƯ HTML ---
    st.markdown("""
        <style>
            /* Nhúng Font chữ Google (Roboto) cho hiện đại */
            @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@400;500;700&display=swap');
            
            html, body, [class*="css"] {
                font-family: 'Roboto', sans-serif;
            }

            /* Nền trang web */
            .stApp {
                background-color: #f0f2f5; /* Màu xám nhạt của Facebook/Bootstrap */
                background-image: url("https://img.freepik.com/free-vector/white-abstract-background-design_23-2148825582.jpg");
                background-size: cover;
            }

            /* Ẩn các thành phần thừa của Streamlit */
            header {visibility: hidden;}
            section[data-testid="stSidebar"] {display: none;}
            
            /* TẠO KHUNG LOGIN (CARD) */
            div[data-testid="stForm"] {
                background: rgba(255, 255, 255, 0.95); /* Trắng mờ */
                padding: 50px 40px;
                border-radius: 20px;
                box-shadow: 0 15px 35px rgba(0,0,0,0.1); /* Đổ bóng sâu 3D */
                width: 100%;
                border: 1px solid white;
            }

            /* CHỈNH Ô NHẬP LIỆU (INPUT) GIỐNG BOOTSTRAP */
            div[data-testid="stTextInput"] {
                margin-bottom: 15px;
            }
            div[data-testid="stTextInput"] label {
                font-size: 14px;
                color: #555;
                font-weight: 500;
                margin-bottom: 8px;
            }
            div[data-testid="stTextInput"] input {
                border-radius: 8px !important;
                padding: 12px 15px !important;
                border: 1px solid #ced4da !important; /* Viền xám chuẩn Bootstrap */
                color: #495057 !important;
                background-color: #fff !important;
                transition: all 0.2s;
            }
            /* Hiệu ứng khi bấm vào ô nhập */
            div[data-testid="stTextInput"] input:focus {
                border-color: #0e0259 !important; 
                box-shadow: 0 0 0 0.2rem rgba(14, 2, 89, 0.25) !important; /* Hào quang xanh */
            }

            /* CHỈNH NÚT BẤM (BUTTON) */
            div[data-testid="stFormSubmitButton"] button {
                width: 100%;
                background-color: #0e0259 !important;
                color: white !important;
                padding: 12px 20px !important;
                font-size: 16px !important;
                border-radius: 30px !important; /* Bo tròn kiểu viên thuốc */
                border: none !important;
                font-weight: 700 !important;
                letter-spacing: 1px;
                margin-top: 10px;
                box-shadow: 0 4px 10px rgba(14, 2, 89, 0.3);
            }
            div[data-testid="stFormSubmitButton"] button:hover {
                background-color: #1a0b7e !important; /* Sáng hơn chút khi di chuột */
                transform: translateY(-2px); /* Nhẹ nhàng bay lên */
                box-shadow: 0 6px 12px rgba(14, 2, 89, 0.4);
            }

            /* Tinh chỉnh Tab */
            div[data-baseweb="tab-list"] {
                background-color: transparent;
                margin-bottom: 20px;
                border-bottom: 2px solid #e9ecef;
            }
            button[data-baseweb="tab"] {
                font-weight: bold !important;
                font-size: 16px;
            }
        </style>
    """, unsafe_allow_html=True)

    # 2. --- PHẦN LOGIC (PYTHON) - BỘ NÃO ---
    
    # Chia cột để căn giữa màn hình (Cột trái trống - Cột giữa Login - Cột phải trống)
    col1, col2, col3 = st.columns([1, 1, 1]) 

    with col2:
        # Khoảng cách phía trên
        st.write("") 
        st.write("") 
        
        # LOGO CÔNG TY
        if comp['logo_b64_str']:
             st.markdown(f'''
                <div style="text-align: center; margin-bottom: 20px;">
                    <img src="data:image/png;base64,{comp["logo_b64_str"]}" style="width: 100px; height: auto;">
                    <h2 style="color: #0e0259; margin-top: 10px; font-weight: 700;">{comp['name']}</h2>
                    <p style="color: #6c757d; font-size: 14px;">Welcome back! Please login to your account.</p>
                </div>
            ''', unsafe_allow_html=True)
        else:
             st.markdown(f'''<h2 style="text-align: center; color: #0e0259;">{comp['name']}</h2>''', unsafe_allow_html=True)

        # TAB CHUYỂN ĐỔI
        tab_login, tab_reg = st.tabs(["ĐĂNG NHẬP", "ĐĂNG KÝ"])

        with tab_login:
            # Dùng st.form để gom nhóm input và nút bấm
            with st.form("login_form"):
                # Input Python (đã được CSS ở trên làm đẹp)
                u = st.text_input("Username", placeholder="Nhập tên đăng nhập")
                p = st.text_input("Password", type="password", placeholder="Nhập mật khẩu")
                
                st.write("") # Khoảng trắng nhỏ
                
                # Nút bấm Python
                submitted = st.form_submit_button("LOGIN")
                
                # Xử lý Logic khi bấm nút
                if submitted:
                    if not u or not p:
                        st.error("⚠️ Vui lòng nhập đầy đủ thông tin!")
                    else:
                        pw_hash = hash_pass(p)
                        df_users = load_table('users')
                        
                        # Kiểm tra trong Database
                        if not df_users.empty:
                            mask = (df_users['username'] == u) & (df_users['password'] == pw_hash)
                            user_found = df_users.loc[mask]
                            
                            if not user_found.empty:
                                if user_found.iloc[0]['status'] == 'approved':
                                    st.session_state.logged_in = True
                                    st.session_state.user_info = {
                                        "name": user_found.iloc[0]['username'],
                                        "role": user_found.iloc[0]['role']
                                    }
                                    st.success("Login thành công!")
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    st.error("🚫 Tài khoản đang chờ duyệt!")
                            else:
                                st.error("❌ Sai tên đăng nhập hoặc mật khẩu!")
                        else:
                            st.error("⚠️ Lỗi kết nối dữ liệu!")

        with tab_reg:
            with st.form("reg_form"):
                st.markdown("<p style='text-align: center; color: #666;'>Tạo tài khoản mới cho nhân viên</p>", unsafe_allow_html=True)
                nu = st.text_input("New Username", placeholder="Tên đăng nhập mong muốn")
                np = st.text_input("New Password", type="password", placeholder="Mật khẩu mong muốn")
                
                if st.form_submit_button("REGISTER"):
                    if not nu or not np:
                        st.warning("Vui lòng nhập đủ thông tin!")
                    else:
                        try:
                            conn = get_connection()
                            exist = run_query("SELECT id FROM users WHERE username=?", (nu,), fetch_one=True)
                            if exist:
                                st.error("Tên đăng nhập đã tồn tại!")
                            else:
                                add_row_to_table('users', {'username': nu, 'password': hash_pass(np), 'role': 'user', 'status': 'pending'})
                                st.success("✅ Đăng ký thành công! Vui lòng chờ Admin duyệt.")
                        except Exception as e:
                            st.error(f"Lỗi: {e}")

        st.markdown("<div style='text-align: center; margin-top: 30px; color: #adb5bd; font-size: 12px;'>© 2026 Bali Tourist Platform</div>", unsafe_allow_html=True)

def render_admin_notifications():
    st.divider()
    st.markdown("### 🔔 Trung Tâm Thông Báo & Phê Duyệt")
    
    # --- LẤY DỮ LIỆU CẦN DUYỆT ---
    pending_projs = run_query("SELECT * FROM projects WHERE pending_name IS NOT NULL AND pending_name != ''")
    pending_tours = run_query("SELECT * FROM tours WHERE pending_name IS NOT NULL AND pending_name != ''")
    del_tours = run_query("SELECT * FROM tours WHERE request_delete=1")
    req_edit_tours = run_query("SELECT * FROM tours WHERE request_edit_act=1")
    pending_users = run_query("SELECT * FROM users WHERE role='user' AND status='pending'")
    req_invoices = run_query("SELECT * FROM invoices WHERE request_edit=1 AND status='active'")
    
    has_requests = False

    # 1. DUYỆT ĐỔI TÊN DỰ ÁN
    if pending_projs:
        has_requests = True
        st.markdown(f"#### 📝 Đổi tên Dự án ({len(pending_projs)})")
        for p in pending_projs:
            with st.container(border=True):
                st.markdown(f"**Dự án:** `{p['project_name']}` ➡ <span style='color:green'><b>`{p['pending_name']}`</b></span>", unsafe_allow_html=True) # type: ignore
                c_app, c_rej = st.columns(2)
                if c_app.button("✔ Duyệt", key=f"app_ren_{p['id']}", type="primary"): # type: ignore
                    run_query("UPDATE projects SET project_name=?, pending_name=NULL WHERE id=?", (p['pending_name'], p['id']), commit=True) # type: ignore
                    st.rerun()
                if c_rej.button("✖ Hủy", key=f"rej_ren_{p['id']}"): # type: ignore
                    run_query("UPDATE projects SET pending_name=NULL WHERE id=?", (p['id'],), commit=True) # type: ignore
                    st.rerun()

    # 2. DUYỆT ĐỔI TÊN TOUR
    if pending_tours:
        has_requests = True
        st.markdown(f"#### 📦 Đổi tên Tour ({len(pending_tours)})")
        for t in pending_tours:
            with st.container(border=True):
                st.markdown(f"**Tour:** `{t['tour_name']}` ➡ <span style='color:green'><b>`{t['pending_name']}`</b></span>", unsafe_allow_html=True) # type: ignore
                c_app, c_rej = st.columns(2)
                if c_app.button("✔ Duyệt", key=f"app_ren_t_{t['id']}", type="primary"): # type: ignore
                    run_query("UPDATE tours SET tour_name=?, pending_name=NULL WHERE id=?", (t['pending_name'], t['id']), commit=True) # type: ignore
                    st.rerun()
                if c_rej.button("✖ Hủy", key=f"rej_ren_t_{t['id']}"): # type: ignore
                    run_query("UPDATE tours SET pending_name=NULL WHERE id=?", (t['id'],), commit=True) # type: ignore
                    st.rerun()

    # 3. DUYỆT XÓA TOUR
    if del_tours:
        has_requests = True
        st.markdown(f"#### <span style='color:red;'>🗑️ Xóa Tour ({len(del_tours)})</span>", unsafe_allow_html=True)
        for t in del_tours:
            with st.container(border=True):
                st.markdown(f"❌ Yêu cầu xóa Tour: **{t['tour_name']}**") # type: ignore
                c_app, c_rej = st.columns(2)
                if c_app.button("✔ Duyệt xóa", key=f"app_del_t_{t['id']}", type="primary"): # type: ignore
                    run_query("UPDATE tours SET request_delete=2 WHERE id=?", (t['id'],), commit=True) # type: ignore
                    st.success("Đã duyệt! Chờ người dùng xác nhận."); time.sleep(1); st.rerun()
                if c_rej.button("✖ Từ chối", key=f"rej_del_t_{t['id']}"): # type: ignore
                    run_query("UPDATE tours SET request_delete=0 WHERE id=?", (t['id'],), commit=True) # type: ignore
                    st.rerun()

    # 4. DUYỆT SỬA QUYẾT TOÁN (MỚI)
    if req_edit_tours:
        has_requests = True
        st.markdown(f"#### 💸 Sửa Quyết toán ({len(req_edit_tours)})")
        for t in req_edit_tours:
            with st.container(border=True):
                st.write(f"Tour: **{t['tour_name']}**") # type: ignore
                c1, c2 = st.columns(2)
                if c1.button("✔ Duyệt", key=f"app_edit_act_{t['id']}"): # type: ignore
                    run_query("UPDATE tours SET request_edit_act=2 WHERE id=?", (t['id'],), commit=True); st.rerun() # type: ignore
                if c2.button("✖ Từ chối", key=f"rej_edit_act_{t['id']}"): # type: ignore
                    run_query("UPDATE tours SET request_edit_act=0 WHERE id=?", (t['id'],), commit=True); st.rerun() # type: ignore

    # 5. DUYỆT USER
    if pending_users:
        has_requests = True
        st.markdown(f"#### 👤 Đăng ký mới ({len(pending_users)})")
        for u in pending_users:
            with st.container(border=True):
                st.write(f"User: **{u['username']}**") # type: ignore
                c1, c2 = st.columns(2)
                if c1.button("✔ Duyệt", key=f"app_user_{u['id']}"): # type: ignore
                    run_query("UPDATE users SET status='approved' WHERE id=?", (u['id'],), commit=True) # type: ignore
                    st.rerun()
                if c2.button("✖ Xóa", key=f"del_user_{u['id']}"): # type: ignore
                    run_query("DELETE FROM users WHERE id=?", (u['id'],), commit=True) # type: ignore
                    st.rerun()

    # 6. DUYỆT SỬA GIÁ HÓA ĐƠN
    if req_invoices:
        has_requests = True
        st.markdown(f"#### 💰 Sửa giá Hóa đơn ({len(req_invoices)})")
        for r in req_invoices:
            with st.container(border=True):
                st.info(f"HĐ: {r['invoice_number']} | Tiền: {format_vnd(r['total_amount'])}") # type: ignore
                c1, c2 = st.columns(2)
                if c1.button("✔ Duyệt", key=f"app_inv_{r['id']}"): # type: ignore
                    run_query("UPDATE invoices SET edit_count=0, request_edit=0 WHERE id=?", (r['id'],), commit=True) # type: ignore
                    st.success("Đã duyệt!"); time.sleep(0.5); st.rerun()
                if c2.button("✖ Từ chối", key=f"rej_inv_{r['id']}"): # type: ignore
                    run_query("UPDATE invoices SET request_edit=0 WHERE id=?", (r['id'],), commit=True) # type: ignore
                    st.rerun()

    if not has_requests:
        st.success("✅ Hiện không có yêu cầu nào cần duyệt.")

def render_admin_panel(comp):
    with st.expander("⚙️ Admin Panel", expanded=False):
        st.caption("Cập nhật thông tin Công ty")
        with st.form("comp_update"):
            cn = st.text_input("Tên", value=comp['name'])
            ca = st.text_input("Địa chỉ", value=comp['address'])
            cp = st.text_input("Mã Số Thuế", value=comp['phone'])
            ul = st.file_uploader("Logo", type=['png','jpg'])
            if st.form_submit_button("Lưu"):
                update_company_info(cn, ca, cp, ul.read() if ul else None)
                st.success("Xong!"); time.sleep(0.5); st.rerun()
        
        # Chỉ admin chính mới thấy mục xóa
        if (st.session_state.user_info or {}).get('role') == 'admin':
            st.divider()
            st.markdown("##### 🗑️ Quản lý dữ liệu")
            
            c1, c2 = st.columns(2)
            with c1:
                if st.button("Xóa Hóa Đơn", use_container_width=True, help="Xóa TOÀN BỘ dữ liệu Hóa đơn & UNC"):
                    run_query("DELETE FROM invoices", commit=True)
                    run_query("DELETE FROM sqlite_sequence WHERE name='invoices'", commit=True)
                    if os.path.exists(UPLOAD_FOLDER):
                        for f in os.listdir(UPLOAD_FOLDER):
                            if "UNC" not in f and "converted" not in f: 
                                    try: os.remove(os.path.join(UPLOAD_FOLDER, f))
                                    except: pass
                    st.toast("Đã xóa sạch Hóa Đơn!"); time.sleep(1); st.rerun()
                
                if st.button("Xóa Tour", use_container_width=True, help="Xóa TOÀN BỘ dữ liệu Tour (Dự toán và Quyết toán)"):
                    run_query("DELETE FROM tours", commit=True)
                    run_query("DELETE FROM tour_items", commit=True)
                    run_query("DELETE FROM sqlite_sequence WHERE name='tours'", commit=True)
                    run_query("DELETE FROM sqlite_sequence WHERE name='tour_items'", commit=True)
                    st.toast("Đã xóa sạch dữ liệu Tour!"); time.sleep(1); st.rerun()
            
            with c2:
                if st.button("Xóa Booking", use_container_width=True, help="Xóa TOÀN BỘ dữ liệu Booking dịch vụ"):
                    run_query("DELETE FROM service_bookings", commit=True)
                    run_query("DELETE FROM sqlite_sequence WHERE name='service_bookings'", commit=True)
                    st.toast("Đã xóa sạch Booking!"); time.sleep(1); st.rerun()
                
                if st.button("Xóa Khách Hàng", use_container_width=True, help="Xóa TOÀN BỘ dữ liệu Khách hàng"):
                    run_query("DELETE FROM customers", commit=True); run_query("DELETE FROM sqlite_sequence WHERE name='customers'", commit=True)
                    st.toast("Đã xóa sạch Khách hàng!"); time.sleep(1); st.rerun()

            with st.popover("💥 XÓA TOÀN BỘ DỮ LIỆU 💥", use_container_width=True):
                st.error("CẢNH BÁO CỰC KỲ NGUY HIỂM!")
                st.warning("Hành động này sẽ **XÓA SẠCH TOÀN BỘ** dữ liệu kinh doanh (Hóa đơn, Tour, Booking, Khách hàng...). Dữ liệu người dùng và thông tin công ty sẽ được giữ lại. Hành động này không thể hoàn tác.")
                st.warning("Chỉ thực hiện khi bạn muốn bắt đầu lại từ đầu. Bạn có chắc chắn không?")
                if st.button("CÓ, TÔI HIỂU RỦI RO VÀ MUỐN XÓA TẤT CẢ", type="primary"):
                    TABLES_TO_DELETE = [
                        'invoices', 'projects', 'project_links', 'service_bookings', 
                        'customers', 'tours', 'tour_items', 'ocr_learning',
                        'transaction_history',
                        'flight_tickets', 'flight_groups', 'flight_group_links'
                    ]
                    with st.spinner("Đang dọn dẹp hệ thống..."):
                        for table in TABLES_TO_DELETE:
                            run_query(f"DELETE FROM {table}", commit=True)
                            run_query(f"DELETE FROM sqlite_sequence WHERE name='{table}'", commit=True)
                        if os.path.exists(UPLOAD_FOLDER):
                            for f in os.listdir(UPLOAD_FOLDER):
                                try: os.remove(os.path.join(UPLOAD_FOLDER, f))
                                except: pass
                    st.success("Đã xóa toàn bộ dữ liệu kinh doanh và các file đã upload!")
                    time.sleep(2); st.rerun()

        with st.popover("🔄 Đồng bộ lên Google Sheet", use_container_width=True):
            st.warning("⚠️ Hành động này sẽ **ghi đè toàn bộ** dữ liệu trên Google Sheet bằng dữ liệu hiện tại trên máy của bạn. Bạn có chắc chắn không?")
            if st.button("Có, tôi muốn đồng bộ ngay", type="primary"):
                sync_all_data_to_gsheet()

def render_sidebar(comp):
    with st.sidebar:
        if comp['logo_b64_str']: st.markdown(f'<div style="text-align:center; margin-bottom:20px;"><img src="data:image/png;base64,{comp["logo_b64_str"]}" width="120" style="border-radius:10px;"></div>', unsafe_allow_html=True)
        
        user_info = st.session_state.get("user_info")
        if user_info and isinstance(user_info, dict):
            st.success(f"Xin chào **{user_info.get('name', 'User')}** 👋")
        else:
            st.session_state.logged_in = False
            st.rerun()
        
        st.markdown("### 🗂️ Phân Hệ Quản Lý")
        module = st.selectbox("Chọn chức năng:", ["🏠 Trang Chủ", "📅 Lịch Thông Báo", "🔖 Quản Lý Booking", "💰 Kiểm Soát Chi Phí", "💳 Quản Lý Công Nợ", "📦 Quản Lý Tour ", "🧾 Quản Lý Hóa Đơn", "🤝 Quản Lý Khách Hàng", "👥 Quản Lý Nhân Sự", "🔍 Tra cứu thông tin"], label_visibility="collapsed")
        
        menu = None
        if module == "💰 Kiểm Soát Chi Phí":
            menu = st.radio("Menu", ["1. Nhập Hóa Đơn", "2. Báo Cáo Tổng Hợp"])
        
        if st.session_state.user_info and st.session_state.user_info.get('role') in ['admin', 'admin_f1']:
            render_admin_notifications()

        st.divider()

        if st.session_state.user_info and st.session_state.user_info.get('role') in ['admin', 'admin_f1']:
            render_admin_panel(comp)

        if st.button("Đăng xuất", use_container_width=True):
            st.session_state.logged_in = False
            st.rerun()
        with st.popover("🔐 Đổi mật khẩu", use_container_width=True):
            st.markdown("##### Cập nhật mật khẩu")
            with st.form("change_pass"):
                op = st.text_input("Mật khẩu hiện tại", type="password")
                new_p = st.text_input("Mật khẩu mới", type="password")
                cp = st.text_input("Xác nhận mật khẩu mới", type="password")
                if st.form_submit_button("Lưu thay đổi"):
                    c_user = (st.session_state.user_info or {}).get('name', '')
                    db_u = run_query("SELECT * FROM users WHERE username=?", (c_user,), fetch_one=True)
                    if isinstance(db_u, sqlite3.Row) and db_u['password'] == hash_pass(op): # type: ignore
                        if new_p and new_p == cp:
                            run_query("UPDATE users SET password=? WHERE username=?", (hash_pass(new_p), c_user), commit=True)
                            st.success("Đổi mật khẩu thành công! Đăng nhập lại nhé.")
                            time.sleep(1)
                            st.session_state.logged_in = False
                            st.rerun()
                        else:
                            st.error("Mật khẩu mới không khớp!")
                    else:
                        st.error("Mật khẩu cũ sai rồi!")

        # --- KIỂM TRA KẾT NỐI GOOGLE (DEBUG) ---
        st.divider() # type: ignore
        with st.expander("🔌 Kiểm tra kết nối Google"):
            if st.button("Test Kết Nối Ngay", use_container_width=True):
                try:
                    with st.spinner("Đang kết nối Google API..."):
                        gc = get_gspread_client()
                        sh = gc.open_by_key(SPREADSHEET_ID)
                        st.success(f"✅ Sheet OK: {sh.title}")
                        drive = get_drive_service()
                        st.success(f"✅ Drive OK (ID: ...{DRIVE_FOLDER_ID[-5:]})")
                except Exception as e:
                    st.error(f"❌ Lỗi: {str(e)}")
                    st.info("💡 Gợi ý: Kiểm tra file service_account.json hoặc quyền chia sẻ của Sheet/Folder.")
    return module, menu

# --- HÀM HIỂN THỊ SO SÁNH CHI PHÍ (UNC vs HÓA ĐƠN) ---
def render_cost_comparison(code):
    # Lấy tất cả hóa đơn/UNC theo mã
    docs = run_query("SELECT * FROM invoices WHERE cost_code=? AND status='active'", (code,))
    if not docs:
        st.info("Chưa có chứng từ nào liên kết.")
        return 0

    df = pd.DataFrame([dict(r) for r in docs])
    
    # Lọc chi phí đầu vào (IN)
    df_in = df.loc[df['type'] == 'IN'].copy() # type: ignore
    if df_in.empty:
        st.info("Chưa có chi phí đầu vào.")
        return 0

    # Tách Hóa đơn và UNC (Dựa vào số hóa đơn có chứa 'UNC' hay không)
    df_in['Is_UNC'] = df_in['invoice_number'].astype(str).str.contains("UNC", case=False, na=False) # type: ignore
    
    df_bills = df_in.loc[~df_in['Is_UNC']]
    df_uncs = df_in.loc[df_in['Is_UNC']]
    
    total_bills = df_bills['total_amount'].sum()
    total_uncs = df_uncs['total_amount'].sum()
    
    # Hiển thị so sánh
    c1, c2, c3 = st.columns(3)
    c1.metric("Tổng Hóa Đơn (Chi phí)", format_vnd(total_bills), help="Tổng giá trị các hóa đơn đầu vào (Không tính UNC)")
    c2.metric("Tổng UNC (Đã chi)", format_vnd(total_uncs), help="Tổng số tiền đã chuyển khoản (UNC)")
    
    diff = total_uncs - total_bills
    if diff == 0:
        c3.success("✅ Đã khớp")
    elif diff > 0:
        c3.warning(f"⚠️ UNC dư: {format_vnd(diff)}")
    else:
        c3.error(f"⚠️ Thiếu UNC: {format_vnd(abs(diff))}")
        
    # Bảng chi tiết
    t1, t2 = st.tabs(["📄 Danh sách Hóa Đơn", "💸 Danh sách UNC"])
    with t1:
        st.dataframe(df_bills[['date', 'invoice_number', 'seller_name', 'total_amount', 'memo']], 
                     column_config={"total_amount": st.column_config.NumberColumn("Số tiền", format="%d")}, use_container_width=True, hide_index=True)
    with t2:
        st.dataframe(df_uncs[['date', 'invoice_number', 'seller_name', 'total_amount', 'memo']], 
                     column_config={"total_amount": st.column_config.NumberColumn("Số tiền", format="%d")}, use_container_width=True, hide_index=True)
        
    return total_bills

def render_cost_control(menu):
    if menu == "1. Nhập Hóa Đơn":
        # 1. Logic Nhập UNC mặc định là Đầu vào (Nhưng Type IN)
        doc_type = st.radio("📂 Loại chứng từ", ["Ủy nhiệm chi ", "Hóa đơn"], horizontal=True, index=1 if st.session_state.current_doc_type == "Hóa đơn" else 0)
        
        if doc_type != st.session_state.current_doc_type:
            st.session_state.current_doc_type = doc_type
            st.session_state.pdf_data = None
            st.session_state.ready_pdf_bytes = None
            st.session_state.ready_file_name = None
            st.session_state.uploader_key += 1
            st.rerun()

        uploaded_file = st.file_uploader(f"Upload {doc_type} (PDF/Ảnh)", type=["pdf", "png", "jpg", "jpeg"], key=f"up_{st.session_state.uploader_key}")
        
        if uploaded_file and st.session_state.ready_file_name != uploaded_file.name:
            st.session_state.ready_pdf_bytes = None
            st.session_state.ready_file_name = uploaded_file.name
            st.session_state.pdf_data = None
            st.session_state.invoice_view_page = 0
        
        is_ready_to_analyze = False
        is_pdf_origin = False
        
        if uploaded_file:
            file_type = uploaded_file.type
            is_pdf_origin = "pdf" in file_type
            is_ready_to_analyze = True

            c_view, c_action = st.columns([1, 1])
            with c_view:
                if is_pdf_origin:
                    st.info("📄 File PDF Gốc")
                    pdf_img = None
                    total_pages = 0
                    try:
                        uploaded_file.seek(0)
                        with pdfplumber.open(uploaded_file) as pdf:
                            total_pages = len(pdf.pages)
                            if st.session_state.invoice_view_page >= total_pages: st.session_state.invoice_view_page = 0
                            pdf_img = pdf.pages[st.session_state.invoice_view_page].to_image(resolution=200).original
                    except: pass
                    
                    if total_pages > 0:
                        if total_pages > 1:
                            c_p, c_n = st.columns(2)
                            if c_p.button("⬅ Trước", key="btn_inv_prev", use_container_width=True): st.session_state.invoice_view_page = max(0, st.session_state.invoice_view_page - 1); st.rerun()
                            if c_n.button("Sau ➡", key="btn_inv_next", use_container_width=True): st.session_state.invoice_view_page = min(total_pages - 1, st.session_state.invoice_view_page + 1); st.rerun()
                        if pdf_img:
                            st.image(pdf_img, caption=f"Trang {st.session_state.invoice_view_page+1}/{total_pages}", width="stretch")
                else:
                    st.info("🖼️ File Ảnh")
                    st.image(uploaded_file, caption="Ảnh gốc", width="stretch")
                    
            with c_action:
                if not is_pdf_origin and st.session_state.ready_pdf_bytes is None:
                    st.info("👉 Bạn đang dùng File Ảnh. Hệ thống sẽ dùng OCR để quét.")
                    if st.button("🔄 CHUYỂN ĐỔI SANG PDF (ĐỂ LƯU TRỮ)", type="secondary", width="stretch"):
                        with st.spinner("Đang chuyển đổi..."):
                            uploaded_file.seek(0)
                            converted_bytes = convert_image_to_pdf(uploaded_file)
                            if converted_bytes:
                                st.session_state.ready_pdf_bytes = converted_bytes
                                st.success("Đã convert xong!")
                                time.sleep(0.5)
                                st.rerun()

                if is_ready_to_analyze:
                    if st.button(f"🔍 QUÉT THÔNG TIN ({doc_type})", type="primary", width="stretch"):
                        file_to_scan = None
                        is_img_input = not is_pdf_origin
                        if is_img_input: file_to_scan = uploaded_file 
                        else: file_to_scan = uploaded_file 
                        
                        if file_to_scan:
                            file_to_scan.seek(0)
                            data, msg = extract_data_smart(file_to_scan, is_img_input, doc_type)
                            if msg: st.warning(msg)
                            if data is None: st.error("Lỗi hệ thống khi đọc file.")
                            else:
                                data['file_name'] = uploaded_file.name
                                st.session_state.pdf_data = data
                                st.session_state.edit_lock = True
                                st.session_state.local_edit_count = 0
                                
                                if not HAS_OCR and is_img_input:
                                    st.error("❌ Máy chưa cài Tesseract OCR. Không thể đọc số từ ảnh đâu á!")
                                
                                # --- 1. KHÔI PHỤC THÔNG BÁO KHỚP TIỀN ---
                                if doc_type == "Hóa đơn":
                                    diff = abs(data['total'] - (data['pre_tax'] + data['tax']))
                                    if diff < 10: st.success(f"✅ Chuẩn men! Tổng: {format_vnd(data['total'])}")
                                    else: st.warning(f"⚠️ Lệch tiền: {format_vnd(diff)} (Tổng != Tiền hàng + Thuế)")
                                else:
                                    st.success(f"✅ Đã quét UNC! Số tiền: {format_vnd(data['total'])}")
                                st.rerun()

                if st.session_state.pdf_data:
                    d = st.session_state.pdf_data
                    st.divider()
                    
                    # --- LOGIC MÃ CHI PHÍ (COST CODE) - MOVED OUTSIDE FORM ---
                    # Lấy danh sách Tour đang chạy để chọn
                    user_info_cost = st.session_state.get("user_info", {})
                    user_role_cost = user_info_cost.get('role')
                    user_name_cost = user_info_cost.get('name')
                    tour_query = "SELECT tour_name, tour_code FROM tours WHERE status='running'"
                    tour_params = []
                    if user_role_cost == 'sale' and user_name_cost:
                        tour_query += " AND sale_name=?"
                        tour_params.append(user_name_cost)
                    active_tours = run_query(tour_query, tuple(tour_params))
                    tour_choices = {f"[{t['tour_code']}] {t['tour_name']}": t['tour_code'] for t in active_tours} if active_tours else {} # type: ignore
                    tour_choices = {f"📦 TOUR: [{t['tour_code']}] {t['tour_name']}": t['tour_code'] for t in active_tours} if active_tours else {} # type: ignore
                    
                    # Lấy danh sách các mã Cost Code đã tồn tại (từ UNC hoặc Hóa đơn trước đó) để Hóa đơn chọn lại
                    existing_codes_query = run_query("SELECT DISTINCT cost_code FROM invoices WHERE cost_code IS NOT NULL AND cost_code != ''")
                    existing_codes = [r['cost_code'] for r in existing_codes_query] if existing_codes_query else [] # type: ignore
                    
                    # Lấy danh sách Booking Dịch Vụ (Lọc theo sale nếu cần)
                    bk_query = "SELECT name, code FROM service_bookings WHERE status='active'"
                    bk_params = []
                    if user_role_cost == 'sale' and user_name_cost:
                        bk_query += " AND sale_name=?"
                        bk_params.append(user_name_cost)
                    active_bookings = run_query(bk_query, tuple(bk_params))
                    booking_choices = {f"🔖 BOOKING: [{b['code']}] {b['name']}": b['code'] for b in active_bookings} if active_bookings else {} # type: ignore

                    selected_cost_code = ""
                    new_bk_name = None
                    new_bk_code = None
                    
                    st.markdown("##### 🔖 Phân loại & Liên kết chi phí")
                    with st.container(border=True):
                        if doc_type == "Ủy nhiệm chi ":
                            st.info("🔖 Phân loại chi phí")
                            # Logic mới: Luôn yêu cầu chọn Mã (Tour hoặc Booking)
                            link_type = st.radio("Liên kết với:", ["Tour", "Booking Dịch Vụ"], horizontal=True)
                            
                            if link_type == "Tour":
                                if tour_choices:
                                    sel_t = st.selectbox("Chọn Tour:", list(tour_choices.keys()))
                                    selected_cost_code = tour_choices[sel_t]
                                else:
                                    st.warning("Chưa có Tour nào đang chạy.")
                            else:
                                # Booking Dịch Vụ
                                bk_action = st.radio("Thao tác:", ["Chọn Booking có sẵn", "➕ Tạo Booking mới"], horizontal=True, label_visibility="collapsed")
                                
                                if bk_action == "Chọn Booking có sẵn":
                                    if booking_choices:
                                        sel_b = st.selectbox("Chọn Booking:", list(booking_choices.keys()))
                                        selected_cost_code = booking_choices[sel_b]
                                    else:
                                        st.warning("Chưa có Tour nào đang chạy.")
                                        st.warning("Chưa có Booking nào.")
                                else:
                                    # Tự tạo mã Booking lẻ
                                    if "gen_booking_code" not in st.session_state:
                                        st.session_state.gen_booking_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                                    # Tạo mới Booking Dịch Vụ ngay tại đây
                                    c_new_b1, c_new_b2 = st.columns([1, 2])
                                    if "new_bk_code" not in st.session_state:
                                        st.session_state.new_bk_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                                    
                                    c_gen1, c_gen2 = st.columns([1, 3])
                                    c_gen1.text_input("Mã Booking:", value=st.session_state.gen_booking_code, disabled=True)
                                    if c_gen2.button("🔄 Tạo mã khác"):
                                        st.session_state.gen_booking_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                                        st.rerun()
                                    selected_cost_code = st.session_state.gen_booking_code
                                    new_bk_code = c_new_b1.text_input("Mã Booking (Tự động)", value=st.session_state.new_bk_code, disabled=True)
                                    new_bk_name = c_new_b2.text_input("Tên Booking / Dịch vụ", placeholder="VD: Khách lẻ A, Vé máy bay B...")
                                
                        else: # Hóa đơn
                            st.info("🔗 Liên kết chi phí")
                            inv_opt = st.radio("Nguồn gốc:", ["Theo mã UNC/Booking/Tour", "Không có UNC (Tự tạo mã)"], horizontal=True)
                            if inv_opt == "Theo mã UNC/Booking/Tour":
                                # Gộp cả mã Tour và mã Booking lẻ đã có
                                all_avail_codes = sorted(list(set(list(tour_choices.values()) + existing_codes)))
                                if all_avail_codes:
                                    selected_cost_code = st.selectbox("Chọn Mã liên kết:", all_avail_codes)
                                else:
                                    st.warning("Chưa có mã nào để liên kết.")
                            else:
                                if "gen_inv_code" not in st.session_state:
                                    st.session_state.gen_inv_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                                st.text_input("Mã chi phí mới:", value=st.session_state.gen_inv_code, disabled=True)
                                selected_cost_code = st.session_state.gen_inv_code
                                st.caption("Vui lòng nhập tên để tạo mã.")
                    
                    # Initialize variables to avoid unbound errors
                    txn_content = ""; seller = ""; buyer = ""

                    with st.form("inv_form"):
                        # Mặc định UNC là Đầu vào
                        default_idx = 0 
                        
                        # --- PHẦN 1: THÔNG TIN CHUNG ---
                        st.markdown("##### 📝 Thông tin chung")
                        with st.container(border=True):
                            st.text_input("Mã chi phí / Booking:", value=selected_cost_code, disabled=True)
                            st.divider()
                            
                            typ = st.radio("Loại", ["Đầu vào", "Đầu ra"], horizontal=True, index=default_idx)
                            drive_link = st.text_input("🔗 Link Drive (Tùy chọn)")
                            
                            c1, c2 = st.columns(2)
                            if doc_type == "Hóa đơn":
                                memo = st.text_input("Gợi nhớ (Memo)", value=d.get('file_name',''))
                                date = st.text_input("Ngày", value=d['date'])
                                num = c1.text_input("Số hóa đơn", value=d['inv_num'])
                                sym = c2.text_input("Ký hiệu/Mẫu số", value=d['inv_sym'])
                            else:
                                memo = c1.text_input("Gợi nhớ (Tên file)", value=d.get('file_name', ''))
                                date = c2.text_input("Ngày chuyển khoản", value=d['date'])
                                content_val = d.get('content', '')
                                txn_content = st.text_area("Nội dung chuyển khoản (OCR)", value=content_val, height=70)
                                num = ""; sym = ""; buyer = "" 
                        
                        # --- PHẦN 2: BÊN MUA / BÁN ---
                        if doc_type == "Hóa đơn" or doc_type == "Ủy nhiệm chi ":
                            st.markdown("##### 🤝 Đối tượng")
                            with st.container(border=True):
                                if doc_type == "Hóa đơn":
                                    seller = st.text_input("Bên Bán", value=d['seller'])
                                    buyer = st.text_input("Bên Mua", value=d['buyer'])
                                else:
                                    seller = st.text_input("Đơn vị nhận tiền", value=d['seller'])
                        
                        # --- PHẦN 3: TÀI CHÍNH ---
                        st.markdown("##### 💰 Tài chính")
                        with st.container(border=True):
                            if doc_type == "Hóa đơn":
                                pre = st.number_input("Tiền hàng", value=float(d['pre_tax']), disabled=st.session_state.edit_lock, format="%.0f")
                                tax = st.number_input("VAT", value=float(d['tax']), disabled=st.session_state.edit_lock, format="%.0f")
                                total = pre + tax
                            else:
                                st.caption("(Với UNC, chỉ cần nhập Số tiền đã chuyển nha)")
                                pre = 0; tax = 0
                                total = st.number_input("Số tiền đã chuyển", value=float(d['total']), disabled=st.session_state.edit_lock, format="%.0f")

                            is_locked_admin = False
                            # 3. & 5. LOGIC DUYỆT:
                            
                            if st.session_state.local_edit_count == 2:
                                st.markdown('<div style="background:#fff3cd; color:orange; padding:10px; border-radius:5px; margin-bottom:10px;">⚠️ <b>Lưu ý:</b> Nếu chỉnh sửa lần 3 phải gửi admin duyệt.</div>', unsafe_allow_html=True)
                            elif st.session_state.local_edit_count >= 3 and st.session_state.local_edit_count < 5:
                                is_locked_admin = True
                                st.markdown(f'<div style="background:#ffeef7; color:red; padding:10px; border-radius:5px; margin-bottom:10px;">🔒 <b>Chế độ duyệt:</b> Bạn đang sửa lần {st.session_state.local_edit_count}. Cần Admin duyệt.</div>', unsafe_allow_html=True)
                            elif st.session_state.local_edit_count >= 5:
                                st.error("⛔ Đã quá số lần chỉnh sửa cho phép (5 lần).")

                            # 6. HIỂN THỊ TIỀN 1 HÀNG (CSS .money-box đã xử lý)
                            st.write("") 
                            st.markdown(f'<div class="money-box">{format_vnd(total)}</div>', unsafe_allow_html=True)
                            
                            b1, b2 = st.columns(2)
                            
                            if st.session_state.local_edit_count < 5:
                                if b1.form_submit_button("✏️ Sửa giá"):
                                    st.session_state.edit_lock = False
                                    st.rerun()
                            
                            if not st.session_state.edit_lock and b2.form_submit_button("✅ Chốt giá"):
                                new_pre = pre if doc_type == "Hóa đơn" else total
                                st.session_state.pdf_data.update({'pre_tax': new_pre, 'tax': tax, 'total': total})
                                st.session_state.edit_lock = True
                                st.session_state.local_edit_count += 1
                                st.rerun()

                        # Nút Lưu / Gửi Duyệt
                        if is_locked_admin:
                            btn_label = "🚀 GỬI ADMIN DUYỆT"
                        elif st.session_state.local_edit_count >= 5:
                            btn_label = "⛔ ĐÃ KHÓA"
                        else:
                            btn_label = "💾 LƯU CHỨNG TỪ"
                        
                        if st.form_submit_button(btn_label, type="primary", width="stretch", disabled=(st.session_state.local_edit_count >= 5)):
                            if doc_type == "Hóa đơn" and (not date or not num): st.error("Ơ kìa, thiếu ngày hoặc số hóa đơn rồi!")
                            elif doc_type == "Ủy nhiệm chi " and not date: st.error("Thiếu ngày chuyển khoản rồi nè!")
                            elif not st.session_state.edit_lock: st.warning("Bấm 'Chốt giá' trước khi lưu nha!")
                            else:
                                # --- CHUẨN BỊ DỮ LIỆU ---
                                t = 'OUT' if "Đầu ra" in typ else 'IN'
                                save_memo = memo
                                save_num = num
                                
                                if doc_type == "Ủy nhiệm chi ":
                                    save_memo = f"[UNC] {memo} - {txn_content}"
                                    if not save_num: save_num = f"UNC-{datetime.now().strftime('%y%m%d%H%M')}"

                                # --- TẠO TÊN FILE ---
                                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                                clean_name = re.sub(r'[\\/*?:"<>|]', "", uploaded_file.name)
                                final_name = f"{ts}_{clean_name}"
                                if st.session_state.ready_pdf_bytes and not final_name.lower().endswith('.pdf'):
                                    final_name = os.path.splitext(final_name)[0] + ".pdf"

                                # [CODE MỚI] 
                                # 1. Upload file lên Drive (Đã tắt theo yêu cầu - Chỉ lưu dữ liệu)
                                drive_link = ""
                                # if uploaded_file:
                                #     # Xử lý file upload (nếu là ảnh đã convert sang PDF thì dùng bytes)
                                #     if st.session_state.ready_pdf_bytes:
                                #         file_obj = io.BytesIO(st.session_state.ready_pdf_bytes)
                                #         drive_link = upload_to_drive(file_obj, final_name, mimetype='application/pdf')
                                #     else:
                                #         drive_link = upload_to_drive(uploaded_file, final_name)
                                
                                # 2. Chuẩn bị dữ liệu để lưu
                                new_invoice = {
                                    'type': t, 
                                    'date': date,
                                    'invoice_number': save_num,
                                    'invoice_symbol': sym,
                                    'seller_name': seller,
                                    'buyer_name': buyer,
                                    'pre_tax_amount': pre,
                                    'tax_amount': tax,
                                    'total_amount': total,
                                    'file_name': final_name,
                                    'status': 'active',
                                    'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                    'memo': save_memo,
                                    'file_path': drive_link, 
                                    'cost_code': selected_cost_code,
                                    'edit_count': st.session_state.local_edit_count,
                                    'request_edit': 1 if is_locked_admin else 0
                                }
                                
                                # 3. Ghi vào Sheet 'invoices'
                                if add_row_to_table('invoices', new_invoice):
                                    st.success("Đã lưu thành công lên Cloud! 🎉")
                                    
                                    # Reset state
                                    time.sleep(1)
                                    st.session_state.pdf_data = None
                                    st.session_state.uploader_key += 1
                                    st.session_state.ready_pdf_bytes = None
                                    st.session_state.ready_file_name = None
                                    st.session_state.local_edit_count = 0
                                    if "gen_booking_code" in st.session_state: del st.session_state.gen_booking_code
                                    if "gen_inv_code" in st.session_state: del st.session_state.gen_inv_code
                                    if "new_bk_code" in st.session_state: del st.session_state.new_bk_code
                                    if "pending_booking_create" in st.session_state: del st.session_state.pending_booking_create
                                    st.rerun()

        st.divider()
        # --- 4. LỊCH SỬ NHẬP LIỆU (HIỆN TẤT CẢ NHƯNG CÓ NOTE) ---
        with st.expander("Lịch sử nhập liệu", expanded=True):
            rows = run_query("SELECT id, type, invoice_number, total_amount, status, memo, request_edit, edit_count, cost_code FROM invoices ORDER BY id DESC LIMIT 20")
            if rows:
                df = pd.DataFrame([dict(r) for r in rows])
                df['Chọn'] = False 
                
                def get_status_note(row): # type: ignore
                    if row['status'] == 'deleted': # type: ignore
                        return "❌ Đã xóa"
                    note = ""
                    if row['request_edit'] == 1: # type: ignore
                        note += "⏳ Chờ duyệt"
                    if row['edit_count'] > 0: # type: ignore
                        if note: note += " | "
                        note += f"✏️ Sửa {row['edit_count']} lần" # type: ignore
                    
                    if not note:
                        return "✅ Hoạt động"
                    return note.strip(" | ")
                
                df['Trạng thái'] = df.apply(get_status_note, axis=1)
                
                df = df[['Chọn', 'id', 'cost_code', 'type', 'invoice_number', 'total_amount', 'Trạng thái', 'memo']]
                df.columns = ['Chọn', 'ID', 'Mã Chi Phí', 'Loại', 'Số HĐ', 'Tổng Tiền', 'Trạng thái', 'Ghi chú']
                
                df['Tổng Tiền'] = df['Tổng Tiền'].apply(format_vnd)

                edited_df = st.data_editor(
                    df,
                    column_config={
                        "Chọn": st.column_config.CheckboxColumn(required=True),
                        "ID": st.column_config.NumberColumn(disabled=True),
                        "Mã Chi Phí": st.column_config.TextColumn(disabled=True),
                        "Loại": st.column_config.TextColumn(disabled=True),
                        "Số HĐ": st.column_config.TextColumn(disabled=True),
                        "Tổng Tiền": st.column_config.TextColumn(disabled=True),
                        "Trạng thái": st.column_config.TextColumn(disabled=True),
                        "Ghi chú": st.column_config.TextColumn(disabled=True),
                    },
                    hide_index=True,
                    use_container_width=True
                )

                if st.button("🗑️ Xóa các mục đã chọn", type="primary"):
                    selected_ids = edited_df[edited_df['Chọn']]['ID'].tolist()
                    if selected_ids:
                        for i in selected_ids:
                            run_query("UPDATE invoices SET status='deleted' WHERE id=?", (i,), commit=True)
                        st.success(f"Đã xóa {len(selected_ids)} hóa đơn!")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.warning("Bạn chưa chọn mục nào cả.")
            else:
                st.info("Chưa có hóa đơn nào.")
    elif menu == "2. Báo Cáo Tổng Hợp":
        st.title("📊 Báo Cáo Tài Chính")

        all_financial_records = []
        with st.spinner("Đang tổng hợp dữ liệu từ tất cả các phân hệ..."):
            # --- OPTIMIZED DATA FETCHING ---
            # Lọc booking theo sale nếu cần
            user_info_rpt = st.session_state.get("user_info", {})
            user_role_rpt = user_info_rpt.get('role')
            user_name_rpt = user_info_rpt.get('name')

            # 1. Fetch all base data in a few queries
            tour_rpt_query = "SELECT * FROM tours WHERE status != 'deleted'"
            tour_rpt_params = []
            if user_role_rpt == 'sale' and user_name_rpt:
                tour_rpt_query += " AND sale_name=?"
                tour_rpt_params.append(user_name_rpt)
            all_tours = run_query(tour_rpt_query, tuple(tour_rpt_params))
            
            bk_rpt_query = "SELECT * FROM service_bookings WHERE status != 'deleted'"
            bk_rpt_params = []
            if user_role_rpt == 'sale' and user_name_rpt:
                bk_rpt_query += " AND sale_name=?"
                bk_rpt_params.append(user_name_rpt)
            all_bookings = run_query(bk_rpt_query, tuple(bk_rpt_params))

            all_linked_invoices = run_query("SELECT cost_code, type, invoice_number, total_amount FROM invoices WHERE status='active' AND request_edit=0 AND cost_code IS NOT NULL AND cost_code != ''")
            # [NEW] Fetch all transactions for debt calculation
            all_transactions = run_query("SELECT ref_code, type, amount FROM transaction_history")

            # 2. Process data in memory using dictionaries for fast lookups
            invoice_costs_by_code = {}
            for inv in all_linked_invoices:
                code = inv['cost_code']
                if code not in invoice_costs_by_code:
                    invoice_costs_by_code[code] = {'IN_INV': 0, 'IN_UNC': 0}
                if inv['type'] == 'IN':
                    is_unc = 'UNC' in (inv.get('invoice_number') or '') # type: ignore
                    if is_unc:
                        invoice_costs_by_code[code]['IN_UNC'] += inv['total_amount'] # type: ignore
                    else:
                        invoice_costs_by_code[code]['IN_INV'] += inv['total_amount'] # type: ignore
            
            # [NEW] Process transactions to get paid amounts
            paid_amounts = {}
            if all_transactions:
                df_txns = pd.DataFrame([dict(r) for r in all_transactions])
                if not df_txns.empty:
                    df_thu = df_txns[df_txns['type'] == 'THU'].groupby('ref_code')['amount'].sum()
                    df_chi = df_txns[df_txns['type'] == 'CHI'].groupby('ref_code')['amount'].sum() # CHI means refund
                    paid_amounts = (df_thu.subtract(df_chi, fill_value=0)).to_dict()

            # --- Process Tours ---
            if all_tours:
                for tour_row in all_tours:
                    tour = dict(tour_row)
                    # [NEW] Add status to record
                    tour_status = tour.get('status', 'running')
                    revenue, cost = get_tour_financials(tour['id'], tour)
                    if revenue > 0: all_financial_records.append({'date_str': tour['start_date'], 'name': tour['tour_name'], 'code': tour['tour_code'], 'category': 'Tour', 'type': 'thu', 'amount': revenue, 'status': tour_status}) # type: ignore
                    if cost > 0: all_financial_records.append({'date_str': tour['start_date'], 'name': tour['tour_name'], 'code': tour['tour_code'], 'category': 'Tour', 'type': 'chi', 'amount': cost, 'status': tour_status}) # type: ignore

            # --- Process Service Bookings ---
            if all_bookings:
                for booking_row in all_bookings:
                    booking = dict(booking_row)
                    
                    # [FIX] Chuyển đổi định dạng ngày YYYY-MM-DD sang DD/MM/YYYY để đồng bộ
                    try:
                        booking_date_obj = datetime.strptime(str(booking['created_at']).split(" ")[0], '%Y-%m-%d')
                        booking_date_str = booking_date_obj.strftime('%d/%m/%Y')
                    except:
                        booking_date_str = booking['created_at']
                    # [NEW] Add status to record
                    booking_status = booking.get('status', 'active')

                    if booking.get('selling_price', 0) > 0:
                        all_financial_records.append({'date_str': booking_date_str, 'name': booking['name'], 'code': booking['code'], 'category': 'Booking Dịch Vụ', 'type': 'thu', 'amount': booking['selling_price'], 'status': booking_status}) # type: ignore
                    
                    # [FIX] Chỉ tính chi phí từ hóa đơn (IN_INV), không tính UNC để tránh double-count.
                    # UNC là thanh toán cho chi phí, không phải bản thân chi phí.
                    total_cost_booking = invoice_costs_by_code.get(booking['code'], {}).get('IN_INV', 0)
                    if total_cost_booking == 0 and booking.get('net_price', 0) > 0:
                        total_cost_booking = booking['net_price'] # type: ignore
                    if total_cost_booking > 0:
                        all_financial_records.append({'date_str': booking_date_str, 'name': booking['name'], 'code': booking['code'], 'category': 'Booking Dịch Vụ', 'type': 'chi', 'amount': total_cost_booking, 'status': booking_status}) # type: ignore

            # --- Process old Projects & Unlinked Invoices (These queries are already efficient) ---
            project_invoices = run_query("SELECT p.project_name, i.type, i.total_amount, i.date, p.id as project_id FROM projects p JOIN project_links l ON p.id = l.project_id JOIN invoices i ON l.invoice_id = i.id WHERE i.status = 'active' AND i.request_edit = 0")
            if project_invoices:
                for inv in project_invoices:
                    all_financial_records.append({'date_str': inv['date'], 'name': inv['project_name'], 'code': f"PROJ_{inv['project_id']}", 'category': 'Dự án (cũ)', 'type': 'thu' if inv['type'] == 'OUT' else 'chi', 'amount': inv['total_amount'], 'status': 'N/A'}) # type: ignore

            unlinked_invoices = run_query("SELECT * FROM invoices i WHERE i.status = 'active' AND i.request_edit = 0 AND (i.cost_code IS NULL OR i.cost_code = '') AND NOT EXISTS (SELECT 1 FROM project_links pl WHERE pl.invoice_id = i.id)")
            if unlinked_invoices:
                for inv in unlinked_invoices:
                    all_financial_records.append({'date_str': inv['date'], 'name': inv['memo'] or inv['seller_name'] or 'Chi phí chung', 'code': f"INV_{inv['id']}", 'category': 'Chi phí chung', 'type': 'thu' if inv['type'] == 'OUT' else 'chi', 'amount': inv['total_amount'], 'status': 'N/A'}) # type: ignore

        if not all_financial_records:
            st.info("Chưa có dữ liệu tài chính để báo cáo.")
        else:
            df = pd.DataFrame(all_financial_records)
            df['date'] = pd.to_datetime(df['date_str'], errors='coerce', dayfirst=True)
            df['status'] = df['status'].fillna('N/A') # Đảm bảo cột status không có giá trị null
            df = df.dropna(subset=['date'])

            # Explicitly create a DatetimeIndex to help Pylance with type inference
            dt_index = pd.DatetimeIndex(df['date'])
            df['year'] = dt_index.year
            df['quarter'] = dt_index.quarter
            df['month_year'] = dt_index.to_period('M').astype(str)
            df['quarter_year'] = df.apply(lambda row: f"Q{row['quarter']}/{row['year']}", axis=1)

            st.markdown("####  Lọc báo cáo")
            c1, c2, c3 = st.columns(3)
            filter_type = c1.selectbox("Lọc theo thời gian:", ["Tháng", "Quý", "Năm"])
            
            options = []
            period_col = ''
            if filter_type == "Tháng":
                options = sorted(df['month_year'].unique(), reverse=True)
                period_col = 'month_year'
            elif filter_type == "Quý":
                options = sorted(df['quarter_year'].unique(), reverse=True)
                period_col = 'quarter_year'
            elif filter_type == "Năm":
                options = sorted(df['year'].unique(), reverse=True)
                period_col = 'year'
                
            selected_period = c2.selectbox(f"Chọn kỳ:", ["Tất cả"] + options)

            # [NEW] Thêm bộ lọc trạng thái
            status_map = {
                "Tất cả trạng thái": None,
                "Đang chạy / Hoạt động": ['running', 'active'],
                "Đã hoàn thành": ['completed']
            }
            selected_status_label = c3.selectbox("Lọc theo trạng thái:", list(status_map.keys()))
            selected_statuses = status_map[selected_status_label]

            # Áp dụng các bộ lọc
            df_filtered = df.copy()
            if selected_period != "Tất cả":
                df_filtered = df_filtered[df_filtered[period_col] == selected_period]
            
            if selected_statuses:
                # Chỉ lọc các mục có trạng thái (Tour/Booking), giữ lại các mục khác (Chi phí chung...)
                mask = df_filtered['status'].isin(selected_statuses) | (df_filtered['status'] == 'N/A')
                df_filtered = df_filtered[mask]

            if not df_filtered.empty:
                agg = df_filtered.pivot_table(index=['category', 'name', 'code'], columns='type', values='amount', aggfunc='sum').fillna(0)
                agg = agg.reset_index()
                
                if 'thu' not in agg.columns: agg['thu'] = 0
                if 'chi' not in agg.columns: agg['chi'] = 0
                agg['lợi nhuận'] = agg['thu'] - agg['chi']
                
                total_thu = agg['thu'].sum()
                total_chi = agg['chi'].sum()
                total_loi_nhuan = agg['lợi nhuận'].sum()
                
                m1, m2, m3 = st.columns(3)
                m1.metric(f"Tổng Thu ({selected_period})", format_vnd(total_thu))
                m2.metric(f"Tổng Chi ({selected_period})", format_vnd(total_chi))
                m3.metric(f"Lợi Nhuận ({selected_period})", format_vnd(total_loi_nhuan), delta=format_vnd(total_loi_nhuan) if total_loi_nhuan != 0 else None)

                st.divider()
                
                st.markdown("#### Chi tiết theo hạng mục")
                # Sort categories by total profit
                category_profit = agg.groupby('category')['lợi nhuận'].sum().sort_values(ascending=False)
                
                for category in category_profit.index:
                    group = agg[agg['category'] == category]
                    with st.expander(f"📂 {category} (Lợi nhuận: {format_vnd(group['lợi nhuận'].sum())})", expanded=True):
                        group = group.sort_values('lợi nhuận', ascending=False)
                        for _, r in group.iterrows():
                            # --- [NEW] Debt calculation & display ---
                            debt_html = ""
                            # Only calculate for Tours and Bookings which have revenue
                            if r['category'] in ['Tour', 'Booking Dịch Vụ'] and r['thu'] > 0:
                                code = r['code']
                                revenue = r['thu']
                                paid = paid_amounts.get(code, 0.0)
                                remaining = revenue - paid
                                
                                if remaining <= 0.1: # Use a small threshold for float comparison
                                    debt_html = f'''<div style="margin-top: 8px; font-size: 0.9em; text-align: right;">
                                        <span style="color: #2e7d32; font-weight: bold;">✅ Đã thanh toán đủ</span>
                                    </div>'''
                                else:
                                    debt_html = f'''<div style="margin-top: 8px; font-size: 0.9em; text-align: right;">
                                        <span style="color: #c62828; font-weight: bold;">Còn phải thu: {format_vnd(remaining)}</span>
                                    </div>'''
                            # --- End of new code ---

                            st.markdown(f"""
                            <div class="report-card" style="padding: 15px; margin-bottom: 10px; border-left: 5px solid {'#28a745' if r['lợi nhuận']>=0 else '#e53935'};">
                                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom: 8px;">
                                    <h5 style="margin:0; padding-right: 10px;">{r['name']}</h5>
                                    <span style="font-size: 0.8em; color: #6c757d; background-color: #f1f3f5; padding: 2px 6px; border-radius: 5px; white-space: nowrap;">CODE: {r['code']}</span>
                                </div>
                                <div style="display:flex; justify-content:space-between; font-size: 0.95em; border-bottom: 1px solid #f1f3f5; padding-bottom: 8px;">
                                    <span>Thu: <b>{format_vnd(r['thu'])}</b></span>
                                    <span>Chi: <b>{format_vnd(r['chi'])}</b></span>
                                    <span style="font-weight: bold; color:{'#1B5E20' if r['lợi nhuận']>=0 else '#c62828'}">Lãi: {format_vnd(r['lợi nhuận'])}</span>
                                </div>
                                {debt_html}
                            </div>
                            """, unsafe_allow_html=True)
            else:
                st.info(f"Không có dữ liệu cho kỳ báo cáo '{selected_period}'.")

def render_debt_management():
    st.title("💳 Quản Lý Công Nợ")
    st.caption("Theo dõi và tổng hợp các khoản phải thu từ khách hàng.")

    tab_lookup, tab_summary = st.tabs(["Tra cứu theo Mã", "Tổng hợp Công nợ"])

    with tab_lookup:
        st.subheader("Tra cứu công nợ theo Mã Tour / Booking")
        
        # --- LẤY DỮ LIỆU ĐỂ TÌM KIẾM (CHỈ HIỆN CÁC MÃ CÒN NỢ) ---
        with st.spinner("Đang tải danh sách còn nợ..."):
            # 1. Lấy tất cả giao dịch và tính toán số tiền đã trả cho mỗi mã
            all_txns_cn = run_query("SELECT ref_code, type, amount FROM transaction_history")
            paid_amounts_cn = {}
            if all_txns_cn:
                df_txns_cn = pd.DataFrame([dict(r) for r in all_txns_cn])
                if not df_txns_cn.empty:
                    df_thu_cn = df_txns_cn[df_txns_cn['type'] == 'THU'].groupby('ref_code')['amount'].sum()
                    df_chi_cn = df_txns_cn[df_txns_cn['type'] == 'CHI'].groupby('ref_code')['amount'].sum()
                    paid_amounts_cn = (df_thu_cn.subtract(df_chi_cn, fill_value=0)).to_dict()

            # 2. Lấy tất cả tour và booking (lọc theo sale nếu cần)
            user_info_cn = st.session_state.get("user_info", {})
            user_role_cn = user_info_cn.get('role')
            user_name_cn = user_info_cn.get('name')

            # [FIX] Lấy tất cả tour/booking chưa bị xóa (bao gồm cả mục đã hoàn thành) để kiểm tra công nợ
            tour_cn_query = "SELECT * FROM tours WHERE COALESCE(status, 'running') NOT IN ('deleted')"
            tour_cn_params = []
            if user_role_cn == 'sale' and user_name_cn:
                tour_cn_query += " AND sale_name=?"
                tour_cn_params.append(user_name_cn)
            all_tours_cn = run_query(tour_cn_query, tuple(tour_cn_params))

            bk_cn_query = "SELECT * FROM service_bookings WHERE COALESCE(status, 'active') NOT IN ('deleted')"
            bk_cn_params = []
            if user_role_cn == 'sale' and user_name_cn:
                bk_cn_query += " AND sale_name=?"
                bk_cn_params.append(user_name_cn)
            all_bookings_cn = run_query(bk_cn_query, tuple(bk_cn_params))

            search_options = {"": "-- Chọn mã để theo dõi --"}

            # 3. Xử lý Tours: Chỉ thêm vào danh sách nếu chưa thu đủ
            if all_tours_cn:
                for t_row in all_tours_cn:
                    tour = dict(t_row)
                    # Tính giá trị hợp đồng
                    final_price = float(tour.get('final_tour_price', 0) or 0)
                    child_price = float(tour.get('child_price', 0) or 0)
                    final_qty = float(tour.get('final_qty', 0) or 0)
                    child_qty = float(tour.get('child_qty', 0) or 0)
                    if final_qty == 0: final_qty = float(tour.get('guest_count', 1))
                    contract_value = (final_price * final_qty) + (child_price * child_qty)
                    
                    paid = paid_amounts_cn.get(tour['tour_code'], 0.0)
                    
                    if contract_value > 0 and contract_value - paid > 0.1:
                        search_options[f"📦 TOUR: [{tour['tour_code']}] {tour['tour_name']}"] = tour['tour_code']

            # 4. Xử lý Bookings: Chỉ thêm vào danh sách nếu chưa thu đủ
            if all_bookings_cn:
                for b_row in all_bookings_cn:
                    booking = dict(b_row)
                    contract_value = float(booking.get('selling_price', 0) or 0)
                    paid = paid_amounts_cn.get(booking['code'], 0.0)
                    if contract_value > 0 and contract_value - paid > 0.1:
                        search_options[f"🔖 BOOKING: [{booking['code']}] {booking['name']}"] = booking['code']

        # --- GIAO DIỆN CHÍNH ---
        col1, col2 = st.columns([1, 2])
        remaining = 0.0

        with col1:
            st.markdown("#### 🔍 Chọn đối tượng")
            selected_label = st.selectbox("Tìm theo Mã Tour / Booking (chỉ hiện mã còn nợ):", list(search_options.keys()), label_visibility="collapsed")
            selected_code = search_options.get(selected_label)

            if selected_code:
                # Reset trạng thái phiếu vừa tạo nếu chuyển mã khác
                if "last_voucher_code" not in st.session_state or st.session_state.last_voucher_code != selected_code:
                    if "last_voucher" in st.session_state: del st.session_state.last_voucher
                    if "last_voucher_pdf" in st.session_state: del st.session_state.last_voucher_pdf
                    st.session_state.last_voucher_code = selected_code

                st.markdown("---")
                st.markdown("#### 📊 Tổng quan công nợ")

                contract_value = 0.0
                # Xác định giá trị hợp đồng
                if "TOUR" in selected_label:
                    tour_info = run_query("SELECT * FROM tours WHERE tour_code=?", (selected_code,), fetch_one=True)
                    if tour_info:
                        t_dict = dict(tour_info)
                        final_price = float(t_dict.get('final_tour_price', 0) or 0)
                        child_price = float(t_dict.get('child_price', 0) or 0)
                        final_qty = float(t_dict.get('final_qty', 0) or 0)
                        child_qty = float(t_dict.get('child_qty', 0) or 0)
                        if final_qty == 0: final_qty = float(t_dict.get('guest_count', 1))
                        contract_value = (final_price * final_qty) + (child_price * child_qty)
                elif "BOOKING" in selected_label:
                    booking_info = run_query("SELECT selling_price FROM service_bookings WHERE code=?", (selected_code,), fetch_one=True)
                    if booking_info:
                        contract_value = float(booking_info['selling_price'] or 0)

                # Lấy tổng đã thu
                paid_data = run_query("SELECT SUM(amount) as total FROM transaction_history WHERE ref_code=? AND type='THU'", (selected_code,), fetch_one=True)
                total_paid = paid_data['total'] if paid_data and paid_data['total'] else 0.0

                # Lấy tổng đã chi (hoàn tiền)
                refund_data = run_query("SELECT SUM(amount) as total FROM transaction_history WHERE ref_code=? AND type='CHI'", (selected_code,), fetch_one=True)
                total_refund = refund_data['total'] if refund_data and refund_data['total'] else 0.0
                
                actual_paid = total_paid - total_refund
                
                remaining = contract_value - actual_paid

                with st.container(border=True):
                    st.metric("Giá trị Hợp đồng/Booking", format_vnd(contract_value))
                    st.metric("Đã thu thực tế", format_vnd(actual_paid))
                    delta_color = "inverse" if remaining > 0 else "off"
                    st.metric("Còn phải thu", format_vnd(remaining), delta=f"-{format_vnd(remaining)}" if remaining > 0 else "✅ Đã thu đủ", delta_color=delta_color)

        with col2:
            if selected_code:
                tab_add, tab_history = st.tabs(["➕ Tạo Phiếu Thu/Chi", "📜 Lịch sử giao dịch"])

                with tab_add:
                    st.markdown("##### Tạo phiếu mới")
                    
                    # Hiển thị nút tải phiếu vừa tạo (nếu có)
                    if "last_voucher" in st.session_state and st.session_state.last_voucher.get('ref_code') == selected_code:
                        lv = st.session_state.last_voucher
                        st.success("✅ Đã lưu phiếu thành công!")
                        
                        # Sử dụng cache PDF nếu có để tránh tạo lại liên tục
                        if "last_voucher_pdf" in st.session_state:
                            pdf_data = st.session_state.last_voucher_pdf
                        else:
                            pdf_data = create_voucher_pdf(lv)
                            st.session_state.last_voucher_pdf = pdf_data

                        st.download_button(
                            label=f"📥 Tải Phiếu {lv['type']} (PDF)",
                            data=pdf_data,
                            file_name=f"Phieu_{lv['type']}_{lv['date'].replace('/','')}.pdf",
                            mime="application/pdf",
                            type="primary"
                        )
                        st.divider()

                    # Form nhập liệu (Luôn hiển thị)
                    k_amt = f"txn_amt_{selected_code}"
                    k_note = f"txn_note_{selected_code}"
                    
                    if k_amt not in st.session_state:
                        if remaining >= 1:
                            st.session_state[k_amt] = "{:,.0f}".format(remaining).replace(",", ".") + " VND"
                        else:
                            st.session_state[k_amt] = ""
                    
                    def fmt_txn_amt_dynamic(key_name):
                        if key_name in st.session_state:
                            val = st.session_state[key_name]
                            try:
                                clean_val = val.replace('.', '').replace(',', '').replace(' VND', '').strip()
                                if clean_val:
                                    v_float = float(clean_val)
                                    st.session_state[key_name] = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                            except: pass

                    c1, c2 = st.columns(2)
                    txn_type = c1.radio("Loại phiếu", ["THU", "CHI (Hoàn tiền)"], horizontal=True, key=f"txn_type_{selected_code}")
                    
                    txn_amount_input = c2.text_input("Số tiền", key=k_amt, on_change=fmt_txn_amt_dynamic, args=(k_amt,), help="Nhập số tiền (VD: 1.000.000)")
                    
                    try:
                        txn_amount = float(txn_amount_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    except: txn_amount = 0.0
                    
                    c3, c4 = st.columns(2)
                    txn_method = c3.selectbox("Hình thức", ["Chuyển khoản", "Tiền mặt"], key=f"txn_method_{selected_code}")
                    txn_note = c4.text_input("Nội dung", placeholder="VD: Cọc lần 1, Thanh toán...", key=k_note)
                    
                    btn_label = "💾 Tạo Phiếu Thu" if txn_type == "THU" else "💾 Tạo Phiếu Chi"
                    if st.button(btn_label, type="primary", use_container_width=True, key=f"btn_save_txn_{selected_code}"):
                        if txn_amount > 0 and txn_note:
                            now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                            run_query(
                                "INSERT INTO transaction_history (ref_code, type, amount, payment_method, note, created_at) VALUES (?, ?, ?, ?, ?, ?)",
                                (selected_code, txn_type, txn_amount, txn_method, txn_note, now_str),
                                commit=True
                            )
                            
                            # Lưu thông tin để tạo PDF
                            v_data = {
                                'ref_code': selected_code,
                                'type': txn_type,
                                'amount': txn_amount,
                                'method': txn_method,
                                'note': txn_note,
                                'date': datetime.now().strftime("%d/%m/%Y"),
                                'issuer': st.session_state.user_info.get('name', '')
                            }
                            st.session_state.last_voucher = v_data
                            
                            # Tạo DOCX ngay và cache lại
                            docx_bytes = create_voucher_docx(v_data)
                            st.session_state.last_voucher_pdf = docx_bytes # Tái sử dụng biến session cũ
                            
                            if k_amt in st.session_state: del st.session_state[k_amt]
                            if k_note in st.session_state: del st.session_state[k_note]
                            st.rerun()
                        else:
                            st.warning("Vui lòng nhập số tiền và nội dung.")

                with tab_history:
                    st.markdown("##### Lịch sử các lần thanh toán")
                    history = run_query("SELECT * FROM transaction_history WHERE ref_code=? ORDER BY id DESC", (selected_code,))
                    
                    if history:
                        # Hiển thị dạng bảng cho gọn
                        df_hist = pd.DataFrame([dict(r) for r in history])
                        
                        # Format dữ liệu hiển thị
                        df_display = df_hist.copy()
                        df_display['created_at'] = pd.to_datetime(df_display['created_at'], errors='coerce').dt.strftime('%d/%m/%Y')
                        df_display['amount'] = df_display['amount'].apply(lambda x: format_vnd(x))
                        
                        df_display = df_display.rename(columns={
                            'created_at': 'Ngày',
                            'type': 'Loại',
                            'amount': 'Số tiền',
                            'payment_method': 'Hình thức',
                            'note': 'Nội dung',
                            'id': 'ID'
                        })
                        
                        st.dataframe(
                            df_display[['ID', 'Ngày', 'Loại', 'Số tiền', 'Hình thức', 'Nội dung']],
                            use_container_width=True,
                            hide_index=True
                        )
                        
                        st.divider()
                        st.markdown("###### 🛠️ Thao tác (Tải phiếu / Xóa)")
                        
                        # Tạo danh sách lựa chọn
                        txn_options = {}
                        for r in history:
                            try: d_lbl = datetime.strptime(r['created_at'], "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%Y")
                            except: d_lbl = r['created_at']
                            label = f"#{r['id']} | {d_lbl} | {r['type']} | {format_vnd(r['amount'])}"
                            txn_options[label] = r

                        selected_txn_label = st.selectbox("Chọn giao dịch:", ["-- Chọn giao dịch --"] + list(txn_options.keys()))
                        
                        if selected_txn_label and selected_txn_label != "-- Chọn giao dịch --":
                            txn = txn_options[selected_txn_label]
                            
                            # Chỉ tạo PDF khi đã chọn (Tối ưu hiệu năng)
                            try: d_str = datetime.strptime(txn['created_at'], "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%Y")
                            except: d_str = txn['created_at']
                            
                            v_data = {
                                'ref_code': selected_code,
                                'type': txn['type'],
                                'amount': txn['amount'],
                                'method': txn['payment_method'],
                                'note': txn['note'],
                                'date': d_str,
                                'issuer': st.session_state.user_info.get('name', '')
                            }
                            docx_bytes = create_voucher_docx(v_data)
                            
                            c_dl, c_del = st.columns([1, 1])
                            with c_dl:
                                st.download_button(
                                    label="📥 Tải Phiếu (Word)",
                                    data=docx_bytes,
                                    file_name=f"Phieu_{txn['type']}_{txn['id']}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_hist_btn_{txn['id']}",
                                    use_container_width=True,
                                    type="primary"
                                )
                            
                            with c_del:
                                if st.button("🗑️ Xóa giao dịch này", key=f"del_hist_btn_{txn['id']}", use_container_width=True):
                                    run_query("DELETE FROM transaction_history WHERE id=?", (txn['id'],), commit=True)
                                    st.success("Đã xóa!")
                                    time.sleep(0.5)
                                    st.rerun()
                    else:
                        st.info("Chưa có lịch sử giao dịch cho mã này.")
            else:
                st.info("👆 Vui lòng chọn một Mã Tour hoặc Mã Booking để xem công nợ.")

    with tab_summary:
        st.subheader("Tổng hợp các khoản phải thu")
        with st.spinner("Đang tính toán công nợ..."):
            # 1. Lấy tất cả giao dịch và tính toán số tiền đã trả cho mỗi mã
            all_txns = run_query("SELECT ref_code, type, amount FROM transaction_history")
            paid_amounts = {}
            if all_txns:
                df_txns = pd.DataFrame([dict(r) for r in all_txns])
                if not df_txns.empty:
                    df_thu = df_txns[df_txns['type'] == 'THU'].groupby('ref_code')['amount'].sum()
                    df_chi = df_txns[df_txns['type'] == 'CHI'].groupby('ref_code')['amount'].sum()
                    paid_amounts = (df_thu.subtract(df_chi, fill_value=0)).to_dict()

            debt_records = []

            # 2. Lấy tất cả tour đang hoạt động và tính công nợ
            user_info_debt = st.session_state.get("user_info", {})
            user_role_debt = user_info_debt.get('role')
            user_name_debt = user_info_debt.get('name')
            # [FIX] Lấy tất cả tour chưa bị xóa (bao gồm cả tour đã hoàn thành) để tổng hợp công nợ
            tour_debt_query = "SELECT * FROM tours WHERE COALESCE(status, 'running') NOT IN ('deleted')"
            tour_debt_params = []
            if user_role_debt == 'sale' and user_name_debt:
                tour_debt_query += " AND sale_name=?"
                tour_debt_params.append(user_name_debt)
            active_tours = run_query(tour_debt_query, tuple(tour_debt_params))
            if active_tours:
                for tour_row in active_tours:
                    tour = dict(tour_row)
                    final_price = float(tour.get('final_tour_price', 0) or 0)
                    child_price = float(tour.get('child_price', 0) or 0)
                    final_qty = float(tour.get('final_qty', 0) or 0)
                    child_qty = float(tour.get('child_qty', 0) or 0)
                    if final_qty == 0: final_qty = float(tour.get('guest_count', 1))
                    contract_value = (final_price * final_qty) + (child_price * child_qty)
 
                    if contract_value > 0:
                        paid = paid_amounts.get(tour['tour_code'], 0.0)
                        remaining = contract_value - paid
                        if remaining > 0.1:
                            debt_records.append({'customer_name': tour.get('customer_name', 'N/A'), 'ref_name': tour['tour_name'], 'ref_code': tour['tour_code'], 'type': 'Tour', 'contract_value': contract_value, 'paid': paid, 'remaining': remaining})
 
            # 3. Lấy tất cả booking lẻ đang hoạt động và tính công nợ
            # [FIX] Lấy tất cả booking chưa bị xóa (bao gồm cả booking đã hoàn thành) để tổng hợp công nợ
            bk_debt_query = "SELECT * FROM service_bookings WHERE COALESCE(status, 'active') NOT IN ('deleted')"
            bk_debt_params = []
            if user_role_debt == 'sale' and user_name_debt:
                bk_debt_query += " AND sale_name=?"
                bk_debt_params.append(user_name_debt)
            active_bookings = run_query(bk_debt_query, tuple(bk_debt_params))
            if active_bookings:
                for booking_row in active_bookings:
                    booking = dict(booking_row)
                    contract_value = float(booking.get('selling_price', 0) or 0)
 
                    if contract_value > 0:
                        paid = paid_amounts.get(booking['code'], 0.0)
                        remaining = contract_value - paid
                        if remaining > 0.1:
                            customer_info = booking.get('customer_info', 'N/A')
                            customer_name = customer_info.split(' - ')[0] if ' - ' in customer_info else customer_info
                            debt_records.append({'customer_name': customer_name, 'ref_name': booking['name'], 'ref_code': booking['code'], 'type': 'Booking', 'contract_value': contract_value, 'paid': paid, 'remaining': remaining})
 
            # 4. Hiển thị kết quả
            if not debt_records:
                st.success("🎉 Không có công nợ nào cần thu.")
            else:
                df_debt = pd.DataFrame(debt_records)
                total_debt = df_debt['remaining'].sum()
                
                st.metric("TỔNG SỐ TIỀN CẦN THU", format_vnd(total_debt))
                
                st.divider()
                st.markdown("#### Danh sách khách hàng đang nợ")
                
                customer_debt = df_debt.groupby('customer_name')['remaining'].sum().reset_index().sort_values('remaining', ascending=False)
                customer_debt.columns = ['Khách hàng', 'Tổng nợ']
                
                st.dataframe(customer_debt, column_config={"Tổng nợ": st.column_config.NumberColumn(format="%d VND")}, use_container_width=True, hide_index=True)
                
                st.divider()
                st.markdown("#### Chi tiết các khoản nợ")
                st.dataframe(
                    df_debt.sort_values('remaining', ascending=False),
                    column_config={ 'customer_name': 'Khách hàng', 'ref_name': 'Tên Tour/Booking', 'ref_code': 'Mã', 'type': 'Loại', 'contract_value': st.column_config.NumberColumn("Giá trị HĐ", format="%d VND"), 'paid': st.column_config.NumberColumn("Đã thu", format="%d VND"), 'remaining': st.column_config.NumberColumn("Còn lại", format="%d VND"), },
                    use_container_width=True, hide_index=True
                )

                # --- TÍNH NĂNG XUẤT EXCEL CÔNG NỢ ---
                st.write("")
                if "debt_xls_data" not in st.session_state: st.session_state.debt_xls_data = None
                
                if st.button("📊 Tạo file Excel báo cáo"):
                    buffer_debt = io.BytesIO()
                    try:
                        with pd.ExcelWriter(buffer_debt, engine='xlsxwriter') as writer:
                            workbook: Any = writer.book
                            worksheet = workbook.add_worksheet('CongNo')
                            
                            # Formats
                            fmt_title = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter', 'font_color': '#B71C1C', 'font_name': 'Times New Roman'})
                            fmt_header = workbook.add_format({'bold': True, 'bg_color': '#FFEBEE', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'font_color': '#B71C1C', 'text_wrap': True, 'font_name': 'Times New Roman'})
                            fmt_text = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                            fmt_money = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_name': 'Times New Roman'})
                            fmt_comp = workbook.add_format({'bold': True, 'font_size': 12, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                            
                            # Company Info
                            comp_data = get_company_data()
                            worksheet.write('A1', comp_data['name'], fmt_comp)
                            worksheet.write('A2', f"ĐC: {comp_data['address']}")
                            worksheet.write('A3', f"MST: {comp_data['phone']}")
                            
                            # Title
                            worksheet.merge_range('A5:G5', "BÁO CÁO CÔNG NỢ KHÁCH HÀNG", fmt_title)
                            worksheet.write('A6', f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y')}")
                            
                            # Headers
                            headers = ['Khách hàng', 'Tên Tour/Booking', 'Mã', 'Loại', 'Giá trị HĐ', 'Đã thu', 'Còn lại']
                            for i, h in enumerate(headers):
                                worksheet.write(7, i, h, fmt_header)
                            
                            # Data
                            df_export = df_debt.sort_values(['customer_name', 'remaining'], ascending=[True, False])
                            
                            row = 8
                            for _, r in df_export.iterrows():
                                worksheet.write(row, 0, r['customer_name'], fmt_text)
                                worksheet.write(row, 1, r['ref_name'], fmt_text)
                                worksheet.write(row, 2, r['ref_code'], fmt_text)
                                worksheet.write(row, 3, r['type'], fmt_text)
                                worksheet.write(row, 4, r['contract_value'], fmt_money)
                                worksheet.write(row, 5, r['paid'], fmt_money)
                                worksheet.write(row, 6, r['remaining'], fmt_money)
                                row += 1
                                
                            # Total row
                            fmt_total = workbook.add_format({'bold': True, 'bg_color': '#FFCDD2', 'border': 1, 'num_format': '#,##0', 'align': 'right', 'font_name': 'Times New Roman'})
                            worksheet.merge_range(row, 0, row, 3, "TỔNG CỘNG", fmt_total)
                            worksheet.write(row, 4, df_export['contract_value'].sum(), fmt_total)
                            worksheet.write(row, 5, df_export['paid'].sum(), fmt_total)
                            worksheet.write(row, 6, df_export['remaining'].sum(), fmt_total)
                            
                            # Column widths
                            worksheet.set_column('A:A', 25)
                            worksheet.set_column('B:B', 35)
                            worksheet.set_column('C:D', 15)
                            worksheet.set_column('E:G', 18)

                        st.session_state.debt_xls_data = buffer_debt.getvalue()
                        st.rerun()
                    except Exception as e:
                        st.error(f"Lỗi tạo file Excel: {e}")

                if st.session_state.debt_xls_data:
                    st.download_button(
                        label="📥 Tải Báo Cáo Công Nợ (Excel)",
                        data=st.session_state.debt_xls_data,
                        file_name=f"BaoCao_CongNo_{datetime.now().strftime('%d%m%Y')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        type="primary"
                    )

def render_booking_management():
    st.title("🔖 Quản Lý Booking")
    st.caption("Quản lý các booking lẻ, booking dịch vụ (Không phải Tour trọn gói)")
    
    # Lấy thông tin user hiện tại để gán cho booking và lọc dữ liệu
    current_user_info = st.session_state.get("user_info", {})
    current_user_name = current_user_info.get('name', 'N/A')
    current_user_role = current_user_info.get('role')

    # --- 2. TÁCH LIÊN KẾT RA 2 PHẦN RIÊNG BIỆT ---
    tab1, tab2, tab3 = st.tabs(["✨ Tạo Booking", "🔗 Chi tiết Booking", "📜 Lịch sử Booking"])
    
    # ---------------- TAB 1: TẠO BOOKING ----------------
    with tab1:
        with st.container(border=True):
            st.markdown("### ➕ Tạo Booking Mới")
            
            # --- GỢI Ý KHÁCH HÀNG ---
            cust_query = "SELECT * FROM customers ORDER BY id DESC"
            cust_params = []
            if current_user_role == 'sale' and current_user_name:
                cust_query = "SELECT * FROM customers WHERE sale_name=? ORDER BY id DESC"
                cust_params.append(current_user_name)
            customers = run_query(cust_query, tuple(cust_params))
            cust_opts = ["-- Khách mới --"]
            if customers:
                cust_opts += [f"{c['name']} | {c['phone']}" for c in customers]
            sel_cust = st.selectbox("🔍 Chọn khách hàng cũ (Gợi ý):", cust_opts, key="bk_cust_suggest")
            
            pre_name, pre_phone = "", ""
            if sel_cust and sel_cust != "-- Khách mới --":
                parts = sel_cust.split(" | ")
                pre_name = parts[0]
                pre_phone = parts[1] if len(parts) > 1 else ""
            
            # Chọn loại dịch vụ
            bk_type = st.radio("Chọn loại dịch vụ:", ["🏨 Khách sạn", "🚌 Vận chuyển", "🧩 Combo / Đa dịch vụ", "🔖 Khác"], horizontal=True)
            st.divider()

            if bk_type == "🏨 Khách sạn":
                st.markdown("##### 🏨 Thông tin lưu trú & Tài chính")
                
                # [NEW] Move dates out to calculate nights
                c_date, c_room = st.columns([2, 1])
                dates = c_date.date_input("Thời gian lưu trú", value=[], help="Chọn ngày nhận và trả phòng", format="DD/MM/YYYY")
                room_count = c_room.number_input("Số lượng phòng", min_value=1, step=1, value=1)
                
                nights = 1
                if len(dates) == 2:
                    nights = (dates[1] - dates[0]).days
                    if nights < 1: nights = 1
                    st.caption(f"Thời gian: {dates[0].strftime('%d/%m')} - {dates[1].strftime('%d/%m')} ({nights} đêm) x {room_count} phòng")
                elif len(dates) == 1:
                    st.caption("Vui lòng chọn ngày trả phòng.")
                else:
                    st.caption("Vui lòng chọn ngày nhận và trả phòng.")

                # Financials
                c1, c2, c3 = st.columns(3)
                
                # [CODE MỚI] Xử lý nhập tiền có định dạng
                if "bk_hotel_net_val" not in st.session_state: st.session_state.bk_hotel_net_val = ""
                if "bk_hotel_sell_val" not in st.session_state: st.session_state.bk_hotel_sell_val = ""

                def fmt_hotel_net():
                    val = st.session_state.bk_hotel_net_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_hotel_net_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass
                
                def fmt_hotel_sell():
                    val = st.session_state.bk_hotel_sell_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_hotel_sell_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass

                net_price_input = c1.text_input("Giá nét / đêm / phòng", key="bk_hotel_net_val", on_change=fmt_hotel_net, help="Nhập số tiền (VD: 1000000)")
                selling_price_input = c2.text_input("Giá bán / đêm / phòng", key="bk_hotel_sell_val", on_change=fmt_hotel_sell, help="Nhập số tiền (VD: 1500000)")
                
                try: net_price_unit = float(net_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: net_price_unit = 0.0
                
                try: selling_price_unit = float(selling_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: selling_price_unit = 0.0
                
                # Tax logic
                tax_option = st.radio("Giá nét đã bao gồm thuế?", ["Đã bao gồm thuế", "Chưa bao gồm thuế"], horizontal=True)
                tax_percent = 0.0
                if tax_option == "Chưa bao gồm thuế":
                    tax_percent = st.number_input("Nhập % Thuế", min_value=0.0, max_value=100.0, step=0.5, format="%.1f")
                
                # Calculations
                net_price_unit_incl_tax = net_price_unit * (1 + tax_percent / 100)
                total_net = net_price_unit_incl_tax * nights * room_count
                total_sell = selling_price_unit * nights * room_count
                total_profit = total_sell - total_net
                
                # Display Table
                st.markdown("###### 📊 Bảng tính chi tiết")
                calc_df = pd.DataFrame({
                    "Loại": ["Giá Nét (Vốn)", "Giá Bán (Doanh thu)"],
                    "Đơn giá": [format_vnd(net_price_unit_incl_tax) + " VND", format_vnd(selling_price_unit) + " VND"],
                    "Số lượng": [f"{nights} đêm x {room_count} phòng", f"{nights} đêm x {room_count} phòng"],
                    "Thành tiền": [format_vnd(total_net) + " VND", format_vnd(total_sell) + " VND"]
                })
                st.dataframe(calc_df, use_container_width=True, hide_index=True)
                
                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 10px; border-radius: 5px; border: 1px solid #c8e6c9; text-align: center;">
                    <span style="color: #2e7d32; font-weight: bold; font-size: 1.1em;">LỢI NHUẬN DỰ KIẾN: {format_vnd(total_profit)} VND</span>
                </div>""", unsafe_allow_html=True)

                st.divider()
                st.text_input("Sales phụ trách", value=current_user_name, disabled=True)
                with st.form("bk_hotel", clear_on_submit=True):
                    h_name = st.text_input("Tên Khách sạn", placeholder="VD: Mường Thanh Luxury")
                    
                    # [NEW] Thêm các trường thông tin chi tiết
                    c_h1, c_h2 = st.columns(2)
                    hotel_code = c_h1.text_input("Mã code khách sạn (Booking ID)", placeholder="VD: 12345678")
                    room_type = c_h2.text_area("Hạng phòng", placeholder="VD: 2 Deluxe, 1 Suite (Xuống dòng nếu nhiều hạng)", height=68)
                    guest_list = st.text_area("Danh sách khách lưu trú", placeholder="VD: Nguyen Van A, Tran Thi B...", height=100)
                    
                    # Dates are already outside
                    
                    c_cust_n, c_cust_p = st.columns(2)
                    cust_name = c_cust_n.text_input("Tên khách hàng (*)", value=pre_name, placeholder="Nhập tên khách")
                    cust_phone = c_cust_p.text_input("Số điện thoại", value=pre_phone, placeholder="Nhập SĐT (Tùy chọn)")

                    new_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                    st.caption(f"Mã Booking dự kiến: {new_code}")
                    if st.form_submit_button("Tạo Booking Khách sạn", type="primary"):
                        if h_name and len(dates) == 2 and cust_name:
                            cust_info = f"{cust_name} - {cust_phone}" if cust_phone else cust_name
                            d_range = f"{dates[0].strftime('%d/%m/%Y')} - {dates[1].strftime('%d/%m/%Y')} ({nights} đêm, {room_count} phòng)"
                            save_customer_check(cust_name, cust_phone, current_user_name)

                            add_row_to_table('service_bookings', {
                                'code': new_code, 'name': f"[KS] {h_name}", 'created_at': datetime.now().strftime("%Y-%m-%d"),
                                'type': 'HOTEL', 'details': f"Lưu trú: {d_range}", 'customer_info': cust_info,
                                'net_price': total_net, # Storing TOTAL
                                'tax_percent': tax_percent,
                                'selling_price': total_sell, # Storing TOTAL
                                'profit': total_profit,
                                'sale_name': current_user_name,
                                'hotel_code': hotel_code,
                                'room_type': room_type,
                                'guest_list': guest_list
                            })
                            # Clear inputs
                            if "bk_hotel_net_val" in st.session_state: del st.session_state.bk_hotel_net_val
                            if "bk_hotel_sell_val" in st.session_state: del st.session_state.bk_hotel_sell_val
                            st.success("Đã tạo!"); time.sleep(0.5); st.rerun()
                        else: st.warning("Vui lòng nhập tên khách sạn, tên khách hàng và chọn đủ ngày đi/về.")

            elif bk_type == "🚌 Vận chuyển":
                trans_type = st.radio("Loại phương tiện:", ["Xe (Ô tô)", "Máy bay", "Tàu hỏa", "Du thuyền"], horizontal=True)
                
                st.divider()
                st.markdown("##### 💰 Thông tin tài chính")
                
                c_qty, c_net, c_sell = st.columns(3)
                qty = c_qty.number_input("Số lượng (Vé/Khách)", min_value=1, value=1, key="trans_qty")
                
                # [CODE MỚI] Xử lý nhập tiền có định dạng cho Vận chuyển
                if "bk_trans_net_val" not in st.session_state: st.session_state.bk_trans_net_val = ""
                if "bk_trans_sell_val" not in st.session_state: st.session_state.bk_trans_sell_val = ""

                def fmt_trans_net():
                    val = st.session_state.bk_trans_net_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_trans_net_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass
                
                def fmt_trans_sell():
                    val = st.session_state.bk_trans_sell_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_trans_sell_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass

                net_price_input = c_net.text_input("Giá nét / vé", key="bk_trans_net_val", on_change=fmt_trans_net, help="Nhập số tiền (VD: 1000000)")
                selling_price_input = c_sell.text_input("Giá bán / vé", key="bk_trans_sell_val", on_change=fmt_trans_sell, help="Nhập số tiền (VD: 1500000)")
                
                try: net_price_unit = float(net_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: net_price_unit = 0.0
                
                try: selling_price_unit = float(selling_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: selling_price_unit = 0.0
                
                tax_option = st.radio("Giá nét đã bao gồm thuế?", ["Đã bao gồm thuế", "Chưa bao gồm thuế"], horizontal=True, key="trans_tax_opt")
                tax_percent = 0.0
                
                if tax_option == "Chưa bao gồm thuế":
                    tax_percent = st.number_input("Nhập % Thuế", min_value=0.0, max_value=100.0, step=0.5, format="%.1f", key="trans_tax_pct")
                
                # Calculations
                net_price_unit_incl_tax = net_price_unit * (1 + tax_percent / 100)
                total_net = net_price_unit_incl_tax * qty
                total_sell = selling_price_unit * qty
                profit = total_sell - total_net

                # Display Table
                st.markdown("###### 📊 Bảng tính chi tiết")
                calc_df = pd.DataFrame({
                    "Loại": ["Giá Nét (Vốn)", "Giá Bán (Doanh thu)"],
                    "Đơn giá": [format_vnd(net_price_unit_incl_tax) + " VND", format_vnd(selling_price_unit) + " VND"],
                    "Số lượng": [qty, qty],
                    "Thành tiền": [format_vnd(total_net) + " VND", format_vnd(total_sell) + " VND"]
                })
                st.dataframe(calc_df, use_container_width=True, hide_index=True)

                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 10px; border-radius: 5px; border: 1px solid #c8e6c9; text-align: center;">
                    <span style="color: #2e7d32; font-weight: bold; font-size: 1.1em;">LỢI NHUẬN DỰ KIẾN: {format_vnd(profit)} VND</span>
                </div>""", unsafe_allow_html=True)

                st.divider()
                st.text_input("Sales phụ trách", value=current_user_name, disabled=True, key="trans_sale")
                with st.form("bk_trans", clear_on_submit=True):
                    details = ""
                    bk_name = ""
                    is_valid = False

                    if trans_type == "Xe (Ô tô)":
                        c1, c2 = st.columns(2)
                        route_from = c1.text_input("Điểm đi")
                        route_to = c2.text_input("Điểm đến")
                        c3, c4, c5 = st.columns(3)
                        car_type = c3.selectbox("Loại xe", ["4S", "7S", "16S", "29S", "35S", "45S"])
                        car_no = c4.text_input("Biển số / Mã xe")
                        t_date = c5.date_input("Ngày đi", format="DD/MM/YYYY")
                        
                        # [NEW] Thêm thông tin tài xế
                        c_drv1, c_drv2 = st.columns(2)
                        driver_name = c_drv1.text_input("Tên tài xế")
                        driver_phone = c_drv2.text_input("SĐT Tài xế")
                        
                        if route_from and route_to:
                            is_valid = True
                            bk_name = f"[XE] {route_from} - {route_to}"
                            details = f"Xe {car_type}: {car_no} | Ngày: {t_date.strftime('%d/%m/%Y')} | SL: {qty}"
                            if driver_name:
                                details += f" | Tài xế: {driver_name}"
                            if driver_phone:
                                details += f" - {driver_phone}"

                    elif trans_type == "Máy bay":
                        c1, c2 = st.columns(2)
                        ticket_code = c1.text_input("Mã vé / Số hiệu")
                        flight_date = c2.date_input("Ngày bay", format="DD/MM/YYYY")
                        flight_route = st.text_input("Hành trình / Hãng bay (Tùy chọn)", placeholder="VD: VN123 HAN-SGN")
                        
                        if ticket_code:
                            is_valid = True
                            desc = flight_route if flight_route else ticket_code
                            bk_name = f"[BAY] {desc}"
                            details = f"Vé: {ticket_code} | Ngày: {flight_date.strftime('%d/%m/%Y')} | SL: {qty}"

                    elif trans_type == "Tàu hỏa":
                        c1, c2 = st.columns(2)
                        ticket_code = c1.text_input("Mã vé / Toa / Ghế")
                        train_date = c2.date_input("Ngày đi", format="DD/MM/YYYY")
                        train_route = st.text_input("Ga đi - Ga đến (Tùy chọn)", placeholder="VD: Hà Nội - Vinh")
                        
                        if ticket_code:
                            is_valid = True
                            desc = train_route if train_route else ticket_code
                            bk_name = f"[TAU] {desc}"
                            details = f"Vé: {ticket_code} | Ngày: {train_date.strftime('%d/%m/%Y')} | SL: {qty}"

                    elif trans_type == "Du thuyền":
                        c1, c2 = st.columns(2)
                        cruise_name = c1.text_input("Tên du thuyền / Tuyến")
                        cruise_date = c2.date_input("Ngày đi", format="DD/MM/YYYY")
                        cabin_type = st.text_input("Loại Cabin / Ghi chú", placeholder="VD: Junior Suite, Balcony...")
                        
                        if cruise_name:
                            is_valid = True
                            bk_name = f"[THUYEN] {cruise_name}"
                            details = f"Cabin: {cabin_type} | Ngày: {cruise_date.strftime('%d/%m/%Y')} | SL: {qty}"

                    st.divider()
                    c_cust_n, c_cust_p = st.columns(2)
                    cust_name = c_cust_n.text_input("Tên khách hàng (*)", value=pre_name, placeholder="Nhập tên khách")
                    cust_phone = c_cust_p.text_input("Số điện thoại", value=pre_phone, placeholder="Nhập SĐT (Tùy chọn)")

                    new_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                    st.caption(f"Mã Booking dự kiến: {new_code}")
                    if st.form_submit_button("Tạo Booking Vận chuyển", type="primary"):
                        if is_valid and cust_name:
                            cust_info = f"{cust_name} - {cust_phone}" if cust_phone else cust_name
                            save_customer_check(cust_name, cust_phone, current_user_name)
                            add_row_to_table('service_bookings', {
                                'code': new_code, 'name': bk_name, 'created_at': datetime.now().strftime("%Y-%m-%d"),
                                'type': 'TRANS', 'details': details, 'customer_info': cust_info,
                                'net_price': total_net,
                                'tax_percent': tax_percent,
                                'selling_price': total_sell, 'profit': profit,
                                'sale_name': current_user_name
                            })
                            # Clear inputs
                            if "bk_trans_net_val" in st.session_state: del st.session_state.bk_trans_net_val
                            if "bk_trans_sell_val" in st.session_state: del st.session_state.bk_trans_sell_val
                            st.success("Đã tạo!"); time.sleep(0.5); st.rerun()
                        else: st.warning("Vui lòng nhập đủ thông tin (Hành trình/Mã vé và Tên khách).")

            elif bk_type == "🧩 Combo / Đa dịch vụ":
                if "combo_list" not in st.session_state: st.session_state.combo_list = []
                c_add, c_list = st.columns([1, 1.5])
                with c_add:
                    st.markdown("##### Thêm dịch vụ con")
                    sub_type = st.selectbox("Loại", ["Khách sạn", "Vận chuyển", "Khác"], key="cb_sub")
                    if sub_type == "Khách sạn":
                        sh_n = st.text_input("Tên KS", key="cb_h_n")
                        c_qty, c_date = st.columns([1, 2])
                        sh_qty = c_qty.number_input("Số lượng phòng", min_value=1, value=1, key="cb_h_q")
                        sh_d = c_date.date_input("Ngày ở", [], key="cb_h_d", format="DD/MM/YYYY")
                        if st.button("Thêm KS") and sh_n and len(sh_d)==2:
                            st.session_state.combo_list.append(f"🏨 {sh_n} - {sh_qty} phòng ({sh_d[0].strftime('%d/%m')} - {sh_d[1].strftime('%d/%m')})"); st.rerun()
                    elif sub_type == "Vận chuyển":
                        tr_mode = st.selectbox("Loại phương tiện", ["Xe", "Máy bay", "Tàu hỏa", "Du thuyền"], key="cb_tr_mode")
                        st_r = st.text_input("Hành trình / Mã vé / Tên tàu", key="cb_t_r")
                        st_d = st.date_input("Ngày", key="cb_t_d", format="DD/MM/YYYY")
                        
                        icon_map = {"Xe": "🚌", "Máy bay": "✈️", "Tàu hỏa": "🚆", "Du thuyền": "🚢"}
                        
                        if st.button("Thêm Vận chuyển") and st_r:
                            st.session_state.combo_list.append(f"{icon_map[tr_mode]} {st_r} ({st_d.strftime('%d/%m')})"); st.rerun()
                    else:
                        so_n = st.text_input("Tên dịch vụ", key="cb_o_n")
                        if st.button("Thêm DV") and so_n:
                            st.session_state.combo_list.append(f"🔖 {so_n}"); st.rerun()
                with c_list:
                    st.markdown("##### Danh sách đã thêm")
                    # [FIX] Dùng list() để ép kiểu rõ ràng, Pylance sẽ hiểu đây là danh sách lặp được
                    safe_combo_list = list(st.session_state.get("combo_list", []))
                    for i, item in enumerate(safe_combo_list): st.text(f"{i+1}. {item}")
                    
                    if st.session_state.get("combo_list") and st.button("Xóa hết", type="secondary"): 
                        st.session_state.combo_list = []
                        st.rerun()
                
                st.divider()
                st.markdown("##### 💰 Thông tin tài chính")
                
                c_qty, c_net, c_sell = st.columns(3)
                qty = c_qty.number_input("Số lượng (Combo/Pax)", min_value=1, value=1, key="combo_qty")
                
                # [CODE MỚI] Xử lý nhập tiền có định dạng cho Combo
                if "bk_combo_net_val" not in st.session_state: st.session_state.bk_combo_net_val = ""
                if "bk_combo_sell_val" not in st.session_state: st.session_state.bk_combo_sell_val = ""

                def fmt_combo_net():
                    val = st.session_state.bk_combo_net_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_combo_net_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass
                
                def fmt_combo_sell():
                    val = st.session_state.bk_combo_sell_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_combo_sell_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass

                net_price_input = c_net.text_input("Giá nét / combo", key="bk_combo_net_val", on_change=fmt_combo_net, help="Nhập số tiền (VD: 1000000)")
                selling_price_input = c_sell.text_input("Giá bán / combo", key="bk_combo_sell_val", on_change=fmt_combo_sell, help="Nhập số tiền (VD: 1500000)")
                
                try: net_price_unit = float(net_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: net_price_unit = 0.0
                
                try: selling_price_unit = float(selling_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: selling_price_unit = 0.0
                
                tax_option = st.radio("Giá nét đã bao gồm thuế?", ["Đã bao gồm thuế", "Chưa bao gồm thuế"], horizontal=True, key="combo_tax_opt")
                tax_percent = 0.0
                
                if tax_option == "Chưa bao gồm thuế":
                    tax_percent = st.number_input("Nhập % Thuế", min_value=0.0, max_value=100.0, step=0.5, format="%.1f", key="combo_tax_pct")
                
                # Calculations
                net_price_unit_incl_tax = net_price_unit * (1 + tax_percent / 100)
                total_net = net_price_unit_incl_tax * qty
                total_sell = selling_price_unit * qty
                profit = total_sell - total_net

                # Display Table
                st.markdown("###### 📊 Bảng tính chi tiết")
                calc_df = pd.DataFrame({
                    "Loại": ["Giá Nét (Vốn)", "Giá Bán (Doanh thu)"],
                    "Đơn giá": [format_vnd(net_price_unit_incl_tax) + " VND", format_vnd(selling_price_unit) + " VND"],
                    "Số lượng": [qty, qty],
                    "Thành tiền": [format_vnd(total_net) + " VND", format_vnd(total_sell) + " VND"]
                })
                st.dataframe(calc_df, use_container_width=True, hide_index=True)

                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 10px; border-radius: 5px; border: 1px solid #c8e6c9; text-align: center;">
                    <span style="color: #2e7d32; font-weight: bold; font-size: 1.1em;">LỢI NHUẬN DỰ KIẾN: {format_vnd(profit)} VND</span>
                </div>""", unsafe_allow_html=True)

                st.divider()
                st.text_input("Sales phụ trách", value=current_user_name, disabled=True, key="combo_sale")
                with st.form("bk_combo", clear_on_submit=True):
                    combo_name = st.text_input("Tên Combo / Gói", placeholder="VD: Combo Đà Nẵng 3N2Đ")
                    # [NEW] Thêm danh sách khách cho Combo
                    guest_list_cb = st.text_area("Danh sách khách (Đoàn)", placeholder="Nhập danh sách khách hàng...", height=100)
                    
                    c_cust_n, c_cust_p = st.columns(2)
                    cust_name = c_cust_n.text_input("Tên khách hàng (*)", value=pre_name, placeholder="Nhập tên khách")
                    cust_phone = c_cust_p.text_input("Số điện thoại", value=pre_phone, placeholder="Nhập SĐT (Tùy chọn)")

                    new_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                    if st.form_submit_button("Lưu Combo", type="primary"):
                        if combo_name and st.session_state.combo_list and cust_name:
                            cust_info = f"{cust_name} - {cust_phone}" if cust_phone else cust_name
                            save_customer_check(cust_name, cust_phone, current_user_name)
                            details_str = " | ".join(st.session_state.combo_list) + f" (SL: {qty})"
                            add_row_to_table('service_bookings', {
                                'code': new_code, 'name': f"[CB] {combo_name}", 'created_at': datetime.now().strftime("%Y-%m-%d"),
                                'type': 'COMBO', 'details': details_str, 'customer_info': cust_info,
                                'net_price': total_net,
                                'tax_percent': tax_percent,
                                'selling_price': total_sell, 'profit': profit,
                                'sale_name': current_user_name,
                                'guest_list': guest_list_cb
                            })
                            # Clear inputs
                            if "bk_combo_net_val" in st.session_state: del st.session_state.bk_combo_net_val
                            if "bk_combo_sell_val" in st.session_state: del st.session_state.bk_combo_sell_val
                            st.session_state.combo_list = []; st.success("Đã tạo!"); time.sleep(0.5); st.rerun()
                        else: st.warning("Cần tên Combo, tên khách hàng và ít nhất 1 dịch vụ.")

            else:
                st.markdown("##### 💰 Thông tin tài chính")
                c_qty, c_net, c_sell = st.columns(3)
                qty = c_qty.number_input("Số lượng", min_value=1, value=1, key="other_qty")
                
                # [CODE MỚI] Xử lý nhập tiền có định dạng cho Khác
                if "bk_other_net_val" not in st.session_state: st.session_state.bk_other_net_val = ""
                if "bk_other_sell_val" not in st.session_state: st.session_state.bk_other_sell_val = ""

                def fmt_other_net():
                    val = st.session_state.bk_other_net_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_other_net_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass
                
                def fmt_other_sell():
                    val = st.session_state.bk_other_sell_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_other_sell_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass

                net_price_input = c_net.text_input("Giá nét / đơn vị", key="bk_other_net_val", on_change=fmt_other_net, help="Nhập số tiền (VD: 1000000)")
                selling_price_input = c_sell.text_input("Giá bán / đơn vị", key="bk_other_sell_val", on_change=fmt_other_sell, help="Nhập số tiền (VD: 1500000)")
                
                try: net_price_unit = float(net_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: net_price_unit = 0.0
                
                try: selling_price_unit = float(selling_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: selling_price_unit = 0.0
                
                tax_option = st.radio("Giá nét đã bao gồm thuế?", ["Đã bao gồm thuế", "Chưa bao gồm thuế"], horizontal=True, key="other_tax_opt")
                tax_percent = 0.0
                
                if tax_option == "Chưa bao gồm thuế":
                    tax_percent = st.number_input("Nhập % Thuế", min_value=0.0, max_value=100.0, step=0.5, format="%.1f", key="other_tax_pct")
                
                # Calculations
                net_price_unit_incl_tax = net_price_unit * (1 + tax_percent / 100)
                total_net = net_price_unit_incl_tax * qty
                total_sell = selling_price_unit * qty
                profit = total_sell - total_net

                # Display Table
                st.markdown("###### 📊 Bảng tính chi tiết")
                calc_df = pd.DataFrame({
                    "Loại": ["Giá Nét (Vốn)", "Giá Bán (Doanh thu)"],
                    "Đơn giá": [format_vnd(net_price_unit_incl_tax) + " VND", format_vnd(selling_price_unit) + " VND"],
                    "Số lượng": [qty, qty],
                    "Thành tiền": [format_vnd(total_net) + " VND", format_vnd(total_sell) + " VND"]
                })
                st.dataframe(calc_df, use_container_width=True, hide_index=True)

                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 10px; border-radius: 5px; border: 1px solid #c8e6c9; text-align: center;">
                    <span style="color: #2e7d32; font-weight: bold; font-size: 1.1em;">LỢI NHUẬN DỰ KIẾN: {format_vnd(profit)} VND</span>
                </div>""", unsafe_allow_html=True)

                st.divider()
                st.text_input("Sales phụ trách", value=current_user_name, disabled=True, key="other_sale")
                with st.form("bk_other", clear_on_submit=True):
                    new_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                    c1, c2 = st.columns([1, 3])
                    c1.text_input("Mã (Auto)", value=new_code, disabled=True)
                    new_name = c2.text_input("Tên Booking / Dịch vụ")
                    
                    c_cust_n, c_cust_p = st.columns(2)
                    cust_name = c_cust_n.text_input("Tên khách hàng (*)", value=pre_name, placeholder="Nhập tên khách")
                    cust_phone = c_cust_p.text_input("Số điện thoại", value=pre_phone, placeholder="Nhập SĐT (Tùy chọn)")

                    if st.form_submit_button("Tạo"):
                        if new_name and cust_name:
                            cust_info = f"{cust_name} - {cust_phone}" if cust_phone else cust_name
                            save_customer_check(cust_name, cust_phone, current_user_name)
                            add_row_to_table('service_bookings', {
                                'code': new_code, 'name': new_name, 'created_at': datetime.now().strftime("%Y-%m-%d"),
                                'type': 'OTHER', 'customer_info': cust_info, 'details': f"SL: {qty}",
                                'net_price': total_net,
                                'tax_percent': tax_percent,
                                'selling_price': total_sell, 'profit': profit,
                                'sale_name': current_user_name
                            })
                            # Clear inputs
                            if "bk_other_net_val" in st.session_state: del st.session_state.bk_other_net_val
                            if "bk_other_sell_val" in st.session_state: del st.session_state.bk_other_sell_val
                            st.success("Đã tạo!"); time.sleep(0.5); st.rerun()
                        else: st.warning("Vui lòng nhập tên dịch vụ và tên khách hàng.")

    # ---------------- TAB 2: KHỚP UNC & HÓA ĐƠN (DỰ ÁN UNC) ----------------
    with tab2:
        st.subheader("🔗 Chi tiết Booking")
        # --- Lọc danh sách booking theo sale ---
        bk_query = "SELECT * FROM service_bookings WHERE status='active'"
        bk_params = []
        if current_user_role == 'sale' and current_user_name:
            bk_query += " AND sale_name=?"
            bk_params.append(current_user_name)
        bk_query += " ORDER BY id DESC"
        bookings = run_query(bk_query, tuple(bk_params))
        
        if bookings:
            bk_map = {f"[{b['code']}] {b['name']}": b['code'] for b in bookings} # type: ignore
            selected_bk_label = st.selectbox("Chọn Booking để xem chi tiết:", list(bk_map.keys()))
            
            if selected_bk_label:
                code = bk_map[selected_bk_label] # type: ignore
                
                bk_info = run_query("SELECT * FROM service_bookings WHERE code=?", (code,), fetch_one=True)
                st.divider()
                st.markdown(f"### 📊 Chi tiết: {selected_bk_label}")
                if isinstance(bk_info, sqlite3.Row):
                    st.markdown("##### 💰 Tổng quan tài chính")
                    fin1, fin2, fin3 = st.columns(3)
                    net_p = bk_info['net_price'] or 0 # type: ignore
                    sell_p = bk_info['selling_price'] or 0 # type: ignore
                    prof_p = bk_info['profit'] or 0 # type: ignore
                    fin1.metric("Giá nét (đã gồm thuế)", format_vnd(net_p))
                    fin2.metric("Giá bán", format_vnd(sell_p))
                    fin3.metric("Lợi nhuận", format_vnd(prof_p))

                    if bk_info['customer_info']:
                        st.markdown(f"**👤 Khách hàng:** {bk_info['customer_info']}")
                    if bk_info['details']:
                        st.info(f"ℹ️ **Thông tin:** {bk_info['details']}")
                
                # Gọi hàm hiển thị so sánh
                render_cost_comparison(code)
                
                # --- NÚT TẢI BOOKING CONFIRMATION (MỚI) ---
                st.write("")
                c_lang, c_dl_btn = st.columns([1, 2])
                sel_lang = c_lang.radio("Ngôn ngữ PDF:", ["Tiếng Việt", "English"], horizontal=True)
                lang_code = 'vi' if sel_lang == "Tiếng Việt" else 'en'
                
                comp_data_cfm = get_company_data()
                pdf_cfm = create_booking_cfm_pdf(dict(bk_info), comp_data_cfm, lang=lang_code)
                c_dl_btn.download_button(
                    label="📥 Tải Booking Confirmation (PDF)",
                    data=pdf_cfm,
                    file_name=f"Booking_CFM_{code}.pdf",
                    mime="application/pdf",
                    type="secondary"
                )
                
                st.divider()
                # Nút hoàn tất & xóa booking
                c_complete, c_delete = st.columns(2)
                if c_complete.button("✅ Hoàn tất Booking", type="primary", use_container_width=True):
                    run_query("UPDATE service_bookings SET status='completed' WHERE code=?", (code,), commit=True)
                    st.success("Đã hoàn tất! Booking đã được chuyển sang tab Lịch sử."); time.sleep(1); st.rerun()

                if c_delete.button("🗑️ Xóa Booking này", use_container_width=True):
                    run_query("UPDATE service_bookings SET status='deleted' WHERE code=?", (code,), commit=True)
                    st.success("Đã xóa!"); time.sleep(0.5); st.rerun()
        else:
            st.info("Chưa có booking nào.")

    # ---------------- TAB 3: LỊCH SỬ BOOKING ----------------
    with tab3:
        st.subheader("📜 Lịch sử Booking đã hoàn tất")
        # --- Lọc danh sách booking theo sale ---
        hist_bk_query = "SELECT * FROM service_bookings WHERE status='completed'"
        hist_bk_params = []
        if current_user_role == 'sale' and current_user_name:
            hist_bk_query += " AND sale_name=?"
            hist_bk_params.append(current_user_name)
        hist_bk_query += " ORDER BY id DESC"
        history_bk = run_query(hist_bk_query, tuple(hist_bk_params))
        if history_bk:
            df_hist = pd.DataFrame([dict(r) for r in history_bk])
            st.dataframe(
                df_hist[['code', 'name', 'created_at', 'type', 'customer_info', 'details', 'net_price', 'selling_price', 'profit']],
                column_config={
                    "code": "Mã Booking",
                    "name": "Tên Booking",
                    "created_at": "Ngày tạo",
                    "type": "Loại",
                    "customer_info": "Khách hàng",
                    "details": "Chi tiết",
                    "net_price": st.column_config.NumberColumn("Giá nét", format="%d"),
                    "selling_price": st.column_config.NumberColumn("Giá bán", format="%d"),
                    "profit": st.column_config.NumberColumn("Lợi nhuận", format="%d"),
                },
                use_container_width=True,
                hide_index=True
            )
        else:
            st.info("Chưa có booking nào hoàn tất.")

def render_tour_management():
    st.title("📦 Quản Lý Tour ")
    
    # Sử dụng Tabs theo yêu cầu
    tab_est, tab_list_srv, tab_act, tab_hist, tab_rpt = st.tabs(["📝 Dự Toán Chi Phí", "📋 Danh sách & Dịch vụ", "💸 Quyết Toán Tour", "📜 Lịch sử Tour", "📈 Tổng Hợp Lợi Nhuận"])
    
    # Lấy thông tin user hiện tại để lọc
    current_user_info_tour = st.session_state.get("user_info", {})
    current_user_name_tour = current_user_info_tour.get('name', 'N/A')
    current_user_role_tour = current_user_info_tour.get('role')

    # Lấy danh sách Tour cho Selectbox dùng chung
    all_tours_query = "SELECT * FROM tours ORDER BY id DESC"
    all_tours_params = []
    if current_user_role_tour == 'sale' and current_user_name_tour:
        all_tours_query = "SELECT * FROM tours WHERE sale_name=? ORDER BY id DESC"
        all_tours_params.append(current_user_name_tour)
    all_tours = run_query(all_tours_query, tuple(all_tours_params))
    running_tours = [t for t in all_tours if t['status'] == 'running']
    tour_options = {f"[{t['tour_code']}] {t['tour_name']} ({t['start_date']})": t['id'] for t in running_tours} if running_tours else {} # type: ignore
    
    # ---------------- TAB 1: DỰ TOÁN CHI PHÍ ----------------
    with tab_est:
        with st.expander("➕ Tạo Thông Tin Đoàn Mới", expanded=False):
            # --- GỢI Ý KHÁCH HÀNG ---
            cust_query_t = "SELECT * FROM customers ORDER BY id DESC"
            cust_params_t = []
            if current_user_role_tour == 'sale' and current_user_name_tour:
                cust_query_t = "SELECT * FROM customers WHERE sale_name=? ORDER BY id DESC"
                cust_params_t.append(current_user_name_tour)
            customers = run_query(cust_query_t, tuple(cust_params_t))
            cust_opts_t = ["-- Khách mới --"] + [f"{c['name']} | {c['phone']}" for c in customers] if customers else ["-- Khách mới --"] # type: ignore
            sel_cust_t = st.selectbox("🔍 Gợi ý khách hàng:", cust_opts_t, key="tour_cust_suggest")
            
            t_pre_name, t_pre_phone = "", ""
            if sel_cust_t and sel_cust_t != "-- Khách mới --":
                parts = sel_cust_t.split(" | ")
                t_pre_name = parts[0]
                t_pre_phone = parts[1] if len(parts) > 1 else ""

            with st.form("create_tour_form", clear_on_submit=True):
                c1, c2 = st.columns(2)
                t_name = c1.text_input("Tên Đoàn")
                t_sale = c2.text_input("Sales phụ trách", value=current_user_name_tour, disabled=True)
                c_cust1, c_cust2 = st.columns(2)
                t_cust_name = c_cust1.text_input("Tên Khách / Đại diện", value=t_pre_name)
                t_cust_phone = c_cust2.text_input("SĐT Khách", value=t_pre_phone)
                c3, c4, c5 = st.columns(3)
                t_start = c3.date_input("Ngày đi", format="DD/MM/YYYY")
                t_end = c4.date_input("Ngày về", format="DD/MM/YYYY")
                t_pax = c5.number_input("Số lượng khách", min_value=1, step=1)
                
                if st.form_submit_button("Tạo Đoàn"):
                    if t_name:
                        save_customer_check(t_cust_name, t_cust_phone, current_user_name_tour)
                        new_tour_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                        add_row_to_table('tours', {
                            'tour_name': t_name, 'sale_name': current_user_name_tour, 'start_date': t_start.strftime('%d/%m/%Y'),
                            'end_date': t_end.strftime('%d/%m/%Y'), 'guest_count': t_pax, 'created_at': datetime.now().strftime('%Y-%m-%d'),
                            'tour_code': new_tour_code, 'customer_name': t_cust_name, 'customer_phone': t_cust_phone
                        })
                        st.success(f"Đã tạo đoàn mới! Mã tour: {new_tour_code}. Hãy chọn ở danh sách bên dưới để làm dự toán.")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("Vui lòng nhập tên đoàn.")

        st.divider()
        st.subheader("Bảng Tính Dự Toán")
        
        selected_tour_label = st.selectbox("Chọn Đoàn để làm dự toán:", list(tour_options.keys()) if tour_options else [], key="sel_tour_est")
        
        if selected_tour_label:
            tour_id = tour_options[selected_tour_label] # type: ignore
            tour_info = next((t for t in all_tours if t['id'] == tour_id), None)
            if not tour_info:
                st.error("Không tìm thấy thông tin tour.")
                st.stop()
            assert tour_info is not None

            # --- TOOLBAR: SỬA / XÓA TOUR ---
            c_ren, c_del = st.columns(2)
            with c_ren:
                with st.popover("✏️ Sửa thông tin", use_container_width=True):
                    with st.form(f"edit_tour_{tour_id}"):
                        en_n = st.text_input("Tên Đoàn", value=tour_info['tour_name']) # type: ignore
                        en_s = st.text_input("Sales", value=tour_info['sale_name']) # type: ignore
                        en_p = st.number_input("Số khách", value=tour_info['guest_count'], min_value=1) # type: ignore
                        if st.form_submit_button("Lưu thay đổi"):
                            if en_n != tour_info['tour_name']: # type: ignore
                                run_query("UPDATE tours SET pending_name=?, sale_name=?, guest_count=? WHERE id=?", (en_n, en_s, en_p, tour_id), commit=True)
                                st.success("Đã cập nhật thông tin & Gửi yêu cầu đổi tên (Chờ Admin duyệt)!"); time.sleep(0.5); st.rerun()
                            else:
                                run_query("UPDATE tours SET sale_name=?, guest_count=? WHERE id=?", (en_s, en_p, tour_id), commit=True)
                                st.success("Đã cập nhật!"); time.sleep(0.5); st.rerun()
            with c_del:
                req_status = tour_info['request_delete'] # type: ignore
                if req_status == 0:
                    with st.popover("🗑️ Yêu cầu xóa", use_container_width=True):
                        st.warning(f"Gửi yêu cầu xóa đoàn: {tour_info['tour_name']}?") # type: ignore
                        if st.button("Gửi yêu cầu", type="primary", use_container_width=True, key=f"req_del_t_{tour_id}"):
                            run_query("UPDATE tours SET request_delete=1 WHERE id=?", (tour_id,), commit=True)
                            st.success("Đã gửi yêu cầu xóa (Chờ Admin duyệt)!"); time.sleep(0.5); st.rerun()
                elif req_status == 1:
                    st.warning("⏳ Đang chờ Admin duyệt xóa...")
                    if st.button("Hủy yêu cầu", key=f"cancel_req_{tour_id}", use_container_width=True): # type: ignore
                        run_query("UPDATE tours SET request_delete=0 WHERE id=?", (tour_id,), commit=True)
                        st.rerun()
                elif req_status == 2:
                    st.success("✅ Admin đã duyệt xóa!")
                    c_conf, c_can = st.columns(2)
                    if c_conf.button("🗑️ Xóa ngay", type="primary", key=f"final_del_{tour_id}"): # type: ignore
                        run_query("DELETE FROM tours WHERE id=?", (tour_id,), commit=True)
                        run_query("DELETE FROM tour_items WHERE tour_id=?", (tour_id,), commit=True)
                        st.success("Đã xóa vĩnh viễn!"); time.sleep(0.5); st.rerun()
                    if c_can.button("Hủy xóa", key=f"keep_tour_{tour_id}"):
                        run_query("UPDATE tours SET request_delete=0 WHERE id=?", (tour_id,), commit=True)
                        st.rerun()

            # Reset edit mode when changing tour
            if st.session_state.current_tour_id_est != tour_id:
                st.session_state.est_edit_mode = False
                st.session_state.current_tour_id_est = tour_id
                if "est_df_temp" in st.session_state: del st.session_state.est_df_temp
                st.session_state.est_editor_key += 1
                
                # [NEW] Reset price values in session state
                t_dict_init = dict(tour_info)
                fp = float(t_dict_init.get('final_tour_price', 0) or 0)
                cp = float(t_dict_init.get('child_price', 0) or 0)
                st.session_state.est_final_price_val = "{:,.0f}".format(fp).replace(",", ".") + " VND"
                st.session_state.est_child_price_val = "{:,.0f}".format(cp).replace(",", ".") + " VND"
            
            # --- IMPORT EXCEL (MỚI - DỰ TOÁN) ---
            with st.expander("📥 Nhập dữ liệu từ Excel (Import)", expanded=False):
                st.caption("💡 File Excel cần có dòng tiêu đề: **Hạng mục, Diễn giải, Đơn vị, Đơn giá, Số lượng, Số lần**")
                
                # Widget upload file
                uploaded_est_file = st.file_uploader("Chọn file Excel dự toán", type=["xlsx", "xls"], key="up_est_tool")
                
                if uploaded_est_file:
                    if st.button("🚀 Đọc file & Điền vào bảng", type="primary"):
                        try:
                            # 1. Đọc file Excel (Tìm dòng tiêu đề tự động)
                            uploaded_est_file.seek(0)
                            df_raw = pd.read_excel(uploaded_est_file, header=None)
                            
                            header_idx = 0
                            detect_kws = ['hạng mục', 'tên hàng', 'diễn giải', 'đơn giá', 'số lượng', 'thành tiền', 'item', 'price', 'qty', 'đvt']
                            
                            # Quét 15 dòng đầu để tìm dòng chứa nhiều từ khóa nhất
                            for i in range(min(15, len(df_raw))):
                                row_vals = [str(x).lower() for x in df_raw.iloc[i].tolist()]
                                if sum(1 for kw in detect_kws if any(kw in val for val in row_vals)) >= 2:
                                    header_idx = i
                                    break
                            
                            uploaded_est_file.seek(0)
                            df_in = pd.read_excel(uploaded_est_file, header=header_idx)
                            
                            # 2. Chuẩn hóa tên cột
                            # Chuyển hết về chữ thường để so sánh
                            df_in.columns = [str(c).lower().strip() for c in df_in.columns]
                            
                            # Định nghĩa các từ khóa (Aliases) cho từng cột DB - Ưu tiên từ trái sang phải
                            col_aliases = {
                                'category': ['hạng mục', 'hang muc', 'tên hàng', 'ten hang', 'tên dịch vụ', 'ten dich vu', 'nội dung', 'noi dung', 'item'],
                                'description': ['diễn giải', 'dien giai', 'chi tiết', 'chi tiet', 'ghi chú', 'ghi chu', 'mô tả', 'mo ta', 'description', 'desc'],
                                'unit': ['đơn vị', 'don vi', 'đvt', 'dvt', 'unit', 'uom'],
                                'quantity': ['số lượng', 'so luong', 'sl', 'qty', 'quantity', 'vol'],
                                'unit_price': ['đơn giá', 'don gia', 'giá', 'gia', 'price', 'unit_price', 'unit price'],
                                'times': ['số lần', 'so lan', 'lần', 'lan', 'times']
                            }
                            
                            # Xác định cột nào trong Excel map vào cột nào trong DB
                            final_col_map = {}
                            for db_col, aliases in col_aliases.items():
                                for alias in aliases:
                                    if alias in df_in.columns:
                                        final_col_map[db_col] = alias
                                        break
                            
                            new_data = []
                            if not final_col_map:
                                st.warning("⚠️ Không tìm thấy các cột thông tin cần thiết (Hạng mục, Đơn giá...). Vui lòng kiểm tra tên cột trong file Excel.")
                            else:
                                for _, row in df_in.iterrows():
                                    item = {}
                                    for db_col, xls_col in final_col_map.items():
                                        val = row[xls_col]
                                        if pd.isna(val):
                                            val = 0 if db_col in ['quantity', 'unit_price', 'times'] else ""
                                        item[db_col] = val
                                    
                                    # Default values
                                    if 'category' not in item: item['category'] = ""
                                    if 'description' not in item: item['description'] = ""
                                    if 'unit' not in item: item['unit'] = ""
                                    
                                    # Safe numeric conversion
                                    def safe_float(v):
                                        try: return float(v)
                                        except: return 0.0
                                    
                                    item['quantity'] = safe_float(item.get('quantity', 1))
                                    item['unit_price'] = safe_float(item.get('unit_price', 0))
                                    item['times'] = safe_float(item.get('times', 1))
                                    if item['times'] == 0: item['times'] = 1
                                    
                                    if str(item['category']).strip() or str(item['description']).strip():
                                        new_data.append(item)

                            if new_data:
                                # 3. Cập nhật vào Session State (Hiển thị lên màn hình)
                                st.session_state.est_df_temp = pd.DataFrame(new_data)
                                st.session_state.est_edit_mode = True # Bật chế độ sửa để hiện nút Lưu
                                st.success(f"Đã đọc thành công {len(new_data)} dòng! Vui lòng kiểm tra bảng bên dưới và bấm LƯU.")
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.warning(f"Không đọc được dữ liệu! (Đã thử dòng {header_idx+1} làm tiêu đề). Vui lòng kiểm tra tên cột.")
                                
                        except Exception as e:
                            st.error(f"Lỗi khi đọc file: {str(e)}")

            # --- Fetch Items (EST) ---
            if "est_df_temp" not in st.session_state:
                existing_items = run_query("SELECT * FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id,))
                if existing_items:
                    df_est = pd.DataFrame([dict(r) for r in existing_items])
                    if 'times' not in df_est.columns: df_est['times'] = 1.0
                    df_est = df_est[['category', 'description', 'unit', 'unit_price', 'quantity', 'times']]
                else:
                    df_est = pd.DataFrame([
                        {"category": "Vận chuyển", "description": "Xe 16 chỗ", "unit": "Xe", "unit_price": 0, "quantity": 1, "times": 1},
                        {"category": "Lưu trú", "description": "Khách sạn 3 sao", "unit": "Phòng", "unit_price": 0, "quantity": 1, "times": 1},
                        {"category": "Ăn uống", "description": "Bữa trưa ngày 1", "unit": "Suất", "unit_price": 0, "quantity": 1, "times": 1},
                    ])
                st.session_state.est_df_temp = df_est

            # Prepare Display Data (Tạo bản sao để hiển thị format đẹp)
            df_display = st.session_state.est_df_temp.copy()
            # [NEW] Tạo số thứ tự (STT) tự động
            df_display.index = pd.RangeIndex(start=1, stop=len(df_display) + 1)
            
            # [MODIFIED] Tính Giá/Pax và ẩn cột Times
            guest_cnt = tour_info['guest_count'] if tour_info['guest_count'] else 1 # type: ignore
            df_display['total_val'] = df_display['quantity'] * df_display['unit_price'] * df_display['times']
            df_display['price_per_pax'] = df_display['total_val'] / guest_cnt
            
            df_display['price_per_pax'] = df_display['price_per_pax'].apply(lambda x: format_vnd(x) + " VND")
            df_display['total_display'] = df_display['total_val'].apply(lambda x: format_vnd(x) + " VND")
            df_display['unit_price'] = df_display['unit_price'].apply(lambda x: format_vnd(x) + " VND") # type: ignore

            st.markdown(f"**Đoàn:** {tour_info['tour_name']} (Mã: {tour_info['tour_code']}) | **Pax:** {tour_info['guest_count']}")
            
            is_disabled = not st.session_state.est_edit_mode

            # --- DATA EDITOR ---
            edited_est = st.data_editor(
                df_display,
                disabled=is_disabled,
                num_rows="dynamic",
                column_config={
                    "_index": st.column_config.NumberColumn("STT", disabled=True),
                    "category": st.column_config.TextColumn("Hạng mục chi phí", required=False),
                    "description": st.column_config.TextColumn("Diễn giải"),
                    "unit": st.column_config.TextColumn("Đơn vị"),
                    "unit_price": st.column_config.TextColumn("Đơn giá (VND)", required=False),
                    "quantity": st.column_config.NumberColumn("Số lượng", min_value=0),
                    "times": st.column_config.NumberColumn("Số lần", min_value=1),
                    "price_per_pax": st.column_config.TextColumn("Giá/Pax", disabled=False),
                    "total_display": st.column_config.TextColumn("Tổng chi phí", disabled=False),
                    "total_val": st.column_config.NumberColumn("Hidden", disabled=True),
                },
                column_order=("category", "description", "unit", "unit_price", "quantity", "times", "price_per_pax", "total_display"),
                use_container_width=True,
                hide_index=False,
                key=f"editor_est_{st.session_state.est_editor_key}"
            )
            
            # --- AUTO-UPDATE CALCULATION ---
            if st.session_state.est_edit_mode:
                # Tự động cập nhật khi dữ liệu thay đổi
                df_new = edited_est.copy()
                
                def clean_vnd_auto(x):
                    if isinstance(x, str):
                        return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    return float(x) if x else 0.0
                
                df_new['unit_price'] = df_new['unit_price'].apply(clean_vnd_auto)
                df_new['quantity'] = pd.to_numeric(df_new['quantity'], errors='coerce').fillna(0)
                if 'times' not in df_new.columns: df_new['times'] = 1
                df_new['times'] = pd.to_numeric(df_new['times'], errors='coerce').fillna(1)

                # [NEW] Logic xử lý sửa Tổng tiền / Giá Pax
                df_new['total_val_edit'] = df_new['total_display'].apply(clean_vnd_auto)
                df_new['pax_val_edit'] = df_new['price_per_pax'].apply(clean_vnd_auto)
                
                df_old: pd.DataFrame = st.session_state.est_df_temp.copy()  # type: ignore
                df_old = df_old.reset_index(drop=True)
                if 'times' not in df_old.columns: df_old['times'] = 1
                
                guest_cnt = tour_info['guest_count'] if tour_info['guest_count'] else 1

                force_rerun_fmt = False

                for pos_idx, (idx, row) in enumerate(df_new.iterrows()):
                    n_unit = row['unit_price']
                    n_qty = row['quantity']
                    n_times = row['times']
                    n_total_edit = row['total_val_edit']
                    n_pax_edit = row['pax_val_edit']
                    
                    o_unit = 0.0
                    o_qty = 0.0
                    o_times = 1.0
                    
                    if pos_idx < len(df_old):
                        try:
                            # Get old values for comparison
                            old_row = df_old.iloc[pos_idx]
                            o_unit = float(old_row['unit_price'])
                            o_qty = float(old_row['quantity'])
                            o_times = float(old_row['times'])
                        except (IndexError, KeyError):
                            pass  # Use default values if access fails
                  
                    unit_changed = abs(n_unit - o_unit) > 0.1
                    
                    old_total = o_unit * o_qty * o_times
                    total_changed = abs(n_total_edit - old_total) > 0.1
                    
                    old_pax = old_total / guest_cnt
                    pax_changed = abs(n_pax_edit - old_pax) > 0.1
                    
                    if unit_changed:
                        pass 
                    elif total_changed:
                        if n_qty * n_times != 0:
                            df_new.at[idx, 'unit_price'] = n_total_edit / (n_qty * n_times)
                    elif pax_changed:
                        new_total = n_pax_edit * guest_cnt
                        if n_qty * n_times != 0:
                            df_new.at[idx, 'unit_price'] = new_total / (n_qty * n_times)
                    
                    # [NEW] Check formatting to force refresh if user typed raw number
                    if str(row['total_display']) != (format_vnd(n_total_edit) + " VND"): force_rerun_fmt = True
                    if str(row['price_per_pax']) != (format_vnd(n_pax_edit) + " VND"): force_rerun_fmt = True
                
                # So sánh với dữ liệu cũ
                cols_check = ['category', 'description', 'unit', 'unit_price', 'quantity', 'times']
                
                # Reset index và fillna để so sánh
                df_new_check = df_new[cols_check].reset_index(drop=True).fillna(0)
                df_old_check = df_old[cols_check].reset_index(drop=True).fillna(0)
                
                has_changes = False
                if len(df_new_check) != len(df_old_check): has_changes = True
                elif not df_new_check.equals(df_old_check): has_changes = True
                
                if has_changes or force_rerun_fmt:
                    st.session_state.est_df_temp = df_new[cols_check]
                    st.rerun()

            # --- TÍNH TOÁN REAL-TIME ---
            total_cost = 0
            if not edited_est.empty:
                # [FIX] Handle case where a cell is None, which becomes the string 'None' after astype(str)
                cleaned_prices_est = edited_est['unit_price'].astype(str).str.replace('.', '', regex=False).str.replace(' VND', '', regex=False).str.strip()
                p_price = cleaned_prices_est.apply(lambda x: float(x) if x and x.lower() != 'none' else 0.0)
                t_times = edited_est['times'].fillna(1) # type: ignore
                total_cost = (edited_est['quantity'] * p_price * t_times).sum()
            
            st.divider()
            
            # --- PHẦN TÍNH LỢI NHUẬN & THUẾ (YÊU CẦU 2: Sắp xếp hàng ngang) ---
            c_cost, c_profit, c_tax = st.columns(3)
            
            with c_cost:
                st.metric("Tổng Chi Phí Dự Toán", format_vnd(total_cost) + " VND") # type: ignore
            with c_profit:
                p_percent = st.number_input("Lợi Nhuận Mong Muốn (%)", value=float(tour_info['est_profit_percent']), step=0.5, key="p_pct", disabled=is_disabled) # type: ignore
            with c_tax:
                t_percent = st.number_input("Thuế VAT Đầu Ra (%)", value=float(tour_info['est_tax_percent']), step=1.0, key="t_pct", disabled=is_disabled) # type: ignore
            
            # Công thức: Giá Bán = Chi Phí + Lợi Nhuận + Thuế
            # Lợi nhuận = Chi Phí * %
            # Thuế = (Chi Phí + Lợi Nhuận) * %
            profit_amt = total_cost * (p_percent / 100)
            base_price = total_cost + profit_amt
            tax_amt = base_price * (t_percent / 100)
            final_price = base_price + tax_amt

            st.markdown(f"""<div class="finance-summary-card">
                <div class="row"><span>Tiền Lợi Nhuận ({p_percent}%):</span> <b>{format_vnd(profit_amt)} VND</b></div>
                <div class="row"><span>Tiền Thuế ({t_percent}%):</span> <b>{format_vnd(tax_amt)} VND</b></div>
                <div class="row total-row"><span>TỔNG GIÁ BÁN DỰ KIẾN:</span> <b>{format_vnd(final_price)} VND</b></div>
                <div class="pax-price">(Giá trung bình/khách: {format_vnd(final_price/tour_info['guest_count'] if tour_info['guest_count'] else 1)} VND)</div>
            </div>
            """, unsafe_allow_html=True)

            # --- THÊM Ô NHẬP GIÁ CHỐT & GIÁ TRẺ EM ---
            st.write("")
            t_dict: Dict[str, Any] = dict(tour_info) if tour_info else {}
            
            # Ensure session state is initialized if not present
            if "est_final_price_val" not in st.session_state:
                cur_final_price = float(t_dict.get('final_tour_price', 0) or 0)
                st.session_state.est_final_price_val = "{:,.0f}".format(cur_final_price).replace(",", ".") + " VND"
            if "est_child_price_val" not in st.session_state:
                cur_child_price = float(t_dict.get('child_price', 0) or 0)
                st.session_state.est_child_price_val = "{:,.0f}".format(cur_child_price).replace(",", ".") + " VND"

            def fmt_est_final_price():
                val = st.session_state.est_final_price_val
                try:
                    v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    st.session_state.est_final_price_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                except: pass

            def fmt_est_child_price():
                val = st.session_state.est_child_price_val
                try:
                    v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    st.session_state.est_child_price_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                except: pass

            c_final_p, c_child_p = st.columns(2)
            with c_final_p:
                # Giá chốt tour
                st.text_input("Giá chốt tour (VND)", key="est_final_price_val", on_change=fmt_est_final_price, disabled=is_disabled, help="Nhập số tiền (VD: 1.000.000)")
                try: final_tour_price_val = float(st.session_state.est_final_price_val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: final_tour_price_val = 0.0

                # Số lượng người lớn
                cur_qty = float(t_dict.get('final_qty', 0))
                if cur_qty == 0: cur_qty = float(t_dict.get('guest_count', 1))
                final_qty_val = st.number_input("Số lượng người lớn", value=cur_qty, min_value=0.0, step=1.0, disabled=is_disabled)

            with c_child_p:
                # Giá trẻ em
                st.text_input("Giá trẻ em (VND)", key="est_child_price_val", on_change=fmt_est_child_price, disabled=is_disabled)
                try: child_price_val = float(st.session_state.est_child_price_val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: child_price_val = 0.0

                cur_child_qty = float(t_dict.get('child_qty', 0))
                child_qty_val = st.number_input("Số lượng trẻ em", value=cur_child_qty, min_value=0.0, step=1.0, disabled=is_disabled)
            
            total_final_manual = (final_tour_price_val * final_qty_val) + (child_price_val * child_qty_val)
            st.markdown(f"""<div style="background-color: #e8f5e9; padding: 15px; border-radius: 10px; margin-top: 10px; border: 1px solid #c8e6c9;"><div style="display:flex; justify-content:space-between; font-size: 1.3em; color: #2e7d32;"><span><b>TỔNG DOANH THU</b></span> <b>{format_vnd(total_final_manual)} VND</b></div></div>""", unsafe_allow_html=True)

            est_profit_manual = total_final_manual - total_cost
            st.markdown(f"""<div style="background-color: #e3f2fd; padding: 15px; border-radius: 10px; margin-top: 10px; border: 1px solid #90caf9;"><div style="display:flex; justify-content:space-between; font-size: 1.3em; color: #1565c0;"><span><b>TỔNG LỢI NHUẬN</b></span> <b>{format_vnd(est_profit_manual)} VND</b></div></div>""", unsafe_allow_html=True)

            # --- EXPORT EXCEL ---
            st.write("")
            df_exp = st.session_state.est_df_temp.copy()
            
            # Chuẩn hóa dữ liệu số
            def clean_price_exp(x): # type: ignore
                if isinstance(x, str):
                    return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip())
                return float(x) if x else 0.0
            
            df_exp['unit_price'] = df_exp['unit_price'].apply(clean_price_exp)
            df_exp['quantity'] = pd.to_numeric(df_exp['quantity'], errors='coerce').fillna(0)
            if 'times' not in df_exp.columns: df_exp['times'] = 1
            df_exp['times'] = pd.to_numeric(df_exp['times'], errors='coerce').fillna(1)
            
            # Tính toán các cột hiển thị giống Web
            df_exp['total_amount'] = df_exp['quantity'] * df_exp['unit_price'] * df_exp['times']
            g_cnt = tour_info['guest_count'] if tour_info['guest_count'] else 1 # type: ignore
            df_exp['price_per_pax'] = df_exp['total_amount'] / g_cnt
            
            # Chọn và đổi tên cột
            df_exp = df_exp[['category', 'description', 'unit', 'unit_price', 'quantity', 'times', 'price_per_pax', 'total_amount']]
            df_exp.columns = ['Hạng mục', 'Diễn giải', 'Đơn vị', 'Đơn giá', 'Số lượng', 'Số lần', 'Giá/Pax', 'Tổng chi phí']

            buffer = io.BytesIO()
            file_ext = "xlsx"
            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            try:
                with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer: # type: ignore
                        # Start table at row 11 (index 10) to leave space for info
                        start_row = 10
                        df_exp.to_excel(writer, index=False, sheet_name='DuToan', startrow=start_row)
                        
                        # --- FORMATTING (Nếu dùng xlsxwriter) ---
                        workbook: Any = writer.book
                        worksheet = writer.sheets['DuToan']
                        
                        # --- STYLES ---
                        company_name_fmt = workbook.add_format({'bold': True, 'font_size': 14, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                        company_info_fmt = workbook.add_format({'font_size': 10, 'italic': True, 'font_color': '#424242', 'font_name': 'Times New Roman'})
                        
                        title_fmt = workbook.add_format({'bold': True, 'font_size': 18, 'align': 'center', 'valign': 'vcenter', 'font_color': '#0D47A1', 'bg_color': '#E3F2FD', 'border': 1, 'font_name': 'Times New Roman'})
                        section_fmt = workbook.add_format({'bold': True, 'font_size': 11, 'font_color': '#E65100', 'underline': True, 'font_name': 'Times New Roman'})
                        
                        header_fmt = workbook.add_format({'bold': True, 'fg_color': '#2E7D32', 'font_color': 'white', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        body_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_size': 10, 'font_name': 'Times New Roman'})
                        body_center_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'align': 'center', 'font_size': 10, 'font_name': 'Times New Roman'})
                        money_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_size': 10, 'font_name': 'Times New Roman'})
                        
                        # Summary Section Styles
                        sum_header_bg_fmt = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'font_color': '#E65100', 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        sum_label_fmt = workbook.add_format({'bold': True, 'align': 'left', 'border': 1, 'bg_color': '#FAFAFA', 'font_name': 'Times New Roman'})
                        sum_val_fmt = workbook.add_format({'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_name': 'Times New Roman'})
                        sum_val_bold_fmt = workbook.add_format({'bold': True, 'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_name': 'Times New Roman'})
                        sum_total_fmt = workbook.add_format({'bold': True, 'bg_color': '#C8E6C9', 'font_color': '#1B5E20', 'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_size': 12, 'font_name': 'Times New Roman'})
                        
                        # --- 1. COMPANY INFO (Rows 0-3) ---
                        if comp['logo_b64_str']:
                            try:
                                logo_data = base64.b64decode(comp['logo_b64_str'])
                                image_stream = io.BytesIO(logo_data)
                                img = Image.open(image_stream)
                                w, h = img.size
                                scale = 60 / h if h > 0 else 0.5
                                image_stream.seek(0)
                                worksheet.insert_image('A1', 'logo.png', {'image_data': image_stream, 'x_scale': scale, 'y_scale': scale, 'x_offset': 5, 'y_offset': 5})
                            except: pass
                        
                        worksheet.write('B1', comp['name'], company_name_fmt)
                        worksheet.write('B2', f"ĐC: {comp['address']}", company_info_fmt)
                        worksheet.write('B3', f"MST: {comp['phone']}", company_info_fmt)
                        
                        # --- 2. TOUR INFO (Rows 4-9) ---
                        worksheet.merge_range('A5:G5', "BẢNG DỰ TOÁN CHI PHÍ TOUR", title_fmt)
                        
                        # Info Data
                        t_info_dict = dict(tour_info) if tour_info else {}
                        t_name = t_info_dict.get('tour_name', '')
                        t_code = t_info_dict.get('tour_code', '')
                        t_sale = t_info_dict.get('sale_name', '')
                        t_start = t_info_dict.get('start_date', '')
                        t_end = t_info_dict.get('end_date', '')
                        t_cust = t_info_dict.get('customer_name', '')
                        t_phone = t_info_dict.get('customer_phone', '')
                        t_guest = t_info_dict.get('guest_count', 0)
                        
                        # Layout Info nicely
                        worksheet.write('A7', "Tên đoàn:", sum_label_fmt)
                        worksheet.merge_range('B7:D7', t_name, sum_val_fmt)
                        worksheet.write('E7', "Mã đoàn:", sum_label_fmt)
                        worksheet.merge_range('F7:G7', t_code, sum_val_fmt)
                        
                        worksheet.write('A8', "Khách hàng:", sum_label_fmt)
                        worksheet.merge_range('B8:D8', f"{t_cust} - {t_phone}", sum_val_fmt)
                        worksheet.write('E8', "Sales:", sum_label_fmt)
                        worksheet.merge_range('F8:G8', t_sale, sum_val_fmt)
                        
                        worksheet.write('A9', "Thời gian:", sum_label_fmt)
                        worksheet.merge_range('B9:D9', f"{t_start} - {t_end}", sum_val_fmt)
                        worksheet.write('E9', "Số khách:", sum_label_fmt)
                        worksheet.merge_range('F9:G9', t_guest, sum_val_fmt)

                        # --- 3. TABLE HEADER & BODY ---
                        # Apply Header
                        for col_num, value in enumerate(df_exp.columns):
                            worksheet.write(start_row, col_num, value, header_fmt)
                        
                        # Apply Body
                        for row in range(len(df_exp)):
                            for col in range(len(df_exp.columns)):
                                val = df_exp.iloc[row, col]
                                # Cols: 0=Cat, 1=Desc, 2=Unit, 3=Price, 4=Qty, 5=PaxPrice, 6=Total
                                if col == 2: fmt = body_center_fmt # Unit centered
                                elif col in [3, 4, 5, 6, 7]: fmt = money_fmt # Money columns
                                else: fmt = body_fmt
                                
                                if pd.isna(val): val = ""
                                worksheet.write(row+start_row+1, col, val, fmt)
                        
                        # --- 4. SUMMARY SECTION ---
                        last_row = start_row + 1 + len(df_exp)
                        sum_row = last_row + 2
                        
                        # --- BẢNG TÍNH GIÁ THÀNH ---
                        worksheet.merge_range(sum_row, 0, sum_row, 3, "PHÂN TÍCH GIÁ THÀNH & LỢI NHUẬN", sum_header_bg_fmt)
                        
                        # Dòng 1: Tổng chi phí
                        worksheet.write(sum_row+1, 0, "1. Tổng chi phí dự toán:", sum_label_fmt)
                        worksheet.merge_range(sum_row+1, 1, sum_row+1, 3, total_cost, sum_val_bold_fmt)
                        
                        # Dòng 2: Lợi nhuận
                        worksheet.write(sum_row+2, 0, "2. Lợi nhuận mong muốn:", sum_label_fmt)
                        worksheet.write(sum_row+2, 1, f"{p_percent:g}%", body_center_fmt)
                        worksheet.merge_range(sum_row+2, 2, sum_row+2, 3, profit_amt, sum_val_fmt)
                        
                        # Dòng 3: Thuế
                        worksheet.write(sum_row+3, 0, "3. Thuế VAT:", sum_label_fmt)
                        worksheet.write(sum_row+3, 1, f"{t_percent:g}%", body_center_fmt)
                        worksheet.merge_range(sum_row+3, 2, sum_row+3, 3, tax_amt, sum_val_fmt)
                        
                        # Dòng 4: Giá tính toán
                        worksheet.write(sum_row+4, 0, "4. Giá bán tính toán:", sum_label_fmt)
                        worksheet.merge_range(sum_row+4, 1, sum_row+4, 3, final_price, sum_total_fmt)
                        
                        # --- BẢNG CHỐT GIÁ BÁN ---
                        # Đặt bên phải bảng giá thành (Cột E, F, G)
                        worksheet.merge_range(sum_row, 4, sum_row, 6, "BẢNG CHỐT GIÁ BÁN THỰC TẾ", sum_header_bg_fmt)
                        
                        # Người lớn
                        worksheet.write(sum_row+1, 4, "Người lớn:", sum_label_fmt)
                        worksheet.write(sum_row+1, 5, final_qty_val, sum_val_fmt) # SL
                        worksheet.write(sum_row+1, 6, final_tour_price_val, sum_val_fmt) # Giá
                        
                        # Trẻ em
                        worksheet.write(sum_row+2, 4, "Trẻ em:", sum_label_fmt)
                        worksheet.write(sum_row+2, 5, child_qty_val, sum_val_fmt) # SL
                        worksheet.write(sum_row+2, 6, child_price_val, sum_val_fmt) # Giá
                        
                        # Tổng doanh thu
                        worksheet.write(sum_row+4, 4, "TỔNG DOANH THU:", sum_label_fmt)
                        worksheet.merge_range(sum_row+4, 5, sum_row+4, 6, total_final_manual, sum_total_fmt)

                        # Column Widths
                        worksheet.set_column('A:A', 25) # Category
                        worksheet.set_column('B:B', 40) # Desc
                        worksheet.set_column('C:C', 10) # Unit
                        worksheet.set_column('D:G', 18) # Numbers
            except Exception as e:
                # If xlsxwriter fails, fall back to a simple CSV export
                buffer.seek(0)
                buffer.truncate()
                df_exp.to_csv(buffer, index=False, encoding='utf-8-sig')
                file_ext = "csv"
                mime_type = "text/csv"
                st.error(f"⚠️ Lỗi khi tạo file Excel: {e}. Đã chuyển sang xuất file CSV.")
                st.info("💡 Gợi ý: Nếu bạn vừa cài thư viện, hãy TẮT HẲN ứng dụng (Ctrl+C tại terminal) và chạy lại lệnh `streamlit run app.py`.")

            clean_t_name = re.sub(r'[\\/*?:"<>|]', "", tour_info['tour_name'] if tour_info else "Tour") # type: ignore
            st.download_button(
                f"📥 Tải Bảng Dự Toán ({file_ext.upper()})", 
                data=buffer.getvalue(), 
                file_name=f"DuToan_{clean_t_name}.{file_ext}", 
                mime=mime_type, 
                use_container_width=True
            )

            # --- Nút Chỉnh sửa / Lưu ---
            if st.session_state.est_edit_mode:
                if st.button("💾 Lưu và chuyển sang phần Danh sách và dịch vụ", type="primary", use_container_width=True):
                    # 1. Update Tour Meta
                    run_query("UPDATE tours SET est_profit_percent=?, est_tax_percent=?, final_tour_price=?, child_price=?, final_qty=?, child_qty=? WHERE id=?", (p_percent, t_percent, final_tour_price_val, child_price_val, final_qty_val, child_qty_val, tour_id), commit=True)
                    
                    # 2. Update Tour Items (Xóa cũ thêm mới)
                    run_query("DELETE FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id,), commit=True)

                    data_to_insert = []
                    query = """INSERT INTO tour_items (tour_id, item_type, category, description, unit, quantity, unit_price, total_amount, times)
                               VALUES (?, 'EST', ?, ?, ?, ?, ?, ?, ?)"""

                    for _, row in edited_est.iterrows():
                        if row['category'] or row['description']: # type: ignore
                            # Xử lý dữ liệu
                            u_price = float(str(row['unit_price']).replace('.', '').replace(' VND', '').strip()) if row['unit_price'] else 0 # type: ignore
                            t_times = row.get('times', 1) # type: ignore
                            if pd.isna(t_times): t_times = 1
                            total_row = row['quantity'] * u_price * t_times # type: ignore
                            
                            # Thêm vào danh sách chờ (chưa ghi ngay)
                            data_to_insert.append((
                                tour_id, 
                                row['category'], 
                                row['description'], 
                                row['unit'], 
                                row['quantity'],  # type: ignore
                                u_price, 
                                total_row, 
                                t_times
                            ))

                    # Ghi tất cả trong 1 lần bắn
                    if data_to_insert:
                        run_query_many(query, data_to_insert)

                    if "est_df_temp" in st.session_state: del st.session_state.est_df_temp
                    st.session_state.est_edit_mode = False
                    st.success("Đã lưu dự toán thành công! Hãy chuyển sang tab 'Danh sách & Dịch vụ'.")
                    time.sleep(1); st.rerun()
            else:
                if st.button("✏️ Chỉnh sửa Dự toán", use_container_width=True):
                    st.session_state.est_edit_mode = True
                    st.rerun()

    # ---------------- TAB MỚI: DANH SÁCH & DỊCH VỤ ----------------
    with tab_list_srv:
        st.subheader("📋 Danh sách Đoàn & Dịch vụ")
        
        def clean_vnd_val(x):
            if isinstance(x, (int, float)): return float(x)
            if isinstance(x, str):
                return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip()) if x.strip() else 0.0
            return 0.0

        # Hàm chuẩn hóa và tính toán lại cột Còn lại
        def recalc_remaining(df, total_col='total_amount', dep_col='deposit', rem_col='remaining'):
            df[total_col] = df[total_col].apply(clean_vnd_val)
            df[dep_col] = df[dep_col].apply(clean_vnd_val)
            df[rem_col] = df[total_col] - df[dep_col]
            return df

        selected_tour_ls_label = st.selectbox("Chọn Đoàn:", list(tour_options.keys()) if tour_options else [], key="sel_tour_ls")
        
        if selected_tour_ls_label:
            tour_id_ls = tour_options[selected_tour_ls_label]
            tour_info_ls = next((t for t in all_tours if t['id'] == tour_id_ls), None)
            
            if tour_info_ls:
                tour_info_ls = dict(tour_info_ls)
                
                # --- SESSION STATE INIT FOR LIST & SERVICES (AUTO CALCULATION) ---
                if "current_tour_id_ls" not in st.session_state or st.session_state.current_tour_id_ls != tour_id_ls or "ls_incurred_temp" not in st.session_state:
                    st.session_state.current_tour_id_ls = tour_id_ls
                    
                    # Load Hotels
                    hotels = run_query("SELECT * FROM tour_hotels WHERE tour_id=?", (tour_id_ls,))
                    df_h = pd.DataFrame([dict(r) for r in hotels]) if hotels else pd.DataFrame(columns=['hotel_name', 'address', 'phone', 'total_rooms', 'room_type', 'total_amount', 'deposit'])
                    if df_h.empty: df_h = pd.DataFrame(columns=['hotel_name', 'address', 'phone', 'total_rooms', 'room_type', 'total_amount', 'deposit'])
                    for col in ['total_amount', 'deposit']:
                        if col not in df_h.columns: df_h[col] = 0.0
                    st.session_state.ls_hotels_temp = df_h[['hotel_name', 'address', 'phone', 'total_rooms', 'room_type', 'total_amount', 'deposit']]

                    # Load Restaurants
                    rests = run_query("SELECT * FROM tour_restaurants WHERE tour_id=?", (tour_id_ls,))
                    if not rests:
                        est_meals = run_query("SELECT description FROM tour_items WHERE tour_id=? AND item_type='EST' AND (category LIKE '%Ăn%' OR description LIKE '%Bữa%' OR description LIKE '%Ăn%')", (tour_id_ls,))
                        if est_meals:
                            df_r = pd.DataFrame([{'date': '', 'meal_name': m['description'], 'restaurant_name': '', 'address': '', 'phone': '', 'menu': '', 'total_amount': 0, 'deposit': 0} for m in est_meals])
                        else:
                            df_r = pd.DataFrame(columns=['date', 'meal_name', 'restaurant_name', 'address', 'phone', 'menu', 'total_amount', 'deposit'])
                    else:
                        df_r = pd.DataFrame([dict(r) for r in rests])
                    
                    for col in ['total_amount', 'deposit']:
                        if col not in df_r.columns: df_r[col] = 0.0
                    if 'date' not in df_r.columns: df_r['date'] = ''
                    st.session_state.ls_rests_temp = df_r[['date', 'meal_name', 'restaurant_name', 'address', 'phone', 'menu', 'total_amount', 'deposit']]

                    # Load Sightseeings
                    sights = run_query("SELECT * FROM tour_sightseeings WHERE tour_id=?", (tour_id_ls,))
                    df_s = pd.DataFrame([dict(r) for r in sights]) if sights else pd.DataFrame(columns=['date', 'name', 'address', 'quantity', 'total_amount', 'deposit', 'note'])
                    if df_s.empty: df_s = pd.DataFrame(columns=['date', 'name', 'address', 'quantity', 'total_amount', 'deposit', 'note'])
                    if 'date' not in df_s.columns: df_s['date'] = ''
                    if 'total_amount' not in df_s.columns: df_s['total_amount'] = 0.0
                    if 'deposit' not in df_s.columns: df_s['deposit'] = 0.0
                    st.session_state.ls_sight_temp = df_s[['date', 'name', 'address', 'quantity', 'total_amount', 'deposit', 'note']]

                    # Load Incurred Costs
                    incurred = run_query("SELECT * FROM tour_incurred_costs WHERE tour_id=?", (tour_id_ls,))
                    df_inc = pd.DataFrame([dict(r) for r in incurred]) if incurred else pd.DataFrame(columns=['name', 'unit', 'quantity', 'price', 'total_amount', 'deposit', 'note'])
                    if df_inc.empty: df_inc = pd.DataFrame(columns=['name', 'unit', 'quantity', 'price', 'total_amount', 'deposit', 'note'])
                    st.session_state.ls_incurred_temp = df_inc[['name', 'unit', 'quantity', 'price', 'total_amount', 'deposit', 'note']]

                # 1. DANH SÁCH ĐOÀN
                st.markdown("##### 1. Danh sách đoàn")
                
                # --- [NEW] THÔNG TIN BÀN GIAO & ĐÓN TIỄN ---
                with st.expander("🚌 Thông Tin Bàn Giao & Đón Tiễn (Điều hành)", expanded=False):
                    with st.form(f"handover_form_{tour_id_ls}"):
                        c_h1, c_h2, c_h3 = st.columns(3)
                        pk_loc = c_h1.text_input("Điểm đón", value=tour_info_ls.get('pickup_location', ''))
                        pk_time = c_h2.text_input("Thời gian đón", value=tour_info_ls.get('pickup_time', ''))
                        fl_code = c_h3.text_area("Chuyến bay/Tàu", value=tour_info_ls.get('flight_code', ''), height=68)
                        
                        c_d1, c_d2, c_d3, c_d4 = st.columns(4)
                        drv_name = c_d1.text_input("Tên lái xe", value=tour_info_ls.get('driver_name', ''))
                        drv_phone = c_d2.text_input("SĐT Lái xe", value=tour_info_ls.get('driver_phone', ''))
                        car_plate = c_d3.text_input("Biển số xe", value=tour_info_ls.get('car_plate', ''))
                        car_type = c_d4.text_input("Loại xe", value=tour_info_ls.get('car_type', ''))
                        
                        c_g1, c_g2 = st.columns(2)
                        gd_name = c_g1.text_input("Tên HDV", value=tour_info_ls.get('guide_name', ''))
                        gd_phone = c_g2.text_input("SĐT HDV", value=tour_info_ls.get('guide_phone', ''))

                        if st.form_submit_button("💾 Lưu thông tin vận hành"):
                            run_query("""UPDATE tours SET 
                                pickup_location=?, pickup_time=?, flight_code=?, 
                                driver_name=?, driver_phone=?, car_plate=?, car_type=?, 
                                guide_name=?, guide_phone=? WHERE id=?""",
                                (pk_loc, pk_time, fl_code, drv_name, drv_phone, car_plate, car_type, gd_name, gd_phone, tour_id_ls), commit=True)
                            st.success("Đã cập nhật thông tin điều hành!"); time.sleep(0.5); st.rerun()

                    st.markdown("##### 📅 Lịch trình chi tiết")
                    try:
                        s_raw = tour_info_ls.get('start_date', '')
                        e_raw = tour_info_ls.get('end_date', '')
                        
                        def try_parse_date(d_str):
                            if not d_str: return None
                            for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y'):
                                try: return datetime.strptime(str(d_str).strip(), fmt)
                                except: continue
                            return None

                        s_d = try_parse_date(s_raw)
                        e_d = try_parse_date(e_raw)

                        if s_d and e_d:
                            num_days = (e_d - s_d).days + 1
                            if num_days < 1: num_days = 1
                        else:
                            num_days = 1
                            s_d = datetime.now()
                    except: 
                        num_days = 1
                        s_d = datetime.now()
                    
                    exist_itin = run_query("SELECT * FROM tour_itineraries WHERE tour_id=? ORDER BY day_index", (tour_id_ls,))
                    itin_map = {r['day_index']: r['content'] for r in exist_itin} if exist_itin else {}
                    
                    itin_data = []
                    for i in range(num_days):
                        d_str = (s_d + pd.Timedelta(days=i)).strftime('%d/%m/%Y')
                        itin_data.append({
                            "day_label": f"Ngày {i+1} ({d_str})",
                            "content": itin_map.get(i, ""),
                            "day_index": i
                        })
                    
                    df_itin = pd.DataFrame(itin_data)
                    edited_itin = st.data_editor(
                        df_itin,
                        column_config={
                            "day_label": st.column_config.TextColumn("Ngày", disabled=True),
                            "content": st.column_config.TextColumn("Nội dung lịch trình", width="large"),
                            "day_index": st.column_config.NumberColumn("Hidden", disabled=True)
                        },
                        column_order=("day_label", "content"),
                        use_container_width=True,
                        hide_index=True,
                        key=f"itin_ed_{tour_id_ls}"
                    )
                    
                    if st.button("💾 Lưu lịch trình", key=f"save_itin_{tour_id_ls}"):
                        run_query("DELETE FROM tour_itineraries WHERE tour_id=?", (tour_id_ls,), commit=True)
                        data_itin = [(tour_id_ls, r['day_index'], r['content']) for _, r in edited_itin.iterrows()]
                        if data_itin:
                            run_query_many("INSERT INTO tour_itineraries (tour_id, day_index, content) VALUES (?,?,?)", data_itin)
                        st.success("Đã lưu lịch trình!"); time.sleep(0.5); st.rerun()

                guests = run_query("SELECT * FROM tour_guests WHERE tour_id=?", (tour_id_ls,))
                df_guests = pd.DataFrame([dict(r) for r in guests]) if guests else pd.DataFrame(columns=['name', 'dob', 'hometown', 'cccd', 'type'])
                if df_guests.empty:
                    df_guests = pd.DataFrame(columns=['name', 'dob', 'hometown', 'cccd', 'type'])
                else:
                    df_guests = df_guests[['name', 'dob', 'hometown', 'cccd', 'type']]

                edited_guests = st.data_editor(
                    df_guests,
                    num_rows="dynamic",
                    key="guest_editor",
                    column_config={
                        "name": st.column_config.TextColumn("Họ và tên", required=True),
                        "dob": st.column_config.TextColumn("Ngày sinh"),
                        "hometown": st.column_config.TextColumn("Quê quán"),
                        "cccd": st.column_config.TextColumn("Số CCCD"),
                        "type": st.column_config.SelectboxColumn("Phân loại", options=["Khách", "Nội bộ", "HDV"], required=True)
                    },
                    use_container_width=True
                )

                # 2. DANH SÁCH PHÒNG KHÁCH SẠN
                st.markdown("##### 2. Danh sách phòng Khách sạn")
                df_hotels = st.session_state.ls_hotels_temp.copy()
                
                # Tính toán hiển thị ban đầu
                df_hotels = recalc_remaining(df_hotels)

                edited_hotels = st.data_editor(
                    df_hotels,
                    num_rows="dynamic",
                    key="hotel_editor",
                    column_config={
                        "hotel_name": st.column_config.TextColumn("Tên Khách sạn", required=True),
                        "address": "Địa chỉ",
                        "phone": "SĐT",
                        "total_rooms": st.column_config.TextColumn("Tổng số phòng"),
                        "room_type": st.column_config.TextColumn("Loại phòng"),
                        "total_amount": st.column_config.NumberColumn("Tổng tiền", format="%d VND"),
                        "deposit": st.column_config.NumberColumn("Đã ứng/cọc", format="%d VND"),
                        "remaining": st.column_config.NumberColumn("Còn lại (Guide trả)", format="%d VND", disabled=True)
                    },
                    use_container_width=True
                )
                
                # Xử lý cập nhật tự động
                cols_h = ['hotel_name', 'address', 'phone', 'total_rooms', 'room_type', 'total_amount', 'deposit']
                # Làm sạch dữ liệu vừa nhập
                edited_hotels = recalc_remaining(edited_hotels)
                
                  # So sánh với dữ liệu cũ (chỉ so các cột nhập liệu để tránh lặp vô tận do cột tính toán)
                if not edited_hotels[cols_h].equals(st.session_state.ls_hotels_temp[cols_h].map(lambda x: x if not isinstance(x, float) else x)):
                    st.session_state.ls_hotels_temp = edited_hotels
                    st.rerun()
                # 3. MENU NHÀ HÀNG
                st.markdown("##### 3. Menu nhà hàng")
                df_rests = st.session_state.ls_rests_temp.copy()
                
                df_rests = recalc_remaining(df_rests)

                edited_rests = st.data_editor(
                    df_rests,
                    num_rows="dynamic",
                    key="rest_editor",
                    column_config={
                        "date": st.column_config.TextColumn("Ngày"),
                        "meal_name": st.column_config.TextColumn("Bữa ăn (Dự toán)", required=True),
                        "restaurant_name": "Tên nhà hàng",
                        "address": "Địa chỉ",
                        "phone": "SĐT",
                        "menu": st.column_config.TextColumn("Thực đơn", width="large"),
                        "total_amount": st.column_config.NumberColumn("Tổng tiền", format="%d VND"),
                        "deposit": st.column_config.NumberColumn("Đã ứng/cọc", format="%d VND"),
                        "remaining": st.column_config.NumberColumn("Còn lại (Guide trả)", format="%d VND", disabled=True)
                    },
                    column_order=("date", "meal_name", "restaurant_name", "address", "phone", "menu", "total_amount", "deposit", "remaining"),
                    use_container_width=True
                )
                
                cols_r = ['date', 'meal_name', 'restaurant_name', 'address', 'phone', 'menu', 'total_amount', 'deposit']
                edited_rests = recalc_remaining(edited_rests)
                
                if not edited_rests[cols_r].equals(st.session_state.ls_rests_temp[cols_r].map(lambda x: x if not isinstance(x, float) else x)):
                    st.session_state.ls_rests_temp = edited_rests
                    st.rerun()
              # 4. ĐIỂM THAM QUAN (MỚI)
                st.markdown("##### 4. Điểm tham quan")
                df_sightseeings = st.session_state.ls_sight_temp.copy()
                
                df_sightseeings = recalc_remaining(df_sightseeings)

                edited_sightseeings = st.data_editor(
                    df_sightseeings,
                    num_rows="dynamic",
                    key="sightseeing_editor",
                    column_config={
                        "date": st.column_config.TextColumn("Ngày"),
                        "name": st.column_config.TextColumn("Tên địa điểm", required=True),
                        "address": "Địa chỉ",
                        "quantity": st.column_config.NumberColumn("Số lượng", min_value=0),
                        "total_amount": st.column_config.NumberColumn("Tổng tiền", format="%d VND"),
                        "deposit": st.column_config.NumberColumn("Đã cọc", format="%d VND"),
                        "remaining": st.column_config.NumberColumn("Còn lại", format="%d VND", disabled=True),
                        "note": st.column_config.TextColumn("Lưu ý")
                    },
                    column_order=("date", "name", "address", "quantity", "total_amount", "deposit", "remaining", "note"),
                    use_container_width=True
                )
                
                cols_s = ['date', 'name', 'address', 'quantity', 'total_amount', 'deposit', 'note']
                edited_sightseeings = recalc_remaining(edited_sightseeings)
                
                if not edited_sightseeings[cols_s].equals(st.session_state.ls_sight_temp[cols_s].map(lambda x: x if not isinstance(x, float) else x)):
                    st.session_state.ls_sight_temp = edited_sightseeings
                    st.rerun()
              # 5. CHI PHÍ PHÁT SINH (Đã đổi thứ tự lên trên)
                st.divider()
                st.markdown("##### 5. Chi phí phát sinh (Nước, Sim, Banner...)")
                df_incurred = st.session_state.ls_incurred_temp.copy()
                
                # Clean numbers và tính toán
                df_incurred['price'] = df_incurred['price'].apply(clean_vnd_val)
                df_incurred['quantity'] = df_incurred['quantity'].apply(clean_vnd_val)
                df_incurred['deposit'] = df_incurred['deposit'].apply(clean_vnd_val)
                df_incurred['total_amount'] = df_incurred['price'] * df_incurred['quantity']
                df_incurred['remaining'] = df_incurred['total_amount'] - df_incurred['deposit']

                edited_incurred = st.data_editor(
                    df_incurred,
                    num_rows="dynamic",
                    key="incurred_editor",
                    column_config={
                        "name": st.column_config.TextColumn("Tên chi phí", required=True),
                        "unit": st.column_config.TextColumn("ĐVT"),
                        "quantity": st.column_config.NumberColumn("Số lượng", min_value=0, format="%.0f"),
                        "price": st.column_config.NumberColumn("Đơn giá", format="%d VND"),
                        "total_amount": st.column_config.NumberColumn("Thành tiền", format="%d VND", disabled=True),
                        "deposit": st.column_config.NumberColumn("Đã ứng/cọc", format="%d VND"),
                        "remaining": st.column_config.NumberColumn("Còn lại", format="%d VND", disabled=True),
                        "note": st.column_config.TextColumn("Ghi chú")
                    },
                    column_order=["name", "unit", "quantity", "price", "total_amount", "deposit", "remaining", "note"],
                    use_container_width=True
                )
                
                # Tính toán lại sau khi edit
                edited_incurred['price'] = edited_incurred['price'].apply(clean_vnd_val)
                edited_incurred['quantity'] = edited_incurred['quantity'].apply(clean_vnd_val)
                edited_incurred['deposit'] = edited_incurred['deposit'].apply(clean_vnd_val)
                edited_incurred['total_amount'] = edited_incurred['price'] * edited_incurred['quantity']
                edited_incurred['remaining'] = edited_incurred['total_amount'] - edited_incurred['deposit']
                
                cols_inc = ['name', 'unit', 'quantity', 'price', 'total_amount', 'deposit', 'note']
                # So sánh với dữ liệu cũ
                if not edited_incurred[cols_inc].equals(st.session_state.ls_incurred_temp[cols_inc].map(lambda x: x if not isinstance(x, float) else x)):
                     st.session_state.ls_incurred_temp = edited_incurred[cols_inc]
                     st.rerun()
                st.write("")
                # 6. CHECKLIST BÀN GIAO (Đã đổi thứ tự xuống dưới)
                st.markdown("##### 6. Checklist bàn giao hồ sơ HDV")
                checklist_items = ["Chương trình đóng mộc", "Danh sách đóng mộc", "Bảo hiểm du lịch", "Thực đơn đóng mộc", "Vé máy bay", "Xác nhận khu du lịch/nhà hàng (Nếu có)", "Hợp đồng hướng dẫn"]
                
                current_checklist = dict(tour_info_ls).get('handover_checklist', '')
                checked_items = current_checklist.split(',') if current_checklist else []
                
                cols_chk = st.columns(2)
                new_checked_list = []
                all_checked = True
                
                for i, item in enumerate(checklist_items):
                    is_checked = item in checked_items
                    if cols_chk[i % 2].checkbox(item, value=is_checked, key=f"chk_ho_{tour_id_ls}_{i}"):
                        new_checked_list.append(item)
                    else:
                        all_checked = False

                st.write("")
                st.markdown("##### Tạm ứng cho HDV")
                
                # [FIX] Tự động tính tổng tiền còn lại để làm Tạm ứng
                def calc_rem_total(df):
                    if df.empty: return 0.0
                    def clean_val(x):
                        if isinstance(x, (int, float)): return float(x)
                        try: return float(str(x).replace('.', '').replace(' VND', '').strip())
                        except: return 0.0
                    
                    t = df['total_amount'].apply(clean_val)
                    d = df['deposit'].apply(clean_val)
                    return (t - d).sum()

                rem_h = calc_rem_total(st.session_state.ls_hotels_temp)
                rem_r = calc_rem_total(st.session_state.ls_rests_temp)
                rem_s = calc_rem_total(st.session_state.ls_sight_temp)
                
                # Tính riêng cho Incurred vì cần tính lại từ price * qty
                df_inc_c = st.session_state.ls_incurred_temp.copy()
                q_i = pd.to_numeric(df_inc_c['quantity'], errors='coerce').fillna(0)
                p_i = pd.to_numeric(df_inc_c['price'], errors='coerce').fillna(0)
                d_i = pd.to_numeric(df_inc_c['deposit'], errors='coerce').fillna(0)
                rem_i = ((q_i * p_i) - d_i).sum()
                
                total_rem_all = rem_h + rem_r + rem_s + rem_i
                
                tam_ung = float(total_rem_all)
                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 15px; border-radius: 10px; margin-top: 10px; border: 1px solid #c8e6c9;"><div style="display:flex; justify-content:space-between; font-size: 1.3em; color: #2e7d32;"><span><b>TẠM ỨNG CHO HDV</b></span> <b>{format_vnd(tam_ung)} VND</b></div></div>""", unsafe_allow_html=True)
                
                st.write("")
                st.write("")
                c_save, c_export = st.columns([1, 2])
                
                if c_save.button("💾 Lưu và chờ quyết toán", type="primary", use_container_width=True):
                    if not all_checked:
                        st.error("⛔ Bạn chưa hoàn thành Checklist bàn giao! Vui lòng kiểm tra đủ các mục trước khi lưu.")
                    else:
                        # Lưu Danh sách đoàn
                        run_query("DELETE FROM tour_guests WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_guests.empty:
                            data_guests = [(tour_id_ls, r['name'], r['dob'], r['hometown'], r['cccd'], r['type']) for _, r in edited_guests.iterrows() if r['name']]
                            if data_guests: run_query_many("INSERT INTO tour_guests (tour_id, name, dob, hometown, cccd, type) VALUES (?,?,?,?,?,?)", data_guests)
                        
                        # Lưu Khách sạn
                        run_query("DELETE FROM tour_hotels WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_hotels.empty:
                            data_hotels = [(tour_id_ls, r['hotel_name'], r['address'], r['phone'], r['total_rooms'], r['room_type'], r['total_amount'], r['deposit']) for _, r in edited_hotels.iterrows() if r['hotel_name']]
                            if data_hotels: run_query_many("INSERT INTO tour_hotels (tour_id, hotel_name, address, phone, total_rooms, room_type, total_amount, deposit) VALUES (?,?,?,?,?,?,?,?)", data_hotels)

                        # Lưu Nhà hàng
                        run_query("DELETE FROM tour_restaurants WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_rests.empty:
                            data_rests = [(tour_id_ls, r['meal_name'], r['restaurant_name'], r['address'], r['phone'], r['menu'], r['total_amount'], r['deposit'], r['date']) for _, r in edited_rests.iterrows() if r['meal_name']]
                            if data_rests: run_query_many("INSERT INTO tour_restaurants (tour_id, meal_name, restaurant_name, address, phone, menu, total_amount, deposit, date) VALUES (?,?,?,?,?,?,?,?,?)", data_rests)

                        # Lưu Điểm tham quan
                        run_query("DELETE FROM tour_sightseeings WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_sightseeings.empty:
                            data_sightseeings = [(tour_id_ls, r['name'], r['address'], r['quantity'], r['note'], r['date'], r['total_amount'], r['deposit']) for _, r in edited_sightseeings.iterrows() if r['name']]
                            if data_sightseeings: run_query_many("INSERT INTO tour_sightseeings (tour_id, name, address, quantity, note, date, total_amount, deposit) VALUES (?,?,?,?,?,?,?,?)", data_sightseeings)

                        # Lưu Chi phí phát sinh
                        run_query("DELETE FROM tour_incurred_costs WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_incurred.empty:
                            data_inc = [(tour_id_ls, r['name'], r['unit'], r['quantity'], r['price'], r['total_amount'], r['deposit'], r['note']) for _, r in edited_incurred.iterrows() if r['name']]
                            if data_inc: run_query_many("INSERT INTO tour_incurred_costs (tour_id, name, unit, quantity, price, total_amount, deposit, note) VALUES (?,?,?,?,?,?,?,?)", data_inc)

                        # Lưu Checklist
                        checklist_str = ",".join(new_checked_list)
                        run_query("UPDATE tours SET handover_checklist=? WHERE id=?", (checklist_str, tour_id_ls), commit=True)
                        
                        st.success("✅ Đã lưu hồ sơ và checklist thành công! Tour đang chờ quyết toán.")
                        time.sleep(1); st.rerun()
                
                with c_export:
                    # --- XUẤT FILE TỔNG HỢP (BÀN GIAO + THỰC ĐƠN) ---
                    buffer_combined = io.BytesIO()
                    with pd.ExcelWriter(buffer_combined, engine='xlsxwriter') as writer:
                        workbook: Any = writer.book
                        # ws = workbook.add_worksheet("ThucDon")
                        
                        # Formats
                        fmt_comp = workbook.add_format({'bold': True, 'font_size': 12, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                        fmt_info = workbook.add_format({'font_size': 10, 'italic': True, 'font_name': 'Times New Roman'})
                        fmt_title = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter', 'font_color': '#E65100', 'border': 0, 'font_name': 'Times New Roman'})
                        fmt_header = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_color': '#E65100', 'font_name': 'Times New Roman'})
                        fmt_text = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        # ==========================================
                        # SHEET 1: BÀN GIAO (BAN_GIAO_HDV)
                        # ==========================================
                        ws_bg = workbook.add_worksheet("BAN_GIAO_HDV")
                        
                        # --- FORMATS (BÀN GIAO) ---
                        fmt_title_bg = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter', 'font_color': '#0D47A1', 'border': 0, 'font_name': 'Times New Roman'})
                        fmt_comp_bg = workbook.add_format({'bold': True, 'font_size': 11, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                        fmt_header_bg = workbook.add_format({'bold': True, 'bg_color': '#E0F7FA', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        fmt_label_bg = workbook.add_format({'bold': True, 'bg_color': '#F5F5F5', 'border': 1, 'align': 'left', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        fmt_text_bg = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        fmt_center_bg = workbook.add_format({'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        fmt_section_bg = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'font_color': '#E65100', 'align': 'left', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        money_fmt_bg = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_name': 'Times New Roman'})

                        # Helper to safely parse float from potential strings
                        def safe_float_exp(x):
                            if isinstance(x, (int, float)): return float(x)
                            try: return float(str(x).replace('.', '').replace(',', '').replace(' VND', '').strip())
                            except: return 0.0

                        # --- DATA PREP ---
                        t = dict(tour_info_ls)
                        
                        # --- LAYOUT BÀN GIAO ---
                        ws_bg.merge_range('A1:F1', comp['name'].upper(), fmt_comp_bg)
                        ws_bg.merge_range('A2:F2', "PHIẾU BÀN GIAO TOUR / TOUR ORDER", fmt_title_bg)
                        
                        # SECTION A
                        row = 3
                        ws_bg.merge_range(row, 0, row, 5, "A. THÔNG TIN ĐOÀN", fmt_section_bg)
                        row += 1
                        ws_bg.write(row, 0, "Tên đoàn:", fmt_label_bg)
                        ws_bg.merge_range(row, 1, row, 2, t.get('tour_name', ''), fmt_text_bg)
                        ws_bg.write(row, 3, "Mã Tour:", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, t.get('tour_code', ''), fmt_center_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "Số lượng:", fmt_label_bg)
                        ws_bg.merge_range(row, 1, row, 2, f"{t.get('guest_count', 0)} khách", fmt_text_bg)
                        ws_bg.write(row, 3, "Thời gian:", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, f"{t.get('start_date','')} - {t.get('end_date','')}", fmt_center_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "Điểm đón:", fmt_label_bg)
                        ws_bg.write(row, 1, t.get('pickup_location', ''), fmt_text_bg)
                        ws_bg.write(row, 2, "Giờ đón:", fmt_label_bg)
                        ws_bg.write(row, 3, t.get('pickup_time', ''), fmt_text_bg)
                        ws_bg.write(row, 4, "Chuyến bay:", fmt_label_bg)
                        ws_bg.write(row, 5, t.get('flight_code', ''), fmt_text_bg)

                        # SECTION B: LỊCH TRÌNH (DỜI TỪ D LÊN B VÀ CHIA THEO NGÀY)
                        row += 2
                        ws_bg.merge_range(row, 0, row, 5, "B. LỊCH TRÌNH CHI TIẾT", fmt_section_bg)
                        row += 1
                            
                        # Lấy dữ liệu lịch trình từ DB
                        itins_xls = run_query("SELECT * FROM tour_itineraries WHERE tour_id=? ORDER BY day_index", (tour_id_ls,))
                        itin_map_xls = {r['day_index']: r['content'] for r in itins_xls} if itins_xls else {}

                        # Tự động tạo dòng theo ngày
                        try:
                            s_date = datetime.strptime(t.get('start_date', ''), '%d/%m/%Y')
                            e_date = datetime.strptime(t.get('end_date', ''), '%d/%m/%Y')
                            delta = (e_date - s_date).days + 1
                            
                            if delta > 0:
                                for i in range(delta):
                                    curr_date = s_date + pd.Timedelta(days=i)
                                    date_str = curr_date.strftime('%d/%m')
                                    content_str = itin_map_xls.get(i, "")
                                    ws_bg.write(row, 0, f"Ngày {i+1} ({date_str})", fmt_label_bg)
                                    ws_bg.merge_range(row, 1, row, 5, content_str, fmt_text_bg)
                                    row += 1
                        except:
                            pass

                        # SECTION C: NHÂN SỰ (DỜI TỪ B XUỐNG C)
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "C. THÔNG TIN NHÂN SỰ & VẬN CHUYỂN", fmt_section_bg)
                        row += 1
                        headers_b = ["Vai trò", "Họ và tên", "Điện thoại", "Ghi chú / Biển số", "", ""]
                        for i, h in enumerate(headers_b): 
                            if h: ws_bg.write(row, i, h, fmt_header_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "Hướng dẫn viên", fmt_center_bg)
                        ws_bg.write(row, 1, t.get('guide_name', ''), fmt_text_bg)
                        ws_bg.write(row, 2, t.get('guide_phone', ''), fmt_center_bg)
                        ws_bg.write(row, 3, "", fmt_text_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "Lái xe", fmt_center_bg)
                        ws_bg.write(row, 1, t.get('driver_name', ''), fmt_text_bg)
                        ws_bg.write(row, 2, t.get('driver_phone', ''), fmt_center_bg)
                        ws_bg.write(row, 3, f"{t.get('car_plate', '')} ({t.get('car_type', '')})", fmt_text_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "Điều hành/Sale", fmt_center_bg)
                        ws_bg.write(row, 1, t.get('sale_name', ''), fmt_text_bg)
                        ws_bg.write(row, 2, "", fmt_center_bg)
                        ws_bg.write(row, 3, "", fmt_text_bg)

                        # SECTION C
                        row += 2
                        ws_bg.merge_range(row, 0, row, 5, "C. CHI TIẾT DỊCH VỤ", fmt_section_bg)
                        
                        # 1. Khách sạn
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "1. Lưu trú (Khách sạn)", fmt_label_bg)
                        row += 1
                        ws_bg.write(row, 0, "Tên KS", fmt_header_bg)
                        ws_bg.write(row, 1, "Liên hệ", fmt_header_bg)
                        ws_bg.write(row, 2, "Phòng/Loại", fmt_header_bg)
                        ws_bg.write(row, 3, "Tổng tiền", fmt_header_bg)
                        ws_bg.write(row, 4, "Đã cọc", fmt_header_bg)
                        ws_bg.write(row, 5, "Còn lại", fmt_header_bg)
                        
                        df_hotels_exp = st.session_state.ls_hotels_temp
                        if not df_hotels_exp.empty:
                            for _, h in df_hotels_exp.iterrows():
                                total = safe_float_exp(h.get('total_amount', 0))
                                dep = safe_float_exp(h.get('deposit', 0))
                                rem = total - dep
                                row += 1
                                ws_bg.write(row, 0, h['hotel_name'], fmt_text_bg)
                                ws_bg.write(row, 1, f"{h['address']}\n{h['phone']}", fmt_text_bg)
                                ws_bg.write(row, 2, f"{h['total_rooms']} ({h['room_type']})", fmt_center_bg)
                                ws_bg.write(row, 3, total, money_fmt_bg)
                                ws_bg.write(row, 4, dep, money_fmt_bg)
                                ws_bg.write(row, 5, rem, money_fmt_bg)
                        else:
                            row += 1; ws_bg.merge_range(row, 0, row, 5, "(Chưa có thông tin)", fmt_center_bg)

                        # 2. Nhà hàng
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "2. Ẩm thực (Nhà hàng)", fmt_label_bg)
                        row += 1
                        ws_bg.write(row, 0, "Bữa ăn", fmt_header_bg)
                        ws_bg.write(row, 1, "Nhà hàng", fmt_header_bg)
                        ws_bg.write(row, 2, "Liên hệ", fmt_header_bg)
                        ws_bg.write(row, 3, "Tổng tiền", fmt_header_bg)
                        ws_bg.write(row, 4, "Đã cọc", fmt_header_bg)
                        ws_bg.write(row, 5, "Còn lại", fmt_header_bg)
                        
                        df_rests_exp = st.session_state.ls_rests_temp
                        if not df_rests_exp.empty:
                            for _, r in df_rests_exp.iterrows():
                                total = safe_float_exp(r.get('total_amount', 0))
                                dep = safe_float_exp(r.get('deposit', 0))
                                rem = total - dep
                                row += 1
                                ws_bg.write(row, 0, r['meal_name'], fmt_text_bg)
                                ws_bg.write(row, 1, r['restaurant_name'], fmt_text_bg)
                                ws_bg.write(row, 2, f"{r['address']}\n{r['phone']}", fmt_text_bg)
                                ws_bg.write(row, 3, total, money_fmt_bg)
                                ws_bg.write(row, 4, dep, money_fmt_bg)
                                ws_bg.write(row, 5, rem, money_fmt_bg)
                        else:
                            row += 1; ws_bg.merge_range(row, 0, row, 5, "(Chưa có thông tin)", fmt_center_bg)

                        # 3. Điểm tham quan
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "3. Điểm tham quan", fmt_label_bg)
                        row += 1
                        ws_bg.write(row, 0, "Tên địa điểm", fmt_header_bg)
                        ws_bg.write(row, 1, "Địa chỉ", fmt_header_bg)
                        ws_bg.write(row, 2, "Số lượng", fmt_header_bg)
                        ws_bg.write(row, 3, "Tổng tiền", fmt_header_bg)
                        ws_bg.write(row, 4, "Còn lại", fmt_header_bg)
                        ws_bg.write(row, 5, "Lưu ý", fmt_header_bg)
                        
                        df_sightseeings_exp = st.session_state.ls_sight_temp
                        if not df_sightseeings_exp.empty:
                            for _, s in df_sightseeings_exp.iterrows():
                                total = safe_float_exp(s.get('total_amount', 0))
                                dep = safe_float_exp(s.get('deposit', 0))
                                rem = total - dep
                                row += 1
                                ws_bg.write(row, 0, s['name'], fmt_text_bg)
                                ws_bg.write(row, 1, s['address'], fmt_text_bg)
                                ws_bg.write(row, 2, s['quantity'], fmt_center_bg)
                                ws_bg.write(row, 3, total, money_fmt_bg)
                                ws_bg.write(row, 4, rem, money_fmt_bg)
                                ws_bg.write(row, 5, s['note'], fmt_text_bg)
                        else:
                            row += 1; ws_bg.merge_range(row, 0, row, 5, "(Chưa có thông tin)", fmt_center_bg)

                        # 4. Chi phí phát sinh (MỚI)
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "4. Chi phí phát sinh (Nước, Sim, Banner...)", fmt_label_bg)
                        row += 1
                        ws_bg.write(row, 0, "Tên chi phí", fmt_header_bg)
                        ws_bg.write(row, 1, "ĐVT", fmt_header_bg)
                        ws_bg.write(row, 2, "Số lượng", fmt_header_bg)
                        ws_bg.write(row, 3, "Tổng tiền", fmt_header_bg)
                        ws_bg.write(row, 4, "Đã cọc", fmt_header_bg)
                        ws_bg.write(row, 5, "Còn lại", fmt_header_bg)
                        
                        df_inc_exp = st.session_state.ls_incurred_temp
                        if not df_inc_exp.empty:
                            for _, inc in df_inc_exp.iterrows():
                                try:
                                    qty = safe_float_exp(inc.get('quantity', 0))
                                    price = safe_float_exp(inc.get('price', 0))
                                    total = qty * price
                                    dep = safe_float_exp(inc.get('deposit', 0))
                                    rem = total - dep
                                except: total=0; dep=0; rem=0; qty=0

                                row += 1
                                ws_bg.write(row, 0, inc['name'], fmt_text_bg)
                                ws_bg.write(row, 1, inc['unit'], fmt_center_bg)
                                ws_bg.write(row, 2, qty, fmt_center_bg)
                                ws_bg.write(row, 3, total, money_fmt_bg)
                                ws_bg.write(row, 4, dep, money_fmt_bg)
                                ws_bg.write(row, 5, rem, money_fmt_bg)
                        else:
                            row += 1; ws_bg.merge_range(row, 0, row, 5, "(Không có)", fmt_center_bg)

                        # --- [FIX] TÍNH TOÁN TỔNG KẾT (Làm sạch dữ liệu trước khi tính) ---
                        def get_clean_sum(df, col_name):
                            if df.empty or col_name not in df.columns: return 0.0
                            def clean_val(x):
                                if isinstance(x, (int, float)): return float(x)
                                try: return float(str(x).replace('.', '').replace(' VND', '').strip())
                                except: return 0.0
                            return df[col_name].apply(clean_val).sum()

                        # 1. Tính Tổng chi phí (Total Amount)
                        t_h = get_clean_sum(st.session_state.ls_hotels_temp, 'total_amount')
                        t_r = get_clean_sum(st.session_state.ls_rests_temp, 'total_amount')
                        t_s = get_clean_sum(st.session_state.ls_sight_temp, 'total_amount')
                        
                        # Tính riêng cho Incurred (vì cần nhân quantity * price)
                        df_inc_calc = st.session_state.ls_incurred_temp.copy()
                        df_inc_calc['price'] = pd.to_numeric(df_inc_calc['price'], errors='coerce').fillna(0)
                        df_inc_calc['quantity'] = pd.to_numeric(df_inc_calc['quantity'], errors='coerce').fillna(0)
                        t_i = (df_inc_calc['price'] * df_inc_calc['quantity']).sum()

                        grand_total = t_h + t_r + t_s + t_i

                        # 2. Tính Đã cọc (Deposit)
                        d_h = get_clean_sum(st.session_state.ls_hotels_temp, 'deposit')
                        d_r = get_clean_sum(st.session_state.ls_rests_temp, 'deposit')
                        d_s = get_clean_sum(st.session_state.ls_sight_temp, 'deposit')
                        d_i = get_clean_sum(st.session_state.ls_incurred_temp, 'deposit')
                        
                        grand_deposit = d_h + d_r + d_s + d_i

                        # 3. Còn lại (HDV cần thanh toán cho NCC)
                        grand_remaining = grand_total - grand_deposit
                        
                        # 4. Quyết toán (Còn lại - Tạm ứng)
                        # tam_ung đã được tính ở UI và truyền vào đây
                        balance = grand_remaining - tam_ung

                        # SECTION D: TỔNG KẾT & TẠM ỨNG
                        row += 2
                        ws_bg.merge_range(row, 0, row, 5, "D. TỔNG KẾT KINH PHÍ", fmt_section_bg)
                        
                        # 1. Tổng chi phí
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "1. TỔNG CHI PHÍ TOUR:", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, grand_total, money_fmt_bg)
                        
                        # 2. Đã cọc (Phân rã)
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "2. ĐÃ CỌC / THANH TOÁN TRƯỚC (CÔNG TY):", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, grand_deposit, money_fmt_bg)
                        
                        # 3. Còn lại
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "3. CÒN LẠI CẦN THANH TOÁN (HDV CHI):", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, grand_remaining, workbook.add_format({'bold': True, 'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'bg_color': '#FFF9C4', 'font_name': 'Times New Roman'}))

                        # [FIX] Chi tiết còn lại (Thay vì chi tiết cọc)
                        r_h = t_h - d_h
                        r_r = t_r - d_r
                        r_s = t_s - d_s
                        r_i = t_i - d_i

                        row += 1
                        ws_bg.write(row, 0, "   - Khách sạn:", fmt_text_bg)
                        ws_bg.merge_range(row, 1, row, 3, r_h, money_fmt_bg)
                        row += 1
                        ws_bg.write(row, 0, "   - Nhà hàng:", fmt_text_bg)
                        ws_bg.merge_range(row, 1, row, 3, r_r, money_fmt_bg)
                        row += 1
                        ws_bg.write(row, 0, "   - Tham quan:", fmt_text_bg)
                        ws_bg.merge_range(row, 1, row, 3, r_s, money_fmt_bg)
                        row += 1
                        ws_bg.write(row, 0, "   - Phát sinh:", fmt_text_bg)
                        ws_bg.merge_range(row, 1, row, 3, r_i, money_fmt_bg)

                        # 4. Tạm ứng cho HDV
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "4. TẠM ỨNG CHO HDV:", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, tam_ung, money_fmt_bg)

                        # 5. Quyết toán
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "5. QUYẾT TOÁN (THU LẠI / CHI THÊM):", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, balance, workbook.add_format({'bold': True, 'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_color': '#D32F2F', 'font_size': 11, 'font_name': 'Times New Roman'}))

                        # FOOTER: CHỮ KÝ
                        row += 3
                        fmt_sig_title = workbook.add_format({'bold': True, 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        fmt_sig_name = workbook.add_format({'italic': True, 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        
                        ws_bg.write(row, 0, "NGƯỜI LẬP PHIẾU", fmt_sig_title)
                        ws_bg.merge_range(row, 1, row, 2, "KẾ TOÁN", fmt_sig_title)
                        ws_bg.write(row, 3, "GIÁM ĐỐC", fmt_sig_title)
                        ws_bg.merge_range(row, 4, row, 5, "HƯỚNG DẪN VIÊN", fmt_sig_title)
                        
                        row += 1
                        ws_bg.write(row, 0, "(Ký, họ tên)", fmt_sig_name)
                        ws_bg.merge_range(row, 1, row, 2, "(Ký, họ tên)", fmt_sig_name)
                        ws_bg.write(row, 3, "(Ký, họ tên)", fmt_sig_name)
                        ws_bg.merge_range(row, 4, row, 5, "(Ký, họ tên)", fmt_sig_name)
                        
                        # Space for signature
                        row += 5
                        
                        # Names
                        ws_bg.write(row, 0, t.get('sale_name', ''), fmt_sig_title)
                        ws_bg.merge_range(row, 4, row, 5, t.get('guide_name', ''), fmt_sig_title)

                        ws_bg.set_column('A:A', 20)
                        ws_bg.set_column('B:F', 18)

                        # ==========================================
                        # SHEET: DANH SÁCH ĐOÀN (DanhSachDoan)
                        # ==========================================
                        ws_ds = workbook.add_worksheet("DanhSachDoan")
                        
                        # 1. Company Info
                        ws_ds.merge_range('A1:F1', comp['name'].upper(), fmt_comp_bg)
                        ws_ds.merge_range('A2:F2', "DANH SÁCH ĐOÀN / GUEST LIST", fmt_title_bg)
                        
                        # 2. Tour Info
                        ws_ds.write('A4', "Tên đoàn:", fmt_label_bg)
                        ws_ds.merge_range('B4:C4', t.get('tour_name', ''), fmt_text_bg)
                        ws_ds.write('D4', "Mã Tour:", fmt_label_bg)
                        ws_ds.merge_range('E4:F4', t.get('tour_code', ''), fmt_center_bg)
                        
                        ws_ds.write('A5', "Thời gian:", fmt_label_bg)
                        ws_ds.merge_range('B5:C5', f"{t.get('start_date','')} - {t.get('end_date','')}", fmt_center_bg)
                        ws_ds.write('D5', "Số khách:", fmt_label_bg)
                        ws_ds.merge_range('E5:F5', f"{t.get('guest_count', 0)} khách", fmt_center_bg)

                        # 3. Table Header
                        row_ds = 7
                        headers_ds = ["STT", "Họ và tên", "Ngày sinh", "Quê quán", "Số CCCD", "Phân loại"]
                        for i, h in enumerate(headers_ds):
                            ws_ds.write(row_ds, i, h, fmt_header_bg)
                        
                        # 4. Data
                        if not edited_guests.empty:
                            for i, (idx, row_g) in enumerate(edited_guests.iterrows()):
                                row_ds += 1
                                ws_ds.write(row_ds, 0, i + 1, fmt_center_bg)
                                ws_ds.write(row_ds, 1, row_g.get('name', ''), fmt_text_bg)
                                ws_ds.write(row_ds, 2, row_g.get('dob', ''), fmt_center_bg)
                                ws_ds.write(row_ds, 3, row_g.get('hometown', ''), fmt_text_bg)
                                ws_ds.write(row_ds, 4, row_g.get('cccd', ''), fmt_center_bg)
                                ws_ds.write(row_ds, 5, row_g.get('type', ''), fmt_center_bg)
                        
                        # Column Widths
                        ws_ds.set_column('A:A', 5)
                        ws_ds.set_column('B:B', 25)
                        ws_ds.set_column('C:C', 15)
                        ws_ds.set_column('D:D', 20)
                        ws_ds.set_column('E:F', 15)

                        # ==========================================
                        # SHEET 2: THỰC ĐƠN (ThucDon)
                        # ==========================================
                        ws_menu = workbook.add_worksheet("ThucDon")
                        
                        # Formats (Menu)
                        fmt_comp_menu = workbook.add_format({'bold': True, 'font_size': 12, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                        fmt_info_menu = workbook.add_format({'font_size': 10, 'italic': True, 'font_name': 'Times New Roman'})
                        fmt_title_menu = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter', 'font_color': '#E65100', 'border': 0, 'font_name': 'Times New Roman'})
                        fmt_header_menu = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_color': '#E65100', 'font_name': 'Times New Roman'})
                        fmt_text_menu = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        
                        # Company Info
                        comp_data = get_company_data()
                        ws_menu.write('A1', comp_data['name'], fmt_comp_menu)
                        ws_menu.write('A2', f"ĐC: {comp_data['address']}", fmt_info_menu)
                        ws_menu.write('A3', f"MST: {comp_data['phone']}", fmt_info_menu)
                        
                        # Title
                        ws_menu.merge_range('A5:C5', f"DANH SÁCH THỰC ĐƠN TOUR: {tour_info_ls['tour_name']}", fmt_title_menu)
                        
                        # Table Header
                        headers_menu = ["Thông tin nhà hàng", "Bữa ăn / Thời gian", "Thực đơn"]
                        for i, h in enumerate(headers_menu):
                            ws_menu.write(6, i, h, fmt_header_menu)
                            
                        # Data
                        row_menu = 7
                        if not df_rests.empty:
                            df_rests_exp = df_rests.fillna('')
                            for _, r in df_rests_exp.iterrows():
                                # Gộp thông tin: Tên, Địa chỉ, Liên hệ
                                info_parts = [str(r[k]) for k in ['restaurant_name', 'address', 'phone'] if str(r[k]).strip()]
                                info_str = "\n".join(info_parts)
                                
                                ws_menu.write(row_menu, 0, info_str, fmt_text_menu)
                                ws_menu.write(row_menu, 1, r['meal_name'], fmt_text_menu)
                                ws_menu.write(row_menu, 2, r['menu'], fmt_text_menu)
                                row_menu += 1
                        
                        # Column widths
                        ws_menu.set_column('A:A', 40) # Thông tin nhà hàng
                        ws_menu.set_column('B:B', 25) # Bữa ăn / Thời gian
                        ws_menu.set_column('C:C', 50) # Thực đơn

                    st.download_button("📥 Xuất Hồ Sơ Bàn Giao & Thực Đơn (Excel)", buffer_combined.getvalue(), f"HoSo_BanGiao_{tour_info_ls['tour_code']}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

    # ---------------- TAB 2: QUYẾT TOÁN ----------------
    with tab_act:
        st.subheader("💸 Quyết Toán ")
        
        selected_tour_act_label = st.selectbox("Chọn Đoàn quyết toán:", list(tour_options.keys()) if tour_options else [], key="sel_tour_act")
        
        if selected_tour_act_label:
            tour_id_act = tour_options[selected_tour_act_label] # type: ignore
            tour_info_act = next((t for t in all_tours if t['id'] == tour_id_act), None)
            if not tour_info_act:
                st.error("Không tìm thấy thông tin tour.")
                st.stop()
            assert tour_info_act is not None
            
            # --- Lấy Dự toán để so sánh ---
            est_items = run_query("SELECT SUM(total_amount) as total FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id_act,), fetch_one=True)
            # If the query returns a row and the 'total' is not None (SQL SUM can return NULL), use it. Otherwise, default to 0.
            est_total_cost = est_items['total'] if est_items and est_items['total'] is not None else 0
            # Tính lại giá bán chốt (Dựa trên % đã lưu)
            p_pct = tour_info_act['est_profit_percent'] # type: ignore
            t_pct = tour_info_act['est_tax_percent'] # type: ignore
            est_profit_val = est_total_cost * (p_pct / 100)
            est_final_sale = (est_total_cost + est_profit_val) * (1 + t_pct/100)
            
            # [UPDATED] Lấy Tổng doanh thu từ bên Dự toán (Giá chốt * SL)
            t_act_dict_calc = dict(tour_info_act)
            final_price_est = float(t_act_dict_calc.get('final_tour_price', 0) or 0)
            child_price_est = float(t_act_dict_calc.get('child_price', 0) or 0)
            final_qty_est = float(t_act_dict_calc.get('final_qty', 0) or 0)
            child_qty_est = float(t_act_dict_calc.get('child_qty', 0) or 0)
            if final_qty_est == 0: final_qty_est = float(t_act_dict_calc.get('guest_count', 1))
            total_revenue_est = (final_price_est * final_qty_est) + (child_price_est * child_qty_est)
            
            if total_revenue_est > 0:
                est_final_sale = total_revenue_est
            else:
                est_profit_val = est_total_cost * (p_pct / 100)
                est_final_sale = (est_total_cost + est_profit_val) * (1 + t_pct/100)
            
            st.info(f"TỔNG DOANH THU: {format_vnd(est_final_sale)} VND")

            # --- [UPDATED] PHÂN TÍCH CHI PHÍ ---
            st.divider()
            st.markdown("### 📊 Phân tích Chi phí")
            
            linked_docs = run_query("SELECT * FROM invoices WHERE cost_code=? AND status='active'", (tour_info_act['tour_code'],)) # type: ignore
            df_linked = pd.DataFrame([dict(r) for r in linked_docs]) if linked_docs else pd.DataFrame()
            
            total_unc = 0
            total_inv = 0
            df_unc = pd.DataFrame()
            df_inv = pd.DataFrame()

            if not df_linked.empty:
                unc_mask = df_linked['invoice_number'].astype(str).str.contains("UNC", case=False, na=False) # type: ignore
                df_unc = df_linked.loc[unc_mask]
                total_unc = df_unc['total_amount'].sum()
                
                inv_mask = (df_linked['type'] == 'IN') & (~unc_mask)
                df_inv = df_linked.loc[inv_mask]
                total_inv = df_inv['total_amount'].sum()

            c_unc_t, c_inv_t = st.columns(2)
            with c_unc_t:
                st.markdown(f"#### 💸 1. Chi phí UNC: {format_vnd(total_unc)}")
                if not df_unc.empty:
                    # [UPDATED] Format tiền tệ Việt Nam có dấu chấm và chữ VND
                    df_unc_show = df_unc.copy()
                    df_unc_show['total_show'] = df_unc_show['total_amount'].apply(lambda x: format_vnd(x) + " VND") # type: ignore
                    st.dataframe(df_unc_show[['date', 'invoice_number', 'memo', 'total_show']],
                                 column_config={
                                     "date": "Ngày", 
                                     "invoice_number": "Số chứng từ", 
                                     "memo": "Nội dung", 
                                     "total_show": "Thành tiền"
                                 },
                                 use_container_width=True, hide_index=True)
                else: st.caption("Chưa có UNC.")
            
            with c_inv_t:
                st.markdown(f"#### 📄 2. Hóa đơn đầu vào: {format_vnd(total_inv)}")
                if not df_inv.empty:
                    # [UPDATED] Format tiền tệ Việt Nam có dấu chấm và chữ VND
                    df_inv_show = df_inv.copy()
                    df_inv_show['total_show'] = df_inv_show['total_amount'].apply(lambda x: format_vnd(x) + " VND") # type: ignore
                    st.dataframe(df_inv_show[['date', 'invoice_number', 'seller_name', 'total_show']], 
                                 column_config={"date": "Ngày", "invoice_number": "Số hóa đơn", "seller_name": "Đơn vị bán", "total_show": "Thành tiền"}, 
                                 use_container_width=True, hide_index=True)
                else: st.caption("Chưa có hóa đơn đầu vào.")

            # [CODE MỚI] Lấy dữ liệu Dự toán để so sánh
            est_items_ref = run_query("SELECT category, description, total_amount FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id_act,))
            est_lookup = {}
            if est_items_ref:
                for r in est_items_ref:
                    key = (str(r['category']).strip().lower(), str(r['description']).strip().lower()) # type: ignore
                    est_lookup[key] = float(r['total_amount'] or 0) # type: ignore
            
            with st.expander("👀 Bảng Dự Toán (Để đối chiếu)", expanded=False):
                if est_items_ref:
                    df_est_ref = pd.DataFrame([dict(r) for r in est_items_ref])
                    df_est_ref['total_amount'] = df_est_ref['total_amount'].apply(lambda x: format_vnd(x)) # type: ignore
                    st.dataframe(df_est_ref, column_config={"category": "Hạng mục", "description": "Diễn giải", "total_amount": "Dự toán"}, use_container_width=True, hide_index=True)
                else: st.info("Chưa có dữ liệu dự toán.")

            # --- Fetch Items (ACT) with Session State ---
            if "current_tour_id_act" not in st.session_state: st.session_state.current_tour_id_act = None
            if st.session_state.current_tour_id_act != tour_id_act:
                if "act_df_temp" in st.session_state: del st.session_state.act_df_temp
                st.session_state.current_tour_id_act = tour_id_act

            if "act_df_temp" not in st.session_state:
                act_items = run_query("SELECT * FROM tour_items WHERE tour_id=? AND item_type='ACT'", (tour_id_act,))
                if act_items:
                    df_act = pd.DataFrame([dict(r) for r in act_items])
                    if 'times' not in df_act.columns: df_act['times'] = 1.0
                    df_act = df_act[['category', 'description', 'unit', 'unit_price', 'quantity', 'times']]
                else:
                     # Gợi ý: Nếu chưa có item ACT, load item EST để sửa cho nhanh
                     est_items_raw = run_query("SELECT * FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id_act,))
                     if est_items_raw:
                         df_act = pd.DataFrame([dict(r) for r in est_items_raw])
                         if 'times' not in df_act.columns: df_act['times'] = 1.0
                         df_act = df_act[['category', 'description', 'unit', 'unit_price', 'quantity', 'times']]
                     else:
                         df_act = pd.DataFrame([{"category": "", "description": "", "unit": "", "quantity": 0, "unit_price": 0, "times": 1}])
                st.session_state.act_df_temp = df_act

            # Prepare Display Data
            df_act_display = st.session_state.act_df_temp.copy()
            guest_cnt_act = tour_info_act['guest_count'] if tour_info_act['guest_count'] else 1 # type: ignore
            
            # Calculate numeric totals
            # Ensure numeric types
            df_act_display['quantity'] = pd.to_numeric(df_act_display['quantity'], errors='coerce').fillna(0)
            df_act_display['unit_price'] = pd.to_numeric(df_act_display['unit_price'], errors='coerce').fillna(0)
            df_act_display['times'] = pd.to_numeric(df_act_display['times'], errors='coerce').fillna(1)

            # Formula: Total = Unit * Qty * Times
            df_act_display['total_val'] = df_act_display['quantity'] * df_act_display['unit_price'] * df_act_display['times']
            # Formula: Pax = Total / Guests
            df_act_display['price_per_pax'] = df_act_display['total_val'] / guest_cnt_act
            
            # Format strings
            df_act_display['price_per_pax'] = df_act_display['price_per_pax'].apply(lambda x: format_vnd(x) + " VND")
            df_act_display['total_display'] = df_act_display['total_val'].apply(lambda x: format_vnd(x) + " VND") # type: ignore
            df_act_display['unit_price'] = df_act_display['unit_price'].apply(lambda x: format_vnd(x) + " VND") # type: ignore

            # [CODE MỚI] Tính toán so sánh (Dự toán vs Thực tế)
            def get_est_val(row): # type: ignore
                k = (str(row['category']).strip().lower(), str(row['description']).strip().lower()) # type: ignore
                return est_lookup.get(k, 0.0)
            
            df_act_display['est_val'] = df_act_display.apply(get_est_val, axis=1)
            df_act_display['diff_val'] = df_act_display['est_val'] - df_act_display['total_val']
            df_act_display['est_display'] = df_act_display['est_val'].apply(lambda x: format_vnd(x) + " VND")
            df_act_display['diff_display'] = df_act_display['diff_val'].apply(lambda x: format_vnd(x) + " VND")

            # --- LOGIC KHÓA / DUYỆT QUYẾT TOÁN ---
            req_act_status = tour_info_act['request_edit_act'] # type: ignore
            has_act_data = False
            check_act = run_query("SELECT id FROM tour_items WHERE tour_id=? AND item_type='ACT' LIMIT 1", (tour_id_act,))
            if check_act: has_act_data = True

            is_act_editable = False
            if current_user_role_tour in ['admin', 'admin_f1']:
                is_act_editable = True

            st.divider()
            st.markdown("#### ✍️ 3.Quyết toán")
            edited_act = st.data_editor(
                df_act_display,
                num_rows="dynamic",
                column_config={
                    "category": st.column_config.TextColumn("Hạng mục chi phí", required=False),
                    "description": st.column_config.TextColumn("Diễn giải"),
                    "unit": st.column_config.TextColumn("Đơn vị"),
                    "unit_price": st.column_config.TextColumn("Đơn giá (VND)", required=False),
                    "quantity": st.column_config.NumberColumn("Số lượng", min_value=0),
                    "times": st.column_config.NumberColumn("Số lần", min_value=1),
                    "price_per_pax": st.column_config.TextColumn("Giá/Pax", disabled=True),
                    "total_display": st.column_config.TextColumn("Thực tế (VND)", disabled=True),
                    "est_display": st.column_config.TextColumn("Dự toán (VND)", disabled=True),
                    "diff_display": st.column_config.TextColumn("Chênh lệch", disabled=True),
                    "total_val": st.column_config.NumberColumn("Hidden", disabled=True),
                    "est_val": st.column_config.NumberColumn("Hidden", disabled=True),
                    "diff_val": st.column_config.NumberColumn("Hidden", disabled=True),
                },
                disabled=not is_act_editable, # Khóa nếu không được phép sửa
                column_order=("category", "description", "unit", "unit_price", "quantity", "times", "price_per_pax", "total_display", "est_display", "diff_display"),
                use_container_width=True,
                hide_index=True,
                key="editor_act"
            )
            
            # --- AUTO-UPDATE CALCULATION (ACTUAL) ---
            if is_act_editable:
                # Tự động cập nhật khi dữ liệu thay đổi
                df_new_act = edited_act.copy()
                
                def clean_vnd_act_auto(x):
                    if isinstance(x, str):
                        return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    return float(x) if x else 0.0
                
                df_new_act['unit_price'] = df_new_act['unit_price'].apply(clean_vnd_act_auto)
                df_new_act['quantity'] = pd.to_numeric(df_new_act['quantity'], errors='coerce').fillna(0)
                if 'times' not in df_new_act.columns: df_new_act['times'] = 1
                df_new_act['times'] = pd.to_numeric(df_new_act['times'], errors='coerce').fillna(1)
                
                # So sánh với dữ liệu cũ
                cols_check_act = ['category', 'description', 'unit', 'unit_price', 'quantity', 'times']
                df_old_act = st.session_state.act_df_temp.copy()
                if 'times' not in df_old_act.columns: df_old_act['times'] = 1
                
                # Reset index và fillna để so sánh
                df_new_check_act = df_new_act[cols_check_act].reset_index(drop=True).fillna(0)
                df_old_check_act = df_old_act[cols_check_act].reset_index(drop=True).fillna(0)
                
                if len(df_new_check_act) != len(df_old_check_act) or not df_new_check_act.equals(df_old_check_act):
                    st.session_state.act_df_temp = df_new_act[cols_check_act]
                    st.rerun()

            act_total_cost = 0
            if not edited_act.empty:
                # Parse unit_price
                # [FIX] Handle case where a cell is None, which becomes the string 'None' after astype(str)
                cleaned_prices_act = edited_act['unit_price'].astype(str).str.replace('.', '', regex=False).str.replace(' VND', '', regex=False).str.strip()
                p_price_act = cleaned_prices_act.apply(lambda x: float(x) if x and x.lower() != 'none' else 0.0)
                # Ensure 'times' column exists and is numeric before accessing it
                # Use .get() with a default Series to handle cases where 'times' might be missing
                times_col_act = edited_act.get('times', pd.Series([1.0] * len(edited_act), index=edited_act.index)).fillna(1).astype(float) # type: ignore
                act_total_cost = (edited_act['quantity'] * p_price_act * times_col_act).sum()
            # TỔNG CHI PHÍ THỰC TẾ = Hóa đơn + Phát sinh (Nhập tay)
            final_act_cost = act_total_cost + total_inv

            # --- TỔNG KẾT QUYẾT TOÁN ---
            st.divider()
            st.markdown("### ⚖️ Tổng kết & Đối chiếu")
            
            c_sum1, c_sum2, c_sum3 = st.columns(3)
            c_sum1.metric("Tổng Chi phí (HĐ + Phát sinh)", format_vnd(final_act_cost), help="Tổng chi phí thực tế của tour")
            c_sum2.metric("Tổng UNC (Đã thanh toán)", format_vnd(total_unc), help="Tổng tiền đã chi ra từ tài khoản")
            
            diff = total_unc - final_act_cost
            if diff == 0:
                c_sum3.success("✅ Đã khớp (UNC = Chi phí)")
            elif diff > 0:
                c_sum3.warning(f"⚠️ UNC dư: {format_vnd(diff)}")
            else:
                c_sum3.error(f"⚠️ Thiếu UNC: {format_vnd(abs(diff))}")
            
            # Lợi nhuận = Tổng doanh thu (Dự toán) - Tổng chi
            final_profit = est_final_sale - final_act_cost
            
            st.markdown(f"""<div class="profit-summary-card">
                <h3>TỔNG DOANH THU - TỔNG CHI = LỢI NHUẬN</h3>
                <div class="formula">{format_vnd(est_final_sale)} - {format_vnd(final_act_cost)} = <span class="result">{format_vnd(final_profit)} VND</span></div>
            </div>
            """, unsafe_allow_html=True)

            # --- EXPORT EXCEL (ACT) ---
            st.write("")
            # Prepare Data for Export
            df_exp_act = edited_act.copy()
            if 'times' not in df_exp_act.columns: df_exp_act['times'] = 1
            df_exp_act['times'] = df_exp_act.get('times', pd.Series([1.0] * len(df_exp_act), index=df_exp_act.index)).fillna(1).astype(float)

            # Clean numbers

            def clean_num_act(x): # type: ignore
                if isinstance(x, str):
                    return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip())
                return float(x) if x else 0.0
            
            df_exp_act['unit_price'] = df_exp_act['unit_price'].apply(clean_num_act)
            df_exp_act['quantity'] = pd.to_numeric(df_exp_act['quantity'], errors='coerce').fillna(0)
            df_exp_act['total_amount'] = df_exp_act['quantity'] * df_exp_act['unit_price'] * df_exp_act['times']
            df_exp_act['price_per_pax'] = df_exp_act['total_amount'] / guest_cnt_act
            
            # --- COMPARISON LOGIC ---
            # [CODE MỚI] Sử dụng lại est_lookup đã tạo ở trên để tính cột Dự toán và Chênh lệch cho Excel
            def get_est_val_exp(row): # type: ignore
                k = (str(row['category']).strip().lower(), str(row['description']).strip().lower()) # type: ignore
                return est_lookup.get(k, 0.0)

            df_exp_act['est_amount'] = df_exp_act.apply(get_est_val_exp, axis=1)
            df_exp_act['diff_amount'] = df_exp_act['est_amount'] - df_exp_act['total_amount'] # type: ignore
            
            def classify_item(row):
                if row['diff_amount'] < 0: return "Vượt chi"
                elif row['diff_amount'] > 0: return "Tiết kiệm"
                return ""

            df_exp_act['Ghi chú'] = df_exp_act.apply(classify_item, axis=1)

            # Rename
            df_exp_act = df_exp_act.rename(columns={
                'category': 'Hạng mục', 
                'description': 'Diễn giải', 
                'unit': 'Đơn vị', 
                'unit_price': 'Đơn giá', 
                'quantity': 'Số lượng', 
                'times': 'Số lần',
                'price_per_pax': 'Giá/Pax',
                'total_amount': 'Thực tế',
                'est_amount': 'Dự toán',
                'diff_amount': 'Chênh lệch'
            })
            
            # [REQUEST 1] Bỏ cột 'Số lần' -> Keep it
            cols_to_export = ['Hạng mục', 'Diễn giải', 'Đơn vị', 'Đơn giá', 'Số lượng', 'Số lần', 'Giá/Pax', 'Dự toán', 'Thực tế', 'Chênh lệch', 'Ghi chú']
            df_exp_act_filtered = df_exp_act[cols_to_export]

            # [REQUEST 2] Tách thành 2 bảng: Chi phí trong dự toán và chi phí phát sinh
            df_in_est = df_exp_act_filtered[df_exp_act_filtered['Dự toán'] > 0].copy()
            df_extra_cost = df_exp_act_filtered[df_exp_act_filtered['Dự toán'] == 0].copy()

            buffer_act = io.BytesIO()
            file_ext_act = "xlsx"
            mime_type_act = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            
            try:
                with pd.ExcelWriter(buffer_act, engine='xlsxwriter') as writer:
                    workbook: Any = writer.book
                    worksheet = workbook.add_worksheet('QuyetToan')
                    
                    # Styles (Copied and adapted)
                    company_name_fmt = workbook.add_format({'bold': True, 'font_size': 14, 'font_color': '#D84315', 'font_name': 'Times New Roman'}) # Orange for Act
                    company_info_fmt = workbook.add_format({'font_size': 10, 'italic': True, 'font_color': '#424242', 'font_name': 'Times New Roman'})
                    title_fmt = workbook.add_format({'bold': True, 'font_size': 18, 'align': 'center', 'valign': 'vcenter', 'font_color': '#BF360C', 'bg_color': '#FBE9E7', 'border': 1, 'font_name': 'Times New Roman'})
                    
                    header_fmt = workbook.add_format({'bold': True, 'fg_color': '#D84315', 'font_color': 'white', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                    body_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_size': 10, 'font_name': 'Times New Roman'})
                    body_center_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'align': 'center', 'font_size': 10, 'font_name': 'Times New Roman'})
                    money_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_size': 10, 'font_name': 'Times New Roman'})
                    
                    # Summary Styles
                    sum_header_bg_fmt = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'font_color': '#E65100', 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                    sum_label_fmt = workbook.add_format({'bold': True, 'align': 'left', 'border': 1, 'bg_color': '#FAFAFA', 'font_name': 'Times New Roman'})
                    sum_val_fmt = workbook.add_format({'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_name': 'Times New Roman'})
                    sum_val_bold_fmt = workbook.add_format({'bold': True, 'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_name': 'Times New Roman'})
                    
                    # [CODE MỚI] Format màu đỏ cho dòng âm
                    alert_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_size': 10, 'font_color': '#D32F2F', 'font_name': 'Times New Roman'})
                    alert_money_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_size': 10, 'font_color': '#D32F2F', 'font_name': 'Times New Roman'})

                    # [CODE MỚI] Format cho tiêu đề các bảng chi phí
                    section_title_fmt = workbook.add_format({'bold': True, 'font_size': 12, 'font_color': '#004D40', 'bg_color': '#E0F2F1', 'border': 1, 'align': 'center', 'font_name': 'Times New Roman'})

                    # 1. Company Info
                    if comp['logo_b64_str']:
                        try:
                            logo_data = base64.b64decode(comp['logo_b64_str'])
                            image_stream = io.BytesIO(logo_data)
                            img = Image.open(image_stream)
                            w, h = img.size
                            scale = 60 / h if h > 0 else 0.5
                            image_stream.seek(0)
                            worksheet.insert_image('A1', 'logo.png', {'image_data': image_stream, 'x_scale': scale, 'y_scale': scale, 'x_offset': 5, 'y_offset': 5})
                        except: pass
                    
                    worksheet.write('B1', comp['name'], company_name_fmt)
                    worksheet.write('B2', f"ĐC: {comp['address']}", company_info_fmt)
                    worksheet.write('B3', f"MST: {comp['phone']}", company_info_fmt)
                    
                    # 2. Tour Info
                    worksheet.merge_range('A5:I5', "BẢNG QUYẾT TOÁN CHI PHÍ TOUR", title_fmt)
                    
                    t_info_dict = {k: tour_info_act[k] for k in tour_info_act.keys()}
                    worksheet.write('A7', "Tên đoàn:", sum_label_fmt)
                    worksheet.merge_range('B7:D7', t_info_dict.get('tour_name',''), sum_val_fmt)
                    worksheet.write('E7', "Mã đoàn:", sum_label_fmt)
                    worksheet.merge_range('F7:I7', t_info_dict.get('tour_code',''), sum_val_fmt)
                    
                    worksheet.write('A8', "Khách hàng:", sum_label_fmt)
                    worksheet.merge_range('B8:D8', f"{t_info_dict.get('customer_name','')} - {t_info_dict.get('customer_phone','')}", sum_val_fmt)
                    worksheet.write('E8', "Sales:", sum_label_fmt)
                    worksheet.merge_range('F8:I8', t_info_dict.get('sale_name',''), sum_val_fmt)
                    
                    worksheet.write('A9', "Thời gian:", sum_label_fmt)
                    worksheet.merge_range('B9:D9', f"{t_info_dict.get('start_date','')} - {t_info_dict.get('end_date','')}", sum_val_fmt)
                    worksheet.write('E9', "Số khách:", sum_label_fmt)
                    worksheet.merge_range('F9:I9', t_info_dict.get('guest_count',0), sum_val_fmt)

                    # 3. Table Header & Body (MODIFIED)
                    current_row = 10 # Bắt đầu từ dòng 11

                    # --- Bảng 1: Chi phí trong dự toán ---
                    if not df_in_est.empty:
                        worksheet.merge_range(current_row, 0, current_row, len(df_in_est.columns)-1, "CHI PHÍ TRONG DỰ TOÁN", section_title_fmt)
                        current_row += 1
                        for col_num, value in enumerate(df_in_est.columns):
                            worksheet.write(current_row, col_num, value, header_fmt)
                        for row_idx in range(len(df_in_est)):
                            diff_val = df_in_est.iloc[row_idx, 7] # Chênh lệch
                            is_negative = isinstance(diff_val, (int, float)) and diff_val < 0
                            for col_idx in range(len(df_in_est.columns)):
                                val = df_in_est.iloc[row_idx, col_idx]
                                if col_idx == 2: fmt = body_center_fmt
                                elif col_idx in [3, 4, 5, 6, 7, 8, 9]: fmt = money_fmt
                                else: fmt = body_fmt
                                if is_negative:
                                    if col_idx in [3, 4, 5, 6, 7, 8, 9]: fmt = alert_money_fmt
                                    else: fmt = alert_fmt
                                if pd.isna(val): val = ""
                                worksheet.write(current_row + 1 + row_idx, col_idx, val, fmt)
                        current_row += len(df_in_est) + 1

                    # Thêm dòng trống
                    current_row += 1

                    # --- Bảng 2: Chi phí phát sinh ngoài dự toán ---
                    if not df_extra_cost.empty:
                        worksheet.merge_range(current_row, 0, current_row, len(df_extra_cost.columns)-1, "CHI PHÍ PHÁT SINH NGOÀI DỰ TOÁN", section_title_fmt)
                        current_row += 1
                        for col_num, value in enumerate(df_extra_cost.columns):
                            worksheet.write(current_row, col_num, value, header_fmt)
                        for row_idx in range(len(df_extra_cost)):
                            # Chi phí phát sinh luôn là âm (vượt chi)
                            is_negative = True
                            for col_idx in range(len(df_extra_cost.columns)):
                                val = df_extra_cost.iloc[row_idx, col_idx]
                                if col_idx == 2: fmt = body_center_fmt
                                elif col_idx in [3, 4, 5, 6, 7, 8, 9]: fmt = money_fmt
                                else: fmt = body_fmt
                                if is_negative:
                                    if col_idx in [3, 4, 5, 6, 7, 8, 9]: fmt = alert_money_fmt
                                    else: fmt = alert_fmt
                                if pd.isna(val): val = ""
                                worksheet.write(current_row + 1 + row_idx, col_idx, val, fmt)
                        current_row += len(df_extra_cost) + 1
                    
                    # 4. Summary
                    sum_row = current_row + 1
                    
                    worksheet.merge_range(sum_row, 0, sum_row, 3, "TỔNG KẾT QUYẾT TOÁN", sum_header_bg_fmt)
                    
                    # [CODE MỚI] Hiển thị đầy đủ thông tin tài chính
                    # 1. Tổng doanh thu
                    worksheet.write(sum_row+1, 0, "1. Tổng doanh thu:", sum_label_fmt)
                    worksheet.merge_range(sum_row+1, 1, sum_row+1, 3, est_final_sale, sum_val_bold_fmt)
                    
                    # 2. Tổng chi phí (Bảng kê + Hóa đơn ngoài)
                    worksheet.write(sum_row+2, 0, "2. Tổng chi phí thực tế:", sum_label_fmt)
                    worksheet.merge_range(sum_row+2, 1, sum_row+2, 3, final_act_cost, sum_val_bold_fmt)
                    
                    # 3. Lợi nhuận
                    worksheet.write(sum_row+3, 0, "3. Lợi nhuận thực tế:", sum_label_fmt)
                    profit_fmt = workbook.add_format({'bold': True, 'num_format': '#,##0', 'align': 'right', 'border': 1, 'bg_color': '#C8E6C9', 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                    worksheet.merge_range(sum_row+3, 1, sum_row+3, 3, final_profit, profit_fmt)
                    
                    # Note nhỏ về chi phí ngoài
                    if total_inv > 0:
                        worksheet.write(sum_row+4, 0, f"(Bao gồm {format_vnd(total_inv)} hóa đơn phát sinh ngoài bảng kê)", workbook.add_format({'italic': True, 'font_size': 9, 'font_name': 'Times New Roman'}))
                    
                    # Column Widths
                    worksheet.set_column('A:A', 25)
                    worksheet.set_column('B:B', 40)
                    worksheet.set_column('C:C', 10)
                    worksheet.set_column('D:I', 15)

            except Exception as e:
                # If xlsxwriter fails, fall back to a simple CSV export
                st.error(f"⚠️ Lỗi khi tạo file Excel: {e}. Đã chuyển sang xuất file CSV.")
                buffer_act.seek(0)
                buffer_act.truncate()
                df_exp_act_filtered.to_csv(buffer_act, index=False, encoding='utf-8-sig')
                file_ext_act = "csv"
                mime_type_act = "text/csv"

            clean_t_name_act = re.sub(r'[\\/*?:"<>|]', "", tour_info_act['tour_name'] if tour_info_act else "Tour") # type: ignore
            st.download_button(
                f"📥 Tải Bảng Quyết Toán ({file_ext_act.upper()})", 
                data=buffer_act.getvalue(), 
                file_name=f"QuyetToan_{clean_t_name_act}.{file_ext_act}", 
                mime=mime_type_act, 
                use_container_width=True
            )

            def save_act_logic():
                run_query("DELETE FROM tour_items WHERE tour_id=? AND item_type='ACT'", (tour_id_act,), commit=True)
                data_to_insert = []
                query = """INSERT INTO tour_items (tour_id, item_type, category, description, unit, quantity, unit_price, total_amount, times)
                           VALUES (?, 'ACT', ?, ?, ?, ?, ?, ?, ?)"""

                for _, row in edited_act.iterrows():
                    if row['category'] or row['description']: # type: ignore
                        u_price = float(str(row['unit_price']).replace('.', '').replace(' VND', '').strip()) if row['unit_price'] else 0 # type: ignore
                        # Handle times safely
                        t_times = row.get('times', 1) # type: ignore
                        if pd.isna(t_times): t_times = 1
                        total_row = row['quantity'] * u_price * t_times # type: ignore

                        data_to_insert.append((
                            tour_id_act,
                            row['category'],
                            row['description'],
                            row['unit'],
                            row['quantity'],
                            u_price, # type: ignore
                            total_row,
                            t_times
                        ))

                if data_to_insert:
                    run_query_many(query, data_to_insert)

            if is_act_editable:
                if st.button("💾 LƯU QUYẾT TOÁN", type="primary", use_container_width=True):
                    save_act_logic()
                    st.success("Đã lưu quyết toán!"); time.sleep(1); st.rerun()
            else:
                st.info("🔒 Chỉ Admin mới được chỉnh sửa quyết toán.")
            
            st.divider()
            if st.button("✅ HOÀN THÀNH TOUR (Chuyển vào Lịch sử)", type="primary", use_container_width=True, key="complete_tour_btn"):
                run_query("UPDATE tours SET status='completed' WHERE id=?", (tour_id_act,), commit=True)
                st.success("Đã hoàn thành tour! Tour đã được chuyển sang tab Lịch sử.")
                time.sleep(1)

                st.rerun()

    # ---------------- TAB 4: LỊCH SỬ TOUR ----------------
    with tab_hist:
        st.subheader("📜 Lịch sử Tour đã hoàn thành")
        completed_tours = [t for t in all_tours if t['status'] == 'completed']
        
        if completed_tours:
            df_hist = pd.DataFrame([dict(t) for t in completed_tours])
            st.dataframe(
                df_hist[['tour_code', 'tour_name', 'start_date', 'end_date', 'guest_count', 'sale_name']],
                column_config={
                    "tour_code": "Mã Tour",
                    "tour_name": "Tên Tour",
                    "start_date": "Ngày đi",
                    "end_date": "Ngày về",
                    "guest_count": "Số khách",
                    "sale_name": "Sales"
                },
                use_container_width=True,
                hide_index=True
            )
            
            st.divider()
            st.write("🛠️ Thao tác:")
            hist_opts = {f"[{t['tour_code']}] {t['tour_name']}": t['id'] for t in completed_tours} # type: ignore
            sel_hist = st.selectbox("Chọn tour để xem lại hoặc mở lại:", list(hist_opts.keys()), key="sel_hist_tour")
            if sel_hist:
                tid_hist = hist_opts[sel_hist] # type: ignore
                if st.button("🔓 Mở lại Tour (Chuyển về Đang chạy)", key="reopen_tour_btn"):
                    run_query("UPDATE tours SET status='running' WHERE id=?", (tid_hist,), commit=True)
                    st.success("Đã mở lại tour! Kiểm tra lại bên tab Quyết toán.")
                    time.sleep(1)
                    st.rerun()
        else:
            st.info("Chưa có tour nào trong lịch sử.")

    # ---------------- TAB 3: TỔNG HỢP LỢI NHUẬN ----------------
    with tab_rpt:
        st.subheader("📈 Tổng Hợp Lợi Nhuận & Doanh Số")
        
        # Lọc theo thời gian
        rpt_df = pd.DataFrame([dict(r) for r in all_tours])
        if not rpt_df.empty:
            rpt_df['dt'] = pd.to_datetime(rpt_df['start_date'], format='%d/%m/%Y', errors='coerce') # type: ignore
            rpt_df = rpt_df.dropna(subset=['dt'])
            
            rpt_df['Month'] = rpt_df['dt'].apply(lambda x: x.strftime('%m/%Y'))
            rpt_df['Quarter'] = rpt_df['dt'].apply(lambda x: f"Q{(x.month-1)//3+1}/{x.year}")
            rpt_df['Year'] = rpt_df['dt'].apply(lambda x: x.strftime('%Y'))
            
            # --- PRE-FETCH DATA FOR PERFORMANCE ---
            all_items = run_query("SELECT tour_id, item_type, total_amount FROM tour_items")
            items_map = {} 
            if all_items:
                for item in all_items:
                    tid = item['tour_id']
                    itype = item['item_type']
                    amt = item['total_amount'] or 0
                    if tid not in items_map: items_map[tid] = {'EST': 0, 'ACT': 0}
                    items_map[tid][itype] += amt
            
            # Tính toán chỉ số cho từng tour
            results = []
            for _, t in rpt_df.iterrows():
                tid = t['id'] # type: ignore
                costs = items_map.get(tid, {'EST': 0, 'ACT': 0})
                est_cost = costs['EST']
                act_cost = costs['ACT']
                
                p_pct = t.get('est_profit_percent', 0) or 0
                t_pct = t.get('est_tax_percent', 0) or 0

                # Tính doanh thu (Ưu tiên giá chốt tay)
                final_price_manual = float(t.get('final_tour_price', 0) or 0)
                child_price_manual = float(t.get('child_price', 0) or 0)
                final_qty = float(t.get('final_qty', 0) or 0)
                child_qty = float(t.get('child_qty', 0) or 0)
                if final_qty == 0: final_qty = float(t.get('guest_count', 1))
                
                manual_revenue = (final_price_manual * final_qty) + (child_price_manual * child_qty)
                
                if manual_revenue > 0:
                    final_sale = manual_revenue
                else:
                    profit_est_val = est_cost * (p_pct/100)
                    final_sale = (est_cost + profit_est_val) * (1 + t_pct/100)

                net_revenue = final_sale / (1 + t_pct/100) if (1 + t_pct/100) != 0 else final_sale
                
                real_profit = net_revenue - act_cost
                
                results.append({
                    **t.to_dict(),
                    "Tên Đoàn": t['tour_name'], # type: ignore
                    "Sales": t['sale_name'], # type: ignore
                    "Ngày đi": t['start_date'], # type: ignore
                    "Doanh Thu Thuần": net_revenue,
                    "Chi Phí TT": act_cost,
                    "Lợi Nhuận TT": real_profit,
                })
            
            res_df = pd.DataFrame(results)

            # --- UI CONTROLS ---
            c_type, c_period, c_val = st.columns(3)
            report_type = c_type.selectbox("Loại báo cáo:", ["Theo Tour (Chi tiết)", "Theo Sales (Tổng hợp)"])
            period_type = c_period.selectbox("Xem theo:", ["Tháng", "Quý", "Năm"])
            
            period_options = []
            period_col = 'Month'
            if period_type == "Tháng":
                period_col = 'Month'
                period_options = sorted(res_df['Month'].unique(), reverse=True)
            elif period_type == "Quý":
                period_col = 'Quarter'
                period_options = sorted(res_df['Quarter'].unique(), reverse=True)
            else:
                period_col = 'Year'
                period_options = sorted(res_df['Year'].unique(), reverse=True)
            
            selected_period = c_val.selectbox("Chọn thời gian:", ["Tất cả"] + period_options)
            
            # Filter
            if selected_period != "Tất cả":
                res_df = res_df[res_df[period_col] == selected_period]
            
            if report_type == "Theo Tour (Chi tiết)":
                res_df['Tỷ suất LN'] = res_df.apply(lambda x: (x['Lợi Nhuận TT']/x['Doanh Thu Thuần']*100) if x['Doanh Thu Thuần'] else 0, axis=1)
                
                c_sum1, c_sum2 = st.columns(2)
                c_sum1.metric("Tổng Lợi Nhuận", format_vnd(res_df['Lợi Nhuận TT'].sum()))
                c_sum2.metric("Tổng Doanh Thu", format_vnd(res_df['Doanh Thu Thuần'].sum()))
                
                st.dataframe(
                    res_df[['Tên Đoàn', 'Sales', 'Ngày đi', 'Doanh Thu Thuần', 'Chi Phí TT', 'Lợi Nhuận TT', 'Tỷ suất LN']],
                    column_config={
                        "Doanh Thu Thuần": st.column_config.NumberColumn(format="%d VND"),
                        "Chi Phí TT": st.column_config.NumberColumn(format="%d VND"),
                        "Lợi Nhuận TT": st.column_config.NumberColumn(format="%d VND"),
                        "Tỷ suất LN": st.column_config.NumberColumn(format="%.2f %%"),
                    },
                    use_container_width=True,
                    hide_index=True
                )
                
                # Chuẩn bị dữ liệu xuất Excel
                df_export = res_df[['Tên Đoàn', 'Sales', 'Ngày đi', 'Doanh Thu Thuần', 'Chi Phí TT', 'Lợi Nhuận TT', 'Tỷ suất LN']].copy()
                file_name_rpt = f"BaoCao_LoiNhuan_Tour_{selected_period.replace('/', '_')}.xlsx"
            else: # Theo Sales
                df_sales = res_df.groupby('Sales').agg({
                    'Doanh Thu Thuần': 'sum',
                    'Chi Phí TT': 'sum',
                    'Lợi Nhuận TT': 'sum',
                    'id': 'count'
                }).reset_index()
                df_sales.columns = ["Nhân viên Sales", "Doanh Thu Thuần", "Chi Phí TT", "Lợi Nhuận TT", "Số Tour"]
                df_sales['Tỷ suất LN'] = df_sales.apply(lambda x: (x['Lợi Nhuận TT']/x['Doanh Thu Thuần']*100) if x['Doanh Thu Thuần'] else 0, axis=1)
                df_sales = df_sales.sort_values('Lợi Nhuận TT', ascending=False)
                
                st.markdown(f"##### 🏆 Bảng xếp hạng Sales ({selected_period})")
                if not df_sales.empty:
                    best = df_sales.iloc[0]
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Top Sales", best['Nhân viên Sales'], delta=format_vnd(best['Lợi Nhuận TT']))
                    c2.metric("Tổng Doanh Số", format_vnd(df_sales['Doanh Thu Thuần'].sum()))
                    c3.metric("Tổng Lợi Nhuận", format_vnd(df_sales['Lợi Nhuận TT'].sum()))
                    
                    st.bar_chart(df_sales.set_index("Nhân viên Sales")[['Doanh Thu Thuần', 'Lợi Nhuận TT']])
                
                st.dataframe(
                    df_sales,
                    column_config={
                        "Doanh Thu Thuần": st.column_config.NumberColumn(format="%d VND"),
                        "Chi Phí TT": st.column_config.NumberColumn(format="%d VND"),
                        "Lợi Nhuận TT": st.column_config.NumberColumn(format="%d VND"),
                        "Tỷ suất LN": st.column_config.NumberColumn(format="%.2f %%"),
                        "Số Tour": st.column_config.NumberColumn(format="%d"),
                    },
                    use_container_width=True,
                    hide_index=True
                )
                
                # Chuẩn bị dữ liệu xuất Excel
                df_export = df_sales.copy()
                file_name_rpt = f"BaoCao_DoanhSo_Sales_{selected_period.replace('/', '_')}.xlsx"

            # --- TÍNH NĂNG XUẤT EXCEL ---
            st.write("")
            buffer_rpt = io.BytesIO()
            with pd.ExcelWriter(buffer_rpt, engine='xlsxwriter') as writer:
                df_export.to_excel(writer, index=False, sheet_name='Report')
                workbook: Any = writer.book
                worksheet = writer.sheets['Report']
                
                # Định dạng
                header_fmt = workbook.add_format({'bold': True, 'fg_color': '#2E7D32', 'font_color': 'white', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                body_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                money_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_name': 'Times New Roman'})
                pct_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '0.00"%"', 'font_name': 'Times New Roman'})
                
                # Áp dụng định dạng header
                for col_num, value in enumerate(df_export.columns):
                    worksheet.write(0, col_num, value, header_fmt)
                
                # Áp dụng định dạng body
                for row_idx in range(len(df_export)):
                    for col_idx in range(len(df_export.columns)):
                        val = df_export.iloc[row_idx, col_idx]
                        col_name = df_export.columns[col_idx]
                        
                        fmt = body_fmt
                        if col_name in ['Doanh Thu Thuần', 'Chi Phí TT', 'Lợi Nhuận TT']: fmt = money_fmt
                        elif col_name == 'Tỷ suất LN': fmt = pct_fmt
                        
                        if pd.isna(val): val = ""
                        worksheet.write(row_idx + 1, col_idx, val, fmt)
                
                worksheet.set_column('A:A', 25)
                worksheet.set_column('B:Z', 18)

            st.download_button("📥 Xuất báo cáo Excel", buffer_rpt.getvalue(), file_name_rpt, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.info("Chưa có dữ liệu tour.")

def render_invoice_management():
    st.title("🧾 Quản Lý Hóa Đơn")
    
    tab_reverse, tab_profit = st.tabs(["🧮 Tính Hóa Đơn Ngược", "💰 Tính Lợi Nhuận"])
    
    with tab_reverse:
        st.subheader("🧾 Tính Hóa Đơn Ngược")
        st.markdown("""<div style="background-color: #e3f2fd; padding: 15px; border-radius: 10px; margin-bottom: 20px; border-left: 5px solid #2196F3;">
            <b>💡 Công cụ này giúp bạn:</b><br>
            • Tính ngược từ tổng tiền (đã bao gồm VAT và phí phục vụ) về giá gốc<br>
            • Hỗ trợ nhiều dòng phát sinh (mỗi dòng có % phí phục vụ riêng)<br>
            • Xuất Excel theo form bảng hóa đơn để gửi đối chiếu
            </div>""", unsafe_allow_html=True)

        if "inv_total" not in st.session_state: st.session_state.inv_total = "0"
        if "inv_main_qty" not in st.session_state: st.session_state.inv_main_qty = 1
        if "inv_service_pct_main" not in st.session_state: st.session_state.inv_service_pct_main = 5.0
        if "inv_vat_pct" not in st.session_state: st.session_state.inv_vat_pct = 8.0
        if "inv_extra_rows" not in st.session_state:
            # Khởi tạo DataFrame rỗng cho danh sách phát sinh với cột Phí phục vụ %
            st.session_state.inv_extra_rows = pd.DataFrame(columns=["description", "unit", "quantity", "unit_price", "service_pct", "vat_pct"])

        def fmt_inv_total():
            val = st.session_state.inv_total
            try:
                v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                st.session_state.inv_total = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
            except:
                pass

        col_t1, col_t2 = st.columns([2, 1])
        with col_t1:
            st.text_input(
                "💰 Tổng tiền thanh toán (đã gồm VAT + phí)",
                key="inv_total",
                on_change=fmt_inv_total,
                help="Nhập tổng tiền cuối cùng trên hóa đơn (VD: 8.994.000)"
            )
        with col_t2:
            main_qty = st.number_input(
                "🛏️ Số lượng DV chính",
                min_value=1,
                value=st.session_state.inv_main_qty,
                step=1,
                key="inv_main_qty",
                help="Ví dụ khách sạn 3 đêm thì nhập 3"
            )

        c_pct1, c_pct2 = st.columns(2)
        with c_pct1:
            service_pct_main = st.number_input(
                "🛎️ Phí phục vụ - Dịch vụ chính (%)",
                min_value=0.0,
                max_value=100.0,
                value=st.session_state.inv_service_pct_main,
                step=0.5,
                key="inv_service_pct_main"
            )
        with c_pct2:
            vat_pct = st.number_input(
                "📋 VAT (%)",
                min_value=0.0,
                max_value=100.0,
                value=st.session_state.inv_vat_pct,
                step=0.5,
                key="inv_vat_pct"
            )

        # --- BẢNG NHẬP CHI PHÍ PHÁT SINH ---
        st.write("")
        st.markdown("#### 📋 Danh sách Chi phí Phát sinh")
        st.caption("💡 Nhập các chi phí phát sinh kèm theo (nếu có). Mỗi dòng có thể có % Phí phục vụ riêng. Hệ thống sẽ tách VAT riêng cho nhóm này.")
        
        df_incurred = st.session_state.inv_extra_rows.copy()
        if df_incurred.empty:
            df_incurred = pd.DataFrame({
                "description": [""],
                "unit": [""],
                "quantity": [0.0],
                "unit_price": [0.0],
                "service_pct": [5.0],
                "vat_pct": [float(vat_pct)]
            })
        
        # Đảm bảo cột service_pct tồn tại
        if "service_pct" not in df_incurred.columns:
            df_incurred["service_pct"] = 5.0
        # Đảm bảo cột vat_pct tồn tại
        if "vat_pct" not in df_incurred.columns:
            df_incurred["vat_pct"] = float(vat_pct)
        
        # Format cột unit_price để hiển thị đẹp (100.000 VND)
        df_incurred_display = df_incurred.copy()
        df_incurred_display['unit_price'] = df_incurred_display['unit_price'].apply(
            lambda x: format_vnd(float(x) if x else 0) + " VND"
        )
        
        edited_incurred = st.data_editor(
            df_incurred_display,
            num_rows="dynamic",
            column_config={
                "description": st.column_config.TextColumn("Diễn giải", required=True, width="large"),
                "unit": st.column_config.TextColumn("Đơn vị", width="small"),
                "quantity": st.column_config.NumberColumn("Số lượng", min_value=0, format="%.0f"),
                "unit_price": st.column_config.TextColumn("Đơn giá", help="Nhập số tiền (VD: 100000 hoặc 100.000)"),
                "service_pct": st.column_config.NumberColumn("Phí DV (%)", min_value=0, max_value=100, format="%.1f %%", help="Phí phục vụ cho dòng này"),
                "vat_pct": st.column_config.NumberColumn("VAT (%)", min_value=0, max_value=100, format="%.1f %%", help="Ví dụ: rượu/nước ngọt có đường nhập 10"),
            },
            use_container_width=True,
            hide_index=True,
            key="inv_incurred_editor"
        )
        
        # Clean dữ liệu: Chuyển unit_price từ text về số
        edited_incurred_clean = edited_incurred.copy()
        edited_incurred_clean['unit_price'] = edited_incurred_clean['unit_price'].apply(
            lambda x: float(str(x).replace('.', '').replace(',', '').replace(' VND', '').strip()) if x else 0.0
        )
        if 'vat_pct' not in edited_incurred_clean.columns:
            edited_incurred_clean['vat_pct'] = float(vat_pct)
        edited_incurred_clean['vat_pct'] = pd.to_numeric(edited_incurred_clean['vat_pct'], errors='coerce').fillna(float(vat_pct))
        
        # Cập nhật vào session state (lưu dạng số)
        st.session_state.inv_extra_rows = edited_incurred_clean.copy()

        st.write("")
        if st.button("🧮 Tính toán", type="primary", use_container_width=True):
            try:
                total_str = st.session_state.inv_total.replace('.', '').replace(',', '').replace(' VND', '').strip()
                total_amount = float(total_str) if total_str else 0.0
                df_inv = pd.DataFrame(columns=["group", "amount", "vat_pct"])

                if total_amount <= 0:
                    st.error("⚠️ Vui lòng nhập tổng tiền hợp lệ!")
                else:
                    main_service_rate = 1 + (service_pct_main / 100)

                    # --- BƯỚC 1: Tính tổng chi phí phát sinh (chưa gồm phí DV và VAT) ---
                    df_extra = edited_incurred_clean.copy()
                    df_extra = df_extra[df_extra['description'].str.strip().astype(bool)] # Lọc dòng rỗng
                    df_extra['quantity'] = pd.to_numeric(df_extra['quantity'], errors='coerce').fillna(0)
                    df_extra['unit_price'] = pd.to_numeric(df_extra['unit_price'], errors='coerce').fillna(0)
                    df_extra['service_pct'] = pd.to_numeric(df_extra['service_pct'], errors='coerce').fillna(5.0)
                    df_extra['vat_pct'] = pd.to_numeric(df_extra['vat_pct'], errors='coerce').fillna(float(vat_pct))
                    df_extra['subtotal'] = df_extra['quantity'] * df_extra['unit_price']
                    
                    # Tính phí dịch vụ cho từng dòng
                    df_extra['service_fee'] = df_extra['subtotal'] * (df_extra['service_pct'] / 100.0)
                    df_extra['subtotal_with_service'] = df_extra['subtotal'] + df_extra['service_fee']
                    # Với VAT hỗn hợp, cần quy đổi từng dòng theo VAT riêng
                    df_extra['total_with_vat'] = df_extra['subtotal_with_service'] * (1 + df_extra['vat_pct'] / 100.0)
                    
                    total_incurred_base = df_extra['subtotal'].sum()
                    total_incurred_service_fee = df_extra['service_fee'].sum()
                    total_incurred_with_service = df_extra['subtotal_with_service'].sum()
                    total_incurred_with_vat = df_extra['total_with_vat'].sum()
                    
                    # --- BƯỚC 2: Tính ngược từ tổng tiền về 2 nhóm ---
                    # Do phát sinh có thể có nhiều mức VAT, tính trực tiếp trên tổng sau VAT
                    main_vat_rate = 1 + (vat_pct / 100.0)
                    total_main_with_service = (total_amount - total_incurred_with_vat) / main_vat_rate
                    if total_main_with_service < 0:
                        st.error("⚠️ Tổng phát sinh (gồm VAT) đang lớn hơn tổng thanh toán. Vui lòng kiểm tra lại dữ liệu.")
                        st.stop()

                    base_price_main = total_main_with_service / main_service_rate
                    main_service_fee = base_price_main * (service_pct_main / 100.0)
                    main_unit_price = base_price_main / float(main_qty) if float(main_qty) > 0 else base_price_main

                    # --- BƯỚC 3: Xây dựng bảng hóa đơn ---
                    invoice_rows = []
                    
                    # NHÓM 1: Dịch vụ chính
                    invoice_rows.append({
                        "name": "Dịch vụ chính",
                        "unit": "Đêm",
                        "qty": float(main_qty),
                        "unit_price": main_unit_price,
                        "amount": base_price_main,
                        "vat_pct": float(vat_pct),
                        "group": "main"
                    })
                    if main_service_fee > 0:
                        invoice_rows.append({
                            "name": f"Phí dịch vụ {service_pct_main:g}%",
                            "unit": "Lần",
                            "qty": 1.0,
                            "unit_price": main_service_fee,
                            "amount": main_service_fee,
                            "vat_pct": float(vat_pct),
                            "group": "main"
                        })
                    
                    # NHÓM 2: Chi phí phát sinh (hiển thị từng dòng + phí DV riêng)
                    for idx, row in df_extra.iterrows():
                        if row['subtotal'] > 0:
                            invoice_rows.append({
                                "name": row['description'],
                                "unit": row['unit'],
                                "qty": float(row['quantity']),
                                "unit_price": float(row['unit_price']),
                                "amount": float(row['subtotal']),
                                "vat_pct": float(row['vat_pct']),
                                "group": "incurred"
                            })
                            # Thêm dòng phí DV cho mỗi chi phí phát sinh (nếu có)
                            if row['service_fee'] > 0:
                                invoice_rows.append({
                                    "name": f"  ↳ Phí DV {row['service_pct']:g}%",
                                    "unit": "Lần",
                                    "qty": 1.0,
                                    "unit_price": float(row['service_fee']),
                                    "amount": float(row['service_fee']),
                                    "vat_pct": float(row['vat_pct']),
                                    "group": "incurred"
                                })

                    # --- BƯỚC 4: Tính VAT riêng cho từng nhóm ---
                    df_inv = pd.DataFrame(invoice_rows)
                
                    def calc_vat_by_group(df, group_name):
                        """Tính VAT cho một nhóm (theo VAT từng dòng)."""
                        group_df = df[df['group'] == group_name].copy()
                        group_subtotal = group_df['amount'].sum()
                        group_df['vat_amount'] = group_df['amount'] * (group_df['vat_pct'] / 100.0)
                        group_vat = group_df['vat_amount'].sum()
                        group_df['total_amount'] = group_df['amount'] + group_df['vat_amount']
                        return group_df, group_subtotal, group_vat
                
                    df_main, subtotal_main, vat_main = calc_vat_by_group(df_inv, "main")
                    df_incurred_final, subtotal_incurred, vat_incurred = calc_vat_by_group(df_inv, "incurred")
                
                    # Gộp lại
                    df_result = pd.concat([df_main, df_incurred_final], ignore_index=True)
                
                    # Tổng cộng
                    sub_total_check = df_result["amount"].sum()
                    vat_total_check = df_result["vat_amount"].sum()
                    grand_total_check = df_result["total_amount"].sum()

                    st.session_state.inv_last_result = {
                        "df": df_result.copy(),
                        "sub_total": sub_total_check,
                        "vat_total": vat_total_check,
                        "grand_total": grand_total_check,
                        "vat_pct": vat_pct,
                        # Thêm thông tin phân nhóm
                        "subtotal_main": subtotal_main,
                        "vat_main": vat_main,
                        "subtotal_incurred": subtotal_incurred,
                        "vat_incurred": vat_incurred,
                    }

                    st.markdown("### 📊 Kết Quả Tính Toán")
                
                    # --- HIỂN THỊ THEO NHÓM ---
                    st.markdown("#### 🏨 Nhóm 1: Dịch vụ Chính")
                    view_df_main = df_main.copy().reset_index(drop=True)
                    view_df_main.index = view_df_main.index + 1
                    view_df_main["Số lượng"] = view_df_main["qty"]
                    view_df_main["Đơn giá"] = view_df_main["unit_price"].apply(lambda x: format_vnd(x))
                    view_df_main["Thành tiền"] = view_df_main["amount"].apply(lambda x: format_vnd(x))
                    view_df_main["% VAT"] = view_df_main["vat_pct"].apply(lambda x: f"{x:g}%")
                    view_df_main["Tiền thuế"] = view_df_main["vat_amount"].apply(lambda x: format_vnd(x))
                    view_df_main["Tổng cộng"] = view_df_main["total_amount"].apply(lambda x: format_vnd(x))
                    st.dataframe(
                        view_df_main[["name", "unit", "Số lượng", "Đơn giá", "Thành tiền", "% VAT", "Tiền thuế", "Tổng cộng"]],
                        column_config={
                            "name": "Tên hàng hóa, dịch vụ",
                            "unit": "Đơn vị tính",
                        },
                        use_container_width=True
                    )
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Tổng trước VAT", format_vnd(subtotal_main) + " VND")
                    c2.metric("VAT", format_vnd(vat_main) + " VND")
                    c3.metric("Tổng sau VAT", format_vnd(subtotal_main + vat_main) + " VND")
                
                    # Hiển thị nhóm phát sinh nếu có
                    if not df_incurred_final.empty:
                        st.write("")
                        st.markdown("#### 💰 Nhóm 2: Chi Phí Phát Sinh")
                        view_df_incurred = df_incurred_final.copy().reset_index(drop=True)
                        view_df_incurred.index = view_df_incurred.index + 1
                        view_df_incurred["Số lượng"] = view_df_incurred["qty"]
                        view_df_incurred["Đơn giá"] = view_df_incurred["unit_price"].apply(lambda x: format_vnd(x))
                        view_df_incurred["Thành tiền"] = view_df_incurred["amount"].apply(lambda x: format_vnd(x))
                        view_df_incurred["% VAT"] = view_df_incurred["vat_pct"].apply(lambda x: f"{x:g}%")
                        view_df_incurred["Tiền thuế"] = view_df_incurred["vat_amount"].apply(lambda x: format_vnd(x))
                        view_df_incurred["Tổng cộng"] = view_df_incurred["total_amount"].apply(lambda x: format_vnd(x))
                        st.dataframe(
                            view_df_incurred[["name", "unit", "Số lượng", "Đơn giá", "Thành tiền", "% VAT", "Tiền thuế", "Tổng cộng"]],
                            column_config={
                                "name": "Tên hàng hóa, dịch vụ",
                                "unit": "Đơn vị tính",
                            },
                            use_container_width=True
                        )
                        c1, c2, c3 = st.columns(3)
                        c1.metric("Tổng trước VAT", format_vnd(subtotal_incurred) + " VND")
                        c2.metric("VAT", format_vnd(vat_incurred) + " VND")
                        c3.metric("Tổng sau VAT", format_vnd(subtotal_incurred + vat_incurred) + " VND")
                
                    # Tổng kết
                    st.write("")
                    st.markdown("---")
                    m1, m2, m3 = st.columns(3)
                    m1.metric("💵 TỔNG TRƯỚC VAT", format_vnd(sub_total_check) + " VND")
                    m2.metric("📋 TỔNG VAT", format_vnd(vat_total_check) + " VND")
                    m3.metric("💰 TỔNG THANH TOÁN", format_vnd(grand_total_check) + " VND")

                    st.info("💡 Không áp dụng làm tròn hóa đơn: hệ thống giữ nguyên kết quả tính theo công thức hiện tại.")

            except Exception as e:
                st.error(f"❌ Lỗi tính toán: {str(e)}")
                import traceback
                st.code(traceback.format_exc())

        if "inv_last_result" in st.session_state:
            st.write("")
            if st.button("📥 Xuất Excel hóa đơn", use_container_width=True):
                try:
                    rs = st.session_state.inv_last_result
                    df_exp = rs["df"].copy()
                    buffer_inv = io.BytesIO()

                    with pd.ExcelWriter(buffer_inv, engine='xlsxwriter') as writer:
                        df_x = df_exp.copy()
                        df_x.insert(0, "STT", range(1, len(df_x) + 1))
                        df_x = df_x[["STT", "name", "unit", "qty", "unit_price", "amount", "vat_pct", "vat_amount", "total_amount"]]
                        df_x.columns = ["STT", "Tên hàng hóa, dịch vụ", "Đơn vị tính", "Số lượng", "Đơn giá", "Thành tiền", "% VAT", "Tiền thuế", "Tổng cộng"]
                        df_x.to_excel(writer, sheet_name="HoaDon", index=False, startrow=4)

                        wb: Any = writer.book
                        ws = writer.sheets["HoaDon"]

                        title_fmt = wb.add_format({'bold': True, 'font_size': 14, 'align': 'center', 'font_name': 'Times New Roman'})
                        head_fmt = wb.add_format({'bold': True, 'border': 1, 'align': 'center', 'bg_color': '#E3F2FD', 'font_name': 'Times New Roman'})
                        text_fmt = wb.add_format({'border': 1, 'font_name': 'Times New Roman'})
                        num_fmt = wb.add_format({'border': 1, 'num_format': '#,##0.00', 'font_name': 'Times New Roman'})
                        total_fmt = wb.add_format({'bold': True, 'border': 1, 'num_format': '#,##0.00', 'bg_color': '#E8F5E9', 'font_name': 'Times New Roman'})

                        ws.merge_range('A1:I1', 'BẢNG TÍNH HÓA ĐƠN', title_fmt)
                        ws.write('A2', 'Tổng trước VAT', text_fmt)
                        ws.write('B2', float(rs['sub_total']), num_fmt)
                        ws.write('D2', 'Tổng VAT', text_fmt)
                        ws.write('E2', float(rs['vat_total']), num_fmt)
                        ws.write('G2', 'Tổng thanh toán', text_fmt)
                        ws.write('H2', float(rs['grand_total']), total_fmt)

                        for c in range(9):
                            ws.write(4, c, df_x.columns[c], head_fmt)

                        for r in range(len(df_x)):
                            excel_row = 5 + r
                            ws.write_number(excel_row, 0, float(df_x.iloc[r, 0]), text_fmt)
                            ws.write(excel_row, 1, df_x.iloc[r, 1], text_fmt)
                            ws.write(excel_row, 2, df_x.iloc[r, 2], text_fmt)
                            ws.write_number(excel_row, 3, float(df_x.iloc[r, 3]), num_fmt)
                            ws.write_number(excel_row, 4, float(df_x.iloc[r, 4]), num_fmt)
                            ws.write_number(excel_row, 5, float(df_x.iloc[r, 5]), num_fmt)
                            ws.write(excel_row, 6, f"{float(df_x.iloc[r, 6]):g}%", text_fmt)
                            ws.write_number(excel_row, 7, float(df_x.iloc[r, 7]), num_fmt)
                            ws.write_number(excel_row, 8, float(df_x.iloc[r, 8]), num_fmt)

                        sum_row = 5 + len(df_x)
                        ws.merge_range(sum_row, 0, sum_row, 4, "TỔNG CỘNG", total_fmt)
                        ws.write_number(sum_row, 5, float(rs['sub_total']), total_fmt)
                        ws.write(sum_row, 6, "", total_fmt)
                        ws.write_number(sum_row, 7, float(rs['vat_total']), total_fmt)
                        ws.write_number(sum_row, 8, float(rs['grand_total']), total_fmt)

                        ws.set_column('A:A', 6)
                        ws.set_column('B:B', 42)
                        ws.set_column('C:C', 12)
                        ws.set_column('D:I', 16)

                    st.download_button(
                        "⬇️ Tải file hóa đơn (.xlsx)",
                        data=buffer_inv.getvalue(),
                        file_name="HoaDon_TinhNguoc.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"❌ Lỗi xuất Excel: {str(e)}")

        with st.expander("📖 Giải thích công thức tính toán"):
                st.markdown("""
            ### Công thức đang dùng (tách VAT theo nhóm)
    
            **BƯỚC 1: Tính tổng trước VAT**
            ```
            Tổng trước VAT = Tổng thanh toán / (1 + VAT%)
            ```
    
            **BƯỚC 2: Tính chi phí phát sinh (mỗi dòng có % phí DV riêng)**
            ```
            Với mỗi dòng phát sinh i:
              - Thành tiền_i = Đơn giá_i × Số lượng_i
              - Phí DV_i = Thành tiền_i × Phí DV %_i
              - Tổng có phí_i = Thành tiền_i + Phí DV_i
                            - VAT_i = Tổng có phí_i × VAT %_i
                            - Tổng sau VAT_i = Tổng có phí_i + VAT_i
            
            Tổng phát sinh gốc = Σ(Thành tiền_i)
            Tổng phí DV phát sinh = Σ(Phí DV_i)
            Tổng phát sinh có phí = Σ(Tổng có phí_i)
                        Tổng phát sinh sau VAT = Σ(Tổng sau VAT_i)
            ```
    
            **BƯỚC 3: Tính dịch vụ chính (có phí DV)**
            ```
                        Tổng DV chính có phí = (Tổng thanh toán - Tổng phát sinh sau VAT) / (1 + VAT dịch vụ chính%)
            Tổng DV chính gốc = Tổng DV chính có phí / (1 + Phí DV chính%)
            Đơn giá DV chính = Tổng DV chính gốc / Số lượng
            ```
    
            **BƯỚC 4: Tính VAT riêng cho từng nhóm**
            ```
            VAT Nhóm 1 (DV chính) = (DV chính gốc + Phí DV chính) × VAT%
                        VAT Nhóm 2 (Phát sinh) = Σ(Tổng có phí_i × VAT %_i)
            
            Tổng VAT = VAT Nhóm 1 + VAT Nhóm 2
            ```
    
            **💡 Lợi ích:**  
            - Mỗi dòng phát sinh có thể có % phí phục vụ riêng (linh hoạt)
            - Mỗi dòng phát sinh có thể có VAT riêng (vd 8%, 10%)
            - Tách riêng VAT cho dịch vụ chính và chi phí phát sinh  
            - Dễ đối chiếu với hóa đơn có nhiều nhóm dịch vụ  
            - Phù hợp với quy định thuế hiện hành
            """)
    
    with tab_profit:
        st.subheader("💰 Tính Lợi Nhuận")
        st.markdown("""<div style="background-color: #e8f5e9; padding: 15px; border-radius: 10px; margin-bottom: 20px; border-left: 5px solid #4CAF50;">
            <b>💡 Công cụ này giúp bạn:</b><br>
            • Theo dõi các hóa đơn đầu ra (thu) và hóa đơn đầu vào (chi)<br>
            • Tính toán lợi nhuận tự động từ hiệu số thu - chi<br>
            • Ghi chú rõ ràng mã/số hóa đơn để dễ đối chiếu<br>
            • Quản lý theo tháng/năm để dễ dàng theo dõi
        </div>""", unsafe_allow_html=True)
    
        # === KHỞI TẠO SESSION STATE ===
        if "profit_output_invoices" not in st.session_state:
            st.session_state.profit_output_invoices = pd.DataFrame(columns=["period", "project", "invoice_no", "description", "amount"])
        if "profit_input_invoices" not in st.session_state:
            st.session_state.profit_input_invoices = pd.DataFrame(columns=["period", "project", "invoice_no", "description", "amount"])
        
        # Migration: tách cột Số HĐ và Diễn giải
        if 'invoice_no' not in st.session_state.profit_output_invoices.columns:
            st.session_state.profit_output_invoices['invoice_no'] = ""
        if 'invoice_no' not in st.session_state.profit_input_invoices.columns:
            st.session_state.profit_input_invoices['invoice_no'] = ""
        if 'project' not in st.session_state.profit_output_invoices.columns:
            st.session_state.profit_output_invoices['project'] = "Tổng tháng"
        if 'project' not in st.session_state.profit_input_invoices.columns:
            st.session_state.profit_input_invoices['project'] = "Tổng tháng"
        
        # === CHỌN THÁNG/NĂM ===
        from datetime import datetime
        current_date = datetime.now()
        
        col_m, col_y = st.columns(2)
        with col_m:
            sel_month = st.selectbox(
                "📅 Tháng",
                options=list(range(1, 13)),
                index=current_date.month - 1,
                format_func=lambda x: f"Tháng {x:02d}",
                key="profit_sel_month"
            )
        with col_y:
            year_opts = list(range(current_date.year - 5, current_date.year + 2))
            sel_year = st.selectbox(
                "📅 Năm",
                options=year_opts,
                index=year_opts.index(current_date.year),
                key="profit_sel_year"
            )
        
        selected_period = f"{sel_month:02d}/{sel_year}"
        st.info(f"🗓️ Đang làm việc với tháng: **{selected_period}**")

        # Hiển thị nhanh danh sách dự án đang có trong tháng chọn
        projects_in_month_top = set()
        if not st.session_state.profit_output_invoices.empty:
            projects_in_month_top.update(
                st.session_state.profit_output_invoices.loc[
                    st.session_state.profit_output_invoices['period'] == selected_period, 'project'
                ].astype(str).str.strip().tolist()
            )
        if not st.session_state.profit_input_invoices.empty:
            projects_in_month_top.update(
                st.session_state.profit_input_invoices.loc[
                    st.session_state.profit_input_invoices['period'] == selected_period, 'project'
                ].astype(str).str.strip().tolist()
            )
        projects_in_month_top = sorted([p for p in projects_in_month_top if p])
        if projects_in_month_top:
            st.caption(f"📌 Dự án trong tháng {selected_period}: {', '.join(projects_in_month_top)}")
        else:
            st.caption(f"📌 Dự án trong tháng {selected_period}: Chưa có")

        # Đặt tên báo cáo ở đầu tab
        if "profit_report_name" not in st.session_state:
            st.session_state.profit_report_name = "Báo cáo lợi nhuận"
        if "profit_meta_edit_mode" not in st.session_state:
            st.session_state.profit_meta_edit_mode = False
        report_name = st.text_input(
            "🏷️ Tên báo cáo",
            key="profit_report_name",
            help="Tên hiển thị trên file Excel",
            disabled=not st.session_state.profit_meta_edit_mode
        )

        # Tạo nhanh dự án trong tháng đang chọn
        if "profit_new_project_name" not in st.session_state:
            st.session_state.profit_new_project_name = ""
        if "profit_new_project_name_clear" not in st.session_state:
            st.session_state.profit_new_project_name_clear = False
        if st.session_state.profit_new_project_name_clear:
            st.session_state.profit_new_project_name = ""
            st.session_state.profit_new_project_name_clear = False

        col_add_proj_name, col_add_proj_btn = st.columns([4, 1])
        with col_add_proj_name:
            st.text_input(
                "📌 Tên dự án cần thêm",
                key="profit_new_project_name",
                placeholder="VD: Tour Hạ Long 20/08 hoặc Booking Lẻ Anh A"
            )
        with col_add_proj_btn:
            st.write("")
            if st.button("➕ Thêm dự án", use_container_width=True, key="profit_add_project_btn"):
                new_project = str(st.session_state.profit_new_project_name or "").strip()
                if not new_project:
                    st.warning("Vui lòng nhập tên dự án trước khi thêm.")
                else:
                    def append_project_seed(df_src: pd.DataFrame) -> tuple[pd.DataFrame, bool]:
                        df_work = df_src.copy()
                        if df_work.empty:
                            df_work = pd.DataFrame(columns=["period", "project", "invoice_no", "description", "amount"])
                        if 'project' not in df_work.columns:
                            df_work['project'] = "Tổng tháng"
                        if 'invoice_no' not in df_work.columns:
                            df_work['invoice_no'] = ""

                        mask = (
                            df_work['period'].astype(str).eq(selected_period)
                            & df_work['project'].astype(str).str.strip().eq(new_project)
                        )
                        if mask.any():
                            return df_work, False

                        new_row = pd.DataFrame([
                            {"period": selected_period, "project": new_project, "invoice_no": "", "description": "", "amount": 0.0}
                        ])
                        return pd.concat([df_work, new_row], ignore_index=True), True

                    out_full_new, out_added = append_project_seed(st.session_state.profit_output_invoices)
                    in_full_new, in_added = append_project_seed(st.session_state.profit_input_invoices)

                    st.session_state.profit_output_invoices = out_full_new
                    st.session_state.profit_input_invoices = in_full_new

                    if out_added or in_added:
                        if "profit_selected_project_by_period" not in st.session_state:
                            st.session_state.profit_selected_project_by_period = {}
                        st.session_state.profit_selected_project_by_period[selected_period] = new_project
                        st.session_state.profit_current_scope_output = None
                        st.session_state.profit_current_scope_input = None
                        st.session_state.profit_new_project_name_clear = True
                        st.success(f"Đã thêm dự án '{new_project}' cho tháng {selected_period}.")
                        st.rerun()
                    else:
                        st.info("Dự án này đã tồn tại trong tháng hiện tại.")

        # Chọn dự án đang nhập (mỗi lần nhập là cho 1 dự án)
        if "profit_selected_project_by_period" not in st.session_state:
            st.session_state.profit_selected_project_by_period = {}

        period_projects = set()
        if not st.session_state.profit_output_invoices.empty:
            period_projects.update(
                st.session_state.profit_output_invoices.loc[
                    st.session_state.profit_output_invoices['period'] == selected_period, 'project'
                ].astype(str).str.strip().tolist()
            )
        if not st.session_state.profit_input_invoices.empty:
            period_projects.update(
                st.session_state.profit_input_invoices.loc[
                    st.session_state.profit_input_invoices['period'] == selected_period, 'project'
                ].astype(str).str.strip().tolist()
            )
        period_projects = {p for p in period_projects if p}

        if not period_projects:
            period_projects = {"Dự án 1"}

        project_options_sorted = sorted(period_projects)
        remembered_project = st.session_state.profit_selected_project_by_period.get(selected_period)
        if remembered_project not in project_options_sorted:
            remembered_project = project_options_sorted[0]
            st.session_state.profit_selected_project_by_period[selected_period] = remembered_project

        selected_project = st.selectbox(
            "🧩 Dự án đang nhập",
            options=project_options_sorted,
            index=project_options_sorted.index(remembered_project),
            key=f"profit_selected_project_{selected_period}"
        )
        st.session_state.profit_selected_project_by_period[selected_period] = selected_project
        st.caption("Dữ liệu bạn nhập bên dưới sẽ thuộc đúng dự án này. Tổng kết tháng sẽ tự cộng tất cả dự án.")

        def is_profit_row_valid(df: pd.DataFrame) -> pd.Series:
            """Giữ dòng có mô tả hoặc có số tiền khác 0 để tránh mất dữ liệu khi nhập từng bước."""
            if df.empty:
                return pd.Series([], dtype=bool)
            inv_ok = df['invoice_no'].astype(str).str.strip().ne("") if 'invoice_no' in df.columns else pd.Series([False] * len(df), index=df.index)
            desc_ok = df['description'].astype(str).str.strip().ne("")
            amt = pd.to_numeric(df['amount'], errors='coerce').fillna(0)
            amt_ok = amt.ne(0)
            return inv_ok | desc_ok | amt_ok

        def clean_profit_amount(x):
            if isinstance(x, (int, float)):
                return float(x)
            if x is None:
                return 0.0
            raw = str(x).strip()
            if not raw:
                return 0.0
            try:
                return float(raw.replace('.', '').replace(',', '').replace(' VND', '').strip())
            except Exception:
                return 0.0

        if "profit_current_scope_output" not in st.session_state:
            st.session_state.profit_current_scope_output = None
        if "profit_current_scope_input" not in st.session_state:
            st.session_state.profit_current_scope_input = None
        if "profit_output_temp" not in st.session_state:
            st.session_state.profit_output_temp = pd.DataFrame(columns=["period", "project", "invoice_no", "description", "amount"])
        if "profit_input_temp" not in st.session_state:
            st.session_state.profit_input_temp = pd.DataFrame(columns=["period", "project", "invoice_no", "description", "amount"])
        if 'invoice_no' not in st.session_state.profit_output_temp.columns:
            st.session_state.profit_output_temp['invoice_no'] = ""
        if 'invoice_no' not in st.session_state.profit_input_temp.columns:
            st.session_state.profit_input_temp['invoice_no'] = ""

        selected_scope = f"{selected_period}::{selected_project}"

        # Đồng bộ temp theo period + project
        if st.session_state.profit_current_scope_output != selected_scope:
            df_out_seed = st.session_state.profit_output_invoices.copy()
            df_out_seed = (
                df_out_seed[
                    (df_out_seed['period'] == selected_period)
                    & (df_out_seed['project'].astype(str).str.strip() == selected_project)
                ]
                if not df_out_seed.empty
                else pd.DataFrame(columns=["period", "project", "invoice_no", "description", "amount"])
            )
            if df_out_seed.empty:
                df_out_seed = pd.DataFrame([{"period": selected_period, "project": selected_project, "invoice_no": "", "description": "", "amount": 0.0}])
            st.session_state.profit_output_temp = df_out_seed[['period', 'project', 'invoice_no', 'description', 'amount']].reset_index(drop=True)
            st.session_state.profit_current_scope_output = selected_scope

        if st.session_state.profit_current_scope_input != selected_scope:
            df_in_seed = st.session_state.profit_input_invoices.copy()
            df_in_seed = (
                df_in_seed[
                    (df_in_seed['period'] == selected_period)
                    & (df_in_seed['project'].astype(str).str.strip() == selected_project)
                ]
                if not df_in_seed.empty
                else pd.DataFrame(columns=["period", "project", "invoice_no", "description", "amount"])
            )
            if df_in_seed.empty:
                df_in_seed = pd.DataFrame([{"period": selected_period, "project": selected_project, "invoice_no": "", "description": "", "amount": 0.0}])
            st.session_state.profit_input_temp = df_in_seed[['period', 'project', 'invoice_no', 'description', 'amount']].reset_index(drop=True)
            st.session_state.profit_current_scope_input = selected_scope
        
        st.markdown("---")
        
        # === BẢNG HÓA ĐƠN ĐẦU RA (THU) ===
        st.markdown("### 📤 Hóa Đơn Đầu Ra (Thu)")
        st.caption("💡 Sau khi nhập số tiền, nhấn **Enter** hoặc click ra ngoài ô. Số sẽ tự động định dạng thành 1.000.000")
        
        # Dùng temp state theo period để tránh mất dữ liệu khi nhập lần đầu
        df_period_output = st.session_state.profit_output_temp.copy()
        if 'project' not in df_period_output.columns:
            df_period_output['project'] = "Tổng tháng"
        if 'invoice_no' not in df_period_output.columns:
            df_period_output['invoice_no'] = ""
        
        # Format cho display
        df_display_output = df_period_output.copy().reset_index(drop=True)
        df_display_output['amount_display'] = df_display_output['amount'].apply(
            lambda x: format_vnd(float(x) if x else 0)
        )
        
        edited_output = st.data_editor(
            df_display_output[['period', 'invoice_no', 'description', 'amount_display']].reset_index(drop=True),
            num_rows="dynamic",
            column_config={
                "period": st.column_config.TextColumn("Tháng", default=selected_period, disabled=True, width="small"),
                "invoice_no": st.column_config.TextColumn("Số HĐ", required=False, width="medium", help="VD: HD001"),
                "description": st.column_config.TextColumn("Diễn giải", required=False, width="large", help="VD: Bán hàng cho khách A"),
                "amount_display": st.column_config.TextColumn("Số tiền (VND)", width="medium", help="Nhập: 5000000 hoặc 5.000.000")
            },
            use_container_width=True,
            hide_index=True,
            key="editor_output_profit"
        )
        
        # Clean data từ editor
        edited_output['amount'] = edited_output['amount_display'].apply(clean_profit_amount)
        edited_output['project'] = selected_project
        edited_output = edited_output[['period', 'project', 'invoice_no', 'description', 'amount']]
        edited_output['period'] = edited_output['period'].fillna(selected_period)
        edited_output['project'] = selected_project
        edited_output['invoice_no'] = edited_output['invoice_no'].fillna('').astype(str).str.strip()

        # So sánh với temp cũ, nếu đổi thì lưu và rerun như tab Dự Toán
        old_out = st.session_state.profit_output_temp[['period', 'project', 'invoice_no', 'description', 'amount']].reset_index(drop=True).fillna('')
        new_out = edited_output[['period', 'project', 'invoice_no', 'description', 'amount']].reset_index(drop=True).fillna('')
        if len(old_out) != len(new_out) or not old_out.equals(new_out):
            st.session_state.profit_output_temp = edited_output.copy()
            df_other_output = st.session_state.profit_output_invoices[
                ~(
                    (st.session_state.profit_output_invoices['period'] == selected_period)
                    & (st.session_state.profit_output_invoices['project'].astype(str).str.strip() == selected_project)
                )
            ]
            st.session_state.profit_output_invoices = pd.concat([df_other_output, edited_output], ignore_index=True)
            st.rerun()

        # Lọc cho mục đích tính toán
        edited_output_clean = edited_output[is_profit_row_valid(edited_output)].copy()
        
        # Tính tổng thu dự án đang nhập (chỉ tính dòng có nội dung)
        total_revenue_project = edited_output_clean['amount'].sum()
        
        col_refresh1, col_metric1 = st.columns([1, 5])
        with col_refresh1:
            if st.button("🔄", key="refresh_output", help="Làm mới để định dạng số tiền"):
                st.rerun()
        with col_metric1:
            st.metric("💰 Tổng Thu (Dự án đang chọn)", format_vnd(total_revenue_project) + " VND")
        
        st.markdown("---")
        
        # === BẢNG HÓA ĐƠN ĐẦU VÀO (CHI) ===
        st.markdown("### 📥 Hóa Đơn Đầu Vào (Chi)")
        st.caption("💡 Sau khi nhập số tiền, nhấn **Enter** hoặc click ra ngoài ô. Số sẽ tự động định dạng thành 1.000.000")
        
        # Dùng temp state theo period để tránh mất dữ liệu khi nhập lần đầu
        df_period_input = st.session_state.profit_input_temp.copy()
        if 'project' not in df_period_input.columns:
            df_period_input['project'] = "Tổng tháng"
        if 'invoice_no' not in df_period_input.columns:
            df_period_input['invoice_no'] = ""
        
        # Format cho display
        df_display_input = df_period_input.copy().reset_index(drop=True)
        df_display_input['amount_display'] = df_display_input['amount'].apply(
            lambda x: format_vnd(float(x) if x else 0)
        )
        
        edited_input = st.data_editor(
            df_display_input[['period', 'invoice_no', 'description', 'amount_display']].reset_index(drop=True),
            num_rows="dynamic",
            column_config={
                "period": st.column_config.TextColumn("Tháng", default=selected_period, disabled=True, width="small"),
                "invoice_no": st.column_config.TextColumn("Số HĐ", required=False, width="medium", help="VD: PO001"),
                "description": st.column_config.TextColumn("Diễn giải", required=False, width="large", help="VD: Mua hàng từ NCC B"),
                "amount_display": st.column_config.TextColumn("Số tiền (VND)", width="medium", help="Nhập: 3000000 hoặc 3.000.000")
            },
            use_container_width=True,
            hide_index=True,
            key="editor_input_profit"
        )
        
        # Clean data từ editor
        edited_input['amount'] = edited_input['amount_display'].apply(clean_profit_amount)
        edited_input['project'] = selected_project
        edited_input = edited_input[['period', 'project', 'invoice_no', 'description', 'amount']]
        edited_input['period'] = edited_input['period'].fillna(selected_period)
        edited_input['project'] = selected_project
        edited_input['invoice_no'] = edited_input['invoice_no'].fillna('').astype(str).str.strip()

        # So sánh với temp cũ, nếu đổi thì lưu và rerun như tab Dự Toán
        old_in = st.session_state.profit_input_temp[['period', 'project', 'invoice_no', 'description', 'amount']].reset_index(drop=True).fillna('')
        new_in = edited_input[['period', 'project', 'invoice_no', 'description', 'amount']].reset_index(drop=True).fillna('')
        if len(old_in) != len(new_in) or not old_in.equals(new_in):
            st.session_state.profit_input_temp = edited_input.copy()
            df_other_input = st.session_state.profit_input_invoices[
                ~(
                    (st.session_state.profit_input_invoices['period'] == selected_period)
                    & (st.session_state.profit_input_invoices['project'].astype(str).str.strip() == selected_project)
                )
            ]
            st.session_state.profit_input_invoices = pd.concat([df_other_input, edited_input], ignore_index=True)
            st.rerun()

        # Lọc cho mục đích tính toán
        edited_input_clean = edited_input[is_profit_row_valid(edited_input)].copy()
        
        # Tính tổng chi dự án đang nhập (chỉ tính dòng có nội dung)
        total_expense_project = edited_input_clean['amount'].sum()

        # Tính tổng tháng = tổng của tất cả dự án trong tháng
        df_month_output = st.session_state.profit_output_invoices.copy()
        df_month_input = st.session_state.profit_input_invoices.copy()
        df_month_output = df_month_output[df_month_output['period'] == selected_period] if not df_month_output.empty else pd.DataFrame(columns=['period', 'project', 'invoice_no', 'description', 'amount'])
        df_month_input = df_month_input[df_month_input['period'] == selected_period] if not df_month_input.empty else pd.DataFrame(columns=['period', 'project', 'invoice_no', 'description', 'amount'])
        df_month_output_valid = df_month_output[is_profit_row_valid(df_month_output)] if not df_month_output.empty else pd.DataFrame(columns=['period', 'project', 'invoice_no', 'description', 'amount'])
        df_month_input_valid = df_month_input[is_profit_row_valid(df_month_input)] if not df_month_input.empty else pd.DataFrame(columns=['period', 'project', 'invoice_no', 'description', 'amount'])

        total_revenue_month = df_month_output_valid['amount'].sum() if not df_month_output_valid.empty else 0
        total_expense_month = df_month_input_valid['amount'].sum() if not df_month_input_valid.empty else 0
        
        col_refresh2, col_metric2 = st.columns([1, 5])
        with col_refresh2:
            if st.button("🔄", key="refresh_input", help="Làm mới để định dạng số tiền"):
                st.rerun()
        with col_metric2:
            st.metric("💸 Tổng Chi (Dự án đang chọn)", format_vnd(total_expense_project) + " VND")

        # === CẢNH BÁO TRÙNG DỮ LIỆU TRONG DỰ ÁN ===
        dup_out_inv = (
            edited_output_clean[
                edited_output_clean['invoice_no'].astype(str).str.strip().ne("")
            ]['invoice_no']
            .astype(str)
            .str.strip()
            .value_counts()
        )
        dup_out_inv = dup_out_inv[dup_out_inv > 1]

        dup_in_inv = (
            edited_input_clean[
                edited_input_clean['invoice_no'].astype(str).str.strip().ne("")
            ]['invoice_no']
            .astype(str)
            .str.strip()
            .value_counts()
        )
        dup_in_inv = dup_in_inv[dup_in_inv > 1]

        out_amount_series = pd.to_numeric(edited_output_clean['amount'], errors='coerce').fillna(0)
        dup_out_amt = out_amount_series[out_amount_series.ne(0)].value_counts()
        dup_out_amt = dup_out_amt[dup_out_amt > 1]

        in_amount_series = pd.to_numeric(edited_input_clean['amount'], errors='coerce').fillna(0)
        dup_in_amt = in_amount_series[in_amount_series.ne(0)].value_counts()
        dup_in_amt = dup_in_amt[dup_in_amt > 1]

        out_inv_set = set(
            edited_output_clean[edited_output_clean['invoice_no'].astype(str).str.strip().ne("")]['invoice_no']
            .astype(str)
            .str.strip()
            .tolist()
        )
        in_inv_set = set(
            edited_input_clean[edited_input_clean['invoice_no'].astype(str).str.strip().ne("")]['invoice_no']
            .astype(str)
            .str.strip()
            .tolist()
        )
        dup_cross_inv = sorted(out_inv_set.intersection(in_inv_set))

        if (not dup_out_inv.empty) or (not dup_in_inv.empty) or (not dup_out_amt.empty) or (not dup_in_amt.empty) or dup_cross_inv:
            st.markdown("#### ⚠️ Cảnh báo trùng dữ liệu trong dự án")

            if not dup_out_inv.empty:
                vals = ", ".join([f"{k} ({int(v)} lần)" for k, v in dup_out_inv.head(8).items()])
                st.warning(f"[Thu] Trùng Số HĐ: {vals}")

            if not dup_in_inv.empty:
                vals = ", ".join([f"{k} ({int(v)} lần)" for k, v in dup_in_inv.head(8).items()])
                st.warning(f"[Chi] Trùng Số HĐ: {vals}")

            if not dup_out_amt.empty:
                vals = ", ".join([f"{format_vnd(clean_profit_amount(k))} ({int(v)} lần)" for k, v in dup_out_amt.head(8).items()])
                st.warning(f"[Thu] Trùng Số tiền: {vals}")

            if not dup_in_amt.empty:
                vals = ", ".join([f"{format_vnd(clean_profit_amount(k))} ({int(v)} lần)" for k, v in dup_in_amt.head(8).items()])
                st.warning(f"[Chi] Trùng Số tiền: {vals}")

            if dup_cross_inv:
                vals = ", ".join(dup_cross_inv[:10])
                st.warning(f"Số HĐ xuất hiện ở cả Thu và Chi trong cùng dự án: {vals}")
            st.caption("Hệ thống chỉ cảnh báo để bạn rà soát, không tự động xóa dữ liệu.")
        
        # === BÁO CÁO DỰ ÁN TRONG THÁNG ===
        st.markdown("---")
        st.markdown("#### 🧩 Báo Cáo Dự Án Trong Tháng")

        def build_project_amount_table(df_src: pd.DataFrame, amount_col_name: str) -> pd.DataFrame:
            if df_src.empty:
                return pd.DataFrame(columns=['project', amount_col_name])
            projects = sorted(df_src['project'].astype(str).fillna('Tổng tháng').unique())
            rows = []
            for proj in projects:
                proj_mask = df_src['project'].astype(str) == proj
                proj_values = pd.Series(pd.to_numeric(df_src.loc[proj_mask, 'amount'], errors='coerce'))
                proj_amt = proj_values.fillna(0).sum()
                rows.append({'project': proj, amount_col_name: proj_amt})
            return pd.DataFrame(rows)

        out_by_project = build_project_amount_table(df_month_output_valid, 'Tổng Thu')
        in_by_project = build_project_amount_table(df_month_input_valid, 'Tổng Chi')

        if not out_by_project.empty or not in_by_project.empty:
            df_project_summary = pd.merge(out_by_project, in_by_project, on='project', how='outer').fillna(0)
            df_project_summary['Lợi Nhuận'] = df_project_summary['Tổng Thu'] - df_project_summary['Tổng Chi']
            df_project_summary = df_project_summary.sort_values('Lợi Nhuận', ascending=False).reset_index(drop=True)

            df_project_show = df_project_summary.copy()
            df_project_show = df_project_show.rename(columns={'project': 'Dự án'})
            df_project_show['Tổng Thu'] = df_project_show['Tổng Thu'].apply(lambda x: format_vnd(x) + " VND")
            df_project_show['Tổng Chi'] = df_project_show['Tổng Chi'].apply(lambda x: format_vnd(x) + " VND")
            df_project_show['Lợi Nhuận'] = df_project_show['Lợi Nhuận'].apply(lambda x: format_vnd(x) + " VND")
            st.dataframe(df_project_show, use_container_width=True, hide_index=True)
        else:
            df_project_summary = pd.DataFrame(columns=['project', 'Tổng Thu', 'Tổng Chi', 'Lợi Nhuận'])
            st.info("Chưa có dữ liệu dự án trong tháng này.")

        # === TÍNH LỢI NHUẬN ===
        st.markdown("---")
        st.markdown(f"### 📊 Kết Quả - Tháng {selected_period}")

        profit_amount = total_revenue_month - total_expense_month

        c1, c2, c3 = st.columns(3)
        c1.metric("💰 Tổng Thu tháng", format_vnd(total_revenue_month) + " VND")
        c2.metric("💸 Tổng Chi tháng", format_vnd(total_expense_month) + " VND")
        c3.metric("📈 Lợi Nhuận", format_vnd(profit_amount) + " VND", 
                 delta="Lãi" if profit_amount >= 0 else "Lỗ")
        
        # === CHI TIẾT ===
        st.markdown("---")
        st.markdown("#### 📋 Chi Tiết Hóa Đơn")
        
        col_detail1, col_detail2 = st.columns(2)
        
        with col_detail1:
            st.markdown("**📤 Đầu Ra (Thu):**")
            detail_output = df_month_output_valid.copy()
            if not detail_output.empty:
                for _, row in detail_output.iterrows():
                    inv = str(row.get('invoice_no', '') or '').strip()
                    desc = str(row.get('description', '') or '').strip()
                    inv_text = f"{inv} - " if inv else ""
                    st.write(f"• [{row['project']}] {inv_text}{desc}: {format_vnd(row['amount'])} VND")
            else:
                st.info("Chưa có hóa đơn đầu ra")
        
        with col_detail2:
            st.markdown("**📥 Đầu Vào (Chi):**")
            detail_input = df_month_input_valid.copy()
            if not detail_input.empty:
                for _, row in detail_input.iterrows():
                    inv = str(row.get('invoice_no', '') or '').strip()
                    desc = str(row.get('description', '') or '').strip()
                    inv_text = f"{inv} - " if inv else ""
                    st.write(f"• [{row['project']}] {inv_text}{desc}: {format_vnd(row['amount'])} VND")
            else:
                st.info("Chưa có hóa đơn đầu vào")
        
        # === TỔNG HỢP TẤT CẢ CÁC THÁNG ===
        df_all_output_full = st.session_state.profit_output_invoices
        df_all_input_full = st.session_state.profit_input_invoices
        
        # Lọc bỏ dòng rỗng (mô tả rỗng và số tiền = 0)
        df_all_output_valid = df_all_output_full[is_profit_row_valid(df_all_output_full)] if not df_all_output_full.empty else pd.DataFrame()
        df_all_input_valid = df_all_input_full[is_profit_row_valid(df_all_input_full)] if not df_all_input_full.empty else pd.DataFrame()
        
        if not df_all_output_valid.empty or not df_all_input_valid.empty:
            st.markdown("---")
            st.markdown("#### 📈 Tổng Hợp Tất Cả Các Tháng")
            
            # Lấy danh sách tất cả các period
            all_periods_set = set()
            if not df_all_output_valid.empty:
                all_periods_set.update(df_all_output_valid['period'].unique())
            if not df_all_input_valid.empty:
                all_periods_set.update(df_all_input_valid['period'].unique())
            
            # Tạo summary table
            summary_rows = []
            for p in sorted(all_periods_set, reverse=True):
                p_revenue = df_all_output_valid[df_all_output_valid['period'] == p]['amount'].sum() if not df_all_output_valid.empty else 0
                p_expense = df_all_input_valid[df_all_input_valid['period'] == p]['amount'].sum() if not df_all_input_valid.empty else 0
                p_profit = p_revenue - p_expense
                
                summary_rows.append({
                    "Tháng/Năm": p,
                    "Tổng Thu": format_vnd(p_revenue) + " VND",
                    "Tổng Chi": format_vnd(p_expense) + " VND",
                    "Lợi Nhuận": format_vnd(p_profit) + " VND"
                })
            
            if summary_rows:
                df_summary = pd.DataFrame(summary_rows)
                st.dataframe(df_summary, use_container_width=True, hide_index=True)
        
        # === NÚT THAO TÁC ===
        st.markdown("---")
        st.markdown("### 🛠️ Thao Tác")

        # Ghi chú khi xuất Excel
        if "profit_note_thu" not in st.session_state:
            st.session_state.profit_note_thu = "Ghi chú Thu"
        if "profit_note_chi" not in st.session_state:
            st.session_state.profit_note_chi = "Ghi chú Chi"

        c_note_thu, c_note_chi = st.columns([1, 1])
        with c_note_thu:
            note_thu = st.text_input(
                "📝 Ghi chú Thu",
                key="profit_note_thu",
                disabled=not st.session_state.profit_meta_edit_mode
            )
        with c_note_chi:
            note_chi = st.text_input(
                "📝 Ghi chú Chi",
                key="profit_note_chi",
                disabled=not st.session_state.profit_meta_edit_mode
            )

        col_meta_btn1, col_meta_btn2 = st.columns(2)
        with col_meta_btn1:
            if st.session_state.profit_meta_edit_mode:
                if st.button("💾 Lưu thông tin báo cáo", use_container_width=True, key="profit_meta_save_btn"):
                    st.session_state.profit_meta_edit_mode = False
                    st.success("✅ Đã lưu thông tin báo cáo")
                    st.rerun()
        with col_meta_btn2:
            if not st.session_state.profit_meta_edit_mode:
                if st.button("✏️ Chỉnh sửa thông tin báo cáo", use_container_width=True, key="profit_meta_edit_btn"):
                    st.session_state.profit_meta_edit_mode = True
                    st.rerun()
        
        col_btn1, col_btn2 = st.columns(2)
        
        with col_btn1:
            if st.button("📥 Xuất Excel", use_container_width=True, type="primary"):
                try:
                    # Chuẩn bị 2 bảng riêng: Thu và Chi
                    df_exp_out = df_month_output_valid[['period', 'project', 'invoice_no', 'description', 'amount']].copy() if not df_month_output_valid.empty else pd.DataFrame(columns=['period', 'project', 'invoice_no', 'description', 'amount'])
                    df_exp_in = df_month_input_valid[['period', 'project', 'invoice_no', 'description', 'amount']].copy() if not df_month_input_valid.empty else pd.DataFrame(columns=['period', 'project', 'invoice_no', 'description', 'amount'])

                    if not df_exp_out.empty or not df_exp_in.empty:
                        buffer_exp = io.BytesIO()
                        with pd.ExcelWriter(buffer_exp, engine='xlsxwriter') as writer:
                            wb = writer.book
                            ws = wb.add_worksheet('Lợi Nhuận')

                            # Formats chung
                            title_fmt = wb.add_format({'bold': True, 'font_size': 14, 'align': 'center', 'font_name': 'Times New Roman'})
                            text_fmt = wb.add_format({'border': 1, 'font_name': 'Times New Roman'})
                            num_fmt = wb.add_format({'border': 1, 'num_format': '#,##0', 'font_name': 'Times New Roman'})
                            head_thu_fmt = wb.add_format({'bold': True, 'border': 1, 'align': 'center', 'bg_color': '#C8E6C9', 'font_name': 'Times New Roman'})
                            head_chi_fmt = wb.add_format({'bold': True, 'border': 1, 'align': 'center', 'bg_color': '#FFCDD2', 'font_name': 'Times New Roman'})
                            section_thu_fmt = wb.add_format({'bold': True, 'border': 1, 'align': 'left', 'bg_color': '#E8F5E9', 'font_name': 'Times New Roman'})
                            section_chi_fmt = wb.add_format({'bold': True, 'border': 1, 'align': 'left', 'bg_color': '#FFEBEE', 'font_name': 'Times New Roman'})
                            note_thu_fmt = wb.add_format({'italic': True, 'border': 1, 'font_color': '#1B5E20', 'bg_color': '#F1F8E9', 'font_name': 'Times New Roman'})
                            note_chi_fmt = wb.add_format({'italic': True, 'border': 1, 'font_color': '#B71C1C', 'bg_color': '#FFEBEE', 'font_name': 'Times New Roman'})

                            report_title = str(report_name).strip() if report_name else "Báo cáo lợi nhuận"

                            # Title
                            ws.merge_range('A1:F1', f'{report_title} - {selected_period}', title_fmt)
                            ws.write('A2', f'Tổng Thu: {format_vnd(total_revenue_month)} VND', text_fmt)
                            ws.write('B2', f'Tổng Chi: {format_vnd(total_expense_month)} VND', text_fmt)
                            ws.write('C2', f'Lợi Nhuận: {format_vnd(profit_amount)} VND', section_thu_fmt)

                            # Bảng Thu
                            current_row = 3
                            ws.merge_range(current_row, 0, current_row, 5, "I. DANH SÁCH THU", section_thu_fmt)
                            current_row += 1
                            out_cols = ['Tháng', 'Dự án', 'Loại', 'Số HĐ', 'Diễn giải', 'Số tiền']
                            for c_idx, c_name in enumerate(out_cols):
                                ws.write(current_row, c_idx, c_name, head_thu_fmt)
                            current_row += 1

                            for _, row in df_exp_out.iterrows():
                                ws.write(current_row, 0, str(row['period']), text_fmt)
                                ws.write(current_row, 1, str(row['project']), text_fmt)
                                ws.write(current_row, 2, 'Thu', text_fmt)
                                ws.write(current_row, 3, str(row['invoice_no']), text_fmt)
                                ws.write(current_row, 4, str(row['description']), text_fmt)
                                ws.write_number(current_row, 5, float(row['amount']) if row['amount'] else 0.0, num_fmt)
                                current_row += 1

                            ws.merge_range(current_row, 0, current_row, 5, f"Ghi chú Thu: {note_thu}", note_thu_fmt)
                            current_row += 2

                            # Bảng Chi
                            ws.merge_range(current_row, 0, current_row, 5, "II. DANH SÁCH CHI", section_chi_fmt)
                            current_row += 1
                            in_cols = ['Tháng', 'Dự án', 'Loại', 'Số HĐ', 'Diễn giải', 'Số tiền']
                            for c_idx, c_name in enumerate(in_cols):
                                ws.write(current_row, c_idx, c_name, head_chi_fmt)
                            current_row += 1

                            for _, row in df_exp_in.iterrows():
                                ws.write(current_row, 0, str(row['period']), text_fmt)
                                ws.write(current_row, 1, str(row['project']), text_fmt)
                                ws.write(current_row, 2, 'Chi', text_fmt)
                                ws.write(current_row, 3, str(row['invoice_no']), text_fmt)
                                ws.write(current_row, 4, str(row['description']), text_fmt)
                                ws.write_number(current_row, 5, float(row['amount']) if row['amount'] else 0.0, num_fmt)
                                current_row += 1

                            ws.merge_range(current_row, 0, current_row, 5, f"Ghi chú Chi: {note_chi}", note_chi_fmt)
                            current_row += 2

                            # Bảng tổng hợp dự án trong tháng
                            if not df_project_summary.empty:
                                ws.merge_range(current_row, 0, current_row, 5, "III. TỔNG HỢP DỰ ÁN TRONG THÁNG", section_thu_fmt)
                                current_row += 1
                                proj_cols = ['Dự án', 'Tổng Thu', 'Tổng Chi', 'Lợi Nhuận']
                                for c_idx, c_name in enumerate(proj_cols):
                                    ws.write(current_row, c_idx, c_name, head_thu_fmt)
                                current_row += 1

                                for _, row in df_project_summary.iterrows():
                                    ws.write(current_row, 0, str(row['project']), text_fmt)
                                    ws.write_number(current_row, 1, float(row['Tổng Thu']), num_fmt)
                                    ws.write_number(current_row, 2, float(row['Tổng Chi']), num_fmt)
                                    ws.write_number(current_row, 3, float(row['Lợi Nhuận']), num_fmt)
                                    current_row += 1
                            
                            ws.set_column('A:A', 12)
                            ws.set_column('B:B', 24)
                            ws.set_column('C:C', 10)
                            ws.set_column('D:D', 16)
                            ws.set_column('E:E', 30)
                            ws.set_column('F:F', 18)

                            # Đăng ký sheet với writer để tránh warning
                            writer.sheets['Lợi Nhuận'] = ws

                        safe_name = re.sub(r'[\\/*?:"<>|]', "", (str(report_name).strip() if report_name else "BaoCao_LoiNhuan"))
                        if not safe_name:
                            safe_name = "BaoCao_LoiNhuan"
                        
                        st.download_button(
                            "⬇️ Tải file Excel",
                            data=buffer_exp.getvalue(),
                            file_name=f"{safe_name}_{selected_period.replace('/', '_')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                        st.success("✅ File Excel sẵn sàng!")
                    else:
                        st.warning("⚠️ Không có dữ liệu để xuất!")
                except Exception as e:
                    st.error(f"❌ Lỗi: {str(e)}")
        
        with col_btn2:
            if st.button("🗑️ Xóa dữ liệu tháng này", use_container_width=True, type="secondary"):
                # Xóa data của tháng hiện tại
                df_out_other = df_all_output_full[df_all_output_full['period'] != selected_period]
                df_in_other = df_all_input_full[df_all_input_full['period'] != selected_period]
                
                st.session_state.profit_output_invoices = df_out_other
                st.session_state.profit_input_invoices = df_in_other
                
                st.success(f"✅ Đã xóa dữ liệu tháng {selected_period}!")
                time.sleep(0.5)
                st.rerun()
        
        # === GHI CHÚ ===
        with st.expander("💡 Hướng dẫn sử dụng"):
            st.markdown("""
            **📝 Cách nhập dữ liệu:**
            1. Chọn tháng/năm cần làm việc
            2. Nhấn nút "+" để thêm dòng mới
            3. Nhập dự án/tour/booking, diễn giải và số tiền (VD: 5000000 hoặc 5.000.000)
            4. Dữ liệu tự động lưu khi bạn nhập
            
            **💡 Tính năng:**
            - Dữ liệu được lưu riêng theo từng tháng
            - Tách được nhiều dự án/tour/booking trong cùng 1 tháng
            - Tự động tính lợi nhuận = Thu - Chi
            - Xuất Excel báo cáo chi tiết
            - Xem tổng hợp tất cả các tháng ở phía dưới
            
            **🔢 Công thức:**
            ```
            Lợi Nhuận = Tổng Thu - Tổng Chi
            ```
            """)
    
def render_customer_management():
    st.title("🤝 Quản Lý Khách Hàng")
    
    # Lấy thông tin user hiện tại để lọc
    current_user_info_cust = st.session_state.get("user_info", {})
    current_user_name_cust = current_user_info_cust.get('name', 'N/A')
    current_user_role_cust = current_user_info_cust.get('role')
    
    tab_list, tab_add = st.tabs(["📋 Danh sách khách hàng", "➕ Thêm khách hàng"])
    
    with tab_add:
        with st.form("add_cust_form"):
            st.subheader("Thêm khách hàng mới")
            c1, c2 = st.columns(2)
            name = c1.text_input("Tên khách hàng (*)", placeholder="Nguyễn Văn A")
            phone = c2.text_input("Số điện thoại", placeholder="090...")
            email = c1.text_input("Email", placeholder="abc@gmail.com")
            addr = c2.text_input("Địa chỉ")
            note = st.text_area("Ghi chú")
            
            if st.form_submit_button("Lưu khách hàng", type="primary"):
                if name:
                    add_row_to_table('customers', {
                        'name': name, 'phone': phone, 'email': email, 'address': addr, 'notes': note,
                        'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 'sale_name': current_user_name_cust
                    })
                    st.success("Đã thêm khách hàng mới!"); time.sleep(1); st.rerun()
                else:
                    st.warning("Vui lòng nhập tên khách hàng.")

    with tab_list:
        # Search bar
        search_term = st.text_input("🔍 Tìm kiếm", placeholder="Nhập tên hoặc số điện thoại...")
        
        query = "SELECT * FROM customers"
        params = []

        # Base filter for sales role
        if current_user_role_cust == 'sale':
            query += " WHERE sale_name=?"
            params.append(current_user_name_cust)

        # Additional filter for search term
        if search_term:
            if "WHERE" in query:
                query += " AND (name LIKE ? OR phone LIKE ?)"
            else:
                query += " WHERE name LIKE ? OR phone LIKE ?"
            params.extend([f"%{search_term}%", f"%{search_term}%"])
        query += " ORDER BY id DESC"
        
        customers = run_query(query, tuple(params))
        
        if customers:
            # Display as dataframe for overview
            df_cust = pd.DataFrame([dict(r) for r in customers])
            st.dataframe(
                df_cust[['name', 'phone', 'email', 'address', 'notes']],
                column_config={
                    "name": "Tên khách hàng",
                    "phone": "SĐT",
                    "email": "Email",
                    "address": "Địa chỉ",
                    "notes": "Ghi chú"
                },
                use_container_width=True,
                hide_index=True
            )
            
            st.divider()
            st.markdown("##### 🛠️ Chỉnh sửa thông tin")
            
            cust_options = {f"{c['name']} - {c['phone']}": c['id'] for c in customers} # type: ignore
            selected_cust = st.selectbox("Chọn khách hàng để sửa/xóa:", list(cust_options.keys()))
            
            if selected_cust:
                cid = cust_options[selected_cust] # type: ignore
                c_info = next((c for c in customers if c['id'] == cid), None)
                
                if c_info:
                    with st.form(f"edit_cust_{cid}"):
                        c1, c2 = st.columns(2)
                        n_name = c1.text_input("Tên", value=c_info['name']) # type: ignore
                        n_phone = c2.text_input("SĐT", value=c_info['phone']) # type: ignore
                        n_email = c1.text_input("Email", value=c_info['email']) # type: ignore
                        n_addr = c2.text_input("Địa chỉ", value=c_info['address']) # type: ignore
                        n_note = st.text_area("Ghi chú", value=c_info['notes']) # type: ignore
                        
                        c_save, c_del = st.columns(2)
                        if c_save.form_submit_button("💾 Cập nhật"):
                            run_query("UPDATE customers SET name=?, phone=?, email=?, address=?, notes=? WHERE id=?", 
                                      (n_name, n_phone, n_email, n_addr, n_note, cid), commit=True)
                            st.success("Đã cập nhật!"); time.sleep(0.5); st.rerun()
                        
                        if c_del.form_submit_button("🗑️ Xóa khách hàng"):
                            run_query("DELETE FROM customers WHERE id=?", (cid,), commit=True)
                            st.success("Đã xóa!"); time.sleep(0.5); st.rerun()
        else:
            st.info("Chưa có khách hàng nào.")

def render_hr_management():
    st.title("👥 Quản Lý Nhân Sự & Tài Khoản")
    
    if (st.session_state.user_info or {}).get('role') not in ['admin', 'admin_f1']:
        st.warning("⛔ Khu vực này chỉ dành cho Admin hoặc Admin F1. Vui lòng liên hệ quản trị viên.")
    else:
        tab_list, tab_req = st.tabs(["📋 Danh sách tài khoản", "📝 Duyệt đăng ký mới"])
        
        with tab_list:
            st.subheader("Danh sách tài khoản hệ thống")
            
            # Lấy dữ liệu users
            users = run_query("SELECT id, username, role, status FROM users ORDER BY id ASC")
            if users:
                df_users = pd.DataFrame([dict(r) for r in users])
                original_df = df_users.copy()
                
                # Xác định các quyền có thể gán
                role_options = ["admin", "admin_f1", "user", "sale", "accountant"]
                if (st.session_state.user_info or {}).get('role') == 'admin_f1':
                    role_options = ["admin_f1", "user", "sale", "accountant"] # Admin F1 không thể tạo admin chính

                # Hiển thị bảng
                edited_df = st.data_editor(
                    df_users,
                    column_config={
                        "id": st.column_config.NumberColumn("ID", width="small", disabled=True),
                        "username": st.column_config.TextColumn("Tên đăng nhập", width="medium", disabled=True),
                        "role": st.column_config.SelectboxColumn("Quyền hạn", options=role_options, required=True, width="medium"),
                        "status": st.column_config.SelectboxColumn("Trạng thái", options=["approved", "pending", "blocked"], required=True, width="medium")
                    },
                    use_container_width=True,
                    hide_index=True
                )
                
                if st.button("💾 Lưu thay đổi phân quyền", type="primary"):
                    if not original_df.equals(edited_df):
                        with st.spinner("Đang cập nhật..."):
                            current_user_role = (st.session_state.user_info or {}).get('role')
                            # Iterate through the edited dataframe
                            for index, row in edited_df.iterrows():
                                original_row = original_df.loc[index]# type: ignore
                                # Check if the row has changed
                                if not row.equals(original_row):
                                    user_id = row['id'] # type: ignore
                                    username = row['username'] # type: ignore
                                    new_role = row['role'] # type: ignore
                                    new_status = row['status'] # type: ignore
                                    original_role = original_row['role'] # type: ignore

                                    # Prevent changing the main admin
                                    if username == 'admin':
                                        st.warning("Không thể thay đổi quyền của tài khoản 'admin' chính.")
                                        continue
                                    
                                    # Prevent F1 from editing a full admin
                                    if current_user_role == 'admin_f1' and original_role == 'admin':
                                        st.warning(f"Bạn không có quyền chỉnh sửa tài khoản admin '{username}'.")
                                        continue
                                    
                                    run_query(
                                        "UPDATE users SET role=?, status=? WHERE id=?",
                                        (new_role, new_status, user_id),
                                        commit=True
                                    )
                        st.success("Đã cập nhật thành công!")
                        time.sleep(1); st.rerun()
                    else:
                        st.toast("Không có thay đổi nào.")
                
                st.divider()
                st.markdown("##### 🗑️ Xóa tài khoản")
                # Loại bỏ admin chính ra khỏi danh sách xóa để tránh lỗi
                del_options = [u['username'] for u in users if u['username'] != 'admin'] # type: ignore
                user_to_del = st.selectbox("Chọn tài khoản cần xóa:", del_options, key="sel_del_u")
                
                if st.button("Xác nhận xóa tài khoản", type="primary", key="btn_del_u"):
                    if user_to_del:
                        # Kiểm tra quyền trước khi xóa
                        user_to_del_info = run_query("SELECT role FROM users WHERE username=?", (user_to_del,), fetch_one=True)
                        current_user_role = (st.session_state.user_info or {}).get('role')

                        if current_user_role == 'admin_f1' and user_to_del_info and user_to_del_info['role'] == 'admin': # type: ignore
                            st.error(f"Bạn không có quyền xóa tài khoản admin '{user_to_del}'.")
                        else:
                            run_query("DELETE FROM users WHERE username=?", (user_to_del,), commit=True)
                            st.success(f"Đã xóa tài khoản {user_to_del}!")
                            time.sleep(1); st.rerun()
            else:
                st.info("Chưa có tài khoản nào.")

        with tab_req:
            st.subheader("Yêu cầu đăng ký chờ duyệt")
            pending = run_query("SELECT * FROM users WHERE status='pending'")
            if pending:
                for p in pending:
                    with st.container(border=True):
                        c1, c2, c3 = st.columns([2, 1, 1])
                        c1.write(f"User: **{p['username']}**") # type: ignore
                        if c2.button("✔ Duyệt", key=f"hr_app_{p['id']}", use_container_width=True): # type: ignore
                            run_query("UPDATE users SET status='approved' WHERE id=?", (p['id'],), commit=True) # type: ignore
                            st.success("Đã duyệt!"); time.sleep(0.5); st.rerun()
                        if c3.button("✖ Xóa", key=f"hr_del_{p['id']}", use_container_width=True): # type: ignore
                            run_query("DELETE FROM users WHERE id=?", (p['id'],), commit=True) # type: ignore
                            st.success("Đã xóa!"); time.sleep(0.5); st.rerun()
            else:
                st.info("Hiện không có yêu cầu nào.")

def render_search_module():
    st.title("🔍 Tra cứu thông tin hệ thống")
    
    # Lấy thông tin user hiện tại để lọc
    current_user_info = st.session_state.get("user_info", {})
    current_user_name = current_user_info.get('name', 'N/A')
    current_user_role = current_user_info.get('role')

    query = st.text_input("Nhập từ khóa tìm kiếm", placeholder="Nhập Mã Tour, Số Hóa Đơn, Mã Vé, Mã Chi Phí, hoặc Tên Khách...", help="Hệ thống sẽ tìm trong Tour, Hóa đơn, UNC và Vé máy bay")
        
    if query:
        st.divider()
        term = f"%{query.strip()}%"
        found_any = False
        
        # 1. TÌM TRONG TOUR
        tour_sql = "SELECT * FROM tours WHERE (tour_code LIKE ? OR tour_name LIKE ?)"
        tour_params = [term, term]
        if current_user_role == 'sale':
            tour_sql += " AND sale_name=?"
            tour_params.append(current_user_name)
            
        tours = run_query(tour_sql, tuple(tour_params))
        if tours:
            found_any = True
            st.subheader(f"📦 Tìm thấy {len(tours)} Tour")
            for t in tours:
                with st.expander(f"Tour: {t['tour_name']} (Mã: {t['tour_code']})", expanded=True):
                    c1, c2, c3 = st.columns(3) # type: ignore
                    c1.write(f"**Sales:** {t['sale_name']}") # type: ignore
                    c2.write(f"**Ngày:** {t['start_date']} - {t['end_date']}") # type: ignore
                    c3.write(f"**Khách:** {t['guest_count']}") # type: ignore
                    
                    est = run_query("SELECT SUM(total_amount) as t FROM tour_items WHERE tour_id=? AND item_type='EST'", (t['id'],), fetch_one=True) # type: ignore
                    act = run_query("SELECT SUM(total_amount) as t FROM tour_items WHERE tour_id=? AND item_type='ACT'", (t['id'],), fetch_one=True) # type: ignore
                    est_val = est['t'] if isinstance(est, sqlite3.Row) and est['t'] else 0 # type: ignore
                    act_val = act['t'] if isinstance(act, sqlite3.Row) and act['t'] else 0 # type: ignore
                    
                    st.info(f"💰 Dự toán: {format_vnd(est_val)} | 💸 Quyết toán: {format_vnd(act_val)}")

        # 2. TÌM TRONG KHÁCH HÀNG (MỚI)
        cust_sql = "SELECT * FROM customers WHERE (name LIKE ? OR phone LIKE ?)"
        cust_params = [term, term]
        if current_user_role == 'sale':
            cust_sql += " AND sale_name=?"
            cust_params.append(current_user_name)
            
        custs = run_query(cust_sql, tuple(cust_params))
        if custs:
            found_any = True
            st.subheader(f"👥 Tìm thấy {len(custs)} Khách hàng")
            for c in custs:
                with st.expander(f"Khách hàng: {c['name']} - {c['phone']}", expanded=True):
                    st.write(f"**Email:** {c['email']}")
                    st.write(f"**Địa chỉ:** {c['address']}")
                    st.write(f"**Ghi chú:** {c['notes']}")

        # 3. TÌM TRONG HÓA ĐƠN / UNC
        invs = run_query("SELECT * FROM invoices WHERE invoice_number LIKE ? OR cost_code LIKE ? OR memo LIKE ? ORDER BY date DESC", (term, term, term))
        if invs:
            found_any = True
            st.subheader(f"💰 Tìm thấy {len(invs)} Hóa đơn / UNC")
            
            for inv in invs:
                icon = "💸" if "UNC" in (inv['invoice_number'] or "") else "📄"
                i_num = inv['invoice_number'] if inv['invoice_number'] else "(Không số)" # type: ignore
                label = f"{icon} {inv['date']} | {i_num} | {format_vnd(inv['total_amount'])} | {inv['memo']}" # type: ignore
                
                with st.expander(label):
                    c_info, c_file = st.columns([1, 1])
                    with c_info:
                        st.markdown(f"**Bên bán:** {inv['seller_name']}") # type: ignore
                        st.markdown(f"**Bên mua:** {inv['buyer_name']}") # type: ignore
                        st.markdown(f"**Tổng tiền:** {format_vnd(inv['total_amount'])}") # type: ignore
                        st.markdown(f"**Mã chi phí:** `{inv['cost_code']}`") # type: ignore
                        st.caption(f"Trạng thái: {inv['status']}") # type: ignore
                    
                    with c_file:
                        file_path = inv['file_path'] # type: ignore
                        if file_path and os.path.exists(file_path):
                            # The 'file_path' from the database is a Google Drive link, not a local path.
                            # The original code to check os.path.exists(file_path) and open it is incorrect.
                            # We should just provide the link.
                            st.link_button("🔗 Mở file trên Google Drive", file_path, use_container_width=True)

        if not found_any:
            st.warning("📭 Không tìm thấy dữ liệu nào phù hợp.")

def main():
    if not st.session_state.logged_in:
        render_login_page(comp)
        return

    module, menu = render_sidebar(comp)

    # --- HEADER CHÍNH ---
    l_html = f'<img src="data:image/png;base64,{comp["logo_b64_str"]}" class="company-logo-img">' if comp['logo_b64_str'] else ''
    st.markdown(f'''
    <div class="company-header-container">
        {l_html}
        <div class="company-info-text">
            <h1>{comp['name']}</h1>
            <p>📍 {comp['address']}</p>
            <p>MST: {comp['phone']}</p>
        </div>
    </div>
    ''', unsafe_allow_html=True)

    if module == "🏠 Trang Chủ":
        render_dashboard()
    elif module == "📅 Lịch Thông Báo":
        render_notification_calendar()
    elif module == "🔖 Quản Lý Booking":
        render_booking_management()
    elif module == "💰 Kiểm Soát Chi Phí":
        render_cost_control(menu)
    elif module == "💳 Quản Lý Công Nợ":
        render_debt_management()
    elif module == "📦 Quản Lý Tour ":
        render_tour_management()
    elif module == "🧾 Quản Lý Hóa Đơn":
        render_invoice_management()
    elif module == "🤝 Quản Lý Khách Hàng":
        render_customer_management()
    elif module == "👥 Quản Lý Nhân Sự":
        render_hr_management()
    elif module == "🔍 Tra cứu thông tin":
        render_search_module()

if __name__ == "__main__":
    main()
